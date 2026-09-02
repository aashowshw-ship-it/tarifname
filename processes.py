from __future__ import annotations

import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from email import policy
from email.parser import BytesParser
from html import unescape
import xml.etree.ElementTree as ET
import unicodedata
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader, PdfWriter
from PIL import Image, ImageOps

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = f"{{{_W_NS}}}"
ET.register_namespace("w", _W_NS)


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip()).upper()


def _block_text(block: ET.Element) -> str:
    parts: list[str] = []
    for node in block.iter():
        if node.tag == _W + "t" and node.text:
            parts.append(node.text)
        elif node.tag in {_W + "tab"}:
            parts.append("\t")
        elif node.tag in {_W + "br", _W + "cr"}:
            parts.append("\n")
    return "".join(parts).strip()


def _find_heading_index(body_children: list[ET.Element], heading: str) -> int:
    target = _norm(heading)
    for idx, child in enumerate(body_children):
        if child.tag == _W + "sectPr":
            continue
        if _norm(_block_text(child)) == target:
            return idx
    raise ValueError(f"Tarifname Word dosyasında '{heading}' başlığı bulunamadı.")


def _slice_docx(data: bytes, *, start_idx: int, end_idx: int | None) -> bytes:
    """DOCX paket yapısını koruyarak document.xml gövdesini seçilen aralığa indirger."""
    src = io.BytesIO(data)
    out = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin:
        document_xml = zin.read("word/document.xml")
        root = ET.fromstring(document_xml)
        body = root.find(_W + "body")
        if body is None:
            raise ValueError("Word belgesinin gövdesi okunamadı.")
        children = list(body)
        sect_pr = next((c for c in children if c.tag == _W + "sectPr"), None)
        content = [c for c in children if c.tag != _W + "sectPr"]
        selected = content[start_idx:end_idx]
        if not selected:
            raise ValueError("Word belgesinde seçilen bölüm boş.")
        for child in list(body):
            body.remove(child)
        for child in selected:
            body.append(child)
        if sect_pr is not None:
            body.append(sect_pr)
        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

        with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                payload = new_xml if item.filename == "word/document.xml" else zin.read(item.filename)
                zout.writestr(item, payload)
    return out.getvalue()



_TEMPLATE_COLORS = {"FF0000", "0000FF"}


def _run_color(run: ET.Element) -> str:
    rpr = run.find(_W + "rPr")
    if rpr is None:
        return ""
    color = rpr.find(_W + "color")
    if color is None:
        return ""
    return (color.attrib.get(_W + "val") or "").upper()


def _paragraph_has_visible_content(paragraph: ET.Element) -> bool:
    for node in paragraph.iter():
        local = node.tag.rsplit("}", 1)[-1]
        if local == "t" and (node.text or "").strip():
            return True
        if local in {"drawing", "object", "pict"}:
            return True
    return False


def strip_template_colored_text(data: bytes, *, colors: set[str] | None = None) -> bytes:
    """Kırmızı/mavi şablon açıklamalarını DOCX'ten OOXML düzeyinde kaldırır.

    Siyah teknik içerik, başlıklar, numaralandırma ve belge paket yapısı korunur.
    """
    target = {c.upper() for c in (colors or _TEMPLATE_COLORS)}
    src = io.BytesIO(data)
    out = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin:
        document_xml = zin.read("word/document.xml")
        root = ET.fromstring(document_xml)
        parent_map = {child: parent for parent in root.iter() for child in parent}
        touched_paragraphs: set[ET.Element] = set()
        for run in list(root.iter(_W + "r")):
            if _run_color(run) not in target:
                continue
            parent = parent_map.get(run)
            if parent is None:
                continue
            paragraph = next((a for a in parent_map_chain(run, parent_map) if a.tag == _W + "p"), None)
            if paragraph is not None:
                touched_paragraphs.add(paragraph)
            parent.remove(run)

        # Yalnız renkli şablon içeriğinden ibaret kalan paragrafları kaldır; böylece
        # PDF sayfa sayımı gerçek başvuru metnine göre yapılır.
        for paragraph in list(touched_paragraphs):
            if _paragraph_has_visible_content(paragraph):
                continue
            parent = parent_map.get(paragraph)
            if parent is not None:
                try:
                    parent.remove(paragraph)
                except ValueError:
                    pass

        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                payload = new_xml if item.filename == "word/document.xml" else zin.read(item.filename)
                zout.writestr(item, payload)
    return out.getvalue()


def parent_map_chain(node: ET.Element, parent_map: dict[ET.Element, ET.Element]):
    current = node
    while current in parent_map:
        current = parent_map[current]
        yield current


def remove_word_line_numbering(data: bytes) -> bytes:
    """EPATS PDF'lerinde sol kenarda görünen Word otomatik satır numaralarını kaldırır.

    Yalnız section property içindeki w:lnNumType düğümünü siler. Font, stil,
    paragraf, numaralandırma, header/footer, margin ve teknik metne dokunmaz.
    """
    src = io.BytesIO(data)
    out = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin:
        document_xml = zin.read("word/document.xml")
        root = ET.fromstring(document_xml)
        changed = False
        for sect_pr in root.iter(_W + "sectPr"):
            for node in list(sect_pr):
                if node.tag == _W + "lnNumType":
                    sect_pr.remove(node)
                    changed = True
        new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True) if changed else document_xml
        with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                payload = new_xml if item.filename == "word/document.xml" else zin.read(item.filename)
                zout.writestr(item, payload)
    return out.getvalue()



def remove_word_header_page_numbers(data: bytes) -> bytes:
    """EPATS alt belgelerinde üstbilgideki PAGE alanlarını kaldırır.

    Teknik gövdeye, stillere ve diğer üstbilgi içeriğine dokunmaz. Bazı
    LibreOffice sürümlerinde PAGE alanının ``1X / 2X`` gibi artefakt üretmesini
    önlemek içindir.
    """
    src = io.BytesIO(data)
    out = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin:
        replacements: dict[str, bytes] = {}
        for name in zin.namelist():
            if not (name.startswith("word/header") and name.endswith(".xml")):
                continue
            raw = zin.read(name)
            try:
                root = ET.fromstring(raw)
            except Exception:
                continue
            parent_map = {child: parent for parent in root.iter() for child in parent}
            changed = False
            for instr in list(root.iter(_W + "instrText")):
                code = re.sub(r"\s+", " ", (instr.text or "")).strip().upper()
                if not code.startswith("PAGE") or code.startswith("PAGEREF"):
                    continue
                ancestors = list(parent_map_chain(instr, parent_map))
                # Textbox/drawing içindeki sayfa alanında tüm dış drawing run'ını kaldır.
                outer_run = None
                for anc in ancestors:
                    if anc.tag == _W + "r" and any(
                        x.tag.rsplit("}", 1)[-1] in {"drawing", "pict", "AlternateContent"}
                        for x in anc.iter()
                    ):
                        outer_run = anc
                if outer_run is not None:
                    par = parent_map.get(outer_run)
                    if par is not None:
                        try:
                            par.remove(outer_run)
                            changed = True
                            continue
                        except ValueError:
                            pass
                # Düz PAGE field için begin..end run grubunu kaldır.
                paragraph = next((a for a in ancestors if a.tag == _W + "p"), None)
                if paragraph is None:
                    continue
                runs = [c for c in list(paragraph) if c.tag == _W + "r"]
                target = next((i for i, r in enumerate(runs) if instr in list(r.iter(_W + "instrText"))), -1)
                if target < 0:
                    continue
                start = target
                while start > 0 and not any(x.attrib.get(_W + "fldCharType") == "begin" for x in runs[start].iter(_W + "fldChar")):
                    start -= 1
                end = target
                while end < len(runs) - 1 and not any(x.attrib.get(_W + "fldCharType") == "end" for x in runs[end].iter(_W + "fldChar")):
                    end += 1
                for run in runs[start:end + 1]:
                    try:
                        paragraph.remove(run)
                        changed = True
                    except ValueError:
                        pass
            if changed:
                replacements[name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                zout.writestr(item, replacements.get(item.filename, zin.read(item.filename)))
    return out.getvalue()


def _paragraph_numbering(paragraph: ET.Element) -> tuple[str, str] | None:
    ppr = paragraph.find(_W + "pPr")
    if ppr is None:
        return None
    numpr = ppr.find(_W + "numPr")
    if numpr is None:
        return None
    num_id_node = numpr.find(_W + "numId")
    ilvl_node = numpr.find(_W + "ilvl")
    if num_id_node is None:
        return None
    return (
        num_id_node.attrib.get(_W + "val", ""),
        ilvl_node.attrib.get(_W + "val", "0") if ilvl_node is not None else "0",
    )


def count_claims_from_docx(data: bytes) -> int:
    """İSTEMLER bölümündeki gerçek istem sayısını Word liste yapısından belirler."""
    with zipfile.ZipFile(io.BytesIO(data), "r") as zin:
        root = ET.fromstring(zin.read("word/document.xml"))
        numbering_root = ET.fromstring(zin.read("word/numbering.xml")) if "word/numbering.xml" in zin.namelist() else None
    num_to_abstract: dict[str, str] = {}
    if numbering_root is not None:
        for num in numbering_root.findall(_W + "num"):
            num_id = num.attrib.get(_W + "numId", "")
            abs_node = num.find(_W + "abstractNumId")
            if num_id and abs_node is not None:
                num_to_abstract[num_id] = abs_node.attrib.get(_W + "val", "")
    body = root.find(_W + "body")
    if body is None:
        return 0
    content = [c for c in list(body) if c.tag != _W + "sectPr"]
    try:
        claims_idx = _find_heading_index(content, "İSTEMLER")
    except ValueError:
        claims_idx = -1
    claim_nodes = content[claims_idx + 1 :] if claims_idx >= 0 else content

    # Şablonda ilk gerçek numaralı paragraf bağımsız istemin numId'sini verir;
    # alt unsur listeleri farklı numId kullanır. Aynı numId'deki ilvl=0 paragraflar
    # bağımsız + bağımlı istemlerin sayısıdır.
    claim_num_id = ""
    for node in claim_nodes:
        if node.tag != _W + "p" or not _block_text(node).strip():
            continue
        numbering = _paragraph_numbering(node)
        if numbering and numbering[1] == "0":
            claim_num_id = numbering[0]
            break
    if claim_num_id:
        claim_abstract = num_to_abstract.get(claim_num_id, "")
        if claim_abstract:
            count = 0
            for node in claim_nodes:
                if node.tag != _W + "p" or not _block_text(node).strip():
                    continue
                numbering = _paragraph_numbering(node)
                if not numbering or numbering[1] != "0":
                    continue
                if num_to_abstract.get(numbering[0], "") == claim_abstract:
                    count += 1
            if count:
                return count
        count = sum(
            1
            for node in claim_nodes
            if node.tag == _W + "p"
            and _block_text(node).strip()
            and _paragraph_numbering(node) == (claim_num_id, "0")
        )
        if count:
            return count

    # Liste numaralandırması kaybolmuş dosyalar için metinsel geri dönüş.
    text = "\n".join(_block_text(n) for n in claim_nodes if n.tag == _W + "p")
    explicit = re.findall(r"(?mi)^\s*(\d+)\s*[.\-)]+\s+", text)
    if explicit:
        return len(set(explicit))
    dependent = re.findall(r"(?mi)^\s*İstem\s+\d+(?:['’`]?e|['’`]?a)\s+uygun", text)
    if dependent:
        return 1 + len(dependent)
    # Bazı Word/PDF dönüşümlerinde liste numarası metne düşmez ancak her bağımlı
    # istem ayrı paragrafta "İstem X..." diye başlar. Aynı paragrafı iki kez sayma.
    claimish = [
        _block_text(n).strip() for n in claim_nodes
        if n.tag == _W + "p" and _block_text(n).strip()
    ]
    dep_count = sum(1 for t in claimish if re.match(r"(?i)^istem\s+\d+", t))
    if dep_count:
        return 1 + dep_count
    return 0


def pdf_page_count(data: bytes) -> int:
    return len(PdfReader(io.BytesIO(data)).pages)


def epats_document_metrics(
    specification_data: bytes,
    pdfs: dict[str, bytes],
    *,
    specification_name: str = "Tarifname.docx",
) -> dict[str, Any]:
    suffix = Path(specification_name).suffix.lower()
    claims = 0
    if suffix in {".docx", ".doc"}:
        docx = specification_data if suffix == ".docx" else _libreoffice_to_docx(specification_data, specification_name)
        cleaned = strip_template_colored_text(docx)
        split_docs = split_patent_docx(cleaned, clean_template_colors=False)
        claims = count_claims_from_docx(split_docs["Istemler.docx"])
    elif suffix == ".pdf" and pdfs.get("Istemler.pdf"):
        claims = count_claims_from_pdf(pdfs["Istemler.pdf"])
    spec_pages = pdf_page_count(pdfs["Tarifname.pdf"]) if pdfs.get("Tarifname.pdf") else 0
    figures_pages = pdf_page_count(pdfs["Sekiller.pdf"]) if pdfs.get("Sekiller.pdf") else 0
    abstract_ok = bool(pdfs.get("Ozet.pdf")) and pdf_page_count(pdfs["Ozet.pdf"]) > 0
    return {
        "specification_pages": spec_pages,
        "claim_count": claims,
        "abstract_present": abstract_ok,
        "figures_pages": figures_pages,
        "codes": {
            "specification": f"T-{spec_pages}" if spec_pages else "T-?",
            "claims": f"İ-{claims}" if claims else "İ-?",
            "abstract": "Ö" if abstract_ok else "Ö-?",
            "figures": f"Ş-{figures_pages}" if figures_pages else "Ş-",
        },
    }


def _docx_source_text(data: bytes) -> str:
    """DOCX metnini paragraf/tablo sırasını ve tablo hücrelerini koruyarak çıkarır."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    table_index = 0
    for child in doc.element.body.iterchildren():
        local = child.tag.rsplit("}", 1)[-1]
        if local == "p":
            p = Paragraph(child, doc)
            text = p.text.strip()
            if text:
                parts.append(text)
        elif local == "tbl":
            table_index += 1
            table = Table(child, doc)
            parts.append(f"[[TABLO {table_index}]]")
            for row in table.rows:
                vals: list[str] = []
                for cell in row.cells:
                    value = re.sub(r"\s*\n\s*", " / ", cell.text or "").strip()
                    # Birleşik hücrelerde python-docx aynı hücreyi birden çok kez döndürebilir.
                    if value and (not vals or _plain_norm(value) != _plain_norm(vals[-1])):
                        vals.append(value)
                    elif not value and not vals:
                        vals.append("")
                if any(v for v in vals):
                    parts.append("\t".join(vals))
    return "\n".join(parts)


def _html_to_text(value: str) -> str:
    value = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", value or "")
    value = re.sub(r"(?i)<br\s*/?>", "\n", value)
    # E-posta içindeki HTML tablolarında hücre sınırlarını TAB olarak koru.
    value = re.sub(r"(?i)</t[dh]\s*>", "\t", value)
    value = re.sub(r"(?i)</p\s*>", "\n", value)
    value = re.sub(r"(?i)</(?:div|tr|li|h[1-6])\s*>", "\n", value)
    value = re.sub(r"(?s)<[^>]+>", " ", value)
    value = unescape(value)
    value = re.sub(r" +", " ", value)
    value = re.sub(r" *\t *", "\t", value)
    value = re.sub(r"\n\s*\n+", "\n", value)
    return value.strip()


def _attachment_text(filename: str, data: bytes) -> str:
    """E-posta eklerindeki desteklenen metin belgelerini güvenli biçimde oku."""
    suffix = Path(filename or "").suffix.lower()
    if suffix in {".docx", ".doc", ".pdf", ".txt", ".md", ".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}:
        try:
            return extract_application_source_text(filename, data)
        except Exception:
            return ""
    return ""



def _strip_quoted_email_html(value: str) -> str:
    """Yalnız güncel e-posta gövdesini bırak; reply/forward zincirini HTML aşamasında kes."""
    html = value or ""
    # Yaygın istemcilerin alıntı blokları. İlk alıntıdan sonrası önceki yazışmadır.
    patterns = (
        r"(?is)<blockquote\b[^>]*>.*$",
        r"(?is)<div\b[^>]*class=[\"'][^\"']*(?:gmail_quote|moz-cite-prefix|yahoo_quoted)[^\"']*[\"'][^>]*>.*$",
        r"(?is)<div\b[^>]*id=[\"'](?:divRplyFwdMsg|appendonsend)[\"'][^>]*>.*$",
    )
    for pat in patterns:
        html = re.sub(pat, "", html, count=1)
    return html


def _current_email_body(value: str) -> str:
    """E-posta gövdesinden yalnız en üstteki/güncel mesajı döndür.

    Outlook/Gmail'in alta eklediği önceki yazışmalar, forwarded/original message
    blokları ve imza bloğu başvuru verisi olarak değerlendirilmez.
    """
    text = (value or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = text.splitlines()
    cut = len(lines)
    strong_markers = (
        re.compile(r"^\s*-{2,}\s*(?:original message|özgün ileti|iletilen ileti|forwarded message)\s*-{2,}\s*$", re.I),
        re.compile(r"^\s*(?:begin forwarded message|iletinin başlangıcı)\s*:?\s*$", re.I),
        re.compile(r"^\s*on\s+.+\s+wrote\s*:\s*$", re.I),
        re.compile(r"^\s*.+\s+tarihinde\s+.+\s+şunu yazdı\s*:\s*$", re.I),
    )
    header_re = re.compile(r"^\s*(?:from|kimden|gönderen|sent|gönderildi|to|kime|cc|subject|konu)\s*:\s*", re.I)
    for i, raw in enumerate(lines):
        line = raw.strip()
        if not line:
            continue
        if any(p.search(line) for p in strong_markers):
            cut = i
            break
        # Outlook reply zinciri çoğunlukla art arda From/Sent/To/Subject başlıklarıyla başlar.
        if header_re.search(line):
            window = [x.strip() for x in lines[i:i+7] if x.strip()]
            header_hits = sum(1 for x in window if header_re.search(x))
            if header_hits >= 2:
                cut = i
                break
    current = "\n".join(lines[:cut]).strip()

    # İmza bloğu başvuru cevabının parçası değildir; yaygın kapanışlardan sonrasını at.
    sig_patterns = (
        r"(?im)^\s*--\s*$",
        r"(?im)^\s*(?:saygılarımla|saygılarımızla|iyi çalışmalar|best regards|kind regards|regards|sincerely)[,!. ]*\s*$",
    )
    sig_cut = len(current)
    for pat in sig_patterns:
        m = re.search(pat, current)
        if m:
            sig_cut = min(sig_cut, m.start())
    current = current[:sig_cut].strip()
    current = re.sub(r"\n\s*\n{2,}", "\n\n", current)
    return current


def _eml_text(data: bytes) -> str:
    """EML'de yalnız güncel mesaj gövdesini oku; eski reply/forward zincirini alma."""
    msg = BytesParser(policy=policy.default).parsebytes(data)
    plain_parts: list[str] = []
    html_parts: list[str] = []
    if msg.is_multipart():
        for part in msg.walk():
            disposition = part.get_content_disposition()
            filename = part.get_filename() or ""
            if disposition == "attachment" or filename:
                # E-posta eki başvuru kaynağı olacaksa kullanıcı ayrıca yükler.
                continue
            ctype = part.get_content_type()
            if ctype == "text/plain":
                try:
                    plain_parts.append(_current_email_body(str(part.get_content())))
                except Exception:
                    payload = part.get_payload(decode=True) or b""
                    plain_parts.append(_current_email_body(payload.decode(part.get_content_charset() or "utf-8", errors="replace")))
            elif ctype == "text/html":
                try:
                    html = _strip_quoted_email_html(str(part.get_content()))
                    html_parts.append(_current_email_body(_html_to_text(html)))
                except Exception:
                    pass
    else:
        try:
            content = msg.get_content()
            if msg.get_content_type() == "text/html":
                body = _html_to_text(_strip_quoted_email_html(str(content)))
            else:
                body = str(content)
            (html_parts if msg.get_content_type() == "text/html" else plain_parts).append(_current_email_body(body))
        except Exception:
            pass
    structured_html = [x for x in html_parts if "\t" in x]
    body_parts = structured_html or plain_parts or html_parts
    return "\n".join(x for x in body_parts if x and x.strip()).strip()


def _ocr_image_source_text(data: bytes) -> str:
    """Form/ekran görüntüsünü sütun ilişkisini mümkün olduğunca koruyarak OCR eder.

    Tesseract'ın düz ``image_to_string`` çıktısı tablo sütunlarını tek satıra
    yapıştırabildiği için hak sahibi/buluş sahibi alanları kayabiliyordu. Burada
    kelime koordinatlarından satırlar yeniden kurulur ve büyük yatay boşluklar
    TAB olarak korunur. Böylece ``Unvanı <TAB> ABC A.Ş. <TAB> VKN <TAB> ...``
    yapısı kural tabanlı ayrıştırıcıya ulaşır.
    """
    try:
        import pytesseract  # type: ignore
    except ImportError as exc:
        raise ValueError("Resimden metin okuyabilmek için yerel OCR bileşeni kurulu değildir.") from exc

    image = Image.open(io.BytesIO(data))
    image = ImageOps.exif_transpose(image)
    # Küçük ekran görüntülerinde OCR doğruluğunu yükselt; oranı değiştirme.
    if image.width < 1800:
        scale = min(2.5, 1800 / max(1, image.width))
        image = image.resize((int(image.width * scale), int(image.height * scale)))
    gray = ImageOps.autocontrast(image.convert("L"))

    ocr = pytesseract.image_to_data(
        gray,
        lang="tur+eng",
        config="--psm 6 -c preserve_interword_spaces=1",
        output_type=pytesseract.Output.DICT,
    )
    grouped: dict[tuple[int, int, int, int], list[dict[str, float | str]]] = {}
    count = len(ocr.get("text", []))
    for i in range(count):
        word = str(ocr["text"][i] or "").strip()
        if not word:
            continue
        try:
            conf = float(ocr.get("conf", ["0"] * count)[i])
        except Exception:
            conf = 0.0
        if conf < 15:
            continue
        key = (
            int(ocr.get("page_num", [1] * count)[i]),
            int(ocr.get("block_num", [0] * count)[i]),
            int(ocr.get("par_num", [0] * count)[i]),
            int(ocr.get("line_num", [0] * count)[i]),
        )
        grouped.setdefault(key, []).append({
            "text": word,
            "left": float(ocr["left"][i]),
            "top": float(ocr["top"][i]),
            "width": float(ocr["width"][i]),
            "height": float(ocr["height"][i]),
        })

    rows: list[tuple[float, str]] = []
    for words in grouped.values():
        words.sort(key=lambda w: float(w["left"]))
        if not words:
            continue
        char_widths = [float(w["width"]) / max(1, len(str(w["text"]))) for w in words]
        heights = [float(w["height"]) for w in words]
        char_widths.sort(); heights.sort()
        med_char = char_widths[len(char_widths)//2] if char_widths else 8.0
        med_h = heights[len(heights)//2] if heights else 18.0
        # Belirgin hücre/sütun boşluğunu TAB yap. Normal kelime arası boşluk kalır.
        gap_threshold = max(28.0, med_char * 4.8, med_h * 1.9)
        chunks: list[str] = []
        prev_right: float | None = None
        for w in words:
            left = float(w["left"])
            if prev_right is not None:
                gap = left - prev_right
                chunks.append("\t" if gap >= gap_threshold else " ")
            chunks.append(str(w["text"]))
            prev_right = left + float(w["width"])
        line = re.sub(r"[ ]{2,}", " ", "".join(chunks)).strip()
        if line:
            rows.append((min(float(w["top"]) for w in words), line))
    rows.sort(key=lambda x: x[0])
    row_text = "\n".join(x[1] for x in rows).strip()
    if row_text:
        return row_text
    # Koordinatlı OCR olağan dışı biçimde boş dönerse güvenli geri dönüş.
    return pytesseract.image_to_string(gray, lang="tur+eng", config="--psm 6").strip()

def extract_application_source_text(filename: str, data: bytes) -> str:
    """Beyan formu/yazı/e-posta gibi başvuru bilgi kaynaklarından metin çıkarır."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        text = _docx_source_text(data)
    elif suffix == ".doc":
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / Path(filename).name
            src.write_bytes(data)
            proc = subprocess.run(["antiword", str(src)], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)
            text = proc.stdout.decode("utf-8", errors="replace") if proc.returncode == 0 else ""
    elif suffix == ".pdf":
        text = "\n".join((page.extract_text() or "") for page in PdfReader(io.BytesIO(data)).pages)
    elif suffix in {".txt", ".md"}:
        text = data.decode("utf-8", errors="replace")
    elif suffix == ".eml":
        text = _eml_text(data)
    elif suffix == ".msg":
        try:
            import extract_msg  # type: ignore
        except ImportError as exc:
            raise ValueError(".msg e-posta dosyasını okuyabilmek için extract-msg paketi kurulu olmalıdır.") from exc
        with tempfile.TemporaryDirectory() as td:
            src = Path(td) / Path(filename).name
            src.write_bytes(data)
            msg = extract_msg.Message(str(src))
            html_body = getattr(msg, "htmlBody", None)
            if isinstance(html_body, (bytes, bytearray)):
                html_body = bytes(html_body).decode("utf-8", errors="replace")
            structured_body = _current_email_body(_html_to_text(_strip_quoted_email_html(str(html_body)))) if html_body else ""
            plain_body = _current_email_body(getattr(msg, "body", "") or "")
            # Yalnız mevcut mesaj gövdesi kullanılır. HTML ve düz metin aynı güncel
            # gövdeyi temsil eder; EVET/HAYIR ve başvuru etiketlerini daha iyi koruyanı seç.
            def _body_score(v: str) -> tuple[int, int]:
                n = _plain_norm(v)
                signals = sum(n.count(x) for x in [
                    "evet", "hayir", "buluscu", "bulus sahibi", "hak sahibi",
                    "erken yayin", "erken yayim", "tubitak", "kosgeb", "proje",
                ])
                return (signals, len(v or ""))
            candidates = [x for x in [structured_body, plain_body] if x.strip()]
            text = max(candidates, key=_body_score) if candidates else ""
            try:
                msg.close()
            except Exception:
                pass
    elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}:
        try:
            # Türkçe + İngilizce OCR yerel çalışır; OpenAI/API kredisi tüketmez.
            # Tablo/sütun koordinatları mümkün olduğunca TAB olarak korunur.
            text = _ocr_image_source_text(data)
        except Exception as exc:
            raise ValueError(f"Resim OCR ile okunamadı: {exc}") from exc
    else:
        raise ValueError("Bilgi kaynağı için desteklenen türler: .docx, .doc, .pdf, .txt, .md, .eml, .msg, .png, .jpg, .jpeg, .webp, .tif, .tiff, .bmp")
    text = (text or "").replace("\x00", " ").strip()
    if not text:
        raise ValueError(f"{filename} dosyasından metin çıkarılamadı.")
    return text



# -----------------------------------------------------------------------------
# AI'SIZ / DETERMINISTIK BASVURU BILGISI CIKARIMI
# -----------------------------------------------------------------------------

_LABEL_SPLIT_RE = re.compile(r"^\s*([^:\t|]{2,80}?)\s*(?::|\t|\|)\s*(.*?)\s*$")


def _plain_norm(value: str) -> str:
    value = (value or "").casefold()
    value = value.translate(str.maketrans({"ı": "i", "ş": "s", "ğ": "g", "ü": "u", "ö": "o", "ç": "c"}))
    value = unicodedata.normalize("NFKD", value)
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    value = re.sub(r"[^a-z0-9]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def _clean_value(value: str) -> str:
    value = re.sub(r"\s+", " ", (value or "").strip(" \t|:;-"))
    return value.strip()


def _looks_like_company(name: str) -> bool:
    n = _plain_norm(name)
    return bool(re.search(r"\b(a s|anonim|ltd|limited|sanayi|ticaret|holding|sirketi|company|inc|corp|llc)\b", n))


def _entity_type(name: str, explicit: str = "") -> str:
    e = _plain_norm(explicit)
    if "tuzel" in e or "kurum" in e or "sirket" in e:
        return "Tüzel kişi"
    if "gercek" in e or "kisi" in e:
        return "Gerçek kişi"
    if _looks_like_company(name):
        return "Tüzel kişi"
    return ""


def _canonical_label(label: str) -> str:
    n = _plain_norm(label)
    if not n:
        return ""

    # Form başlıklarında sık geçen ek/sonekleri normalleştir; anlamı bozacak
    # genel kelimeleri tamamen silmek yerine rol tespitinde esnek davran.
    applicant_role = any(x in n for x in [
        "hak sahibi", "basvuru sahibi", "basvuran", "applicant", "patent sahibi",
    ])
    inventor_role = any(x in n for x in [
        "bulus sahibi", "bulusu yapan", "bulus yapan", "buluscu", "mucit", "inventor",
    ])
    name_terms = [
        "adi soyadi", "ad soyad", "adi ve soyadi", "ad ve soyad", "ad soyadi",
        "adi unvani", "ad unvan", "adi ve unvani", "unvani", "unvan", "isim", "name",
    ]
    identity_terms = [
        "tckn", "t c kimlik", "tc kimlik", "tc kimlik no", "tc kimlik numarasi", "t c kimlik numarasi",
        "vkn", "vergi no", "vergi numarasi", "vergi kimlik no", "vergi kimlik numarasi",
        "kimlik no", "kimlik numarasi", "identity", "tax number",
    ]

    # Rol + alan birleşik etiketleri: "Başvuru Sahibinin Adı/Unvanı",
    # "Buluşu Yapanın Adresi", "Hak Sahibi TCKN/VKN" vb.
    if applicant_role or inventor_role:
        role = "applicant" if applicant_role else "inventor"
        if any(x in n for x in ["adres", "address", "tebligat adres", "ikametgah"]):
            return role + "_address"
        if any(x in n for x in ["ulke", "uyruk", "tabiyet", "country", "milliyet"]):
            return role + "_country"
        if any(x in n for x in ["il ilce", "il ilçe", "il/ilce", "il/ilçe"]):
            return role + "_location"
        if re.search(r"\bilce\b", n):
            return role + "_district"
        if re.search(r"\b(il|sehir|city)\b", n):
            return role + "_city"
        if any(x in n for x in identity_terms):
            return role + "_identity"
        if any(x in n for x in ["dogum tarihi", "birth date"]):
            return role + "_birth_date"
        if any(x in n for x in ["e posta", "eposta", "email", "e mail"]):
            return role + "_email"
        if any(x in n for x in ["telefon", "gsm", "cep telefonu", "phone"]):
            return role + "_phone"
        if any(x in n for x in name_terms):
            return role + "_name"
        if applicant_role and any(x in n for x in ["kisi turu", "tuzel", "gercek", "entity type", "hukuki nitelik"]):
            return "applicant_entity_type"

    aliases = {
        "application_kind": [
            "basvuru turu", "basvuru tipi", "koruma turu", "basvuru sekli", "koruma sekli",
            "patent faydali model tercihi", "patent veya faydali model",
        ],
        "reference": [
            "dp referans", "dp ref", "dp no", "dosya referansi", "dosya referans no", "dosya no",
            "referans no", "referans numarasi", "referans", "is no", "is numarasi", "dosya kodu",
        ],
        "invention_title": [
            "bulus basligi", "bulusun basligi", "bulus adi", "bulusun adi", "buluşun adı",
            "invention title", "baslik",
        ],
        "applicant": [
            "hak sahibi", "hak sahipleri", "basvuru sahibi", "basvuru sahipleri", "basvuran",
            "applicant", "applicants", "hak sahibi bilgileri", "basvuru sahibi bilgileri",
        ],
        "inventor": [
            "bulus sahibi", "bulus sahipleri", "buluscu", "buluscular", "bulusu yapan",
            "bulusu yapanlar", "bulus yapan", "mucit", "mucitler", "inventor", "inventors",
            "bulus sahibi bilgileri", "bulusu yapan bilgileri",
        ],
        "name": name_terms,
        "identity": identity_terms,
        "address": [
            "adres", "adresi", "tebligat adresi", "ikamet adresi", "ikametgah", "merkez adresi",
            "yazisma adresi", "address",
        ],
        "country": ["ulke", "ulkesi", "uyruk", "uyrugu", "tabiyet", "tabiyeti", "milliyet", "milliyeti", "country"],
        "location": ["il ilce", "il/ilce", "il / ilce", "il ilçe", "il/ilçe", "il / ilçe", "il ve ilce", "city district"],
        "city": ["il", "ili", "sehir", "sehri", "city"],
        "district": ["ilce", "ilcesi", "ilçe", "ilçesi", "district"],
        "birth_date": ["dogum tarihi", "doğum tarihi", "birth date", "date of birth"],
        "fax": ["faks", "fax", "faks no", "fax no"],
        "postal_code": ["posta kodu", "postal code", "zip code"],
        "authorized_person": ["yetkili", "yetkili kisi", "yetkili kişi", "temsilci", "contact person"],
        "website": ["web", "web sitesi", "internet adresi", "website"],
        "entity_type": [
            "kisi turu", "sahip turu", "hak sahibi turu", "basvuru sahibi turu", "tuzel gercek kisi",
            "gercek tuzel", "hukuki nitelik", "entity type",
        ],
        "gender": ["cinsiyet", "gender", "sex"],
        "email": ["e posta", "eposta", "e mail", "email", "mail adresi", "elektronik posta"],
        "phone": ["telefon", "telefonu", "telefon no", "telefon numarasi", "gsm", "cep telefonu", "ev telefonu", "is telefonu", "ev is telefonu", "phone"],
        "priority": ["ruchan", "ruchan durumu", "ruchan talebi", "priority", "priority claim"],
        "priority_country": ["ruchan ulkesi", "priority country"],
        "priority_number": ["ruchan numarasi", "ruchan no", "ruchan basvuru no", "priority number"],
        "priority_date": ["ruchan tarihi", "priority date"],
    }
    order = (
        "priority_country", "priority_number", "priority_date", "application_kind", "invention_title",
        "reference", "entity_type", "applicant", "inventor", "name", "identity", "address", "country",
        "location", "city", "district", "birth_date", "gender", "authorized_person", "postal_code", "email", "phone", "fax", "website", "priority",
    )
    for key in order:
        for alias in aliases[key]:
            a = _plain_norm(alias)
            if n == a or n.startswith(a + " ") or n.endswith(" " + a):
                return key
    return ""



def _is_instruction_text(value: str) -> bool:
    n = _plain_norm(value)
    if not n:
        return False
    markers = [
        "not basvuru sahibinin", "birden fazla olmasi durumunda", "ayri ayri duzenlenmelidir",
        "bu alan", "doldurulmalidir", "zorunlu alan", "gerekiyorsa", "aciklama",
        "ornek olarak", "ornek bilgi", "ornek metin",
        "basvuru sahibinin birden fazla", "bulusu yapan birden fazla", "her bir basvuru sahibi",
    ]
    return any(x in n for x in markers)


def _is_role_section_heading(raw_label: str, canonical: str) -> bool:
    n = _plain_norm(raw_label)
    if canonical not in {"applicant", "inventor"}:
        return False
    # Resmî beyan formlarında bölüm başlığı bazen yalnız "... SAHİBİ(LERİ)"
    # biçimindedir ve hemen yanında açıklama/not bulunur. Bu da bölüm başlığıdır.
    if any(x in n for x in [
        "bilgi", "bilgileri", "bilgisi", "information", "detay", "detaylari",
        "sahipleri", "yapanlar", "mucitler", "buluscular",
    ]):
        return True
    # Kısa ve yalnız rol adlarından oluşan başlıkları kabul et;
    # "Buluş Sahibi: Ali Veli" gibi etiket+değer satırını başlık sanma.
    if re.search(r"[:\t|]", raw_label or ""):
        return False
    pure = n.strip()
    return pure in {
        "hak sahibi", "basvuru sahibi", "hak sahibi basvuru sahibi",
        "hak sahipleri", "basvuru sahipleri", "hak sahipleri basvuru sahipleri",
        "bulus sahibi", "bulus sahipleri", "bulusu yapan", "bulusu yapanlar",
        "buluscu", "buluscular", "mucit", "mucitler",
    }



def _is_form_header_bundle(value: str) -> bool:
    """OCR'ın bir tablo başlık satırını tek değer gibi birleştirmesini tanır."""
    n = _plain_norm(value)
    if not n:
        return False
    markers = [
        "unvani", "ad soyad", "hak sahibi adresi", "adres", "sahip turu", "kisi turu",
        "uyruk", "ulke", "tc kimlik", "tckn", "vergi no", "vkn", "e posta", "telefon",
        "dogum tarihi", "cinsiyet", "il ilce", "ilce",
    ]
    hits = sum(1 for m in markers if m in n)
    # Şirket unvanı veya gerçek bir e-posta/adres gibi açık bir veri varsa başlık
    # demek için daha yüksek eşik kullan; aksi halde üç alan adı yeterlidir.
    threshold = 5 if (_looks_like_company(value) or "@" in value) else 3
    return hits >= threshold


def _is_bad_person_value(field: str, value: str, *, applicant: bool) -> bool:
    value = _clean_value(value)
    n = _plain_norm(value)
    if not n:
        return True
    if _is_instruction_text(value) or _is_form_header_bundle(value):
        return True
    if _canonical_label(value):
        return True
    bad = [
        "bilgileri gizlensin", "cevap verilmemesi", "isleme alinacaktir", "cinsiyet",
        "dogum tarihi", "ev is telefonu", "hak sahibi adresi", "sahip turu", "uyruk",
        "tc kimlik vergi no", "e posta telefon", "basvuru sahibi bilgileri", "hak sahibi bilgileri",
        "bulus sahibi bilgileri",
    ]
    if any(x in n for x in bad):
        return True
    if field == "name":
        if not re.search(r"[A-Za-zÇĞİÖŞÜçğıöşü]", value):
            return True
        if "?" in value or "->" in value or "à" in value or "→" in value:
            return True
        if applicant:
            # Gerçek kişi başvuru sahibi de olabilir; yalnız salt form etiketi/uzun
            # açıklama cümlesini ad/unvan olarak kabul etme.
            if len(value) > 220:
                return True
        else:
            # Buluşçu adı normalde kısa bir kişi adıdır. Form soru/cümlelerini dışla.
            if len(value) > 120 or len(value.split()) > 12:
                return True
    return False


def _inline_field_pair(token: str) -> tuple[str, str, str] | None:
    """OCR aynı hücrede ``Etiket Değer`` ürettiyse ikisini ayırır."""
    token = _clean_value(token)
    # Kimlik etiketi kendi içinde "No", "/ Vergi No" gibi sözcükler taşıyabilir.
    # Uzun bir kimlik değeri yoksa bunu asla inline değerli alan sayma.
    if _canonical_label(token) == "identity" and not re.search(r"\d{5,}", token):
        return None
    # Bazı etiketlerin kendisi slash içerir; bunları "etiket + değer" sanma.
    if re.fullmatch(r"(?i)(?:tckn\s*/\s*kimlik|adı\s*/\s*unvanı|adi\s*/\s*unvani|il\s*/\s*ilçe|il\s*/\s*ilce)", token):
        return None
    patterns: list[tuple[str, str]] = [
        (r"^(doğum\s+tarihi|dogum\s+tarihi|birth\s+date)\s+(.+)$", "birth_date"),
        (r"^(t\.?\s*c\.?\s*kimlik(?:\s+no|\s+numarası|\s+numarasi)?|tc\s+kimlik(?:\s+no)?|tckn\s*/\s*kimlik|tckn|vkn|vergi\s+(?:kimlik\s+)?(?:no|numarası|numarasi))\s+(.+)$", "identity"),
        (r"^(adı\s*/\s*unvanı|adi\s*/\s*unvani|adı\s+soyadı|adi\s+soyadi|ad\s+soyad|unvanı|unvani)\s+(.+)$", "name"),
        (r"^(il\s*/\s*ilçe|il\s*/\s*ilce|il\s+ilçe|il\s+ilce)\s+(.+)$", "location"),
        (r"^(ülke|ulke|uyruk|country)\s+(.+)$", "country"),
        (r"^(ilçe|ilce|district)\s+(.+)$", "district"),
        (r"^(il|şehir|sehir|city)\s+(.+)$", "city"),
        (r"^(e[ -]?posta|eposta|e[ -]?mail|email)\s+(.+)$", "email"),
        (r"^(ev\s*/\s*iş\s+telefonu|ev\s*/\s*is\s+telefonu|telefon|gsm|phone)\s+(.+)$", "phone"),
        (r"^(adres|adresi|address)\s+(.+)$", "address"),
        (r"^(sahip\s+türü|sahip\s+turu|kişi\s+türü|kisi\s+turu)\s+(.+)$", "entity_type"),
        (r"^(cinsiyet|gender|sex)\s+(.+)$", "gender"),
    ]
    for pattern, label in patterns:
        m = re.match(pattern, token, flags=re.IGNORECASE)
        if m:
            value = _clean_value(m.group(2))
            if value:
                return (m.group(1), label, value)
    return None


def _line_label_value_pairs(line: str) -> list[tuple[str, str, str]]:
    """Bir görsel/tablo satırındaki bir veya birden çok etiket-değer çiftini çıkar.

    Dönüş: (ham etiket, kanonik etiket, değer). Tablo satırlarında
    "Hak Sahibi | ABC A.Ş. | Ülke | Türkiye" gibi birden çok çift desteklenir.
    """
    line = (line or "").strip()
    if not line or line.startswith("[[TABLO") or line.startswith("[[E-POSTA EKI"):
        return []
    # OCR tablo başlıklarını tek bir değer gibi yapıştırdıysa kişi adı/unvanı
    # sanma. Koordinatlı OCR ile çoğu satır TAB'lı gelir; bu koruma geri dönüş
    # senaryoları içindir.
    if "\t" not in line and _is_form_header_bundle(line):
        return []

    # Önce tablo/hücre ayraçlarını değerlendir. Hücre içindeki " / " normal
    # metin olarak bırakılır; yalnız tab ve belirgin dikey çizgiler ayraçtır.
    tokens = [x.strip() for x in re.split(r"\t+|\s+\|\s+", line) if x.strip()]
    if len(tokens) >= 2:
        pairs: list[tuple[str, str, str]] = []
        i = 0
        while i < len(tokens):
            inline = _inline_field_pair(tokens[i])
            if inline is not None:
                pairs.append(inline)
                i += 1
                continue
            label = _canonical_label(tokens[i])
            if label:
                j = i + 1
                vals: list[str] = []
                while j < len(tokens) and not _canonical_label(tokens[j]):
                    vals.append(tokens[j])
                    j += 1
                value = _clean_value(" / ".join(vals))
                if label in {"applicant", "inventor"} and _is_instruction_text(value):
                    value = ""
                pairs.append((tokens[i], label, value))
                i = max(j, i + 1)
            else:
                i += 1
        if pairs:
            return pairs

    m = _LABEL_SPLIT_RE.match(line)
    if m:
        raw_label = m.group(1)
        label = _canonical_label(raw_label)
        if label:
            value = _clean_value(m.group(2))
            if label in {"applicant", "inventor"} and _is_instruction_text(value):
                value = ""
            return [(raw_label, label, value)]
    # OCR/PDF aynı hücreyi "Doğum Tarihi 23.05.1992" gibi döndürmüş olabilir.
    generic_inline = _inline_field_pair(line)
    if generic_inline is not None:
        return [generic_inline]

    # PDF/antiword çıktısında tablo sütunları bazen ayraçsız tek satıra
    # düşer: "Hak Sahibi ABC A.Ş." gibi. Yalnız açık ve sabit etiket
    # başlangıçlarında ayraçsız geri dönüş uygula.
    inline_patterns = [
        (r"^(hak\s+sahibi|başvuru\s+sahibi|başvuru\s+sahibinin|hak\s+sahibinin)\s+(.+)$", "applicant"),
        (r"^(buluş\s+sahibi|buluşu\s+yapan|buluşçu|mucit)\s+(.+)$", "inventor"),
        (r"^(buluş\s+başlığı|buluşun\s+başlığı|buluş\s+adı)\s+(.+)$", "invention_title"),
        (r"^(başvuru\s+türü|başvuru\s+tipi|koruma\s+türü)\s+(.+)$", "application_kind"),
        (r"^(dp\s*(?:ref(?:erans)?|no)|dosya\s+(?:referansı|no)|referans\s+no)\s+(.+)$", "reference"),
    ]
    for pattern, default_label in inline_patterns:
        mm = re.match(pattern, line, flags=re.IGNORECASE)
        if mm:
            value = _clean_value(mm.group(2))
            if _plain_norm(value) not in {"bilgi", "bilgileri", "bilgisi"}:
                return [(mm.group(1), default_label, value)]

    label = _canonical_label(line)
    if label:
        return [(line, label, "")]
    return []


def reference_from_filename(filename: str) -> str:
    """Tarifname dosya adındaki DP/ofis referansını güvenli biçimde çıkar.

    Örn. Tarifname_181176_rev3.docx -> 181176. Yalnız açık bir numaralı
    referans bulunduğunda döner; tarih/sürüm gibi kısa parçaları seçmez.
    """
    stem = Path(filename or "").stem
    # DP181176, DP-181176, Tarifname_181176, 181176_Tarifname gibi örnekler.
    explicit = re.search(r"(?i)\bdp\s*[-_ ]?\s*(\d{4,12})\b", stem)
    if explicit:
        return explicit.group(1)
    candidates = re.findall(r"(?<!\d)(\d{5,12})(?!\d)", stem)
    if candidates:
        # En uzun aday; eşitlikte soldaki. 5.4.48 gibi sürüm parçaları zaten
        # noktalı/kısa olduğundan bu kalıba girmez.
        return sorted(enumerate(candidates), key=lambda x: (-len(x[1]), x[0]))[0][1]
    return ""


def _add_other_information(result: dict[str, Any], label: str, value: str, source: str) -> None:
    label = _clean_value(label)
    value = _clean_value(value)
    if not value:
        return
    key = (_plain_norm(label), _plain_norm(value), source)
    for row in result.get("other_information", []):
        if (_plain_norm(row.get("label", "")), _plain_norm(row.get("value", "")), row.get("source", "")) == key:
            return
    result["other_information"].append({"label": label or "Diğer bilgi", "value": value, "source": source})


def _extract_contact_information(result: dict[str, Any], text: str, source: str) -> None:
    """Serbest mail/yazıdaki açık iletişim bilgilerini 'diğer bilgiler'e al.

    Telefon için yalnız açık Telefon/Tel/Cep/GSM etiketi kabul edilir; VKN/TCKN,
    dosya numarası veya metindeki başka sayı telefon sanılmaz.
    """
    for email in sorted(set(re.findall(r"(?i)\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", text or ""))):
        _add_other_information(result, "E-posta", email, source)
    phone_labeled = re.compile(
        r"(?i)(?:telefon|tel\.?|cep(?:\s*telefonu)?|gsm|mobil\s*telefon)"
        r"(?:\s*(?:no|numarası|numarasi))?\s*[:|\-]?\s*"
        r"((?:\+?90\s*)?(?:\(?0?\d{3}\)?[ .-]*)?\d{3}[ .-]*\d{2}[ .-]*\d{2})"
    )
    for match in phone_labeled.finditer(text or ""):
        phone = _clean_value(match.group(1))
        if len(re.sub(r"\D", "", phone)) >= 7:
            _add_other_information(result, "Telefon", phone, source)


def _split_people_value(value: str) -> list[str]:
    value = _clean_value(value)
    if not value:
        return []
    # Birden fazla kişi açık biçimde noktalı virgül, satır içi numara veya " ve " ile verilmişse ayır.
    parts = re.split(r"\s*;\s*|\s+\d+[.)]\s+", value)
    if len(parts) == 1 and " ve " in value.casefold() and not _looks_like_company(value):
        candidate = re.split(r"\s+ve\s+", value, flags=re.IGNORECASE)
        if 1 < len(candidate) <= 4 and all(len(x.split()) >= 2 for x in candidate):
            parts = candidate
    return [p.strip() for p in parts if p.strip()]


def _new_person(name: str, source: str, *, applicant: bool) -> dict[str, str]:
    row = {
        "identity": "",
        "name": _clean_value(name),
        "country": "",
        "city": "",
        "district": "",
        "address": "",
        "email": "",
        "phone": "",
        "birth_date": "",
        "source": source,
    }
    if applicant:
        row["entity_type"] = _entity_type(name)
    return row


def _split_city_district(value: str) -> tuple[str, str]:
    value = _clean_value(value)
    if not value:
        return "", ""
    parts = [x.strip() for x in re.split(r"[/|,]", value, maxsplit=1)]
    if len(parts) == 2 and all(parts):
        return parts[0], parts[1]
    return "", value


def _apply_person_field(row: dict[str, str], field: str, value: str, *, applicant: bool) -> None:
    value = _clean_value(value)
    if not value:
        return
    if field == "identity":
        ident = _valid_identity(value)
        if ident:
            row["identity"] = ident
    elif field == "entity_type" and applicant:
        row["entity_type"] = _entity_type(row.get("name", ""), value)
    elif field == "location":
        city, district = _split_city_district(value)
        if city and not row.get("city"):
            row["city"] = city
        if district and not row.get("district"):
            row["district"] = district
    elif field == "district":
        city, district = _split_city_district(value)
        if city and _plain_norm(city) in _TURKEY_PROVINCES and not row.get("city"):
            row["city"] = city
        if not row.get("district"):
            row["district"] = district or value
    elif field == "birth_date":
        # Form başlığıyla komşu hücre numaralarını doğum tarihi sanma.
        if re.search(r"(?<!\d)(?:0?[1-9]|[12]\d|3[01])[./-](?:0?[1-9]|1[0-2])[./-](?:19|20)\d{2}(?!\d)", value):
            row["birth_date"] = value
    elif field in {"address", "country", "city", "email", "phone"}:
        if _is_bad_person_value(field, value, applicant=applicant):
            return
        if not row.get(field):
            row[field] = value


_TURKEY_PROVINCES = {
    _plain_norm(x) for x in """Adana Adıyaman Afyonkarahisar Ağrı Aksaray Amasya Ankara Antalya Ardahan Artvin Aydın Balıkesir Bartın Batman Bayburt Bilecik Bingöl Bitlis Bolu Burdur Bursa Çanakkale Çankırı Çorum Denizli Diyarbakır Düzce Edirne Elazığ Erzincan Erzurum Eskişehir Gaziantep Giresun Gümüşhane Hakkari Hatay Iğdır Isparta İstanbul İzmir Kahramanmaraş Karabük Karaman Kars Kastamonu Kayseri Kırıkkale Kırklareli Kırşehir Kilis Kocaeli Konya Kütahya Malatya Manisa Mardin Mersin Muğla Muş Nevşehir Niğde Ordu Osmaniye Rize Sakarya Samsun Siirt Sinop Sivas Şanlıurfa Şırnak Tekirdağ Tokat Trabzon Tunceli Uşak Van Yalova Yozgat Zonguldak""".split()
}


def _fill_country_from_explicit_location(row: dict[str, str]) -> None:
    if row.get("country"):
        return
    naddr = _plain_norm(row.get("address", ""))
    if "turkiye" in naddr or re.search(r"\bturkey\b", naddr):
        row["country"] = "Türkiye"
        return
    city = _plain_norm(row.get("city", ""))
    if city in _TURKEY_PROVINCES:
        row["country"] = "Türkiye"



def _append_unique_person(target: list[dict[str, str]], row: dict[str, str], *, applicant: bool) -> dict[str, str]:
    name_key = _plain_norm(row.get("name", ""))
    id_key = re.sub(r"\D", "", row.get("identity", ""))
    for existing in target:
        existing_id = re.sub(r"\D", "", existing.get("identity", ""))
        if (id_key and existing_id and id_key == existing_id) or (name_key and name_key == _plain_norm(existing.get("name", ""))):
            for key in ("identity", "country", "city", "district", "address", "email", "phone", "birth_date"):
                if not existing.get(key) and row.get(key):
                    existing[key] = row[key]
            if applicant and not existing.get("entity_type"):
                existing["entity_type"] = row.get("entity_type") or _entity_type(existing.get("name", ""))
            if row.get("source") and row["source"] not in (existing.get("source") or "").split("; "):
                existing["source"] = "; ".join(x for x in [existing.get("source", ""), row["source"]] if x)
            return existing
    target.append(row)
    return row


def _specification_title(specification_text: str) -> str:
    lines = [_clean_value(x) for x in (specification_text or "").splitlines()]
    lines = [x for x in lines if x]
    for idx, line in enumerate(lines[:15]):
        if _plain_norm(line) == "tarifname":
            for cand in lines[idx + 1: idx + 5]:
                n = _plain_norm(cand)
                if n and n not in {"teknik alan", "onceki teknik", "bulusun kisa aciklamasi"} and not cand.startswith("Araştırma raporunun"):
                    return cand
    # Başlık etiketi olan tarifnamelerde geri dönüş.
    for line in lines[:25]:
        m = _LABEL_SPLIT_RE.match(line)
        if m and _canonical_label(m.group(1)) == "invention_title":
            return _clean_value(m.group(2))
    return ""


def _kind_from_text(text: str) -> str:
    n = _plain_norm(text)
    if re.search(r"\bfaydali model\b", n):
        return "Faydalı Model"
    if re.search(r"\bpatent\b", n):
        return "Patent"
    return ""


def _priority_status(value: str) -> str:
    n = _plain_norm(value)
    if not n:
        return "Belirsiz"
    negative = ["yok", "hayir", "talep edilmiyor", "talep edilmeyecek", "bulunmuyor", "mevcut degil", "no"]
    positive = ["var", "evet", "talep ediliyor", "talep edilecek", "mevcut", "yes"]
    if any(x in n for x in negative):
        return "Yok"
    if any(x in n for x in positive):
        return "Var"
    return "Belirsiz"


def _valid_identity(value: str) -> str:
    value = _clean_value(value)
    if not value or _canonical_label(value) or _is_instruction_text(value) or _is_form_header_bundle(value):
        return ""
    digits = re.sub(r"\D", "", value or "")
    if len(digits) in {10, 11}:
        return digits
    # Yabancı kimlik/pasaport gibi alanlar yalnız sayı/harf içeren makul bir
    # değer ise korunur; "Cinsiyet", "Doğum Tarihi" gibi etiketler artık girmez.
    compact = re.sub(r"[^A-Za-z0-9]", "", value)
    if 5 <= len(compact) <= 24 and any(ch.isdigit() for ch in compact):
        return value
    return ""


def _set_singleton(result: dict[str, Any], key: str, value: str, source: str, seen: dict[str, tuple[str, str]]) -> None:
    value = _clean_value(value)
    if not value:
        return
    prior = seen.get(key)
    if prior and _plain_norm(prior[0]) != _plain_norm(value):
        result["conflicts"].append(f"{key}: '{prior[0]}' ({prior[1]}) / '{value}' ({source})")
        return
    seen[key] = (value, source)
    if key == "priority.status":
        result["priority"]["status"] = value
        result["priority"]["source"] = source
    elif key.startswith("priority."):
        result["priority"][key.split(".", 1)[1]] = value
        result["priority"]["source"] = source
    else:
        result[key] = value



def _explicit_yes_no(segment: str) -> str:
    """Varsayılan HAYIR açıklamasını cevap sanmadan gerçek EVET/HAYIR cevabını bulur."""
    segment = segment or ""

    def mapped(token: str) -> str:
        return "Evet" if _plain_norm(token) == "evet" else "Hayır"

    matches = re.findall(r"(?:à|→|->|=>|⇒|⟶)\s*(EVET|HAYIR)\b", segment, flags=re.IGNORECASE)
    if matches:
        return mapped(matches[-1])

    # HTML/OCR bazı ok karakterlerini düz 'a' olarak bırakabiliyor.
    matches = re.findall(r"\)\s*(?:a|:|[-–—>]+)?\s*(EVET|HAYIR)\b", segment, flags=re.IGNORECASE)
    if matches:
        return mapped(matches[-1])

    for line in reversed(segment.splitlines()):
        m = re.search(r"(?:cevap|yanıt|yanit)\s*[:\-]\s*(EVET|HAYIR)\b", line, flags=re.IGNORECASE)
        if m:
            return mapped(m.group(1))

    # Formlarda varsayılan açıklama önce, gerçek cevap en sonda bulunur.
    tokens = list(re.finditer(r"\b(EVET|HAYIR)\b", segment, flags=re.IGNORECASE))
    if len(tokens) >= 2:
        return mapped(tokens[-1].group(1))
    if len(tokens) == 1 and _plain_norm(tokens[0].group(1)) == "evet":
        return "Evet"
    return "Belirsiz"


def _question_segment(text: str, anchors: list[str]) -> str:
    normalized = text or ""
    for anchor in anchors:
        m = re.search(anchor, normalized, flags=re.IGNORECASE | re.DOTALL)
        if not m:
            continue
        start = m.start()
        tail = normalized[start:start + 1800]
        cuts: list[int] = []
        # Bir sonraki numaralı soruda kes.
        nxt = re.search(r"\n\s*\*{0,2}\d+[.)]\*{0,2}\s+", tail[20:])
        if nxt:
            cuts.append(20 + nxt.start())
        # Numarasız/yapıştırılmış e-postalarda bir sonraki bilinen başvuru sorusunu kes.
        next_question = re.search(
            r"\n\s*(?:\*{0,2}\d+[.)]\*{0,2}\s*)?(?:"
            r"başvuru\s+esnasında\s+)?(?:buluşçu\s+bilgileri\s+gizlensin\s+mi|"
            r"buluş\s+sahibi\s+bilgileri\s+gizlensin\s+mi|"
            r"buluş.{0,220}?(?:tübitak|tubitak|kosgeb|kamu\s+kurum).{0,260}?proje\s+kapsamında|"
            r"erken\s+yay(?:ın|ım)\s+(?:talep\s+ediliyor\s+mu|talebi))",
            tail[20:], flags=re.IGNORECASE | re.DOTALL,
        )
        if next_question:
            cuts.append(20 + next_question.start())
        if cuts:
            tail = tail[:min(cuts)]
        return tail
    return ""


def _set_filing_option(result: dict[str, Any], key: str, status: str, source: str) -> None:
    if status not in {"Evet", "Hayır"}:
        return
    row = result["filing_options"][key]
    prior = row.get("status", "Belirsiz")
    if prior in {"Evet", "Hayır"} and prior != status:
        result["conflicts"].append(f"{row.get('label', key)}: '{prior}' ({row.get('source')}) / '{status}' ({source})")
        return
    row["status"] = status
    row["source"] = source
    row["explicit"] = True


def _extract_filing_options(result: dict[str, Any], text: str, source: str) -> None:
    # 1) Buluşçu bilgilerinin gizlenmesi
    seg = _question_segment(text, [
        r"buluşçu\s+bilgileri\s+gizlensin\s+mi",
        r"buluş\s+sahibi\s+bilgileri\s+gizlensin\s+mi",
        r"mucit\s+bilgileri\s+gizlensin\s+mi",
    ])
    if seg:
        _set_filing_option(result, "inventor_hidden", _explicit_yes_no(seg), source)

    # 2) TÜBİTAK/KOSGEB/kamu destekli proje
    seg = _question_segment(text, [
        r"buluş.{0,120}(?:TÜBİTAK|TUBITAK).{0,120}(?:KOSGEB).{0,260}?proje\s+kapsamında",
        r"buluş.{0,180}?kamu\s+kurum.{0,220}?desteklenen\s+bir\s+proje",
    ])
    if seg:
        status = _explicit_yes_no(seg)
        _set_filing_option(result, "public_project", status, source)
        if status == "Evet":
            inst = re.search(r"(?:kurum|destekleyen\s+kurum)\s*[:\-]\s*([^\n;]{2,160})", seg, flags=re.IGNORECASE)
            proj = re.search(r"(?:proje\s*(?:no|numarası|numarasi)|project\s*(?:no|number))\s*[:\-]\s*([^\n;]{2,80})", seg, flags=re.IGNORECASE)
            row = result["filing_options"]["public_project"]
            if inst:
                row["institution"] = _clean_value(inst.group(1))
            if proj:
                row["project_number"] = _clean_value(proj.group(1))

    # 3) Erken yayın
    seg = _question_segment(text, [
        r"erken\s+yayın\s+talep\s+ediliyor\s+mu",
        r"erken\s+yayım\s+talep\s+ediliyor\s+mu",
        r"erken\s+yayın\s+talebi",
        r"erken\s+yayım\s+talebi",
    ])
    if seg:
        _set_filing_option(result, "early_publication", _explicit_yes_no(seg), source)


def _apply_filing_option_defaults(result: dict[str, Any]) -> None:
    # Kullanıcının beyan formu kurallarında cevap verilmemesi halinde üç alan da
    # HAYIR kabul ediliyor. Açık cevap bulunursa her zaman açık cevap üstün gelir.
    for key in ("inventor_hidden", "public_project", "early_publication"):
        row = result["filing_options"][key]
        if row.get("status") not in {"Evet", "Hayır"}:
            row["status"] = "Hayır"
            row["source"] = "Varsayılan (cevap verilmemiş)"
            row["explicit"] = False


def _table_header_field(cell: str) -> str:
    n = _plain_norm(cell)
    if not n:
        return ""
    if any(x in n for x in ["adi soyadi", "ad soyad", "adi unvani", "adi / unvani", "unvani", "unvan"]):
        return "name"
    if any(x in n for x in ["tc kimlik", "tckn", "vergi no", "vergi kimlik", "vkn"]):
        return "identity"
    if "hak sahibi adres" in n or "basvuru sahibi adres" in n or n in {"adres", "adresi", "address"}:
        return "address"
    if "sahip turu" in n or "kisi turu" in n:
        return "entity_type"
    if n in {"uyruk", "ulke", "country"}:
        return "country"
    if n in {"il", "sehir", "city"}:
        return "city"
    if n in {"ilce", "district"}:
        return "district"
    if "il / ilce" in n or "il ilce" in n:
        return "location"
    if any(x in n for x in ["e posta", "eposta", "email", "e mail"]):
        return "email"
    if any(x in n for x in ["telefon", "gsm", "cep"]):
        return "phone"
    if "dogum tarihi" in n:
        return "birth_date"
    if "cinsiyet" in n:
        return "gender"
    return ""


def _structured_role_rows(text: str, source: str) -> tuple[list[dict[str, str]], list[dict[str, str]], set[int]]:
    """Açık rol bölümündeki tablo başlıklarını hemen sonraki veri satırıyla eşler.

    Kişi/kurum modeli kayıt yaratmaz; kayıt yalnız açık Hak/başvuru sahibi veya
    Buluş sahibi/buluşçu tablo yapısından oluşur.
    """
    lines = [(i, x.strip()) for i, x in enumerate((text or "").splitlines())]
    applicants: list[dict[str, str]] = []
    inventors: list[dict[str, str]] = []
    consumed: set[int] = set()
    current_role = ""
    i = 0
    while i < len(lines):
        idx, line = lines[i]
        n = _plain_norm(line)
        if any(x in n for x in ["hak sahibi", "basvuru sahibi", "basvuru sahibinin"]):
            current_role = "applicant"
        if any(x in n for x in ["bulus sahibi", "buluscu", "mucit", "bulusu yapan"]):
            current_role = "inventor"

        cells = [c.strip() for c in re.split(r"\t+|\s+\|\s+", line) if c.strip()]
        fields = [_table_header_field(c) for c in cells]
        field_count = sum(1 for f in fields if f)
        inferred = current_role
        joined = " ".join(_plain_norm(c) for c in cells)
        if any(x in joined for x in ["sahip turu", "vergi no", "vergi kimlik", "hak sahibi adres"]):
            inferred = "applicant"
        if any(x in joined for x in ["dogum tarihi", "cinsiyet"]):
            inferred = "inventor"

        if field_count >= 3 and inferred in {"applicant", "inventor"}:
            j = i + 1
            while j < min(len(lines), i + 5):
                jidx, candidate = lines[j]
                if not candidate or candidate.startswith("[["):
                    j += 1
                    continue
                cn = _plain_norm(candidate)
                if any(x in cn for x in ["not basvuru sahibinin", "cevap verilmemesi"]):
                    j += 1
                    continue
                vals = [c.strip() for c in re.split(r"\t+|\s+\|\s+", candidate)]
                if sum(1 for v in vals if v) < 2 or abs(len(vals) - len(cells)) > 2:
                    j += 1
                    continue
                if len(vals) < len(cells):
                    vals += [""] * (len(cells) - len(vals))
                row = _new_person("", source, applicant=inferred == "applicant")
                for col, field in enumerate(fields):
                    if not field or col >= len(vals):
                        continue
                    value = _clean_value(vals[col])
                    if not value or _is_form_header_bundle(value) or _is_instruction_text(value):
                        continue
                    if field == "name":
                        if not _is_bad_person_value("name", value, applicant=inferred == "applicant"):
                            row["name"] = value
                            if inferred == "applicant":
                                row["entity_type"] = _entity_type(value, row.get("entity_type", ""))
                    elif field == "gender":
                        continue
                    else:
                        _apply_person_field(row, field, value, applicant=inferred == "applicant")
                _fill_country_from_explicit_location(row)
                if row.get("name") or row.get("identity") or row.get("address") or row.get("email"):
                    (applicants if inferred == "applicant" else inventors).append(row)
                    consumed.update({idx, jidx})
                i = j
                break
        i += 1
    return applicants, inventors, consumed


def _reconcile_role_rows(result: dict[str, Any]) -> None:
    """Açık rol tablosuna aykırı hayalet/çapraz kayıtları temizler."""
    apps = [dict(x) for x in result.get("applicants") or []]
    invs = [dict(x) for x in result.get("inventors") or []]
    inv_ids = {
        re.sub(r"\D", "", x.get("identity", ""))
        for x in invs
        if len(re.sub(r"\D", "", x.get("identity", ""))) == 11
    }
    cleaned_apps: list[dict[str, str]] = []
    for row in apps:
        ident = re.sub(r"\D", "", row.get("identity", ""))
        bad_name = not row.get("name") or _is_bad_person_value("name", row.get("name", ""), applicant=True)
        if ident in inv_ids and bad_name:
            row["identity"] = ""
        if bad_name and not row.get("identity") and not row.get("address") and not row.get("email"):
            continue
        cleaned_apps.append(row)
    result["applicants"] = cleaned_apps
    result["inventors"] = invs


def extract_application_information_rule_based(
    source_blocks: list[tuple[str, str]], *, specification_text: str = "", specification_filename: str = ""
) -> dict[str, Any]:
    """Başvuru verilerini OpenAI/API kullanmadan belge yapısı + açık metin kurallarıyla çıkarır.

    Bilgi uydurmaz. Tablo hücreleri, bölüm başlıkları, etiket/değer satırları ve açık
    e-posta cümleleri desteklenir. Bulunamayan alan ön kontrol kapısı tarafından bloke edilir.
    """
    result: dict[str, Any] = {
        "application_kind": "",
        "reference": "",
        "invention_title": "",
        "applicants": [],
        "inventors": [],
        "priority": {"status": "Belirsiz", "country": "", "number": "", "date": "", "source": ""},
        "filing_options": {
            "inventor_hidden": {"label": "Buluşçu bilgileri gizlensin mi?", "status": "Belirsiz", "source": "", "explicit": False},
            "public_project": {"label": "Kamu destekli proje kapsamında mı?", "status": "Belirsiz", "institution": "", "project_number": "", "source": "", "explicit": False},
            "early_publication": {"label": "Erken yayın talep ediliyor mu?", "status": "Belirsiz", "source": "", "explicit": False},
        },
        "other_information": [],
        "conflicts": [],
        "source_files_used": [],
        "field_sources": {},
    }
    seen: dict[str, tuple[str, str]] = {}

    def ensure_person(role: str, source: str, *, force_new: bool = False) -> dict[str, str]:
        target = result["applicants"] if role == "applicant" else result["inventors"]
        if force_new or not target:
            row = _new_person("", source, applicant=role == "applicant")
            target.append(row)
            return row
        return target[-1]

    for source, raw_text in source_blocks:
        text = (raw_text or "").replace("\r", "\n")
        result["source_files_used"].append(source)
        _extract_contact_information(result, text, source)
        _extract_filing_options(result, text, source)

        structured_apps, structured_invs, structured_consumed = _structured_role_rows(text, source)
        for row in structured_apps:
            _append_unique_person(result["applicants"], row, applicant=True)
        for row in structured_invs:
            _append_unique_person(result["inventors"], row, applicant=False)

        current_role = ""
        current_person: dict[str, str] | None = None
        pending_label = ""

        lines = [x.strip(" \r") for x in text.splitlines()]
        for line_idx, line in enumerate(lines):
            if line_idx in structured_consumed:
                continue
            if not line or line.startswith("[[TABLO") or line.startswith("[[E-POSTA EKI"):
                continue

            whole_label = _canonical_label(line)
            if whole_label in {"applicant", "inventor"} and _is_role_section_heading(line, whole_label):
                current_role = whole_label
                current_person = None
                pending_label = ""
                continue

            pairs = _line_label_value_pairs(line)
            if not pairs and pending_label:
                # Başlık bir satır, değer sonraki satır. Ancak sonraki satır başka
                # bir tablo/bölüm başlığıysa onu değer sanma.
                if not _canonical_label(line):
                    pairs = [(pending_label, pending_label, _clean_value(line))]
                    pending_label = ""
            if not pairs:
                nline = _plain_norm(line)
                if "ruchan" in nline and result["priority"]["status"] == "Belirsiz":
                    status = _priority_status(line)
                    if status != "Belirsiz":
                        _set_singleton(result, "priority.status", status, source, seen)
                if not result["application_kind"] and any(x in nline for x in ["patent basvuru", "faydali model basvuru", "faydali model olarak", "patent olarak"]):
                    kind = _kind_from_text(line)
                    if kind:
                        _set_singleton(result, "application_kind", kind, source, seen)
                continue

            for raw_label, label, value in pairs:
                # Bölüm başlığı: "HAK SAHİBİ BİLGİLERİ" / "BULUŞU YAPAN BİLGİLERİ".
                if label in {"applicant", "inventor"} and not value and _is_role_section_heading(raw_label, label):
                    current_role = label
                    current_person = None
                    pending_label = ""
                    continue

                # Tek satır başlık + sonraki satır değer formatı.
                if not value and label in {
                    "application_kind", "reference", "invention_title", "applicant", "inventor",
                    "priority", "priority_country", "priority_number", "priority_date",
                    "applicant_name", "inventor_name", "applicant_identity", "inventor_identity",
                    "applicant_address", "inventor_address", "applicant_country", "inventor_country",
                    "applicant_city", "inventor_city", "applicant_district", "inventor_district",
                    "applicant_email", "inventor_email", "applicant_phone", "inventor_phone",
                    "applicant_birth_date", "inventor_birth_date", "name", "identity", "address", "country", "city",
                    "district", "email", "phone", "birth_date", "entity_type", "gender",
                }:
                    pending_label = label
                    if label.startswith("applicant") or label == "applicant":
                        current_role = "applicant"
                    elif label.startswith("inventor") or label == "inventor":
                        current_role = "inventor"
                    continue

                if label == "application_kind":
                    kind = _kind_from_text(value)
                    if kind:
                        _set_singleton(result, "application_kind", kind, source, seen)
                    continue
                if label == "reference":
                    _set_singleton(result, "reference", value, source, seen)
                    continue
                if label == "invention_title":
                    _set_singleton(result, "invention_title", value, source, seen)
                    continue
                if label == "priority":
                    status = _priority_status(value)
                    if status != "Belirsiz":
                        _set_singleton(result, "priority.status", status, source, seen)
                    current_role, current_person = "priority", None
                    continue
                if label == "priority_country":
                    _set_singleton(result, "priority.country", value, source, seen)
                    current_role, current_person = "priority", None
                    continue
                if label == "priority_number":
                    _set_singleton(result, "priority.number", value, source, seen)
                    current_role, current_person = "priority", None
                    continue
                if label == "priority_date":
                    _set_singleton(result, "priority.date", value, source, seen)
                    current_role, current_person = "priority", None
                    continue

                if label in {"applicant", "inventor"}:
                    current_role = label
                    current_person = None
                    for person_name in _split_people_value(value):
                        row = _new_person(person_name, source, applicant=label == "applicant")
                        target = result["applicants"] if label == "applicant" else result["inventors"]
                        current_person = _append_unique_person(target, row, applicant=label == "applicant")
                    continue

                if label.startswith("applicant_") or label.startswith("inventor_"):
                    role, field = label.split("_", 1)
                    current_role = role
                    target = result["applicants"] if role == "applicant" else result["inventors"]
                    if field != "entity_type" and _is_bad_person_value(field, value, applicant=role == "applicant"):
                        # Form başlığı/açıklama gerçek kişi verisi değildir.
                        continue
                    if field == "name":
                        # Önce TCKN/doğum tarihi gibi alanları gelen aynı kişinin boş adını
                        # yeni satır açmadan tamamla. Ancak dolu ve farklı bir ad varsa yeni kişi başlat.
                        if current_person is not None and current_person in target and not current_person.get("name"):
                            current_person["name"] = _clean_value(value)
                            if role == "applicant" and not current_person.get("entity_type"):
                                current_person["entity_type"] = _entity_type(value)
                        else:
                            if current_person is not None and current_person.get("name") and _plain_norm(current_person.get("name", "")) != _plain_norm(value):
                                current_person = None
                            row = _new_person(value, source, applicant=role == "applicant")
                            current_person = _append_unique_person(target, row, applicant=role == "applicant")
                    else:
                        if current_person is None or current_person not in target:
                            current_person = target[-1] if target else ensure_person(role, source)
                        _apply_person_field(current_person, field, value, applicant=role == "applicant")
                    continue

                if label in {"name", "identity", "address", "country", "location", "city", "district", "email", "phone", "birth_date", "entity_type", "gender"} and current_role in {"applicant", "inventor"}:
                    if label == "gender":
                        if value and not _is_bad_person_value("gender", value, applicant=current_role == "applicant"):
                            prefix = "Hak sahibi" if current_role == "applicant" else "Buluş sahibi"
                            _add_other_information(result, f"{prefix} Cinsiyet", value, source)
                        continue
                    if label != "entity_type" and _is_bad_person_value(label, value, applicant=current_role == "applicant"):
                        continue
                    target = result["applicants"] if current_role == "applicant" else result["inventors"]
                    if label == "name":
                        if current_person is not None and current_person in target and not current_person.get("name"):
                            current_person["name"] = _clean_value(value)
                            if current_role == "applicant" and not current_person.get("entity_type"):
                                current_person["entity_type"] = _entity_type(value)
                        else:
                            if current_person is not None and current_person.get("name") and _plain_norm(current_person.get("name", "")) != _plain_norm(value):
                                current_person = None
                            row = _new_person(value, source, applicant=current_role == "applicant")
                            current_person = _append_unique_person(target, row, applicant=current_role == "applicant")
                    else:
                        if current_person is None or current_person not in target:
                            current_person = target[-1] if target else ensure_person(current_role, source)
                        _apply_person_field(current_person, label, value, applicant=current_role == "applicant")
                    continue

                # Tanınan fakat çekirdek kişi şemasına dahil olmayan alanları kaybetme.
                if label in {"fax", "postal_code", "authorized_person", "website"}:
                    prefix = "Hak sahibi" if current_role == "applicant" else "Buluş sahibi" if current_role == "inventor" else ""
                    nice = {"fax": "Faks", "postal_code": "Posta kodu", "authorized_person": "Yetkili", "website": "Web sitesi"}[label]
                    _add_other_information(result, f"{prefix} {nice}".strip(), value, source)
                    continue

                # Etiketli fakat çekirdek alanlara ait olmayan bilgileri ön kontrolde koru.
                if value:
                    _add_other_information(result, raw_label, value, source)

        # Serbest e-posta/yazı cümlelerinde yalnız açık rol ifadelerine izin ver.
        applicant_patterns = [
            r"(?:hak|başvuru)\s+sahibi(?:\s+olarak)?\s*(?::|-)?\s+(.{3,180}?)(?=\s+(?:olacaktır|olacak|olarak\s+belirlenmiştir|belirlenmiştir|dır|dir)\b|[\n;]|$)",
            r"başvuru\s+(.{3,180}?)\s+(?:adına|üzerinden)\s+yapılacaktır",
        ]
        inventor_patterns = [
            r"(?:buluş\s+sahibi|buluşçu|mucit|buluşu\s+yapan)(?:\s+olarak)?\s*(?::|-)?\s+(.{3,160}?)(?=\s+(?:olacaktır|olacak|olarak\s+belirlenmiştir|belirlenmiştir|dır|dir)\b|[\n;]|$)",
        ]
        if not result["applicants"]:
            for pattern in applicant_patterns:
                match = re.search(pattern, text, flags=re.IGNORECASE)
                if match:
                    for name in _split_people_value(_clean_value(match.group(1)).strip(" ,")):
                        _append_unique_person(result["applicants"], _new_person(name, source, applicant=True), applicant=True)
                    if result["applicants"]:
                        break
        if not result["inventors"]:
            for pattern in inventor_patterns:
                match = re.search(pattern, text, flags=re.IGNORECASE)
                if match:
                    for name in _split_people_value(_clean_value(match.group(1)).strip(" ,")):
                        _append_unique_person(result["inventors"], _new_person(name, source, applicant=False), applicant=False)
                    if result["inventors"]:
                        break

    # DP referansı tarifname dosya adında zaten bulunuyorsa bunu ana kaynak olarak kullan.
    file_ref = reference_from_filename(specification_filename)
    if file_ref:
        if result["reference"] and _plain_norm(result["reference"]) != _plain_norm(file_ref):
            result["conflicts"].append(
                f"DP / dosya referansı: '{result['reference']}' (başvuru kaynağı) / '{file_ref}' (Tarifname dosya adı)"
            )
        elif not result["reference"]:
            result["reference"] = file_ref
            result["field_sources"]["reference"] = "Tarifname dosya adı"

    # Buluş başlığının tek otoritesi başvuru için yüklenen TARİFNAME'dir.
    # Beyan formu/e-posta içindeki eski veya kısa başlık hiçbir zaman çatışma
    # üretmez ve EPATS başlığını değiştiremez.
    spec_title = _specification_title(specification_text)
    if spec_title:
        source_title = result.get("invention_title", "")
        if source_title and _plain_norm(source_title) != _plain_norm(spec_title):
            _add_other_information(result, "Bilgi kaynağındaki buluş başlığı", source_title, result.get("field_sources", {}).get("invention_title", "Başvuru bilgi kaynağı"))
        result["invention_title"] = spec_title
        result["field_sources"]["invention_title"] = "Tarifname"
        # Daha önce singleton izleme tablosuna bilgi kaynağından yazılmış başlık
        # varsa aşağıdaki field_sources döngüsünün bunu ezmesini engelle.
        seen.pop("invention_title", None)

    # Adreste Türkiye açıkça yazıyorsa veya il alanı 81 Türkiye ilinden biriyse ülkeyi doldur.
    for rows in (result["applicants"], result["inventors"]):
        for row in rows:
            _fill_country_from_explicit_location(row)

    for field_name in ("application_kind", "reference", "invention_title"):
        if field_name in seen and field_name not in result["field_sources"]:
            result["field_sources"][field_name] = seen[field_name][1]
    if result["priority"].get("source"):
        result["field_sources"]["priority"] = result["priority"]["source"]

    # Aynı kimlik numarasının farklı adla kullanılmasını işaretle.
    for role_key, label in (("applicants", "Hak sahibi"), ("inventors", "Buluş sahibi")):
        rows = result[role_key]
        by_id: dict[str, dict[str, str]] = {}
        for row in rows:
            digits = re.sub(r"\D", "", row.get("identity", ""))
            if not digits:
                continue
            if digits in by_id and _plain_norm(by_id[digits].get("name", "")) != _plain_norm(row.get("name", "")):
                result["conflicts"].append(f"{label} kimlik {digits} farklı adlarla bulundu.")
            else:
                by_id[digits] = row

    _reconcile_role_rows(result)
    _apply_filing_option_defaults(result)
    result["conflicts"] = list(dict.fromkeys(x for x in result["conflicts"] if x))
    return normalize_application_information(result)


# -----------------------------------------------------------------------------
# HİBRİT / YEREL AI DESTEKLİ BAŞVURU BİLGİSİ ÇIKARIMI
# -----------------------------------------------------------------------------

_LOCAL_AI_DEFAULT_MODEL = "/opt/models/Qwen2.5-0.5B-Instruct-IQ2_XS.gguf"
_LOCAL_AI_EMAIL_RE = re.compile(r"(?i)\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b")


def _local_ai_schema() -> dict[str, Any]:
    person_properties = {
        "identity": {"type": "string"},
        "name": {"type": "string"},
        "country": {"type": "string"},
        "city": {"type": "string"},
        "district": {"type": "string"},
        "address": {"type": "string"},
        "email": {"type": "string"},
        "phone": {"type": "string"},
        "birth_date": {"type": "string"},
        "source": {"type": "string"},
    }
    applicant_properties = dict(person_properties)
    applicant_properties["entity_type"] = {"type": "string"}
    option_properties = {
        "status": {"type": "string"},
        "source": {"type": "string"},
    }
    return {
        "type": "object",
        "properties": {
            "applicants": {
                "type": "array",
                "items": {"type": "object", "properties": applicant_properties, "additionalProperties": False},
                "maxItems": 6,
            },
            "inventors": {
                "type": "array",
                "items": {"type": "object", "properties": person_properties, "additionalProperties": False},
                "maxItems": 12,
            },
            "filing_options": {
                "type": "object",
                "properties": {
                    "inventor_hidden": {"type": "object", "properties": option_properties, "additionalProperties": False},
                    "public_project": {
                        "type": "object",
                        "properties": {
                            **option_properties,
                            "institution": {"type": "string"},
                            "project_number": {"type": "string"},
                        },
                        "additionalProperties": False,
                    },
                    "early_publication": {"type": "object", "properties": option_properties, "additionalProperties": False},
                },
                "additionalProperties": False,
            },
        },
        "required": ["applicants", "inventors", "filing_options"],
        "additionalProperties": False,
    }


def _local_ai_excerpt(source_blocks: list[tuple[str, str]], *, max_chars: int = 3200) -> str:
    """Yerel modele yalnız başvuru kişileri/tercihleriyle ilgili yoğunlaştırılmış metni verir.

    Küçük modelin bağlamını form açıklamalarıyla doldurmamak için ilgili satırların
    çevresi seçilir. Kaynak adı her blokta korunur.
    """
    needles = (
        "hak sahibi", "başvuru sahibi", "başvuru sahib", "buluş sahibi", "buluşçu", "mucit", "buluşu yapan",
        "unvan", "ad soyad", "adı soyadı", "adı/unvan", "tckn", "tc kimlik", "vkn", "vergi",
        "adres", "uyruk", "ülke", "ilçe", "telefon", "e-posta", "eposta", "email", "doğum",
        "gizlensin", "erken yayın", "erken yayım", "tübitak", "kosgeb", "kamu kurum", "proje",
        "@",
    )
    chunks: list[str] = []
    remaining = max_chars
    for source, text in source_blocks:
        if remaining <= 0:
            break
        lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
        selected: set[int] = set()
        for i, line in enumerate(lines):
            n = _plain_norm(line)
            if any(_plain_norm(x) in n for x in needles) or "@" in line:
                for j in range(max(0, i - 2), min(len(lines), i + 4)):
                    selected.add(j)
        if not selected:
            # Yapısı bozulmuş OCR'da anahtar kelimeler kaybolmuş olabilir; küçük
            # bir başlangıç bölümü yine modele verilir.
            selected.update(range(min(len(lines), 35)))
        body = "\n".join(lines[i] for i in sorted(selected))
        block = f"--- KAYNAK: {source} ---\n{body}\n"
        block = block[:remaining]
        if block.strip():
            chunks.append(block)
            remaining -= len(block)
    return "\n".join(chunks).strip()


def build_local_ai_application_prompt(source_blocks: list[tuple[str, str]]) -> str:
    excerpt = _local_ai_excerpt(source_blocks)
    return f"""Patent/faydalı model başvurusu için aşağıdaki kaynaklardan yapılandırılmış bilgi çıkar.

KURALLAR:
- Yalnız kaynakta AÇIKÇA yazan bilgileri kullan. Tahmin etme, tamamlamaya çalışma.
- Form başlıklarını, açıklama/not metnini, 'İmza', 'Yetkili', 'Cinsiyet', 'Doğum Tarihi' gibi ETİKETLERİ kişi adı veya şirket unvanı sanma.
- Hak/başvuru sahibi şirketse TAM ticaret unvanını al (örn. 'TT MOBİL İLETİŞİM HİZMETLERİ ANONİM ŞİRKETİ').
- Buluş sahibinde gerçek kişinin gerçek AD SOYAD bilgisini al. E-posta imza bloğundaki ad ancak kaynakta buluş sahibi/buluşçu olarak açıkça ilişkilendirilmişse kullanılabilir.
- E-posta alanına yalnız gerçek e-posta adresini yaz; 'İmza' veya başka metin ekleme.
- Telefon alanına yalnız telefon numarasını yaz.
- TCKN/VKN alanına yalnız kimlik/vergi numarasını yaz.
- Aynı kişiye ait adres, il, ilçe, ülke, telefon, e-posta ve doğum tarihini aynı kayıtta birleştir.
- source alanına bilginin geldiği KAYNAK adını aynen yaz.
- filing_options için yalnız açık EVET/HAYIR cevaplarını al. Cevap yoksa status boş string olsun.
- Çıktıda yalnız aşağıdaki yapıda JSON olsun; bulunmayan alanı boş bırak:
{{"applicants":[{{"entity_type":"","identity":"","name":"","country":"","city":"","district":"","address":"","email":"","phone":"","source":""}}],"inventors":[{{"identity":"","name":"","country":"","city":"","district":"","address":"","email":"","phone":"","birth_date":"","source":""}}],"filing_options":{{"inventor_hidden":{{"status":"","source":""}},"public_project":{{"status":"","institution":"","project_number":"","source":""}},"early_publication":{{"status":"","source":""}}}}}}

KAYNAKLAR:
{excerpt}
""".strip()


def _parse_first_json_object(text: str) -> dict[str, Any]:
    raw = (text or "").strip()
    if not raw:
        return {}
    try:
        obj = json.loads(raw)
        return obj if isinstance(obj, dict) else {}
    except Exception:
        pass
    start = raw.find("{")
    if start < 0:
        return {}
    depth = 0
    in_string = False
    escaped = False
    for idx in range(start, len(raw)):
        ch = raw[idx]
        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_string = False
            continue
        if ch == '"':
            in_string = True
        elif ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                try:
                    obj = json.loads(raw[start: idx + 1])
                    return obj if isinstance(obj, dict) else {}
                except Exception:
                    return {}
    return {}


def _run_local_ai_cli(prompt: str, schema: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any]]:
    """Kredisiz yerel GGUF modeli llama.cpp ile tek seferlik çalıştırır."""
    if str(os.getenv("LOCAL_AI_DISABLED", "")).strip().lower() in {"1", "true", "yes", "on"}:
        return {}, {"used": False, "available": False, "warning": "Yerel AI ortam değişkeni ile kapalı."}
    binary = shutil.which(os.getenv("LOCAL_AI_COMMAND", "llama"))
    model_path = os.getenv("LOCAL_AI_MODEL_PATH", _LOCAL_AI_DEFAULT_MODEL)
    if not binary:
        return {}, {"used": False, "available": False, "warning": "Yerel AI çalıştırıcısı (llama.cpp) bulunamadı; kurallı çıkarımla devam edildi."}
    if not Path(model_path).is_file():
        return {}, {"used": False, "available": False, "warning": f"Yerel AI modeli bulunamadı: {model_path}; kurallı çıkarımla devam edildi."}

    # Qwen2.5 ChatML. Küçük modelde düşünme/serbest sohbet yerine doğrudan JSON çıkarımı.
    chat_prompt = (
        "<|im_start|>system\n"
        "Sen patent başvuru belgelerinden veri çıkaran dikkatli bir bilgi çıkarım motorusun. "
        "Kaynakta olmayan hiçbir bilgiyi üretme. Yalnız JSON üret.<|im_end|>\n"
        "<|im_start|>user\n" + prompt + "<|im_end|>\n"
        "<|im_start|>assistant\n"
    )
    command = [
        binary, "cli", "-m", model_path,
        "-c", str(int(os.getenv("LOCAL_AI_CONTEXT", "3072"))),
        "-n", str(int(os.getenv("LOCAL_AI_MAX_TOKENS", "1100"))),
        "--temp", "0",
        "--no-display-prompt",
        "--simple-io",
        "-j", json.dumps(schema, ensure_ascii=False, separators=(",", ":")),
        "-p", chat_prompt,
    ]
    try:
        proc = subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=int(os.getenv("LOCAL_AI_TIMEOUT_SECONDS", "180")),
            check=False,
        )
    except Exception as exc:
        return {}, {"used": False, "available": True, "warning": f"Yerel AI çalıştırılamadı: {exc}"}
    if proc.returncode != 0:
        err = proc.stderr.decode("utf-8", errors="replace").strip()
        return {}, {"used": False, "available": True, "warning": "Yerel AI hata verdi; kurallı çıkarımla devam edildi." + (f" ({err[-240:]})" if err else "")}
    output = proc.stdout.decode("utf-8", errors="replace")
    parsed = _parse_first_json_object(output)
    if not parsed:
        return {}, {"used": False, "available": True, "warning": "Yerel AI geçerli JSON üretemedi; kurallı çıkarımla devam edildi."}
    return parsed, {
        "used": True,
        "available": True,
        "model": Path(model_path).name,
        "warning": "",
    }


def _norm_evidence(text: str) -> str:
    return _plain_norm(re.sub(r"\s+", " ", text or ""))


def _find_value_source(value: str, source_blocks: list[tuple[str, str]]) -> str:
    nv = _norm_evidence(value)
    if not nv:
        return ""
    # Çok kısa EVET/HAYIR gibi değerlerde tek başına kaynak aramak anlamsızdır.
    if len(nv) < 4:
        return ""
    for source, text in source_blocks:
        nt = _norm_evidence(text)
        if nv in nt:
            return source
    return ""


def _verified_ai_value(field: str, value: Any, source_blocks: list[tuple[str, str]]) -> tuple[str, str]:
    value = _clean_value(str(value or ""))
    if not value:
        return "", ""
    if field == "identity":
        clean = _valid_identity(value)
        return (clean, _find_value_source(clean, source_blocks)) if clean else ("", "")
    if field == "email":
        m = _LOCAL_AI_EMAIL_RE.search(value)
        if not m:
            return "", ""
        clean = m.group(0)
        return clean, _find_value_source(clean, source_blocks)
    if field == "phone":
        # Kaynaktaki numarayı biçimini değiştirmeden yakala; en az 7 rakam.
        digits = re.sub(r"\D", "", value)
        if len(digits) < 7:
            return "", ""
        for source, text in source_blocks:
            compact = re.sub(r"\D", "", text or "")
            if digits in compact:
                return value, source
        return "", ""
    source = _find_value_source(value, source_blocks)
    if not source:
        return "", ""
    if field == "name" and _is_instruction_text(value):
        return "", ""
    return value, source


def _sanitize_person_row(row: dict[str, str], *, applicant: bool) -> None:
    if row.get("email"):
        m = _LOCAL_AI_EMAIL_RE.search(row.get("email", ""))
        row["email"] = m.group(0) if m else ""
    if row.get("phone"):
        # Başka metin telefon alanına taşmışsa yalnız telefon benzeri parçayı koru.
        matches = re.findall(r"(?<!\d)(?:\+?90\s*)?(?:\(?0?\d{3}\)?[ .-]*)?\d{3}[ .-]*\d{2}[ .-]*\d{2}(?!\d)", row.get("phone", ""))
        if matches:
            row["phone"] = _clean_value(matches[0])
        elif len(re.sub(r"\D", "", row.get("phone", ""))) < 7:
            row["phone"] = ""
    if row.get("name") and _is_bad_person_value("name", row["name"], applicant=applicant):
        row["name"] = ""
    _fill_country_from_explicit_location(row)


def _ai_person_rows(raw: Any, *, applicant: bool, source_blocks: list[tuple[str, str]]) -> list[dict[str, str]]:
    if not isinstance(raw, list):
        return []
    out: list[dict[str, str]] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        row = _new_person("", "", applicant=applicant)
        detected_sources: list[str] = []
        for field in ("identity", "name", "country", "city", "district", "address", "email", "phone", "birth_date"):
            value, src = _verified_ai_value(field, item.get(field), source_blocks)
            if value:
                row[field] = value
                if src and src not in detected_sources:
                    detected_sources.append(src)
        if applicant:
            # Tüzel/gerçek kişi türü tam unvandan güvenli biçimde hesaplanabilir.
            explicit_type = str(item.get("entity_type") or "")
            row["entity_type"] = _entity_type(row.get("name", ""), explicit_type)
        requested_source = str(item.get("source") or "").strip()
        known_names = {name for name, _ in source_blocks}
        if requested_source in known_names and requested_source not in detected_sources:
            detected_sources.append(requested_source)
        row["source"] = "; ".join(detected_sources)
        _sanitize_person_row(row, applicant=applicant)
        if row.get("name") or row.get("identity") or row.get("email"):
            out.append(row)
    return out


def _person_match_score(existing: dict[str, str], candidate: dict[str, str]) -> int:
    score = 0
    e_id = re.sub(r"\D", "", existing.get("identity", ""))
    c_id = re.sub(r"\D", "", candidate.get("identity", ""))
    if e_id and c_id and e_id == c_id:
        score += 100
    if existing.get("email") and candidate.get("email") and _plain_norm(existing["email"]) == _plain_norm(candidate["email"]):
        score += 80
    if existing.get("name") and candidate.get("name") and _plain_norm(existing["name"]) == _plain_norm(candidate["name"]):
        score += 70
    es = {x for x in (existing.get("source") or "").split("; ") if x}
    cs = {x for x in (candidate.get("source") or "").split("; ") if x}
    if es & cs:
        score += 15
    return score


def _merge_ai_people(existing_rows: list[dict[str, str]], ai_rows: list[dict[str, str]], *, applicant: bool) -> list[dict[str, str]]:
    rows = [dict(x) for x in existing_rows]
    for row in rows:
        _sanitize_person_row(row, applicant=applicant)
    for ai_row in ai_rows:
        best: dict[str, str] | None = None
        best_score = -1
        for current in rows:
            score = _person_match_score(current, ai_row)
            if score > best_score:
                best, best_score = current, score
        # Aynı kaynakta tek kayıt varsa AI'nın rol ayrımı, bozuk tablo parser'ından
        # daha değerlidir. Bu yalnız kaynakta doğrulanmış AI değerleri için geçerlidir.
        if best_score < 15 and len(rows) == 1 and len(ai_rows) == 1:
            best = rows[0]
            best_score = 10
        if best is None or best_score < 10:
            rows.append(dict(ai_row))
            continue
        # AI yalnız kaynakta doğrulanan değerleri taşıdığı için boş/bozuk alanı düzeltmesine izin ver.
        if ai_row.get("name"):
            current_name = best.get("name", "")
            should_replace_name = (
                not current_name
                or _is_bad_person_value("name", current_name, applicant=applicant)
                or (
                    applicant
                    and _looks_like_company(ai_row["name"])
                    and _plain_norm(ai_row["name"]) != _plain_norm(current_name)
                    and len(ai_row["name"]) > len(current_name)
                    and bool(set((best.get("source") or "").split("; ")) & set((ai_row.get("source") or "").split("; ")))
                )
            )
            if should_replace_name:
                best["name"] = ai_row["name"]
        for field in ("identity", "country", "city", "district", "address", "email", "phone", "birth_date"):
            if ai_row.get(field) and (not best.get(field) or field in {"email", "phone"}):
                best[field] = ai_row[field]
        if applicant and ai_row.get("entity_type"):
            best["entity_type"] = ai_row["entity_type"]
        sources = [x for x in (best.get("source") or "").split("; ") if x]
        for src in (ai_row.get("source") or "").split("; "):
            if src and src not in sources:
                sources.append(src)
        best["source"] = "; ".join(sources)
        _sanitize_person_row(best, applicant=applicant)

    # Salt açıklama/form başlığı olan ve hiçbir güvenilir kimlik/e-posta taşımayan hayalet satırları temizle.
    cleaned: list[dict[str, str]] = []
    for row in rows:
        _sanitize_person_row(row, applicant=applicant)
        if row.get("name") or row.get("identity") or row.get("email") or row.get("address"):
            cleaned.append(row)
    return cleaned


def _merge_local_ai_information(rule_data: dict[str, Any], ai_data: dict[str, Any], source_blocks: list[tuple[str, str]]) -> dict[str, Any]:
    merged = normalize_application_information(rule_data)
    ai_applicants = _ai_person_rows(ai_data.get("applicants"), applicant=True, source_blocks=source_blocks)
    ai_inventors = _ai_person_rows(ai_data.get("inventors"), applicant=False, source_blocks=source_blocks)
    merged["applicants"] = _merge_ai_people(merged.get("applicants") or [], ai_applicants, applicant=True)
    merged["inventors"] = _merge_ai_people(merged.get("inventors") or [], ai_inventors, applicant=False)

    # Başvuru tercihlerini de hibrit olarak tamamla; ancak yalnız kaynakta soru ile
    # birlikte açık cevap varsa yerel AI sonucunu kabul et. Kurallı parser'ın açık
    # cevabı her zaman korunur.
    raw_opts = ai_data.get("filing_options") if isinstance(ai_data.get("filing_options"), dict) else {}
    options = merged.get("filing_options") or {}
    question_needles = {
        "inventor_hidden": ["gizlensin", "gizlen"],
        "public_project": ["tubitak", "tübitak", "kosgeb", "kamu kurum", "proje kapsam"],
        "early_publication": ["erken yayin", "erken yayın", "erken yayim", "erken yayım"],
    }
    all_text = "\n".join(text for _, text in source_blocks)
    all_norm = _plain_norm(all_text)
    for key, needles in question_needles.items():
        current = options.get(key) or {}
        ai_opt = raw_opts.get(key) if isinstance(raw_opts.get(key), dict) else {}
        status = str(ai_opt.get("status") or "").strip().casefold()
        status_norm = _plain_norm(status)
        mapped = "Evet" if status_norm == "evet" else "Hayır" if status_norm in {"hayir", "hayır"} else ""
        has_question = any(_plain_norm(x) in all_norm for x in needles)
        # current explicit ise değiştirme; AI sadece varsayılanı açık cevaba yükseltebilir.
        if mapped and has_question and not current.get("explicit"):
            src = str(ai_opt.get("source") or "").strip()
            known = {name for name, _ in source_blocks}
            if src not in known:
                # Soruyu içeren ilk kaynağı bul.
                src = next((name for name, text in source_blocks if any(_plain_norm(x) in _plain_norm(text) for x in needles)), "")
            current["status"] = mapped
            current["source"] = src or current.get("source") or "Yerel AI (kaynak doğrulandı)"
            current["explicit"] = True
            if key == "public_project" and mapped == "Evet":
                for f in ("institution", "project_number"):
                    val, _ = _verified_ai_value(f, ai_opt.get(f), source_blocks)
                    if val:
                        current[f] = val
        options[key] = current
    merged["filing_options"] = options
    return normalize_application_information(merged)



# -----------------------------------------------------------------------------
# TARAYICI CPU/WASM NER SONUCU -> BASVURU ALANLARI
# -----------------------------------------------------------------------------
_CPU_NER_EMAIL_RE = re.compile(r"[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}", re.IGNORECASE)
_CPU_NER_PHONE_LABEL_RE = re.compile(
    r"(?:cep(?:\s*telefonu)?|gsm|mobil\s*telefon|telefon|tel\.?)(?:\s*(?:no|numarasi|numarası))?\s*[:|\-]?\s*"
    r"(\+?\d[\d\s().\-/]{5,}\d)", re.IGNORECASE,
)
_CPU_NER_BIRTH_RE = re.compile(
    r"(?:dogum|doğum)\s*tarihi\s*[:|\-]?\s*(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})", re.IGNORECASE,
)
_CPU_NER_ADDRESS_RE = re.compile(
    r"(?:hak\s*sahibi\s*adresi|başvuru\s*sahibi\s*adresi|basvuru\s*sahibi\s*adresi|buluş\s*sahibi\s*adresi|bulus\s*sahibi\s*adresi|adres)"
    r"\s*[:|\-]?\s*([^\n]{5,320})", re.IGNORECASE,
)
_CPU_NER_TCKN_RE = re.compile(
    r"(?:t\.?\s*c\.?\s*kimlik(?:\s*(?:no|numarasi|numarası))?|tckn)\D{0,28}(\d{11})(?!\d)", re.IGNORECASE,
)
_CPU_NER_VKN_RE = re.compile(
    r"(?:vergi(?:\s*kimlik)?(?:\s*(?:no|numarasi|numarası))?|vkn)\D{0,28}(\d{10})(?!\d)", re.IGNORECASE,
)


def _cpu_ner_role(entity: dict[str, Any]) -> str:
    """NER varligini en yakin rol basligina baglar; tahmin icin tum belgeyi kullanmaz."""
    before = _plain_norm(str(entity.get("before") or ""))
    after = _plain_norm(str(entity.get("after") or ""))
    label = str(entity.get("label") or "").upper()
    applicant_terms = ["hak sahibi", "basvuru sahibi", "basvuran", "applicant", "patent sahibi"]
    inventor_terms = ["bulus sahibi", "buluscu", "mucit", "bulusu yapan", "inventor"]

    def _last_pos(text: str, terms: list[str]) -> int:
        return max([text.rfind(_plain_norm(t)) for t in terms] + [-1])

    ap = _last_pos(before, applicant_terms)
    inv = _last_pos(before, inventor_terms)
    # En yakin onceki rol basligi birincil sinyaldir.
    if ap >= 0 or inv >= 0:
        if ap >= 0 and inv < 0:
            return "applicant"
        if inv >= 0 and ap < 0:
            return "inventor"
        if abs(ap - inv) > 8:
            return "applicant" if ap > inv else "inventor"
    # Tablo/OCR'da etiket degerden sonra kalmissa yalniz cok yakin sonraki etikete bak.
    after_head = after[:220]
    ap2 = min([after_head.find(_plain_norm(t)) for t in applicant_terms if after_head.find(_plain_norm(t)) >= 0] + [9999])
    inv2 = min([after_head.find(_plain_norm(t)) for t in inventor_terms if after_head.find(_plain_norm(t)) >= 0] + [9999])
    if min(ap2, inv2) < 9999:
        if ap2 < 9999 and inv2 == 9999:
            return "applicant"
        if inv2 < 9999 and ap2 == 9999:
            return "inventor"
        if abs(ap2 - inv2) > 8:
            return "applicant" if ap2 < inv2 else "inventor"
    # Açık rol başlığı yoksa NER tek başına kayıt yaratmaz.
    # ORG=hak sahibi veya PER=buluşçu varsayımı yanlış rol karışmasına yol açar.
    return ""


def _cpu_ner_nearest_match(pattern: re.Pattern[str], context: str, center: int, *, group: int = 1) -> str:
    candidates: list[tuple[int, str]] = []
    for m in pattern.finditer(context):
        try:
            value = _clean_value(m.group(group))
        except Exception:
            continue
        if not value:
            continue
        mid = (m.start() + m.end()) // 2
        candidates.append((abs(mid - center), value))
    if not candidates:
        return ""
    candidates.sort(key=lambda x: x[0])
    return candidates[0][1]


def _cpu_ner_context_fields(entity: dict[str, Any], *, applicant: bool) -> dict[str, str]:
    before = str(entity.get("before") or "")[-1200:]
    word = _clean_value(str(entity.get("word") or ""))
    after = str(entity.get("after") or "")[:1700]
    context = before + "\n" + word + "\n" + after
    center = len(before) + 1 + len(word) // 2
    row: dict[str, str] = {}

    # Kimlik numarasi yalniz acik TCKN/VKN etiketiyle eslesirse baglanir.
    if applicant:
        vkn = _cpu_ner_nearest_match(_CPU_NER_VKN_RE, context, center)
        tckn = _cpu_ner_nearest_match(_CPU_NER_TCKN_RE, context, center)
        row["identity"] = vkn or tckn
    else:
        row["identity"] = _cpu_ner_nearest_match(_CPU_NER_TCKN_RE, context, center)

    # E-posta icin tum contextte aday olabilir; merkezdeki kisi/kuruma en yakinini al.
    emails: list[tuple[int, str]] = []
    for m in _CPU_NER_EMAIL_RE.finditer(context):
        emails.append((abs(((m.start() + m.end()) // 2) - center), m.group(0)))
    if emails:
        emails.sort(key=lambda x: x[0])
        # Cok uzaktaki imza/onceki blok e-postasini baglama.
        if emails[0][0] <= 950:
            row["email"] = emails[0][1]

    # Telefon yalniz Telefon/Cep/GSM etiketiyle yakalanir; belgedeki ilk sayi asla alinmaz.
    phone = _cpu_ner_nearest_match(_CPU_NER_PHONE_LABEL_RE, context, center)
    if phone:
        row["phone"] = phone
    # Adres yalniz acik Adres/Hak Sahibi Adresi/Bulus Sahibi Adresi etiketiyle baglanir.
    address = _cpu_ner_nearest_match(_CPU_NER_ADDRESS_RE, context, center)
    if address:
        # Ayni satirda sonraki tablo etiketi de gelmisse orada kes.
        address = re.split(
            r"\s+(?=(?:e[ -]?posta|email|telefon|tel\.?|cep|gsm|tckn|tc\s*kimlik|vkn|vergi|doğum\s*tarihi|dogum\s*tarihi|uyruk|ülke|ulke|ilçe|ilce)\b)",
            address, maxsplit=1, flags=re.IGNORECASE,
        )[0].strip(" \t|:;-")
        if len(address) >= 6:
            row["address"] = address
            # Acik Ilce/Il kalibi (örn. Umraniye/Istanbul) varsa bagla.
            locs = list(re.finditer(r"([A-Za-zÇĞİÖŞÜçğıöşü]+)\s*/\s*([A-Za-zÇĞİÖŞÜçğıöşü]+)", address))
            if locs:
                district, city = locs[-1].group(1), locs[-1].group(2)
                if _plain_norm(city) in _TURKEY_PROVINCES:
                    row["city"] = city
                    row["district"] = district
            if not row.get("city"):
                for token in reversed(re.findall(r"[A-Za-zÇĞİÖŞÜçğıöşü]+", address)):
                    if _plain_norm(token) in _TURKEY_PROVINCES:
                        row["city"] = token
                        break
            if "turkiye" in _plain_norm(address) or (row.get("city") and _plain_norm(row["city"]) in _TURKEY_PROVINCES):
                row["country"] = "Türkiye"
    if not applicant:
        birth = _cpu_ner_nearest_match(_CPU_NER_BIRTH_RE, context, center)
        if birth:
            row["birth_date"] = birth
    return row


def _cpu_ner_verified_entity_word(entity: dict[str, Any], source_text: str, *, applicant: bool) -> str:
    word = _clean_value(str(entity.get("word") or "").replace("##", ""))
    if not word or len(word) < 3:
        return ""
    label = str(entity.get("label") or "").upper()
    if applicant and label not in {"ORG", "PER"}:
        return ""
    if not applicant and label != "PER":
        return ""
    if _plain_norm(word) not in _plain_norm(source_text):
        return ""
    if _is_bad_person_value("name", word, applicant=applicant):
        return ""
    return word


def _cpu_ner_ai_payload(ner_data: dict[str, Any], source_blocks: list[tuple[str, str]]) -> dict[str, Any]:
    """CPU NER varliklarini mevcut guvenli AI-merge semasina cevirir."""
    by_source = {name: text for name, text in source_blocks}
    applicants: list[dict[str, str]] = []
    inventors: list[dict[str, str]] = []
    entities = ner_data.get("entities") if isinstance(ner_data, dict) else []
    if not isinstance(entities, list):
        entities = []

    # Ayni kaynak/rol/ad NER chunk overlap nedeniyle birden fazla kez gelebilir.
    combined: dict[tuple[str, str, str], dict[str, str]] = {}
    for entity in entities:
        if not isinstance(entity, dict):
            continue
        source = str(entity.get("source") or "").strip()
        if source not in by_source:
            continue
        score = float(entity.get("score") or 0.0)
        if score < 0.55:
            continue
        role = _cpu_ner_role(entity)
        if role not in {"applicant", "inventor"}:
            continue
        applicant = role == "applicant"
        name = _cpu_ner_verified_entity_word(entity, by_source[source], applicant=applicant)
        if not name:
            continue
        row = _new_person(name, source, applicant=applicant)
        row.update({k: v for k, v in _cpu_ner_context_fields(entity, applicant=applicant).items() if v})
        key = (source, role, _plain_norm(name))
        if key in combined:
            current = combined[key]
            for field, value in row.items():
                if value and not current.get(field):
                    current[field] = value
        else:
            combined[key] = row

    for (_, role, _), row in combined.items():
        if role == "applicant":
            applicants.append(row)
        else:
            inventors.append(row)
    # Filing options generatif AI gerektirmez; mevcut acik EVET/HAYIR parser'i korunur.
    return {"applicants": applicants, "inventors": inventors, "filing_options": {}}


def merge_verified_cpu_ner_application_information(
    rule_data: dict[str, Any], ner_data: dict[str, Any], source_blocks: list[tuple[str, str]]
) -> dict[str, Any]:
    """CPU NER yalnız açık rol bloklarındaki mevcut kaydı düzeltir.

    Yeni hak sahibi/buluşçu satırı yaratmaz. Telefon ve kimlik ancak açık
    Telefon/TCKN/VKN etiketiyle NER bağlamında yakalanmışsa yanlış mevcut değeri
    düzeltebilir; belgedeki serbest sayı asla taşınmaz.
    """
    merged = normalize_application_information(rule_data)
    candidates = _cpu_ner_ai_payload(ner_data if isinstance(ner_data, dict) else {}, source_blocks)
    source_map = {name: text for name, text in source_blocks}

    def _explicit_current_ok(field: str, value: str, source_text: str, *, applicant: bool) -> bool:
        if not value:
            return False
        if field == "phone":
            digits = re.sub(r"\D", "", value)
            return any(digits and digits == re.sub(r"\D", "", m.group(1)) for m in _CPU_NER_PHONE_LABEL_RE.finditer(source_text or ""))
        if field == "identity":
            digits = re.sub(r"\D", "", value)
            pats = [_CPU_NER_VKN_RE, _CPU_NER_TCKN_RE] if applicant else [_CPU_NER_TCKN_RE]
            return any(digits and digits == re.sub(r"\D", "", m.group(1)) for pat in pats for m in pat.finditer(source_text or ""))
        return True

    for key, applicant in (("applicants", True), ("inventors", False)):
        rows = [dict(x) for x in merged.get(key) or []]
        cands = candidates.get(key) if isinstance(candidates.get(key), list) else []
        for row in rows:
            _sanitize_person_row(row, applicant=applicant)
        for cand in cands:
            cname = str(cand.get("name") or "").strip()
            csrcs = {x for x in str(cand.get("source") or "").split("; ") if x}
            if not cname or not csrcs:
                continue
            eligible = [
                r for r in rows
                if csrcs & {x for x in str(r.get("source") or "").split("; ") if x}
            ]
            if len(eligible) != 1:
                continue
            row = eligible[0]
            if not row.get("name") or _is_bad_person_value("name", row.get("name", ""), applicant=applicant):
                row["name"] = cname
                if applicant:
                    row["entity_type"] = _entity_type(cname, row.get("entity_type", ""))

            source_text = "\n".join(source_map.get(x, "") for x in csrcs)
            for field in ("identity", "address", "email", "phone", "birth_date", "country", "city", "district"):
                val = str(cand.get(field) or "").strip()
                if not val:
                    continue
                current = str(row.get(field) or "").strip()
                if not current or (field in {"phone", "identity"} and not _explicit_current_ok(field, current, source_text, applicant=applicant)):
                    row[field] = val
            _sanitize_person_row(row, applicant=applicant)
        merged[key] = rows
    _reconcile_role_rows(merged)
    return normalize_application_information(merged)


def merge_verified_ai_application_information(
    rule_data: dict[str, Any], ai_data: dict[str, Any], source_blocks: list[tuple[str, str]]
) -> dict[str, Any]:
    """Tarayıcı/yerel AI çıktısını yalnız kaynakta doğrulanabilen değerlerle birleştirir."""
    merged = _merge_local_ai_information(rule_data, ai_data if isinstance(ai_data, dict) else {}, source_blocks)
    for row in merged.get("applicants") or []:
        _sanitize_person_row(row, applicant=True)
    for row in merged.get("inventors") or []:
        _sanitize_person_row(row, applicant=False)
    return normalize_application_information(merged)


def extract_application_information_hybrid(
    source_blocks: list[tuple[str, str]], *, specification_text: str = "", specification_filename: str = "",
    local_ai_runner: Any | None = None,
) -> dict[str, Any]:
    """Kurallı parser + kredisiz yerel küçük LLM ile güvenli hibrit çıkarım.

    Yerel AI yalnız kaynakta tekrar doğrulanabilen değerleri doldurabilir. Buluş
    başlığı ve DP referansı yine Tarifname'den/isimden deterministik olarak gelir.
    Yerel model çalışmazsa süreç durmaz; kurallı sonuç ve görünür uyarı döner.
    """
    rule_data = extract_application_information_rule_based(
        source_blocks,
        specification_text=specification_text,
        specification_filename=specification_filename,
    )
    runner = local_ai_runner or _run_local_ai_cli
    prompt = build_local_ai_application_prompt(source_blocks)
    try:
        ai_data, status = runner(prompt, _local_ai_schema())
    except Exception as exc:
        ai_data, status = {}, {"used": False, "available": False, "warning": f"Yerel AI devreye alınamadı: {exc}"}
    result = _merge_local_ai_information(rule_data, ai_data, source_blocks) if ai_data else normalize_application_information(rule_data)
    # AI devreye girmese bile parser'ın e-posta/telefon alanına taşıdığı "İmza" vb.
    # artıklar ön kontrolde görünmesin.
    for row in result.get("applicants") or []:
        _sanitize_person_row(row, applicant=True)
    for row in result.get("inventors") or []:
        _sanitize_person_row(row, applicant=False)
    # Tarifname başlığı ve dosya referansı hibrit birleştirme sonrasında da tek otorite olarak korunur.
    spec_title = _specification_title(specification_text)
    if spec_title:
        result["invention_title"] = spec_title
        result.setdefault("field_sources", {})["invention_title"] = "Tarifname"
    file_ref = reference_from_filename(specification_filename)
    if file_ref:
        result["reference"] = file_ref
        result.setdefault("field_sources", {})["reference"] = "Tarifname dosya adı"
    result["local_ai"] = status
    return result

def build_application_information_prompt(source_blocks: list[tuple[str, str]], *, specification_text: str = "") -> str:
    blocks = []
    for name, text in source_blocks:
        blocks.append(f"--- KAYNAK: {name} ---\n{text[:45000]}")
    if specification_text.strip():
        blocks.append(f"--- TARİFNAME (yalnız başlık/teyit kaynağı) ---\n{specification_text[:18000]}")
    joined = "\n\n".join(blocks)
    return f"""
Aşağıdaki patent/faydalı model başvuru kaynaklarından yalnız açıkça bulunan bilgileri çıkar.
Tahmin etme, şirket unvanını/adresi/kimlik numarasını uydurma, e-posta gönderen kişiyi otomatik olarak buluş sahibi veya hak sahibi sayma.
Bir bilgi birden fazla kaynakta çelişiyorsa en güvenilir açık ifadeyi seç ve conflicts listesinde çelişkiyi belirt.
Tarifname yalnız buluş başlığını teyit etmek için kullanılabilir; tarifnameden hak sahibi/buluş sahibi uydurma.
Bulabildiğin tüm diğer başvuruya ilişkin bilgileri other_information içinde de koru.

Yalnız aşağıdaki JSON şemasında cevap ver:
{{
  "application_kind": "Patent|Faydalı Model|",
  "reference": "",
  "invention_title": "",
  "applicants": [
    {{"entity_type":"Tüzel kişi|Gerçek kişi|", "identity":"", "name":"", "country":"", "city":"", "address":"", "source":""}}
  ],
  "inventors": [
    {{"identity":"", "name":"", "country":"", "city":"", "address":"", "source":""}}
  ],
  "priority": {{"status":"Var|Yok|Belirsiz", "country":"", "number":"", "date":"", "source":""}},
  "other_information": [{{"label":"", "value":"", "source":""}}],
  "conflicts": [""],
  "source_files_used": [""]
}}

KAYNAKLAR:
{joined}
""".strip()


def normalize_application_information(data: dict[str, Any] | None) -> dict[str, Any]:
    data = data if isinstance(data, dict) else {}
    app_kind = str(data.get("application_kind") or "").strip()
    if app_kind.lower().replace("ı", "i") not in {"patent", "faydali model", "faydalı model"}:
        app_kind = ""
    elif app_kind.lower() == "patent":
        app_kind = "Patent"
    else:
        app_kind = "Faydalı Model"

    def people(key: str, applicant: bool) -> list[dict[str, str]]:
        out: list[dict[str, str]] = []
        raw = data.get(key)
        if not isinstance(raw, list):
            return out
        for item in raw:
            if not isinstance(item, dict):
                continue
            row = {
                "identity": str(item.get("identity") or "").strip(),
                "name": str(item.get("name") or "").strip(),
                "country": str(item.get("country") or "").strip(),
                "city": str(item.get("city") or "").strip(),
                "district": str(item.get("district") or "").strip(),
                "address": str(item.get("address") or "").strip(),
                "email": str(item.get("email") or "").strip(),
                "phone": str(item.get("phone") or "").strip(),
                "birth_date": str(item.get("birth_date") or "").strip(),
                "source": str(item.get("source") or "").strip(),
            }
            if applicant:
                row["entity_type"] = str(item.get("entity_type") or "").strip()
            if any(row.values()):
                out.append(row)
        return out

    priority = data.get("priority") if isinstance(data.get("priority"), dict) else {}
    raw_opts = data.get("filing_options") if isinstance(data.get("filing_options"), dict) else {}
    def opt(key: str, label: str, *, project: bool = False) -> dict[str, Any]:
        src = raw_opts.get(key) if isinstance(raw_opts.get(key), dict) else {}
        status = str(src.get("status") or "Belirsiz").strip()
        if status not in {"Evet", "Hayır", "Belirsiz"}:
            status = "Belirsiz"
        source = str(src.get("source") or "").strip()
        explicit = bool(src.get("explicit"))
        if status == "Belirsiz":
            status = "Hayır"
            source = source or "Varsayılan (cevap verilmemiş)"
            explicit = False
        row: dict[str, Any] = {
            "label": label, "status": status, "source": source,
            "explicit": explicit,
        }
        if project:
            row["institution"] = str(src.get("institution") or "").strip()
            row["project_number"] = str(src.get("project_number") or "").strip()
        return row
    filing_options = {
        "inventor_hidden": opt("inventor_hidden", "Buluşçu bilgileri gizlensin mi?"),
        "public_project": opt("public_project", "Kamu destekli proje kapsamında mı?", project=True),
        "early_publication": opt("early_publication", "Erken yayın talep ediliyor mu?"),
    }
    other = data.get("other_information") if isinstance(data.get("other_information"), list) else []
    return {
        "application_kind": app_kind,
        "reference": str(data.get("reference") or "").strip(),
        "invention_title": str(data.get("invention_title") or "").strip(),
        "applicants": people("applicants", True),
        "inventors": people("inventors", False),
        "priority": {
            "status": str(priority.get("status") or "Belirsiz").strip() or "Belirsiz",
            "country": str(priority.get("country") or "").strip(),
            "number": str(priority.get("number") or "").strip(),
            "date": str(priority.get("date") or "").strip(),
            "source": str(priority.get("source") or "").strip(),
        },
        "filing_options": filing_options,
        "other_information": [
            {
                "label": str(x.get("label") or "").strip(),
                "value": str(x.get("value") or "").strip(),
                "source": str(x.get("source") or "").strip(),
            }
            for x in other if isinstance(x, dict) and (x.get("label") or x.get("value"))
        ],
        "conflicts": [str(x).strip() for x in (data.get("conflicts") or []) if str(x).strip()] if isinstance(data.get("conflicts"), list) else [],
        "source_files_used": [str(x).strip() for x in (data.get("source_files_used") or []) if str(x).strip()] if isinstance(data.get("source_files_used"), list) else [],
        "field_sources": {
            str(k): str(v).strip()
            for k, v in (data.get("field_sources") or {}).items()
            if isinstance(data.get("field_sources"), dict) and str(v).strip()
        },
    }


def application_precheck_missing(metadata: dict[str, Any], metrics: dict[str, Any], *, figures_required: bool = False) -> list[str]:
    missing: list[str] = []
    if not str(metadata.get("application_kind") or "").strip():
        missing.append("başvuru türü")
    if not str(metadata.get("invention_title") or "").strip():
        missing.append("buluş başlığı")
    applicants = metadata.get("applicants") or []
    if not applicants:
        missing.append("en az bir hak sahibi / başvuru sahibi")
    else:
        for i, applicant in enumerate(applicants, 1):
            if not str(applicant.get("name") or "").strip():
                missing.append(f"hak sahibi {i} unvan/ad soyad")
            if not str(applicant.get("country") or "").strip():
                missing.append(f"hak sahibi {i} ülke")
            if not str(applicant.get("address") or "").strip():
                missing.append(f"hak sahibi {i} adres")
    inventors = metadata.get("inventors") or []
    if not inventors:
        missing.append("en az bir buluş sahibi")
    else:
        for i, inventor in enumerate(inventors, 1):
            if not str(inventor.get("name") or "").strip():
                missing.append(f"buluş sahibi {i} ad soyad")
            if not str(inventor.get("country") or "").strip():
                missing.append(f"buluş sahibi {i} ülke")
            if not str(inventor.get("address") or "").strip():
                missing.append(f"buluş sahibi {i} adres")
    priority = metadata.get("priority") or {}
    priority_status = str(priority.get("status") or "").strip()
    if priority_status not in {"Var", "Yok"}:
        missing.append("rüçhan durumu (Var/Yok)")
    if priority_status == "Var":
        for field, label in [("country", "rüçhan ülkesi"), ("number", "rüçhan numarası"), ("date", "rüçhan tarihi")]:
            if not str(priority.get(field) or "").strip():
                missing.append(label)
    options = metadata.get("filing_options") or {}
    for key, label in [
        ("inventor_hidden", "buluşçu bilgilerinin gizlenme tercihi"),
        ("public_project", "kamu destekli proje durumu"),
        ("early_publication", "erken yayın tercihi"),
    ]:
        if str((options.get(key) or {}).get("status") or "").strip() not in {"Evet", "Hayır"}:
            missing.append(label)
    pub = options.get("public_project") or {}
    if pub.get("status") == "Evet":
        if not str(pub.get("institution") or "").strip():
            missing.append("kamu destekli proje kurumu")
        if not str(pub.get("project_number") or "").strip():
            missing.append("kamu destekli proje numarası")
    if int(metrics.get("specification_pages") or 0) <= 0:
        missing.append("Tarifname PDF")
    if int(metrics.get("claim_count") or 0) <= 0:
        missing.append("istemler / istem sayısı")
    if not metrics.get("abstract_present"):
        missing.append("özet")
    if figures_required and int(metrics.get("figures_pages") or 0) <= 0:
        missing.append("şekiller")
    return missing

def split_patent_docx(data: bytes, *, clean_template_colors: bool = True) -> dict[str, bytes]:
    """Tek tarifname DOCX'ini EPATS için Tarifname / İstemler / Özet DOCX bölümlerine ayırır."""
    if clean_template_colors:
        data = strip_template_colored_text(data)
    # LibreOffice bazı Word satır numaralarını sol kenarda "5, 10, 15..."
    # şeklinde EPATS PDF'ine taşır. Başvuru paketinde bunlar istenmediği için
    # yalnız bu otomatik section özelliğini kaldırıyoruz.
    data = remove_word_line_numbering(data)
    with zipfile.ZipFile(io.BytesIO(data), "r") as zin:
        root = ET.fromstring(zin.read("word/document.xml"))
    body = root.find(_W + "body")
    if body is None:
        raise ValueError("Word belgesinin gövdesi okunamadı.")
    content = [c for c in list(body) if c.tag != _W + "sectPr"]
    claims_idx = _find_heading_index(content, "İSTEMLER")
    abstract_idx = _find_heading_index(content, "ÖZET")
    if claims_idx <= 0 or abstract_idx <= claims_idx:
        raise ValueError("İSTEMLER / ÖZET başlık sırası beklenen yapıda değil.")
    return {
        "Tarifname.docx": _slice_docx(data, start_idx=0, end_idx=claims_idx),
        "Istemler.docx": _slice_docx(data, start_idx=claims_idx, end_idx=abstract_idx),
        "Ozet.docx": _slice_docx(data, start_idx=abstract_idx, end_idx=None),
    }


def _libreoffice_convert(data: bytes, filename: str, target_ext: str) -> bytes:
    if shutil.which("libreoffice") is None:
        raise RuntimeError("Belge dönüşümü için LibreOffice kurulu değil.")
    target_ext = target_ext.lstrip(".").lower()
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        src = td_path / Path(filename).name
        src.write_bytes(data)
        outdir = td_path / "converted"
        outdir.mkdir()
        proc = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", target_ext, "--outdir", str(outdir), str(src)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120,
        )
        outputs = list(outdir.glob(f"*.{target_ext}"))
        if proc.returncode != 0 or not outputs:
            err = proc.stderr.decode("utf-8", errors="replace").strip()
            raise RuntimeError(f"Belge {target_ext.upper()} biçimine dönüştürülemedi. {err}".strip())
        return outputs[0].read_bytes()


def _libreoffice_to_pdf(data: bytes, filename: str) -> bytes:
    return _libreoffice_convert(data, filename, "pdf")


def _libreoffice_to_docx(data: bytes, filename: str) -> bytes:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return data
    return _libreoffice_convert(data, filename, "docx")


def extract_specification_text(filename: str, data: bytes) -> str:
    """Tarifname kaynağından başlık/şekil kontrolünde kullanılacak düz metni çıkarır."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text).strip()
    if suffix == ".doc":
        # DOC'u DOCX'e dönüştürmek antiword'e göre Unicode/tablo biçimini daha iyi korur.
        docx = _libreoffice_to_docx(data, filename)
        return extract_specification_text(Path(filename).with_suffix(".docx").name, docx)
    if suffix == ".pdf":
        return "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(data)).pages).strip()
    raise ValueError("Tarifname dosyası DOC, DOCX veya PDF olmalıdır.")


def _fitz_span_rgb(color: int) -> tuple[int, int, int]:
    return ((int(color) >> 16) & 255, (int(color) >> 8) & 255, int(color) & 255)


def _pdf_epats_cleanup(data: bytes) -> bytes:
    """PDF'de kırmızı/mavi şablon metnini ve sol marj satır numaralarını görsel olarak kaldırır.

    Sayfa boyutu, fontlar ve kalan içerik yeniden dizilmez. PDF kaynaklarında DOCX gibi
    reflow yapılamayacağından kaldırılan alanlar beyaz boşluk olarak kalabilir.
    """
    import fitz  # pymupdf

    doc = fitz.open(stream=data, filetype="pdf")
    changed = False
    for page in doc:
        page_dict = page.get_text("dict")
        for block in page_dict.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = (span.get("text") or "").strip()
                    if not text:
                        continue
                    bbox = fitz.Rect(span.get("bbox"))
                    r, g, b = _fitz_span_rgb(span.get("color", 0))
                    is_template_color = (r >= 150 and g <= 130 and b <= 130) or (b >= 150 and r <= 130 and g <= 160)
                    is_left_line_number = (
                        bool(re.fullmatch(r"\d{1,3}", text))
                        and bbox.x0 < page.rect.width * 0.10
                        and 25 < bbox.y0 < page.rect.height - 25
                        and int(text) % 5 == 0
                    )
                    is_top_page_artifact = (
                        bool(re.fullmatch(r"\d{1,3}\s*[Xx]", text))
                        and bbox.x0 < page.rect.width * 0.18
                        and bbox.y0 < page.rect.height * 0.12
                    )
                    if is_template_color or is_left_line_number or is_top_page_artifact:
                        page.add_redact_annot(bbox + (-1, -1, 1, 1), fill=(1, 1, 1))
                        changed = True
        if changed:
            try:
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
            except TypeError:
                page.apply_redactions()
    if not changed:
        out = data
    else:
        out = doc.tobytes(garbage=4, deflate=True)
    doc.close()
    return out


def _find_pdf_heading(data: bytes, heading: str) -> tuple[int, float] | None:
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    target = _norm(heading)
    try:
        for page_index, page in enumerate(doc):
            info = page.get_text("dict")
            for block in info.get("blocks", []):
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    text = "".join((sp.get("text") or "") for sp in spans).strip()
                    if _norm(text) == target and spans:
                        y0 = min(float(sp.get("bbox", [0, 0, 0, 0])[1]) for sp in spans)
                        return page_index, y0
    finally:
        doc.close()
    return None


def _copy_pdf_vertical_range(src_doc, out_doc, page_index: int, y0: float, y1: float) -> None:
    import fitz

    page = src_doc[page_index]
    y0 = max(0.0, min(float(y0), page.rect.height))
    y1 = max(0.0, min(float(y1), page.rect.height))
    if y1 - y0 < 2:
        return
    new_page = out_doc.new_page(width=page.rect.width, height=page.rect.height)
    clip = fitz.Rect(0, y0, page.rect.width, y1)
    # İçeriği kaynak sayfadaki aynı koordinatlarda tut; yalnız diğer bölümün
    # bulunduğu alanı taşımayarak biçimi değiştirmeden ayır.
    new_page.show_pdf_page(clip, src_doc, page_index, clip=clip)


def split_patent_pdf(data: bytes) -> dict[str, bytes]:
    """Birleşik Tarifname/İstemler/Özet PDF'ini görsel biçimi bozmadan ayırır."""
    import fitz

    cleaned = _pdf_epats_cleanup(data)
    claims_pos = _find_pdf_heading(cleaned, "İSTEMLER")
    abstract_pos = _find_pdf_heading(cleaned, "ÖZET")
    if not claims_pos or not abstract_pos:
        raise ValueError("PDF içinde İSTEMLER ve ÖZET başlıkları bulunamadı.")
    if claims_pos[0] > abstract_pos[0] or (claims_pos[0] == abstract_pos[0] and claims_pos[1] >= abstract_pos[1]):
        raise ValueError("PDF içindeki İSTEMLER / ÖZET sırası beklenen yapıda değil.")

    src = fitz.open(stream=cleaned, filetype="pdf")
    outputs: dict[str, bytes] = {}
    sections = {
        "Tarifname.pdf": ((0, 0.0), claims_pos),
        "Istemler.pdf": (claims_pos, abstract_pos),
        "Ozet.pdf": (abstract_pos, (len(src) - 1, src[-1].rect.height)),
    }
    try:
        for name, (start, end) in sections.items():
            out = fitz.open()
            start_page, start_y = start
            end_page, end_y = end
            for page_index in range(start_page, end_page + 1):
                page = src[page_index]
                y0 = start_y if page_index == start_page else 0.0
                y1 = end_y if page_index == end_page else page.rect.height
                _copy_pdf_vertical_range(src, out, page_index, y0, y1)
            if len(out) == 0:
                raise ValueError(f"PDF içindeki {name} bölümü boş bulundu.")
            outputs[name] = out.tobytes(garbage=4, deflate=True)
            out.close()
    finally:
        src.close()
    return outputs


def count_claims_from_pdf(data: bytes) -> int:
    text = "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(data)).pages)
    explicit = [int(x) for x in re.findall(r"(?mi)^\s*(\d{1,3})\s*[.\-)]+\s+", text)]
    if explicit:
        # 1..N dizisi varsa maksimum istem numarası en güvenilir sayımdır.
        seq = sorted(set(x for x in explicit if 1 <= x <= 999))
        if seq and seq[0] == 1:
            return max(seq)
    deps = re.findall(r"(?mi)^\s*İstem\s+\d+(?:['’`]?e|['’`]?a)\s+uygun", text)
    if deps:
        return 1 + len(deps)
    return 0


def _ensure_pdf(data: bytes, filename: str) -> bytes:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        PdfReader(io.BytesIO(data))  # temel bütünlük doğrulaması
        return data
    if suffix in {".docx", ".doc"}:
        return _libreoffice_to_pdf(data, Path(filename).name)
    raise ValueError("Şekiller dosyası DOC, DOCX veya PDF olmalıdır.")


def build_epats_application_package(
    specification_data: bytes,
    *,
    specification_name: str = "Tarifname.docx",
    figures_data: bytes | None = None,
    figures_name: str | None = None,
    metadata: dict[str, Any] | None = None,
) -> tuple[bytes, dict[str, bytes]]:
    """EPATS'a yüklenmeye hazır PDF'leri üretip ZIP paket olarak döndürür.

    Tarifname kaynağı DOC/DOCX/PDF olabilir. DOC önce DOCX'e dönüştürülür;
    PDF ise görsel biçimi korunarak İSTEMLER/ÖZET sınırlarından ayrılır.
    """
    suffix = Path(specification_name).suffix.lower()
    pdfs: dict[str, bytes] = {}
    if suffix in {".docx", ".doc"}:
        docx = specification_data if suffix == ".docx" else _libreoffice_to_docx(specification_data, specification_name)
        split_docs = split_patent_docx(docx)
        for docx_name, docx_data in split_docs.items():
            pdf_name = Path(docx_name).with_suffix(".pdf").name
            conversion_docx = remove_word_header_page_numbers(docx_data)
            pdfs[pdf_name] = _pdf_epats_cleanup(_libreoffice_to_pdf(conversion_docx, docx_name))
    elif suffix == ".pdf":
        pdfs.update(split_patent_pdf(specification_data))
    else:
        raise ValueError("Tarifname dosyası DOC, DOCX veya PDF olmalıdır.")

    if figures_data is not None:
        safe_name = figures_name or "Sekiller.docx"
        pdfs["Sekiller.pdf"] = _ensure_pdf(figures_data, safe_name)

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name, payload in pdfs.items():
            zf.writestr(name, payload)
        if metadata:
            zf.writestr("basvuru_verileri.json", json.dumps(metadata, ensure_ascii=False, indent=2).encode("utf-8"))
    return out.getvalue(), pdfs

# -----------------------------------------------------------------------------
# v5.4.57 - CPU/WASM SEMANTIK ROL AI SONUCU -> BASVURU ALANLARI
# -----------------------------------------------------------------------------

def _semantic_block_source_ok(source: str, block_text: str, source_map: dict[str, str]) -> bool:
    """Tarayıcıdan gelen blok gerçekten ilgili yüklenen kaynaktan türetilmiş mi?"""
    full = _plain_norm(source_map.get(source, ""))
    if not full:
        return False
    lines = [
        _plain_norm(x).strip()
        for x in (block_text or "").splitlines()
        if len(_plain_norm(x).strip()) >= 8
    ]
    if not lines:
        return False
    # Pencere oluşturma sırasında satırlar birleştirilmez; anlamlı satırların çoğu
    # tam kaynakta bulunmalı. Böylece tarayıcı sonucu yeni metin uyduramaz.
    hits = sum(1 for x in lines if x in full)
    return hits >= max(1, (len(lines) + 1) // 2)


def _semantic_map_pair_label(label: str, *, applicant: bool) -> str:
    role = "applicant" if applicant else "inventor"
    if label.startswith(role + "_"):
        return label[len(role) + 1:]
    if label in {"name", "identity", "address", "country", "location", "city", "district", "email", "phone", "birth_date", "entity_type"}:
        return label
    if label == role:
        return "name"
    return ""


def _semantic_company_candidate(text: str) -> str:
    values: list[str] = []
    for line in (text or "").splitlines():
        for cell in re.split(r"\t+|\s+\|\s+", line):
            value = _clean_value(cell)
            if not value or _is_instruction_text(value) or _is_form_header_bundle(value):
                continue
            if _looks_like_company(value):
                values.append(value)
    if not values:
        return ""
    values.sort(key=lambda x: (len(x.split()), len(x)), reverse=True)
    return values[0]


def _semantic_probable_person_candidate(text: str) -> str:
    """Yalnız çok sıkı koşullarda etiketsiz kişi adı geri dönüşü.

    Adres, kurum, form başlığı veya soru cümlesi asla kişi adı yapılmaz.
    """
    address_terms = {
        "mah", "mahalle", "cad", "caddesi", "cd", "sok", "sokak", "bulvar", "blv",
        "no", "kat", "daire", "merkezi", "mudurlugu", "müdürlüğü", "turkiye", "türkiye",
    }
    candidates: list[str] = []
    for raw in (text or "").splitlines():
        line = _clean_value(raw)
        if not line or "\t" in raw or ":" in line or "@" in line or re.search(r"\d", line):
            continue
        if _canonical_label(line) or _is_instruction_text(line) or _is_form_header_bundle(line) or _looks_like_company(line):
            continue
        words = [x for x in re.split(r"\s+", line) if x]
        if not (2 <= len(words) <= 5):
            continue
        norm_words = {_plain_norm(x.strip(".,;()[]")) for x in words}
        if norm_words & address_terms:
            continue
        if not all(re.search(r"[A-Za-zÇĞİÖŞÜçğıöşü]", w) for w in words):
            continue
        # Tamamı küçük harfli açıklama cümlelerini dışla; adlarda en az iki kelime
        # büyük harfle/uppercase başlar.
        cap = sum(1 for w in words if w[:1].isupper())
        if cap < 2:
            continue
        candidates.append(line)
    if not candidates:
        return ""
    # Kısa kişi adını uzun açıklamaya tercih et.
    candidates.sort(key=lambda x: (len(x.split()), len(x)))
    return candidates[0]


def _semantic_role_rows_from_block(block_text: str, source: str, *, applicant: bool) -> list[dict[str, str]]:
    """AI'nın yalnız rolünü belirlediği kaynak bloğundan alanları deterministik çıkar.

    AI değer üretmez; ad/unvan, kimlik, adres, e-posta ve telefon yine kaynak
    bloğundaki etiket/değer veya tablo hücrelerinden alınır.
    """
    heading = "HAK SAHİBİ / BAŞVURU SAHİBİ" if applicant else "BULUŞ SAHİBİ / BULUŞÇU"
    forced = heading + "\n" + (block_text or "")
    parsed = extract_application_information_rule_based([(source, forced)])
    key = "applicants" if applicant else "inventors"
    rows = [dict(x) for x in (parsed.get(key) or [])]

    # Farklı form tasarımlarında aynı satırdaki açık etiket/değerleri ayrıca tara.
    fallback = _new_person("", source, applicant=applicant)
    for line in (block_text or "").splitlines():
        for _raw_label, label, value in _line_label_value_pairs(line):
            field = _semantic_map_pair_label(label, applicant=applicant)
            if not field or not value:
                continue
            if field == "name":
                if not _is_bad_person_value("name", value, applicant=applicant):
                    fallback["name"] = _clean_value(value)
                    if applicant:
                        fallback["entity_type"] = _entity_type(fallback["name"], fallback.get("entity_type", ""))
            else:
                _apply_person_field(fallback, field, value, applicant=applicant)

    if applicant and not fallback.get("name"):
        company = _semantic_company_candidate(block_text)
        if company:
            fallback["name"] = company
            fallback["entity_type"] = _entity_type(company)
    if not applicant and not fallback.get("name"):
        person = _semantic_probable_person_candidate(block_text)
        if person:
            fallback["name"] = person

    _fill_country_from_explicit_location(fallback)
    _sanitize_person_row(fallback, applicant=applicant)
    if fallback.get("name") or fallback.get("identity") or fallback.get("email") or fallback.get("address"):
        rows.append(fallback)

    out: list[dict[str, str]] = []
    for row in rows:
        row["source"] = source
        _sanitize_person_row(row, applicant=applicant)
        if row.get("name") or row.get("identity") or row.get("email") or row.get("address"):
            _append_unique_person(out, row, applicant=applicant)
    return out


def _semantic_enrich_from_existing(
    semantic_rows: list[dict[str, str]], existing_rows: list[dict[str, str]], *, applicant: bool
) -> list[dict[str, str]]:
    """Semantik rolü doğru bulunan satırı eski parser'ın güvenli alanlarıyla tamamla."""
    rows = [dict(x) for x in semantic_rows]
    for row in rows:
        best = None
        best_score = -1
        for current in existing_rows:
            score = _person_match_score(current, row)
            if score > best_score:
                best, best_score = current, score
        # Aynı kaynaktan yalnız bir eski ve bir semantik satır varsa alan tamamlama
        # için kaynak eşleşmesi yeterlidir; eski rol/name semantik sonucu değiştiremez.
        if best_score < 15 and len(rows) == 1 and len(existing_rows) == 1:
            rs = {x for x in (row.get("source") or "").split("; ") if x}
            es = {x for x in (existing_rows[0].get("source") or "").split("; ") if x}
            if rs & es:
                best, best_score = existing_rows[0], 10
        if best is None or best_score < 10:
            continue
        for field in ("identity", "country", "city", "district", "address", "email", "phone", "birth_date"):
            if not row.get(field) and best.get(field):
                row[field] = best[field]
        if applicant and not row.get("entity_type"):
            row["entity_type"] = _entity_type(row.get("name", ""), best.get("entity_type", ""))
        _sanitize_person_row(row, applicant=applicant)
    return rows


def merge_verified_cpu_semantic_application_information(
    rule_data: dict[str, Any], semantic_data: dict[str, Any], source_blocks: list[tuple[str, str]]
) -> dict[str, Any]:
    """v5.4.57: form yerleşiminden bağımsız semantik rol AI + kaynak doğrulama.

    Tarayıcı CPU modeli yalnız bir metin bloğunun *rolünü* sınıflandırır. Gerçek
    ad/unvan/telefon/e-posta/kimlik değerleri Python tarafında doğrudan kaynaktan
    çıkarılır ve tam kaynakta tekrar doğrulanır. Böylece AI serbest değer üretemez.
    """
    merged = normalize_application_information(rule_data)
    source_map = {name: text for name, text in source_blocks}
    raw_blocks = semantic_data.get("blocks") if isinstance(semantic_data, dict) else []
    if not isinstance(raw_blocks, list):
        raw_blocks = []

    semantic_apps: list[dict[str, str]] = []
    semantic_invs: list[dict[str, str]] = []
    option_updates: list[tuple[str, dict[str, Any]]] = []

    for item in raw_blocks:
        if not isinstance(item, dict):
            continue
        source = str(item.get("source") or "").strip()
        block_text = str(item.get("text") or "").strip()
        role = str(item.get("role") or "").strip()
        try:
            score = float(item.get("score") or 0.0)
            margin = float(item.get("margin") or 0.0)
        except Exception:
            continue
        if source not in source_map or not _semantic_block_source_ok(source, block_text, source_map):
            continue
        threshold = 0.40 if role == "options" else 0.46
        if score < threshold or (margin < 0.015 and score < 0.62):
            continue

        if role == "applicant":
            rows = _semantic_role_rows_from_block(block_text, source, applicant=True)
            # Her değer tüm orijinal kaynağa karşı tekrar doğrulanır.
            verified = _ai_person_rows(rows, applicant=True, source_blocks=source_blocks)
            for row in verified:
                _append_unique_person(semantic_apps, row, applicant=True)
        elif role == "inventor":
            rows = _semantic_role_rows_from_block(block_text, source, applicant=False)
            verified = _ai_person_rows(rows, applicant=False, source_blocks=source_blocks)
            for row in verified:
                _append_unique_person(semantic_invs, row, applicant=False)
        elif role == "options":
            temp = extract_application_information_rule_based([(source, block_text)])
            option_updates.append((source, temp.get("filing_options") or {}))

    # AI rolü içerisinde kaynakta doğrulanmış bir ad/unvan bulunduysa bu rol için
    # semantik satırlar otoritedir. Eski parser'ın "ronik", "Sultan", açıklama vb.
    # hayalet kayıtları taşınmaz; yalnız eşleşen eski satırdan boş alan tamamlanır.
    if any(x.get("name") for x in semantic_apps):
        merged["applicants"] = _semantic_enrich_from_existing(
            semantic_apps, merged.get("applicants") or [], applicant=True
        )
    elif semantic_apps:
        merged["applicants"] = _merge_ai_people(merged.get("applicants") or [], semantic_apps, applicant=True)

    if any(x.get("name") for x in semantic_invs):
        merged["inventors"] = _semantic_enrich_from_existing(
            semantic_invs, merged.get("inventors") or [], applicant=False
        )
    elif semantic_invs:
        merged["inventors"] = _merge_ai_people(merged.get("inventors") or [], semantic_invs, applicant=False)

    # Başvuru tercihleri AI tarafından üretilmez. AI yalnız ilgili soru bloğunu
    # işaretler; EVET/HAYIR yine blok metninden deterministik okunur.
    opts = merged.get("filing_options") or {}
    for _source, raw_opts in option_updates:
        if not isinstance(raw_opts, dict):
            continue
        for key in ("inventor_hidden", "public_project", "early_publication"):
            cand = raw_opts.get(key) if isinstance(raw_opts.get(key), dict) else {}
            if not cand.get("explicit") or cand.get("status") not in {"Evet", "Hayır"}:
                continue
            cur = opts.get(key) if isinstance(opts.get(key), dict) else {}
            # Mevcut açık cevap varsa koru; varsayılan cevabı ise açık kaynak cevabıyla yükselt.
            if not cur.get("explicit"):
                opts[key] = dict(cand)
    merged["filing_options"] = opts

    for row in merged.get("applicants") or []:
        _sanitize_person_row(row, applicant=True)
    for row in merged.get("inventors") or []:
        _sanitize_person_row(row, applicant=False)
    _reconcile_role_rows(merged)
    return normalize_application_information(merged)
