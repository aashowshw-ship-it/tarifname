from __future__ import annotations

import io
import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def _docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            vals = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            if any(vals):
                parts.append("\t".join(vals))
    return "\n".join(parts)


def _pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _legacy_doc_text(data: bytes, filename: str) -> str:
    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / Path(filename).name
        src.write_bytes(data)
        try:
            p = subprocess.run(
                ["antiword", str(src)],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,
            )
            text = p.stdout.decode("utf-8", errors="replace")
            if text.strip():
                return text
        except Exception:
            pass

        outdir = Path(td) / "converted"
        outdir.mkdir()
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(outdir), str(src)],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=120,
            )
            candidates = list(outdir.glob("*.docx"))
            if candidates:
                return _docx_text(candidates[0].read_bytes())
        except Exception as exc:
            raise ValueError("Eski .doc dosyası okunamadı. Dosyayı .docx olarak kaydedip yeniden yükleyin.") from exc

    raise ValueError("Eski .doc dosyası okunamadı.")


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        text = _docx_text(data)
    elif suffix == ".doc":
        text = _legacy_doc_text(data, filename)
    elif suffix == ".pdf":
        text = _pdf_text(data)
    elif suffix in {".txt", ".md"}:
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError("Desteklenen dosya türleri: .docx, .doc, .pdf, .txt")

    text = text.replace("\x00", " ").strip()
    if not text:
        raise ValueError("Dosyadan metin çıkarılamadı.")
    return text
