from __future__ import annotations

import re
from copy import deepcopy

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

EQ_MARKER_RE = re.compile(r"\[\[(?:EQ|FORMULA)\s*:\s*(.+?)\]\]", re.IGNORECASE | re.DOTALL)


def _math_run(text: str) -> OxmlElement:
    r = OxmlElement("m:r")
    rpr = OxmlElement("m:rPr")
    r.append(rpr)
    t = OxmlElement("m:t")
    t.text = str(text)
    r.append(t)
    return r


def _math_identifier(token: str) -> OxmlElement:
    token = str(token)
    if "_" not in token:
        return _math_run(token)
    base, sub = token.split("_", 1)
    ssub = OxmlElement("m:sSub")
    e = OxmlElement("m:e")
    e.append(_math_run(base))
    sub_el = OxmlElement("m:sub")
    sub_el.append(_math_run(sub))
    ssub.extend([e, sub_el])
    return ssub


def _append_math_expression(parent: OxmlElement, expression: str) -> None:
    """Patent formüllerini Word OMML matematik yapısına çevirir.

    Alt çizgili değişkenler subscript yapılır; `A/B` biçimindeki değişken oranları
    kesir nesnesine çevrilir. Diğer matematik işaretleri OMML matematik run'ında
    korunur. Böylece denklem düz paragraf metni değildir.
    """
    expr = str(expression or "").strip()
    pos = 0
    token_re = re.compile(
        r"(?P<frac>[A-Za-zÇĞİÖŞÜçğıöşüΑ-Ωα-ω][A-Za-zÇĞİÖŞÜçğıöşüΑ-Ωα-ω0-9]*(?:_[A-Za-zÇĞİÖŞÜçğıöşüΑ-Ωα-ω0-9]+)*\s*/\s*[A-Za-zÇĞİÖŞÜçğıöşüΑ-Ωα-ω][A-Za-zÇĞİÖŞÜçğıöşüΑ-Ωα-ω0-9]*(?:_[A-Za-zÇĞİÖŞÜçğıöşüΑ-Ωα-ω0-9]+)*)|(?P<ident>[A-Za-zÇĞİÖŞÜçğıöşüΑ-Ωα-ω][A-Za-zÇĞİÖŞÜçğıöşüΑ-Ωα-ω0-9]*(?:_[A-Za-zÇĞİÖŞÜçğıöşüΑ-Ωα-ω0-9]+)+)"
    )
    for m in token_re.finditer(expr):
        if m.start() > pos:
            parent.append(_math_run(expr[pos:m.start()]))
        if m.group("frac"):
            left, right = re.split(r"\s*/\s*", m.group("frac"), maxsplit=1)
            frac = OxmlElement("m:f")
            num = OxmlElement("m:num")
            den = OxmlElement("m:den")
            num.append(_math_identifier(left))
            den.append(_math_identifier(right))
            frac.extend([num, den])
            parent.append(frac)
        else:
            parent.append(_math_identifier(m.group("ident")))
        pos = m.end()
    if pos < len(expr):
        parent.append(_math_run(expr[pos:]))
    if not len(parent):
        parent.append(_math_run(expr))


def build_inline_omath(expression: str) -> OxmlElement:
    omath = OxmlElement("m:oMath")
    _append_math_expression(omath, expression)
    return omath


def append_text_with_equations(paragraph, text: str) -> None:
    """`[[EQ: ...]]` işaretlerini gerçek inline Word denklemlerine dönüştürür."""
    raw = str(text or "")
    pos = 0
    found = False
    for m in EQ_MARKER_RE.finditer(raw):
        found = True
        if m.start() > pos:
            paragraph.add_run(raw[pos:m.start()])
        paragraph._p.append(build_inline_omath(m.group(1).strip()))
        pos = m.end()
    if pos < len(raw):
        paragraph.add_run(raw[pos:])
    if not found and not paragraph.runs:
        paragraph.add_run(raw)


def add_display_equation(doc, template, prototype_index: int, expression: str):
    """Şablon paragraf arketipi içinde ortalı gerçek Word OMML denklemi oluşturur."""
    p_el = deepcopy(template.paragraphs[prototype_index]._p)
    ppr = p_el.find(qn("w:pPr"))
    for child in list(p_el):
        if child is not ppr:
            p_el.remove(child)
    omath_para = OxmlElement("m:oMathPara")
    omath_para.append(build_inline_omath(expression))
    p_el.append(omath_para)
    doc._element.body.insert(-1, p_el)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    return p
