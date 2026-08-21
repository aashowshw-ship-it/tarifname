from __future__ import annotations

import base64
import io
import json
import os
import re
import subprocess
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import unquote
from urllib.request import Request, urlopen
import html as html_lib

import streamlit as st
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openai import OpenAI
from PIL import Image

try:
    import cairosvg  # SVG müşteri şekillerini teknik içerik değiştirmeden PNG önizlemeye/Word yerleşimine dönüştürmek için
except ImportError:  # pragma: no cover
    cairosvg = None

try:
    import fitz  # PyMuPDF: PDF içindeki şekilleri çıkarmak için
except ImportError:  # pragma: no cover - bağımlılık Render üzerinde requirements ile kurulur
    fitz = None
from pypdf import PdfReader

from rules import APP_VERSION, RULESET_VERSION, ARASTIRMA_RULES, ARASTIRMA_GUNCELLEME_RULES, GORUS_RULES, TARIFNAME_RULES
from template_audit import validate_full_tarifname_template_fidelity
from source_guards import (
    build_source_passage_registry,
    validate_source_passage_audit,
    resolve_tarifname_claim_mode,
    derive_tarifname_output_names,
    validate_final_source_coverage_chain,
    validate_final_raw_source_audit,
)
from word_math import EQ_MARKER_RE, append_text_with_equations as _append_text_with_equations, add_display_equation
from gorus_audit import (
    annotate_quote_locations,
    validate_quote_locations_against_spec,
    extract_cited_original_figure_pages,
    detect_examiner_reasoned_documents,
    detect_ep_xy_documents,
    detect_defense_documents,
    is_ep_search_report,
    validate_gorus_template_fidelity,
    validate_opinion_payload,
    validate_opinion_against_raw_sources,
    validate_gorus_docx_content_flow,
    validate_minimal_tracked_changes,
    validate_ep_prior_art_markup_text,
    validate_ai_quality_audit,
    build_gorus_quality_report,
    render_gorus_docx_smoke_test,
)

BASE_DIR = Path(__file__).resolve().parent
TARIFNAME_TEMPLATE = BASE_DIR / "Tarifname_181176_template.docx"
GORUS_TEMPLATE = BASE_DIR / "Gorus_metni_696809_template.docx"
ARASTIRMA_TEMPLATE = BASE_DIR / "On_Arastirma_Raporu_181612_template.docx"
MODEL = os.getenv("OPENAI_MODEL", "gpt-5.6")
MAX_TEXT_PER_FILE = int(os.getenv("MAX_TEXT_PER_FILE", "180000"))
MAX_TOTAL_TEXT = int(os.getenv("MAX_TOTAL_TEXT", "700000"))
FIGURE_REFERENCE_CONFIDENCE = float(os.getenv("FIGURE_REFERENCE_CONFIDENCE", "0.86"))


# -----------------------------------------------------------------------------
# GENEL KURALLAR
# -----------------------------------------------------------------------------
# Kurallar tek bir kaynaktan (rules.py) yüklenir. Böylece arayüz ve üretim akışları
# arasında kural farkı oluşmaz.


@dataclass
class UploadedAsset:
    name: str
    data: bytes
    mime: str = "application/octet-stream"


# -----------------------------------------------------------------------------
# API / DOSYA YARDIMCILARI
# -----------------------------------------------------------------------------
def get_client() -> OpenAI:
    key = os.getenv("OPENAI_API_KEY", "").strip()
    if not key:
        raise RuntimeError(
            "OPENAI_API_KEY tanımlı değil. Render > servis > Environment bölümünde OPENAI_API_KEY ekleyin."
        )
    return OpenAI(api_key=key)


def extract_json(text: str) -> dict[str, Any]:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start, end = text.find("{"), text.rfind("}")
        if start >= 0 and end > start:
            return json.loads(text[start : end + 1])
        raise ValueError("Yapay zekâ yanıtı geçerli JSON olarak okunamadı.")


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".svg"}
RASTER_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}

def _svg_to_png(data: bytes, *, output_width: int = 2200) -> bytes:
    if cairosvg is None:
        raise ValueError("SVG müşteri şekli işlendi ancak CairoSVG kurulu değil. requirements.txt içindeki CairoSVG bağımlılığını kurun.")
    try:
        # Bazı headless Cairo/Pango kurulumlarında Matplotlib SVG'lerindeki Unicode
        # eksi/en-dash glifleri kare olarak rasterize olabiliyor. Geometri/veri
        # değiştirilmeden yalnız eşdeğer ASCII çizgi karakterine normalize edilir.
        svg_bytes = data.replace("−".encode("utf-8"), b"-").replace("–".encode("utf-8"), b"-")
        return cairosvg.svg2png(bytestring=svg_bytes, output_width=output_width)
    except Exception as exc:
        raise ValueError("SVG müşteri şekli PNG önizlemeye dönüştürülemedi; özgün SVG teknik kaynak olarak korunamadı.") from exc

def _model_ready_image(asset: UploadedAsset) -> UploadedAsset:
    if Path(asset.name).suffix.lower() == ".svg":
        return UploadedAsset(asset.name, _svg_to_png(asset.data), "image/png")
    return asset

def image_content(asset: UploadedAsset) -> dict[str, Any]:
    asset = _model_ready_image(asset)
    b64 = base64.b64encode(asset.data).decode("ascii")
    mime = asset.mime or "image/png"
    return {"type": "input_image", "image_url": f"data:{mime};base64,{b64}", "detail": "high"}


def ask_json(prompt: str, *, web_search: bool = False, images: Iterable[UploadedAsset] | None = None) -> dict[str, Any]:
    client = get_client()
    content: list[dict[str, Any]] = [{"type": "input_text", "text": prompt}]
    for asset in images or []:
        content.append(image_content(asset))
    kwargs: dict[str, Any] = {
        "model": MODEL,
        "input": [{"role": "user", "content": content}],
    }
    if web_search:
        kwargs["tools"] = [{"type": "web_search"}]
        kwargs["tool_choice"] = "required"
    response = client.responses.create(**kwargs)
    return extract_json(response.output_text)


def docx_text(data: bytes) -> str:
    """DOCX metnini, Word denklem düğümleri (m:t) dahil olacak şekilde çıkarır."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            xml = zf.read("word/document.xml")
        root = ET.fromstring(xml)
        paragraphs: list[str] = []
        for node in root.iter():
            if node.tag.rsplit("}", 1)[-1] != "p":
                continue
            pieces: list[str] = []
            for child in node.iter():
                local = child.tag.rsplit("}", 1)[-1]
                if local == "t" and child.text:
                    pieces.append(child.text)
                elif local in {"tab"}:
                    pieces.append("\t")
                elif local in {"br", "cr"}:
                    pieces.append("\n")
            text = "".join(pieces).strip()
            if text:
                paragraphs.append(text)
        if paragraphs:
            return "\n".join(paragraphs)
    except Exception:
        pass

    # Bozuk/alışılmadık DOCX için güvenli geri dönüş.
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            vals = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            if any(vals):
                parts.append("\t".join(vals))
    return "\n".join(parts)


def pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def legacy_doc_text(data: bytes, filename: str) -> str:
    with tempfile.TemporaryDirectory() as td:
        source = Path(td) / Path(filename).name
        source.write_bytes(data)
        try:
            result = subprocess.run(
                ["antiword", str(source)], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60
            )
            text = result.stdout.decode("utf-8", errors="replace")
            if text.strip():
                return text
        except Exception:
            pass
        outdir = Path(td) / "converted"
        outdir.mkdir()
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(outdir), str(source)],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=120,
            )
            files = list(outdir.glob("*.docx"))
            if files:
                return docx_text(files[0].read_bytes())
        except Exception as exc:
            raise ValueError("Eski .doc dosyası okunamadı; .docx olarak kaydedip yükleyin.") from exc
    raise ValueError("Eski .doc dosyası okunamadı.")


def extract_text_from_asset(asset: UploadedAsset) -> str:
    suffix = Path(asset.name).suffix.lower()
    if suffix == ".docx":
        text = docx_text(asset.data)
    elif suffix == ".doc":
        text = legacy_doc_text(asset.data, asset.name)
    elif suffix == ".pdf":
        text = pdf_text(asset.data)
    elif suffix in {".txt", ".md"}:
        text = asset.data.decode("utf-8", errors="replace")
    elif suffix in IMAGE_SUFFIXES:
        text = f"[TEKNİK GÖRSEL DOSYASI: {asset.name}]"
    else:
        text = ""
    return text.replace("\x00", " ").strip()[:MAX_TEXT_PER_FILE]


def assets_from_uploads(files: Iterable[Any] | None) -> list[UploadedAsset]:
    out: list[UploadedAsset] = []
    for f in files or []:
        data = f.getvalue()
        suffix = Path(f.name).suffix.lower()
        if suffix == ".zip":
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                for info in zf.infolist():
                    if info.is_dir() or info.file_size > 30 * 1024 * 1024:
                        continue
                    inner_name = Path(info.filename).name
                    inner_suffix = Path(inner_name).suffix.lower()
                    if inner_suffix not in {".pdf", ".docx", ".doc", ".txt", ".md", ".png", ".jpg", ".jpeg", ".webp", ".svg"}:
                        continue
                    out.append(UploadedAsset(inner_name, zf.read(info)))
        else:
            out.append(UploadedAsset(f.name, data, getattr(f, "type", "application/octet-stream")))
    return out


def combine_asset_text(label: str, assets: list[UploadedAsset]) -> tuple[str, list[UploadedAsset]]:
    blocks: list[str] = []
    images: list[UploadedAsset] = []
    total = 0
    for asset in assets:
        suffix = Path(asset.name).suffix.lower()
        if suffix in IMAGE_SUFFIXES:
            images.append(_model_ready_image(asset))
            blocks.append(f"\n--- {label}: {asset.name} (teknik görsel ayrıca eklenmiştir; dosya adı ve görsel içeriği kaynak envanterine dahildir) ---\n")
            continue
        text = extract_text_from_asset(asset)
        if not text:
            continue
        remain = MAX_TOTAL_TEXT - total
        if remain <= 0:
            break
        text = text[:remain]
        total += len(text)
        blocks.append(f"\n--- {label}: {asset.name} ---\n{text}\n")
    return "".join(blocks), images


def _valid_figure_image(data: bytes, *, min_width: int = 260, min_height: int = 160) -> bool:
    try:
        with Image.open(io.BytesIO(data)) as im:
            width, height = im.size
            return width >= min_width and height >= min_height
    except Exception:
        return False


def extract_embedded_images(asset: UploadedAsset) -> list[UploadedAsset]:
    """DOCX/PDF içindeki büyük görselleri özgün baytlarıyla çıkarır."""
    suffix = Path(asset.name).suffix.lower()
    images: list[UploadedAsset] = []

    if suffix in RASTER_IMAGE_SUFFIXES:
        return [asset] if _valid_figure_image(asset.data, min_width=1, min_height=1) else []

    if suffix == ".svg":
        try:
            png = _svg_to_png(asset.data)
        except Exception:
            return []
        return [UploadedAsset(asset.name, png, "image/png")] if _valid_figure_image(png, min_width=1, min_height=1) else []

    if suffix == ".docx":
        try:
            with zipfile.ZipFile(io.BytesIO(asset.data)) as zf:
                media = sorted(
                    name for name in zf.namelist()
                    if name.startswith("word/media/") and not name.endswith("/")
                )
                for index, name in enumerate(media, 1):
                    data = zf.read(name)
                    ext = Path(name).suffix.lower() or ".png"
                    if ext == ".svg":
                        try:
                            rendered = _svg_to_png(data)
                        except Exception:
                            continue
                        if not _valid_figure_image(rendered):
                            continue
                        images.append(UploadedAsset(f"{Path(asset.name).stem}_sekil_{index}.svg", rendered, "image/png"))
                        continue
                    if not _valid_figure_image(data):
                        continue
                    mime = "image/jpeg" if ext in {".jpg", ".jpeg"} else f"image/{ext.lstrip('.')}"
                    images.append(UploadedAsset(f"{Path(asset.name).stem}_sekil_{index}{ext}", data, mime))
        except Exception:
            return []
        return images

    if suffix == ".pdf" and fitz is not None:
        try:
            pdf = fitz.open(stream=asset.data, filetype="pdf")
            seen: set[int] = set()
            counter = 0
            for page in pdf:
                for info in page.get_images(full=True):
                    xref = int(info[0])
                    if xref in seen:
                        continue
                    seen.add(xref)
                    extracted = pdf.extract_image(xref)
                    data = extracted.get("image", b"")
                    if not data or not _valid_figure_image(data):
                        continue
                    counter += 1
                    ext = "." + str(extracted.get("ext", "png")).lower()
                    mime = "image/jpeg" if ext in {".jpg", ".jpeg"} else f"image/{ext.lstrip('.')}"
                    images.append(UploadedAsset(f"{Path(asset.name).stem}_sekil_{counter}{ext}", data, mime))
            pdf.close()
        except Exception:
            return []
    return images


def _append_word_field(run, field_name: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_name)
    run._r.append(fld)


def _add_figures_page_counter(section) -> None:
    """Şekiller şablonundaki `1 / 3` mantığını dinamik PAGE / NUMPAGES alanlarıyla kurar."""
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Header tekrar üretilirse eski içeriği temizle.
    for child in list(p._element):
        p._element.remove(child)
    r1 = p.add_run()
    _append_word_field(r1, "PAGE")
    p.add_run(" / ")
    r2 = p.add_run()
    _append_word_field(r2, "NUMPAGES")
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold = True


def _figure_dimensions_cm(data: bytes, max_width_cm: float = 16.5, max_height_cm: float = 20.5) -> tuple[float, float]:
    try:
        with Image.open(io.BytesIO(data)) as im:
            w, h = im.size
        if not w or not h:
            raise ValueError
        ratio = h / w
        width = max_width_cm
        height = width * ratio
        if height > max_height_cm:
            height = max_height_cm
            width = height / ratio
        return max(4.0, width), max(2.5, height)
    except Exception:
        return max_width_cm, 10.0



def _figure_reference_context(draft: dict[str, Any], figure_index: int) -> dict[str, Any]:
    descriptions = list(draft.get("figure_descriptions") or [])
    figure_description = descriptions[figure_index - 1] if 0 <= figure_index - 1 < len(descriptions) else ""
    return {
        "figure_index": figure_index,
        "figure_description": figure_description,
        "elements": draft.get("elements") or [],
        "method_steps": draft.get("method_steps") or [],
        "detailed_paragraphs": draft.get("detailed_paragraphs") or [],
        "working_principle": draft.get("working_principle") or "",
    }


def _figure_reference_audit_prompt(draft: dict[str, Any], figure_index: int, language: str) -> str:
    context = _figure_reference_context(draft, figure_index)
    return f"""{TARIFNAME_RULES}

BİRİNCİ görsel denetlenecek patent şeklidir. Varsa sonraki görseller aynı müşteri çizim setinden yalnız fiziksel unsurun tanınmasına yardımcı bağlam olarak verilmiştir; sonraki görsellere ilişkin ayrı audit üretme.
BİRİNCİ şekli nihai tarifnamedeki REFERANS NUMARALARI ve teknik açıklamayla karşılaştır.
Bu aşamada görsel üretme veya değiştirme. Yalnız teknik referans denetimi yap ve JSON döndür.
Şekil sırası: {figure_index}
Tarifname dili: {language}

KRİTİK DENETİM MANTIĞI:
- Şekildeki mevcut numara veya okun doğru olduğunu varsayma.
- Her işaret için dört aşamalı eşleştirme yap: referans işareti → unsur adı → detaylı açıklamadaki teknik tanım/işlev → şekil üzerindeki gerçek fiziksel veya şematik teknik karşılık.
- Kılavuz çizgisi/ok ucu doğrudan ilgili fiziksel unsurda sonlanmalıdır. Boş alan, komşu parça veya genel tertibat doğru hedef değildir.
- Görünüşte temsil edilen zorunlu ana taşıyıcı unsurun referansı atlanmamalıdır; örneğin modülleri barındıran Akıllı SIM Donanım Platformu (1) uygun bir kapsayıcı çerçeve/taşıyıcı olarak gösterilebilir.
- Ok uçları küçük, sade ve şeklin ölçeğiyle orantılı olmalıdır; büyük ok uçları kullanma.
- Yöntem akış şekline kaynak işlem-adımı düzeyinde açık bir döngü vermiyorsa son adımdan önceki bir adıma geri dönüş oku ekleme. Modül geri beslemesini sistem şekli üzerinde modüller arasında göster.
- Bir referans belirli alt parçaya aitse bütün tertibatı gösteremez. Örneğin referans listesinde `9 = Travers` ise 9 yalnız traversin kendisini göstermelidir.
- REFERANS NUMARALARI bölümünde AYRI numaraya ve ayrı ada sahip unsurlar aynı taşıyıcı içinde bulunsa dahi tek ayırt edilemeyen kutu/hedef üzerinde `2-3`, `2/3` veya iki numara birlikte gösterilemez. Her ayrı unsur için ayrı kutucuk, ayrı çağrı alanı veya ayrı kılavuz çizgisi olmalıdır. Böyle bir birleşik gösterim görürsen `merged_reference_groups` alanına numaraları yaz ve status=`needs_edit` yap. Ortak taşıyıcı ve mevcut teknik bağlantılar korunarak yalnız çağrı/referans katmanının ayrıştırılmasını iste.
- Bu şekil tek başına bütün referansları taşımak zorunda değildir; fakat nihai şekil SETİ tamamlandığında referans listesindeki tüm sistem unsurları en az bir şekil üzerinde, yöntem adımları da uygulanabilir yöntem/akış şekillerinde kapsanacaktır.
- Her görünür parçayı zorla numaralandırma. Yalnız tarifnamede gerçek referansla tanımlanmış ve BU ŞEKİLDE fiziksel karşılığı güvenilir biçimde görülen unsuru değerlendir.
- Referanslı ve bu şekilde görünür bir unsur numarasızsa, fiziksel yeri güvenilir biçimde belirlenebiliyorsa action=`add` yap. Belirsizse `unresolved` yaz; uydurma hedef seçme.
- Mevcut numara doğru fakat ok yanlış fiziksel parçaya gidiyorsa action=`correct` yap.
- Doğru numara ve doğru hedef varsa action=`keep` yap.
- Unsur bu şekilde görünmüyorsa action=`omit` yap; sırf tüm referansları kullanmak için ekleme yapma.
- Geçici/yardımcı şekil numaralarını gerçek tarifname referansı gibi kabul etme.
- confidence, teknik fiziksel eşleştirme güvenidir. 0.86 altında add/correct önerme; unresolved olarak bildir.

JSON ŞEMASI:
{{
  "figure_index": {figure_index},
  "figure_description": "",
  "status": "ok/needs_edit/unresolved",
  "existing_reference_marks": [""],
  "annotations": [
    {{
      "reference": "9",
      "name": "Travers",
      "visible": true,
      "action": "keep/correct/add/split/omit",
      "location_description": "Şekilde hedeflenen fiziksel parçanın açık ve ayırt edici tarifi",
      "reason": "",
      "confidence": 0.95
    }}
  ],
  "merged_reference_groups": [["2","3"]],
  "extra_or_temporary_marks": [""],
  "unresolved": [""],
  "edit_instruction": "Yalnız needs_edit durumunda, hangi numara/okun hangi fiziksel parçaya yöneltileceğini kısa ve kesin yaz."
}}

TARİFNAME/ŞEKİL BAĞLAMI:
{json.dumps(context, ensure_ascii=False, indent=2)}
"""


def audit_figure_references(
    asset: UploadedAsset,
    draft: dict[str, Any],
    figure_index: int,
    language: str,
    context_images: list[UploadedAsset] | None = None,
) -> dict[str, Any]:
    visual_context = [asset]
    for candidate in context_images or []:
        if candidate is asset:
            continue
        visual_context.append(candidate)
        if len(visual_context) >= 4:
            break
    audit = ask_json(
        _figure_reference_audit_prompt(draft, figure_index, language),
        images=visual_context,
    )
    audit["figure_index"] = figure_index
    return audit


def _has_nonempty_items(value: Any) -> bool:
    if not value:
        return False
    if isinstance(value, (list, tuple, set)):
        return any(str(x or "").strip() for x in value)
    return bool(str(value).strip())


def _audit_has_unsafe_edit(audit: dict[str, Any]) -> bool:
    if str(audit.get("status", "")).strip().casefold() == "unresolved":
        return True
    if _has_nonempty_items(audit.get("unresolved")):
        return True
    for item in audit.get("annotations") or []:
        action = str(item.get("action", "")).strip().casefold()
        if action in {"add", "correct", "split"}:
            try:
                confidence = float(item.get("confidence", 0.0) or 0.0)
            except (TypeError, ValueError):
                confidence = 0.0
            if confidence < FIGURE_REFERENCE_CONFIDENCE:
                return True
    return False


def _extract_image_generation_result(response: Any) -> bytes:
    for item in getattr(response, "output", []) or []:
        item_type = item.get("type") if isinstance(item, dict) else getattr(item, "type", "")
        if item_type != "image_generation_call":
            continue
        result = item.get("result") if isinstance(item, dict) else getattr(item, "result", None)
        if isinstance(result, list) and result:
            result = result[-1]
        if isinstance(result, str) and result.strip():
            try:
                return base64.b64decode(result)
            except Exception as exc:  # pragma: no cover - API response shape guard
                raise ValueError("Şekil düzeltme çıktısı base64 görsel olarak çözülemedi.") from exc
    raise ValueError("Şekil düzeltme çağrısı görsel çıktı üretmedi.")


def edit_figure_reference_annotations(
    asset: UploadedAsset,
    draft: dict[str, Any],
    figure_index: int,
    language: str,
    audit: dict[str, Any],
) -> UploadedAsset:
    """Özgün şeklin yalnız referans numarası/kılavuz çizgisi katmanını düzeltir."""
    client = get_client()
    context = _figure_reference_context(draft, figure_index)
    prompt = f"""Bu görsel bir patent şeklidir ve özgün müşteri çizimi teknik kaynak olarak bağlayıcıdır.
Yalnız referans numaraları ile bunların kılavuz çizgileri/oklarını ve AYRI REFERANSLI yazılım/modül unsurları için gereken sade çağrı/kutucuk katmanını düzelt. Mekanik/elektronik taşıyıcı geometrisini, parça biçimlerini, delikleri, kesit taramalarını, perspektifi, boyut ilişkilerini, bağlantı oklarını veya teknik kurguyu değiştirme. Yeni teknik parça üretme veya parça silme. Ancak iki ayrı referans kaynak şekil üzerinde tek `2-3`/tek kutuda birleştirilmişse ortak taşıyıcıyı koruyarak yalnız bu referans gösterimini iki ayrı küçük kutucuk/çağrı alanına ayırabilirsin.

REFERANS DENETİMİ:
{json.dumps(audit, ensure_ascii=False, indent=2)}

TARİFNAME BAĞLAMI:
{json.dumps(context, ensure_ascii=False, indent=2)}

UYGULAMA KURALLARI:
- action=keep olan referansı ve doğru hedefini koru.
- action=correct olan numaranın kılavuz çizgisi/okunu location_description içinde tarif edilen fiziksel unsura yönelt.
- action=add olan referansı yalnız tarif edilen fiziksel/şematik unsur kesin görünüyorsa ekle.
- action=split ise birleşik referans gösterimini (örn. 2-3) aynı ortak taşıyıcı içinde iki ayrı ve ayırt edilebilir referans kutucuğu/çağrısına ayır; teknik bağlantıları ve taşıyıcıyı değiştirme.
- action=omit olan referansı sırf referans listesinde var diye bu şekle ekleme.
- Ok/kılavuz çizgisi ucu doğrudan ilgili fiziksel unsur üzerinde sonlansın; boş alana veya genel tertibata yönelmesin.
- Referans belirli bir alt parçaya aitse tüm tertibatı işaretleme.
- Mümkün olduğunca çizgi kesişmelerini azalt; fakat teknik geometrinin hiçbir bölümünü değiştirme.
- Çıktıda yalnız patent şekli bulunsun; açıklama listesi, referans lejandı, başlık veya ek metin ekleme.
- Siyah-beyaz patent çizimi ve özgün en-boy oranı korunmalıdır.
"""
    content = [{"type": "input_text", "text": prompt}, image_content(asset)]
    common = {
        "model": MODEL,
        "input": [{"role": "user", "content": content}],
    }
    # Yüksek girdi sadakati desteklenen hesaplarda müşteri geometrisini daha güçlü korur.
    try:
        response = client.responses.create(
            **common,
            tools=[{"type": "image_generation", "quality": "high", "input_fidelity": "high"}],
            tool_choice="required",
        )
    except Exception:
        # Bazı API/model sürümlerinde input_fidelity araç parametresi bulunmayabilir.
        response = client.responses.create(
            **common,
            tools=[{"type": "image_generation", "quality": "high"}],
            tool_choice="required",
        )
    data = _extract_image_generation_result(response)
    if not _valid_figure_image(data, min_width=1, min_height=1):
        raise ValueError("Şekil düzeltme çıktısı geçerli bir görsel değil.")
    return UploadedAsset(f"{Path(asset.name).stem}_referans_duzeltilmis.png", data, "image/png")


def verify_figure_reference_edit(
    original: UploadedAsset,
    edited: UploadedAsset,
    draft: dict[str, Any],
    figure_index: int,
    language: str,
    audit: dict[str, Any],
) -> dict[str, Any]:
    context = _figure_reference_context(draft, figure_index)
    prompt = f"""İki patent şekli veriliyor. BİRİNCİ görsel özgün müşteri şekli, İKİNCİ görsel yalnız referans numarası/okları düzeltilmiş aday görseldir.
Adayı çok sıkı doğrula. Görsel üretme; yalnız JSON döndür.

KABUL KRİTERLERİ:
1. Özgün müşteri şeklinin teknik geometrisi, parça biçimleri, delikler, kesitler/taramalar, perspektif, bileşen konumları ve teknik kurgu özünde korunmuş olmalıdır. Antialiasing/çizgi keskinliği gibi önemsiz raster farklarını geometri değişikliği sayma.
2. Referanslar dört aşamalı eşleşmeye uymalıdır: referans → unsur adı → teknik tanım → fiziksel karşılık.
3. Her kılavuz çizgisi/ok doğrudan doğru fiziksel unsurda sonlanmalı; boş alanı, komşu parçayı veya genel tertibatı göstermemelidir.
4. Adayda tarifnamede olmayan yeni referans numarası/legend/açıklama eklenmemelidir.
5. Bu şekilde görünmeyen unsurlar sırf numaralandırma amacıyla eklenmemelidir.
6. REFERANS NUMARALARI bölümünde ayrı olan unsurlar tek `2-3`/tek kutu/tek hedefte birleştirilmemelidir; ortak taşıyıcı içinde dahi ayrı kutucuk/çağrı/ok ile ayırt edilebilir olmalıdır.
7. Denetimde unresolved kalan unsur varsa annotations_correct=false yap.

ÖN DENETİM:
{json.dumps(audit, ensure_ascii=False, indent=2)}

TARİFNAME BAĞLAMI:
{json.dumps(context, ensure_ascii=False, indent=2)}

JSON ŞEMASI:
{{
  "geometry_preserved": true,
  "annotations_correct": true,
  "distinct_references_separated": true,
  "wrong_or_missing": [""],
  "extra_reference_marks": [""],
  "confidence": 0.95,
  "notes": ""
}}
"""
    return ask_json(prompt, images=[original, edited])


def prepare_figures_with_reference_audit(
    images: list[UploadedAsset],
    draft: dict[str, Any],
    language: str = "Türkçe",
    progress_callback: Any | None = None,
) -> tuple[list[UploadedAsset], list[dict[str, Any]], list[str]]:
    """Şekilleri tarifnameyle denetler; güvenliyse yalnız referans katmanını düzeltir ve tekrar doğrular."""
    prepared: list[UploadedAsset] = []
    reports: list[dict[str, Any]] = []
    unresolved: list[str] = []
    total = len(images)

    for index, asset in enumerate(images, 1):
        if progress_callback:
            progress_callback(index, total, "audit")
        try:
            audit = audit_figure_references(asset, draft, index, language, context_images=images)
        except Exception as exc:
            message = f"ŞEKİL {index}: referans denetimi çalıştırılamadı ({exc})."
            unresolved.append(message)
            reports.append({"figure_index": index, "final_status": "unresolved", "message": message})
            prepared.append(asset)
            continue

        report: dict[str, Any] = {"figure_index": index, "audit": audit}
        if _audit_has_unsafe_edit(audit):
            message = f"ŞEKİL {index}: en az bir referansın fiziksel karşılığı güvenilir biçimde belirlenemedi."
            unresolved.append(message)
            report.update({"final_status": "unresolved", "message": message})
            reports.append(report)
            prepared.append(asset)
            continue

        status = str(audit.get("status", "ok")).strip().casefold()
        needs_edit = status == "needs_edit" or any(
            str(x.get("action", "")).strip().casefold() in {"add", "correct", "split"}
            for x in audit.get("annotations") or []
        ) or _has_nonempty_items(audit.get("merged_reference_groups"))
        if not needs_edit:
            report["final_status"] = "ok"
            reports.append(report)
            prepared.append(asset)
            continue

        if progress_callback:
            progress_callback(index, total, "edit")
        try:
            edited = edit_figure_reference_annotations(asset, draft, index, language, audit)
        except Exception as exc:
            message = f"ŞEKİL {index}: referans düzeltmesi uygulanamadı ({exc})."
            unresolved.append(message)
            report.update({"final_status": "unresolved", "message": message})
            reports.append(report)
            prepared.append(asset)
            continue

        if progress_callback:
            progress_callback(index, total, "verify")
        try:
            verification = verify_figure_reference_edit(asset, edited, draft, index, language, audit)
        except Exception as exc:
            message = f"ŞEKİL {index}: düzeltme sonrası doğrulama çalıştırılamadı ({exc})."
            unresolved.append(message)
            report.update({"final_status": "unresolved", "message": message})
            reports.append(report)
            prepared.append(asset)
            continue

        try:
            verify_confidence = float(verification.get("confidence", 0.0) or 0.0)
        except (TypeError, ValueError):
            verify_confidence = 0.0
        accepted = (
            bool(verification.get("geometry_preserved"))
            and bool(verification.get("annotations_correct"))
            and bool(verification.get("distinct_references_separated", True))
            and not _has_nonempty_items(verification.get("wrong_or_missing"))
            and not _has_nonempty_items(verification.get("extra_reference_marks"))
            and verify_confidence >= FIGURE_REFERENCE_CONFIDENCE
        )
        report["verification"] = verification
        if accepted:
            report["final_status"] = "corrected"
            prepared.append(edited)
        else:
            message = f"ŞEKİL {index}: otomatik düzeltme, geometri/referans doğrulamasını geçemedi."
            unresolved.append(message)
            report.update({"final_status": "unresolved", "message": message})
            prepared.append(asset)
        reports.append(report)

    # SET BAZINDA REFERANS TAMLIK KAPISI: referans listesindeki her unsur/adım en az bir nihai şekilde görünmelidir.
    expected_elements = {str(x.get("number", "") or "").strip() for x in (draft.get("elements") or []) if str(x.get("number", "") or "").strip()}
    expected_methods = {str(x.get("number", "") or "").strip() for x in (draft.get("method_steps") or []) if str(x.get("number", "") or "").strip()}
    represented: set[str] = set()
    for report in reports:
        audit = report.get("audit") or {}
        for mark in audit.get("existing_reference_marks") or []:
            represented.add(str(mark or "").strip())
        for annotation in audit.get("annotations") or []:
            ref = str(annotation.get("reference", "") or "").strip()
            action = str(annotation.get("action", "") or "").strip().casefold()
            if ref and action in {"keep", "add", "correct", "split"} and annotation.get("visible", True) is not False:
                represented.add(ref)
    missing_elements = sorted(expected_elements - represented)
    missing_methods = sorted(expected_methods - represented)
    if missing_elements:
        unresolved.append("Şekil setinde gösterilmeyen REFERANS NUMARALARI sistem unsurları: " + ", ".join(missing_elements))
    if expected_methods and missing_methods:
        unresolved.append("Şekil/akış setinde gösterilmeyen yöntem adımı referansları: " + ", ".join(missing_methods))

    return prepared, reports, unresolved


def build_figures_docx(images: list[UploadedAsset], language: str = "Türkçe") -> bytes:
    """Referans denetiminden geçmiş müşteri şekillerini dinamik sayfa düzeni ve ŞEKİL N başlıklarıyla Word'e yerleştirir."""
    if not images:
        raise ValueError("Şekiller Word dosyası için BBF içinde veya ayrıca yüklenen dosyalarda kullanılabilir görsel bulunamadı.")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)
        section.header_distance = Cm(0.65)
        _add_figures_page_counter(section)

    # Yaklaşık kullanılabilir dikey alan. Görsel büyüklüğüne göre 1 veya daha fazla şekil aynı sayfaya yerleşebilir.
    usable_height_cm = 24.2
    used_height_cm = 0.0

    for index, asset in enumerate(images, 1):
        width_cm, height_cm = _figure_dimensions_cm(asset.data)
        block_height = height_cm + 1.25  # şekil başlığı ve minimum aralık

        if used_height_cm > 0 and used_height_cm + block_height > usable_height_cm:
            doc.add_page_break()
            used_height_cm = 0.0

        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(2)
        try:
            p_img.add_run().add_picture(io.BytesIO(asset.data), width=Cm(width_cm), height=Cm(height_cm))
        except Exception as exc:
            raise ValueError(f"{asset.name} şekiller dosyasına eklenemedi.") from exc

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(10)
        r = p_cap.add_run(f"{'FIGURE' if _english_spec(language) else 'ŞEKİL'} {index}")
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(11)

        used_height_cm += block_height

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# -----------------------------------------------------------------------------
# DOCX YARDIMCILARI
# -----------------------------------------------------------------------------
def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_paragraph(p, *, bold: bool = False, center: bool = False, italic: bool = False, size: int = 11):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        # Set explicit direct formatting. The 696809 template's Normal style is
        # bold by default, so leaving run.bold=None would accidentally render
        # ordinary opinion paragraphs in bold.
        run.bold = bool(bold)
        run.italic = bool(italic)
    return p


def add_text(doc: Document, text: str, *, bold: bool = False, center: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    p.add_run(text)
    return format_paragraph(p, bold=bold, center=center, italic=italic, size=size)


def add_heading(doc: Document, text: str, *, center: bool = False):
    return add_text(doc, text, bold=True, center=center)


def add_bullet(doc: Document, text: str, *, symbol: str = "•"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.add_run(f"{symbol}\t{text}")
    return format_paragraph(p)


def add_quote(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f'“{text}”')
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def copy_template_paragraph(doc: Document, template: Document, index: int):
    if 0 <= index < len(template.paragraphs):
        doc._element.body.insert(-1, deepcopy(template.paragraphs[index]._p))


def copy_template_paragraph_with_text(doc: Document, template: Document, index: int, text: str):
    """Şablon paragraf biçimini koruyup metni değiştir.

    Metin şablondaki sabit metinle aynıysa paragrafı *aynen* kopyalar. Böylece
    sabit uyarı paragraflarındaki kırmızı/mavi run bölünmeleri, kalınlıklar ve
    diğer run düzeyi biçimler kaybolmaz. Dinamik metinde ise paragraf arketipi
    korunur ve yeni metin ilk run üzerinden yazılır.
    """
    if not (0 <= index < len(template.paragraphs)):
        return
    if str(text) == template.paragraphs[index].text:
        copy_template_paragraph(doc, template, index)
        return
    p_el = deepcopy(template.paragraphs[index]._p)
    text_nodes = list(p_el.iter(qn("w:t")))
    if text_nodes:
        text_nodes[0].text = text
        for node in text_nodes[1:]:
            node.text = ""
    doc._element.body.insert(-1, p_el)


def _english_spec(language: str) -> bool:
    text = str(language or "").strip().casefold().replace("\u0307", "")
    return text in {"ingilizce", "english", "en"}


def add_blank(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    return p


def _copy_list_properties(target, source):
    """Şablondaki gerçek Word numaralandırma/madde işareti ve girintisini kopyala."""
    src_pr = source._p.pPr
    if src_pr is None:
        return
    tgt_pr = target._p.get_or_add_pPr()
    for tag in ("w:numPr", "w:ind"):
        existing = tgt_pr.find(qn(tag))
        if existing is not None:
            tgt_pr.remove(existing)
        src = src_pr.find(qn(tag))
        if src is not None:
            tgt_pr.append(deepcopy(src))


def add_template_list_item(doc: Document, template: Document, prototype_index: int, text: str):
    p = doc.add_paragraph()
    _copy_list_properties(p, template.paragraphs[prototype_index])
    _append_text_with_equations(p, text)
    return format_paragraph(p)


def add_nested_claim_list_item(doc: Document, template: Document, text: str):
    """Ortak taşıyıcı grubundaki alt unsuru gerçek Word alt bullet olarak yazar."""
    p = doc.add_paragraph()
    # Şablondaki tire biçimli gerçek bullet numaralandırmasını kullan.
    _copy_list_properties(p, template.paragraphs[66])
    ppr = p._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:left"), "2136")
    ind.set(qn("w:hanging"), "360")
    _append_text_with_equations(p, text)
    return format_paragraph(p)


def safe_output_name(name: str, default: str) -> str:
    """Return a human-readable, filesystem-safe DOCX download name.

    Browser/URL encoded names such as ``Görüş%20Metni_698891.docx`` or
    ``G%C3%B6r%C3%BC%C5%9F%20Metni_698891.docx`` are decoded before they
    reach Streamlit's ``file_name`` parameter. Turkish characters and normal
    spaces are deliberately preserved; the downloaded file must not contain
    literal ``%20``/``%C3...`` fragments.
    """
    raw = str(name or default).strip()

    # Decode once or twice so a previously double-encoded filename is also
    # repaired, while avoiding an unbounded decode loop.
    for _ in range(2):
        decoded = unquote(raw)
        if decoded == raw:
            break
        raw = decoded

    raw = raw.replace("\u00a0", " ")
    raw = re.sub(r"[\r\n\t]+", " ", raw)
    raw = re.sub(r"\s+", " ", raw).strip()

    # Strip any user-supplied path and keep only the filename. Handle both
    # Windows and POSIX separators regardless of the server OS.
    raw = raw.replace("\\", "/").split("/")[-1].strip()
    if not raw:
        raw = default
    if not raw.lower().endswith(".docx"):
        raw += ".docx"
    return raw


# -----------------------------------------------------------------------------
# TARİFNAME MODÜLÜ
# -----------------------------------------------------------------------------
TARIFNAME_DRAFT_SCHEMA = r"""
{
  "title":"",
  "technical_field":"",
  "prior_art_general_paragraphs":[""],
  "literature_paragraphs":[""],
  "short_description_intro":"",
  "objectives":[""],
  "unumbered_invention_definition":"",
  "unumbered_invention_features":[""],
  "figure_descriptions":[""],
  "elements":[{"number":"<source reference or default 1>","name":"","description":""}],
  "method_steps":[{"number":"<source reference or default 1001>","text":""}],
  "detailed_paragraphs":[""],
  "formulas":[{"label":"","expression":"","explanation":""}],
  "tables":[{"caption":"","headers":[""],"rows":[[""]]}],
  "experimental_results":[""],
  "alternatives":[""],
  "working_principle":"",
  "system_claim":null,
  "dependent_system_claims":[""],
  "method_claim":null,
  "dependent_method_claims":[""],
  "abstract":"",
  "source_coverage_map":[{"fact_id":"T001","covered":true,"sections":["BULUŞUN DETAYLI AÇIKLAMASI"],"evidence":""}],
  "coverage_audit":{
    "prior_art_complete":true,
    "reference_table_complete":true,
    "formulas_complete":true,
    "tables_complete":true,
    "experimental_results_complete":true,
    "alternatives_complete":true,
    "claims_consistent":true,
    "reference_names_clear":true,
    "reference_order_valid":true,
    "how_test_passed":true,
    "core_difference_present":true,
    "scope_not_overlimited":true,
    "dependent_claims_non_redundant":true,
    "dependent_claim_dependencies_valid":true,
    "example_dimensions_not_claim_limited":true,
    "product_claim_language_valid":true,
    "abstract_single_paragraph_sentence":true,
    "source_attribution_removed":true,
    "all_technical_facts_covered":true,
    "software_carrier_valid":true,
    "detail_intro_sentence_case":true,
    "notes":[""]
  }
}
"""


def tarifname_extraction_prompt(
    source_text: str,
    technical_supplement_text: str = "",
    example_structure_text: str = "",
    language: str = "Türkçe",
) -> str:
    return f"""{TARIFNAME_RULES}
ÇIKTI TARİFNAME DİLİ: {language}. Kaynak hangi dilde olursa olsun teknik envanteri kaynak terimlerine sadık kalarak çıkar.
Aşağıdaki kaynakları yalnızca yapı ve kapsam envanteri çıkarmak için incele. Teknik metni yeniden icat etme,
kısaltma nedeniyle önemli bilgi kaybettirme ve örnek tarifnamelerin teknik içeriğini kullanma. Gömülü şekil, grafik, diyagram, ısı haritası ve görsel teknik sonuçları da kaynak olarak incele; yalnız metin çıkarımına güvenme.

EN ÜST TAMLIK KAPISI:
- BBF ve ek teknik kaynaklardaki HER teknik bilgiyi atomik `technical_facts` maddelerine ayır. Bir cümlede iki ayrı teknik bilgi varsa iki ayrı madde yap.
- Teknik avantajları ve ayırt edici yönleri de ayrı maddeler yap; örneğin fiyat-performans dengesi, gecikme, pil ömrü, cihaz/işletim sistemi/operatör bağımsızlığı gibi kaynakta açık sonuçları atlama.
- İdari/form alanlarını `excluded_nontechnical_items` altında ayır; kişi/sicil/ödül/imza, boş idari alanlar ve yalnız araştırma anahtar kelimeleri tarifname technical_facts listesine girmez.
- `coverage_checklist` genel başlık listesi değil, technical_facts maddelerinin tarifnamede korunacağı içerik gruplarını özetleyen yardımcı listedir.

KAYNAK HİYERARŞİSİ:
1. BBF: temel teknik kaynak.
2. Ek teknik müşteri belgeleri: yalnızca açık teknik dayanak olarak kullanılabilir.
3. Örnek tarifnameler: yalnızca unsur, yöntem adımı, istem ve biçim kurgusunu görmek içindir; teknik içerikleri kullanılamaz.

JSON dışında hiçbir şey yazma.
ŞEMA:
{{
 "title":"",
 "technical_field":"",
 "prior_art_inventory":["BBF'deki her ayrı önceki teknik konu ve kısıt"],
 "technical_problems":[""],
 "technical_solution":[""],
 "technical_effects":[""],
 "technical_facts":[{{"id":"T001","category":"alan/problem/çözüm/unsur/işlev/akış/avantaj/alternatif/kullanım/ayırt_edici_yön/görsel","statement":"Kaynakta açıkça verilen tek bir atomik teknik bilgi","mandatory":true}}],
 "excluded_nontechnical_items":["Form talimatı, kişi/sicil/ödül bilgisi, boş idari alan veya yalnız araştırma anahtar kelimesi gibi tarifnameye taşınmaması gereken içerik"],
 "elements":[{{"number":"<müşteri referansı; yoksa 1>","name":"","function":"","source":"BBF/ek teknik belge"}}],
 "method_steps":[{{"number":"<müşteri referansı; yoksa 1001>","text":"","stage":"","essential":true}}],
 "formulas":[{{"label":"","expression":"","variables":[""],"role":"zorunlu/tercihli"}}],
 "tables":[{{"caption":"","headers":[""],"rows":[[""]]}}],
 "experimental_results":[""],
 "alternatives":[""],
 "use_cases":[""],
 "figures":[""],
 "figure_reference_audit":[{{"figure":"Şekil 1","reference_marks":["1"],"method_marks":["S101"],"symbolic_reference_marks":["UW"],"temporary_marks":[],"possible_wrong_targets":["9 numaralı ok travers yerine başka bir parçaya yöneliyor olabilir"],"missing_reference_candidates":[""],"notes":""}}],
 "claim_core":[""],
 "parallel_step_groups":[{{"summary":"","step_numbers":["1007","1008"],"recommended_claim_location":"ana istem/tek bağımlı istem"}}],
 "stage_distinctions":[{{"step_numbers":["1001","1006"],"difference":""}}],
 "has_system_basis":true,
 "has_method_basis":true,
 "recommended_claim_mode":"Yalnızca yöntem/Sistem ve yöntem/Yalnızca sistem",
 "recommended_claim_mode_reason":"",
 "coverage_checklist":["Kaynakta bulunan ve taslakta mutlaka korunması gereken her içerik grubu"]
}}

BBF:
---
{source_text}
---

EK TEKNİK BELGELER/NOTLAR:
---
{technical_supplement_text}
---

ÖRNEK TARİFNAMELER - YALNIZCA KURGU:
---
{example_structure_text}
---
"""

def tarifname_extraction_quality_prompt(
    source_text: str,
    technical_supplement_text: str,
    extracted: dict[str, Any],
    source_passage_registry: list[dict[str, str]],
    language: str = "Türkçe",
) -> str:
    return f"""{TARIFNAME_RULES}
ÇIKTI TARİFNAME DİLİ: {language}.
Bu aşamada TARİFNAME YAZMA. Ham BBF ve ek teknik belgeleri SIFIRDAN yeniden oku ve mevcut yapılandırılmış envanteri denetle.

EN ÜST KURAL: BBF ve ek teknik kaynaklardaki her teknik bilgi atomik `technical_facts` listesinde bulunmalıdır. Mevcut envantere güvenme; kaynakta olup listede olmayan teknik alan, problem, çözüm, unsur, işlev, işlem akışı, avantaj, teknik etki, kullanım, alternatif, ayırt edici yön veya görsel/şema bilgisi varsa EKLE. Özellikle kaynakta açıkça geçen performans/enerji/gecikme/fiyat-performans sonuçları ile cihaz, işletim sistemi veya operatör altyapısı bağımsızlığı gibi ifadeleri atlama.
Kişi/sicil/ödül/imza, form talimatları, boş idari alanlar ve yalnız araştırma anahtar kelimeleri `excluded_nontechnical_items` içinde kalmalıdır; bunları technical_facts yapma.
Her technical fact tek bir atomik teknik anlam taşısın, `id` alanları T001, T002... biçiminde benzersiz ve sıralı olsun. Teknik olarak sınıflandırılan HER fact tarifnameye aktarılması zorunlu olduğundan `mandatory=true` olmalıdır; teknik bir fact için `mandatory=false` kullanma.
Mevcut elements/method_steps/formulas/tables/alternatives/use_cases/figure audit alanlarını kaynakla karşılaştır ve eksik teknik içerik varsa tamamla; kaynakta olmayan bilgi ekleme.

HAM KAYNAK PASAJ KAPISI: Aşağıdaki SOURCE_PASSAGE_REGISTRY yerel kod tarafından ham dosyalardan deterministik çıkarılmıştır. HER passage_id `source_passage_audit` içinde TAM BİR KEZ yer almalıdır. Teknik pasaj classification=`technical` olmalı ve en az bir geçerli technical_facts id'sine bağlanmalıdır. Gerçekten idari/form niteliğindeki pasaj classification=`nontechnical` olabilir ancak reason boş bırakılamaz. Teknik içerikli pasajı nontechnical işaretleyerek atlamak yasaktır.

JSON dışında hiçbir şey yazma. Mevcut JSON yapısını koruyarak düzeltilmiş TAM envanteri döndür ve `source_passage_audit` alanını ekle.

HAM BBF:
---
{source_text}
---

EK TEKNİK BELGELER:
---
{technical_supplement_text}
---

SOURCE_PASSAGE_REGISTRY:
{json.dumps(source_passage_registry, ensure_ascii=False, indent=2)}

MEVCUT ENVANTER:
{json.dumps(extracted, ensure_ascii=False, indent=2)}

ZORUNLU EK ALAN ŞEMASI:
"source_passage_audit":[{{"passage_id":"B0001","classification":"technical/nontechnical","fact_ids":["T001"],"reason":""}}]
"""


def _validate_technical_fact_inventory(extracted: dict[str, Any]) -> None:
    facts = extracted.get("technical_facts") or []
    if not facts:
        raise ValueError("BBF teknik bilgi envanteri boş; technical_facts oluşturulmadan tarifname hazırlanamaz.")
    ids: list[str] = []
    for fact in facts:
        fid = str(fact.get("id", "") or "").strip()
        statement = str(fact.get("statement", "") or "").strip()
        if not fid or not statement:
            raise ValueError("BBF technical_facts envanterinde kimliği veya teknik açıklaması boş madde bulundu.")
        if fact.get("mandatory", True) is not True:
            raise ValueError(
                f"BBF technical_facts envanterinde {fid} teknik bilgi olmasına rağmen mandatory=false işaretlenmiş. "
                "Tarifname akışında bütün teknik bilgiler zorunlu kapsama tabidir."
            )
        ids.append(fid)
    if len(ids) != len(set(ids)):
        raise ValueError("BBF technical_facts envanterinde tekrarlanan fact_id bulundu.")


def tarifname_literature_prompt(extracted: dict[str, Any], count: int, jurisdiction: str, language: str = "Türkçe") -> str:
    return f"""Aşağıdaki buluş için tam olarak {count} teknik olarak yakın patent dokümanı araştır. Web araması kullan.
Nihai tarifname dili: {language}. title_en alanında patentin doğrulanmış özgün İngilizce başlığını koru; title_tr alanını yalnız Türkçe çıktı için yardımcı çeviri olarak doldur.
Doküman uydurma; yayın/başvuru numarasını, İngilizce başlığını, yayın tarihini ve kaynak bağlantısını doğrula.
Tercih edilen ülke/veri tabanı: {jurisdiction or 'global'}.
Her doküman için teknik konusu ile buluşta bulunmayan temel teknik farkı açıkla. Bu aşamada tarifname metni yazma.
JSON dışında yazma.
ŞEMA:
{{"documents":[{{"application_number":"","title_en":"","title_tr":"","publication_date":"","jurisdiction":"","summary":"","difference":"","source_url":""}}]}}
BULUŞ ENVANTERİ:
{json.dumps(extracted, ensure_ascii=False, indent=2)}"""


def tarifname_drafting_prompt(
    extracted: dict[str, Any],
    claim_mode: str,
    literature: list[dict[str, Any]],
    source_text: str,
    technical_supplement_text: str,
    example_structure_text: str,
    language: str = "Türkçe",
) -> str:
    language_instruction = (
        "Nihai tarifnamenin TAMAMI İngilizce olmalıdır. Başlıklar, teknik açıklama, istemler ve özet İngilizce yazılmalı; Türkçe patent kalıpları bırakılmamalıdır. TECHNICAL FIELD ilk paragrafı yalnız 'The invention relates to ... .' cümlesinden oluşmalı, ikinci paragraf ayrı olarak 'In particular, the invention relates to ... .' ile başlamalıdır. İngilizce bağımsız istemlerde 'comprising:'; bağımlı istemlerde 'The system according to claim X, wherein...' veya 'The method according to claim X, wherein...' yapısını kullan. Yazılım ağırlıklı buluşlarda 'an electronic device', 'a processing unit' veya 'software executed on an electronic device' gibi geniş donanımsal taşıyıcı kullan."
        if _english_spec(language) else
        "Nihai tarifnamenin TAMAMI Türkçe olmalıdır ve mevcut Türkçe tarifname kuralları aynen uygulanmalıdır."
    )
    return f"""{TARIFNAME_RULES}
{language_instruction}
Aşağıdaki kaynaklara dayanarak seçilen dilde patent tarifnamesinin TAM metnini oluştur.
İstem yapısı: {claim_mode}

KRİTİK TALİMATLAR:
- BBF'nin bütün teknik bilgilerini kullan. Uzun önceki teknik, formüller, tablolar, deneysel sonuçlar, alternatifler ve referans tablosu atlanamaz.
- Yapılandırılmış envanter yalnızca yardımcıdır. Çelişki halinde ham BBF ve açık teknik müşteri belgeleri esas alınır.
- Örnek tarifnamelerden yalnızca kurguyu öğren; teknik bilgi aktarma.
- technical_field iki paragraf olmalıdır ve paragraflar \n\n ile ayrılmalıdır. Türkçe çıktıda ilk paragraf yalnız “Buluş, ... ile ilgilidir.”, ikinci paragraf “Buluş, özellikle ...” ile; İngilizce çıktıda ilk paragraf yalnız “The invention relates to ... .”, ikinci paragraf “In particular, the invention relates to ... .” ile başlamalıdır. Seçilen istem modu “Sistem ve yöntem” ise Türkçe ilk paragrafın sonu özellikle “... sistemi ve yöntemi ile ilgilidir.” olmalıdır; yalnız sistemde “... sistemi ile ilgilidir.”, yalnız yöntemde “... yöntemi ile ilgilidir.” yapısı kullanılmalıdır.
- ÖNCEKİ TEKNİK'teki aynı anlatımın devamı olan “Özellikle...”, “Bununla birlikte...”, “Bu nedenle...” gibi cümleleri ayrı paragraf yapma. Patent literatürü dokümanları ise ayrı ayrı paragraf olsun.
- Türkçe tarifnamede her patent literatürü paragrafında doğrulanmış İngilizce başlık ile Türkçe başlık karşılığı birlikte yazılsın. Türkçe literatür paragrafı bağlayıcı taslak dilini izlesin: “Literatürde yapılan araştırmalar sonucu ... numaralı, İngilizce başlığı ‘...’ ve Türkçe karşılığı ‘...’ olan patent dokümanına rastlanmıştır. Söz konusu başvuru/doküman ... ile ilgilidir. Ancak bahsedilen başvuruda/dokümanda ... ile ilgili bir emareye rastlanmamıştır.” “Buluşta ise ...” biçiminde karşılaştırmalı görüş/savunma dili kullanılmasın. İngilizce tarifnamede özgün İngilizce patent başlığı kullanılsın; Türkçe başlık karşılığı nihai İngilizce metne eklenmesin.
- BULUŞUN DETAYLI AÇIKLAMASI'nda numaralı sistem/cihaz unsurlarını tek tek ayrı paragraf yapma; bütün unsur açıklamalarını teknik akış içinde tek sürekli paragrafta topla. Sistem unsuru-yöntem adımı ilişkisini açıklamak için “İşlem Adımı / Gerçekleştiren Unsur / Açıklama” türü tablo oluşturma. Bu ilişkiyi modül (1), sonraki modül (2) ve ilgili yöntem adımı (1001, 1002...) arasındaki veri/işlev bağlantısını gösteren doğal teknik paragraf olarak yaz. Yalnız ham kaynakta gerçekten sayısal/deneysel veri tablosu olan tabloları tables alanında koru. Gerçekten ayrı bir yapılanma/alternatif/yöntem/çalışma prensibi ayrıca paragraf olabilir.
- Ana istemde zorunlu teknik çekirdeği kapsayıcı biçimde ver. Aynı işlemin birinci/ikinci/k'ıncı tekrarlarını ana istemde gereksiz yere ayrı satırlara bölme. Bu ayrıntıları, aynı alt akışa aitse tek bağımlı istemde topla.
- Buluş ağırlıklı olarak yazılım/algoritma/modül/birimlerden oluşuyorsa bağımsız istemleri soyut yazılım olarak bırakma. Kaynakta özel donanım zorunlu değilse geniş bir donanımsal taşıyıcı kullan. Türkçede “bir elektronik cihaz üzerinde koşturulan yazılım vasıtasıyla ...”, İngilizcede “software executed on an electronic device ...” veya eşdeğer teknik taşıyıcı dili kullanılabilir. Gereksiz sunucu, cep telefonu/phone veya kişisel bilgisayar daraltması yapma; özel donanım uydurma.
- AYNI referanssız elektronik işlem birimi/cihaz taşıyıcısı üzerinde AYNI çalışma ilişkisine sahip birden fazla ARDIŞIK yürütülebilir yazılım/modül/kontrolör/arayüz/yığın varsa taşıyıcı ifadesini her birinde tekrar etmek zorunda değilsin. Türkçe sistem isteminde `elements` listesine bir grup nesnesi koyabilirsin: `{{"lead":"bir elektronik işlem birimi üzerinde koşturulan yazılım vasıtasıyla çalışan ve;","subelements":["... modülü (2),","... kontrolörü (3),"]}}`. Lead referans taşımaz; her subelement ayrı bir yeni referanslı unsuru ilk-tanım sırasıyla tanımlar. VERİTABANI, bellek, veri deposu, profil tablosu veya salt veri yapısı kaynak açıkça yürütülebilir yazılım/modül olduğunu söylemiyorsa bu ortak grubun altına ALINMAZ; ayrı unsur yazılır. Bu grup yalnız aynı taşıyıcı gerçekten bütün alt unsurlar için ortaksa kullanılır ve Word'de gerçek iç içe bullet olarak yazılır.
- İstemleri yalnız hedeflenen sonuç veya fonksiyonla bırakma. Özellikle bağımsız istemde teknikte uzman kişinin “nasıl gerçekleştiriliyor?” sorusuna cevap verecek şekilde, kaynakta açık dayanağı bulunduğu ölçüde işlemi yapan teknik unsur/taşıyıcıyı, kullanılan girdiyi veya önceki unsurdan gelen veriyi, teknik işlem/mekanizmayı ve ortaya çıkan teknik çıktının sonraki unsurla bağlantısını yaz. “tespit eden / dönüştüren / optimize eden / classifying / transforming / determining” gibi sonuç bildiren fiiller kaynak mekanizmayı açıklıyorsa tek başına yeterli sayılmaz. Buna karşılık tercihli uygulama ayrıntılarıyla ana istemi gereksiz daraltma. Yazılım/modül unsurlarını İngilizce claim sırasını taklit ederek `X modülü (N), ... yapan bir modül` biçiminde kurma; Türkçe istemde önce kaynak destekli teknik işlev/mekanizma yazılır, unsur adı ve `(N)` referansı bu işlevi tanımlayan sıfat-fiil yapısının sonunda gelir: `... verilerini birlikte değerlendirerek ... değerini hesaplayan X modülü (N),`.
- Kaynakta açık matematiksel bağıntı/formül varsa `formulas[].expression` alanında formül metnini koru. Aynı bağıntının bağımlı istemde açıkça yazılması gerekiyorsa düz `x = ...` metni kullanma; bağıntıyı `[[EQ: x = ...]]` işaretleyicisi içinde yaz. Word üreticisi bunu gerçek OMML denklem nesnesine dönüştürecektir. Formül zorunlu teknik çekirdek değilse ana istemi gereksiz daraltma; bağımlı istem/detaylı açıklamada tut.
- Bağımlı istemleri kaynakta geçen her ayrıntı için çoğaltma. Yalnız ana isteme gerçek teknik daraltma/geri çekilme konumu sağlayan seçilmiş özellikleri kullan; istem bağımlılığı ana donanımsal taşıyıcıyı zaten taşıyorsa alt istemde elektronik cihaz/yazılım ifadesini gereksiz yere tekrar etme.
- Eğitim/genel aşama ile test aşamasındaki paralel akışları aynı mantıkla fakat ayrı teknik aşamalar olarak kur.
- REFERANS NUMARALARI bölümünde müşteri tarafından sistem/cihaz unsurları veya yöntem işlem adımları için verilmiş açık referansları AYNEN koru; 10, 20..., S101..., M1... veya başka bir referans ailesini sırf standartlaştırmak için değiştirme. Sistem/cihaz modüllerinde hiç referans yoksa kaynak sırasıyla 1, 2, 3... ver. Yöntem işlem adımlarında hiç referans yoksa varsayılan 1001, 1002, 1003... ailesini kullan. Kısmen numaralandırılmış kaynakta mevcut müşteri işaretlerini koru, yalnız boş kalanlara çakışmayacak varsayılan referans ata. Word'deki yöntem referans satırı `1001. ...` biçiminde başlar; bu satırın içinde sistem/cihaz unsur işaretleri `(1)`, `(2)` vb. yazılmaz. Parantezli unsur referansları BULUŞUN DETAYLI AÇIKLAMASI bölümünden itibaren başlar.
- “Yöntemin gerçekleştirdiği işlem adımları aşağıdaki gibidir:” bölümü için method_steps tam ve tutarlı olsun. Kaynaktaki yöntem referansları varsa aynen korunsun; yalnız kaynakta hiç yöntem referansı yoksa 1001’den başlayan varsayılan sıra oluşturulsun. REFERANS NUMARALARI, detaylı açıklamadaki yöntem listesi ve bağımsız yöntem isteminde aynı referanslı adımın teknik metni birebir aynı olsun. Detaylı açıklamadaki ara maddeler virgülle, son madde noktayla bitsin. Bağımsız yöntem istemindeki ara adımlar virgülle bitsin, son adım noktalamasız bitsin.
- Yalnızca yöntem modunda system_claim null olmalıdır. Yalnızca sistem modunda method_claim null olmalıdır.
- Her bağımlı istem ana isteme göre gerçek bir daraltma sağlamalıdır.
- Patent literatürü yalnızca ÖNCEKİ TEKNİK bölümünde kullanılsın.
- Kullanıcıya sunulan tarifname metninde “BBF”, “buluş bildirim formu” veya kaynak dokümana atıf yapan benzer ifadeler kesinlikle bulunmasın; teknik bilgi doğrudan buluş anlatımı olarak yazılsın.
- YAPILANDIRILMIŞ ENVANTER içindeki `technical_facts` listesinin HER maddesi nihai tarifnamede uygun yerde korunmalıdır. Teknik olarak sınıflandırılmış bir fact için `mandatory=false` kullanılamaz. Taslak JSON'daki `source_coverage_map` alanında her fact_id için `covered=true`, karşılık bulunduğu bölüm(ler) ve tarifname taslağında birebir geçen en az 20 karakterlik kısa kanıt alıntısı verilmelidir. Bir teknik fact isteme uygun değilse isteme zorla koyma; teknik alan/önceki teknik/kısa açıklama/detaylı açıklama/çalışma prensibi/alternatif/özet içinde uygun yere koy. Hiçbir teknik fact karşılıksız bırakılamaz.
- Sistem ve yöntem istemleri birlikte oluşturuluyorsa başlık seçilen dile uygun olarak “... Sistemi ve Yöntemi” veya “... System and Method” yapısını taşısın.
- REFERANS NUMARALARI bölümünde unsur adlarında yalnızca ilk kelimenin ilk harfi büyük olsun; standart teknik kısaltmaları koru. Unsur adları cümle içinde geçtiğinde küçük harfle başlat.
- Türkçe çıktıda “Buluşun bir gerçekleştirilmesinde” ifadesini kullanma; gerekli yerde “Buluşun bir yapılanmasında” yaz ve “Mevcut buluş” kullanma. İngilizce çıktıda doğal patent dili kullan.
- BULUŞUN DETAYLI AÇIKLAMASI giriş cümlesinde buluş adını başlıktaki büyük harf düzeniyle tekrar etme. Türkçe cümle içinde normal küçük harf düzeni kullan; SIM, eSIM, NFC, API gibi teknik kısaltmaları aynen koru.
- Önceki teknik bölümüne ham kaynakta verilen bütün teknik arka plan, eksiklik ve problem anlatımını aktar; literatür paragrafları bunların yerine geçmez.
- ŞEKİLLERİN KISA AÇIKLAMASI kısa ve işlevsel olsun; gerekli değilse yöntem adımı numara aralığını şekil açıklamasında tekrarlama.
- Bağımlı istemlerde Türkçe çıktıda “Önceki istemlerden herhangi birine” kalıbını, İngilizce çıktıda belirsiz “any preceding claim” zincirlerini varsayılan olarak kullanma; ek özellik hangi ana unsur veya işlem adımının ayrıntısıysa doğrudan onu tanımlayan gerekli isteme bağla.
- objectives alanında Türkçe çıktı için amaç gövdesi “... sağlamaktır.” gibi tam yüklemle bitsin. İngilizce çıktı için her objective baştan sona tam bir cümle olarak yazılsın, örneğin “The main objective of the invention is to ... .”
- Müşteri şekillerini teknik kaynak olarak aynen esas al. Görseldeki gerçek referans işaretlerini sayısal unsur, yöntem adımı, sembolik referans ve geçici şekil numarası olarak ayır. Gömülü grafik/ısı haritası/diyagram üzerindeki teknik sonuçları tamlık kontrolünde dikkate al. Geçici şekil numarasını yeni unsur referansı yapma. Şekildeki mevcut ok/numarayı otomatik olarak doğru kabul etme; nihai şekil aşamasında referans → unsur adı → detaylı açıklamadaki teknik tanım → fiziksel karşılık eşleştirmesi yapılacaktır. Belirli alt parçaya ait referansı tüm tertibat olarak yorumlama ve her görünür parçayı zorla numaralandırma.
- REFERANS NUMARALARI bölümünde ayrı numaraya sahip iki unsur müşteri/BBF şekli üzerinde `2-3` gibi tek ve ayırt edilemeyen bir kutu/hedefte gösterilmişse bunu nihai şekle aynen taşıma. Müşteri şeklinin ortak taşıyıcısını ve teknik ilişkilerini koruyarak ayrı kutucuk/çağrı/oklarla her referansı ayrı göster. Şekiller Word isteniyorsa nihai `elements` referanslarının tamamı şekil SETİNDE en az bir kez görünmelidir; sistem+yöntem şekilleri hazırlanıyorsa yöntem referansları da akış şekillerinde tam kapsanmalıdır.
- BBF referans/BOM tablosunu kaynak envanteri olarak çıkar; fakat “Diğer parçalar/Diğer elemanlar” gibi belirsiz üst başlıkları nihai patent unsuru yapma. Altında açıkça tanımlanan gerçek parçaları teknik adlarıyla kullan.
- Sistem şeması yalnız “sistem” sözcüğüne özgü değildir; cihaz, ürün, tertibat, düzenek ve yapılanma istemleri de aynı ürün istem dil kurallarına tabidir. Yöntem dışındaki bağımlı istemler “olmasıdır.” veya “içermesidir.” ile bitmelidir.
- Ana istemde bir referanslı unsuru ilk kez tanımlarken henüz tanımlanmamış sonraki referanslı unsurları kullanma. İlk/ana taşıyıcı unsuru kendi yapısı ve işleviyle tanımla; sonra diğer unsurları sırayla daha önce tanımlanmış unsurlara bağla. Kural olarak her claim bullet yalnız bir yeni referanslı unsur tanımlasın.
- Ürün/yapılanma isteminde “somun flanşının gövdeye bağlanması”, “parçanın oluşturulması” gibi işlem isimleştirmesi kullanma; “gövdeye bağlanan somun flanşı”, “... yapısına sahip parça” gibi unsur merkezli dil kullan.
- Bir unsur yalnız nerede bulunduğuyla bırakılmasın; kaynak destekliyorsa teknik işlevi de yazılsın. Örneğin sızdırmazlık elemanının bağlantı bölgesinde akışkan sızdırmazlığını sağladığı belirtilsin.
- Aynı olmayan fiziksel unsurları “ve/veya” ile tek unsur gibi birleştirme. Gerçek alternatifleri teknik kimlikleri ve işlevleriyle açık ayrı maddelerde tanımla.
- “vidalanan/kaynaklanan/yapıştırılan” ve belirli çap/diş/ölçü gibi daraltıcı ifadeleri yalnız zorunlu teknik çekirdek veya farklılaştırıcı mekanizma ise ana istemde tut. Değilse kaynakla uyumlu daha geniş bağlantı dili kullan.
- Ana istemde tanımlanan bir özelliği başka bullet'ta tekrar etme. Sonraki unsur yalnız kendi ilişkisi ve işleviyle tanımlansın.
- Sistem ve yöntem bağımlı istemlerinin HER BİRİNİ semantik olarak ana/üst istemle ve önceki bağımlı istemlerle karşılaştır; aynı teknik özelliği farklı kelimelerle tekrar eden bağımlı istem üretme. Her alt istem gerçek ek teknik sınırlama getirmelidir. Bir istem silinir/değişirse sonraki bağımlılık numaralarını yeniden kur. Türkçe bağımlı YÖNTEM istemleri eylem sonucu ile bitmez; tek ek adım varsa `işlem adımını içermesidir.`, birden fazla ek adım varsa `işlem adımlarını içermesidir.` şeklinde kapanır.
- Örnek ölçü/çap/diş değerlerini zorunlu değilse istemlere taşıma; detaylı açıklamada örnek yapılanma olarak koru ve kaynak destekliyorsa farklı ölçülere uygulanabilirliği açıkla.
- Referans adı koruma kapsamını gereksiz daraltmasın: özel bir örnek (örn. O-ring) daha genel kaynak destekli teknik işlevin gerçekleştirmesiyse unsur adını genel teknik kavramla (örn. sızdırmazlık elemanı) kur; özel örneği detaylı açıklamada parantez içinde ver.
- Her teknik ayrıntıya zorla referans verme; yapıştırıcı/malzeme/kaplama gibi özellikler ayrı referans gerektirmiyorsa numarasız olarak detaylı açıklama ve uygun bağımlı istemde kullanılabilir.
- Kullanıcıya görünen tarifnamede “BBF’de”, “müşteri tarafından iletilen teknik çizimde”, “müşteri bilgilerine göre”, “ek teknik belgede” gibi kaynak atıfları bulunmasın.
- Türkçe özet tek paragraf ve tek cümle olsun; referans numarası kullanma.

JSON dışında hiçbir şey yazma.
ÇIKTI ŞEMASI:
{TARIFNAME_DRAFT_SCHEMA}

SİSTEM İSTEMİ ŞEMASI (varsa):
{{"preamble":"","elements":["tek düz unsur maddesi",{{"lead":"bir elektronik işlem birimi üzerinde koşturulan yazılım vasıtasıyla çalışan ve;","subelements":["ayrı referanslı alt unsur (2)","ayrı referanslı alt unsur (3)"]}},"sonraki düz unsur maddesi"],"closing":"içermesidir."}}
YÖNTEM İSTEMİ ŞEMASI (varsa):
{{"preamble":"","steps":[""],"closing":"işlem adımlarını içermesidir."}}

YAPILANDIRILMIŞ ENVANTER:
{json.dumps(extracted, ensure_ascii=False, indent=2)}

ONAYLANAN PATENT LİTERATÜRÜ:
{json.dumps(literature, ensure_ascii=False, indent=2)}

HAM BBF:
---
{source_text}
---

EK TEKNİK BELGELER/NOTLAR:
---
{technical_supplement_text}
---

ÖRNEK TARİFNAMELER - YALNIZCA KURGU:
---
{example_structure_text}
---
"""

def tarifname_quality_prompt(
    source_text: str,
    technical_supplement_text: str,
    extracted: dict[str, Any],
    draft: dict[str, Any],
    claim_mode: str,
    literature: list[dict[str, Any]],
    language: str = "Türkçe",
    validation_feedback: str = "",
) -> str:
    language_instruction = (
        "Çıktı dili İngilizcedir. Tüm bölüm metinleri, istemler ve özet İngilizce kalmalı; TECHNICAL FIELD ve İngilizce claim kalıplarını kontrol et."
        if _english_spec(language) else
        "Çıktı dili Türkçedir. Türkçe tarifname ve istem kalıplarını kontrol et."
    )
    return f"""{TARIFNAME_RULES}
{language_instruction}
Aşağıdaki tarifname taslağını kaynaklarla SATIR SATIR ve `technical_facts` bazında karşılaştır ve eksik/yanlış hususları düzelterek tam JSON'u yeniden üret.
Bu bir özetleme görevi değildir. Kaynakta olup taslakta bulunmayan her teknik bilgi geri eklenmelidir. `technical_facts` içindeki HER madde zorunludur ve her biri için `source_coverage_map` kaydı oluştur; bölüm ve kanıt metni boş bırakılamaz; evidence alanı tarifname taslağında birebir geçen en az 20 karakterlik bir alıntı olmalıdır. Genel “tam” beyanı yeterli değildir.

ÖNCEKİ OTOMATİK DOĞRULAMA GERİ BİLDİRİMİ (varsa):
{validation_feedback or "Yok"}
Bu geri bildirimde hata varsa JSON'u buna göre düzelt; yalnız hatayı yamamakla kalma, ilgili kuralın teknik mantığını bütün taslakta yeniden kontrol et.

ZORUNLU KONTROL LİSTESİ:
1. BBF'deki önceki teknik anlatımının tamamı korunmuş mu?
2. Referans tablosundaki bütün unsurlar ve yöntem adımları var mı?
3. Farklı numaralı adımlar yanlışlıkla aynı metinle mi yazılmış? Aynı veri farklı aşamada kullanılıyorsa aşama farkı açık mı?
4. “Yöntemin gerçekleştirdiği işlem adımları” tam liste ve referans tablosuyla uyumlu mu?
5. Ana istem zorunlu çekirdeği kapsıyor mu; paralel tekrarları gereksiz yere tek tek sayıyor mu?
6. Aynı alt akışa ait paralel analizler ve çıktılar gerekiyorsa tek bağımlı istemde mi?
7. Eğitim/genel ve test aşamaları paralel fakat ayrı olarak mı kurulmuş?
8. Formüller, değişken açıklamaları, tablolar, deneysel sonuçlar, alternatifler ve teknik etkiler eksiksiz mi?
9. Seçilen istem modu ({claim_mode}) ile başlık, açıklama ve istemler tutarlı mı? Sistem ve yöntem ise başlık seçilen dile uygun biçimde “Sistemi ve Yöntemi” veya “System and Method” yapısını taşıyor mu?
10. Bağımlı istemler gerçek daraltma sağlıyor mu?
11. Kullanıcıya sunulan metinde “BBF” veya “buluş bildirim formu” gibi kaynak atfı kalmış mı? Kalmışsa doğrudan buluş anlatımına dönüştür.
12. “Buluşun bir gerçekleştirilmesinde” veya “Mevcut buluş” kalıbı var mı? Varsa “Buluşun bir yapılanmasında” / “Buluş” diline dönüştür.
13. REFERANS NUMARALARI unsur adları yalnızca ilk kelime büyük olacak biçimde mi? Cümle içindeki unsur adları küçük harfle mi başlıyor?
14. Önceki teknik kaynakta verilen bütün müşteri teknik arka planını ve eksikliklerini içeriyor mu?
15. Teknik alan seçilen dilde iki kademeli mi? Türkçede ilk paragraf yalnız “Buluş, ... ile ilgilidir.”, ikinci paragraf “Buluş, özellikle ...” ile; İngilizcede ilk paragraf yalnız “The invention relates to ... .”, ikinci paragraf “In particular, the invention relates to ... .” ile başlamalıdır. İstem modu Sistem ve yöntem ise Türkçe ilk paragraf mutlaka “... sistemi ve yöntemi ile ilgilidir.” şeklinde bitmelidir; yalnız sistem/yöntem modlarında buna uygun tek tür adı kullanılmalıdır.
16. ÖNCEKİ TEKNİK'te “Özellikle...”, “Bununla birlikte...”, “Bu nedenle...” gibi aynı anlatımın devamları gereksiz yere ayrı paragraf yapılmış mı? Yapılmışsa birleştir.
17. Patent literatürü başlıkları seçilen dile uygun mu? Türkçe çıktıda İngilizce özgün başlık + Türkçe karşılığı kullanılmalı ve paragraf “Literatürde yapılan araştırmalar sonucu ... rastlanmıştır. Söz konusu başvuru/doküman ... ile ilgilidir. Ancak bahsedilen başvuruda/dokümanda ... ile ilgili bir emareye rastlanmamıştır.” taslak kalıbını izlemelidir. “Buluşta ise ...” dili kullanılmamalıdır. İngilizce çıktıda özgün İngilizce başlık kullanılmalı.
18. BULUŞUN DETAYLI AÇIKLAMASI'nda numaralı unsurlar gereksiz yere ayrı ayrı paragraflara bölünmüş mü? Bölündüyse tek sürekli unsur paragrafında birleştir.
19. Detaylı açıklamadaki yöntem madde listesinde ara maddeler virgül, son madde nokta ile bitiyor mu? Kaynakta yöntem referansları verilmişse aynen korunmuş mu; kaynakta hiç referans yoksa 1001, 1002... varsayılan sırası kullanılmış mı? Bağımsız yöntem istemindeki son işlem adımı noktalamasız mı?
20. Şekil açıklamaları kısa mı ve gerekli olmayan yöntem adımı numara aralıklarını tekrarlamıyor mu?
21. Müşteri şekillerindeki gerçek unsur/yöntem/sembolik referanslar REFERANS NUMARALARI ile uyumlu mu? Geçici şekil numaraları yeni referans olarak uydurulmuş mu?
22. Şekilde kullanılan UW, UW_F, UW_PL, UW_R, UW_M gibi sembolik referansların tarifnamede açık karşılığı var mı?
23. Müşterinin sistem/cihaz ve yöntem referans işaretleri aynen korunmuş mu? Kaynakta referans verilmeyen sistem/cihaz modülleri 1, 2, 3...; referans verilmeyen yöntem adımları 1001, 1002... varsayılanıyla tamamlanmış mı?
24. “İşlem Adımı / Gerçekleştiren Unsur / Açıklama” türü açıklama tablosu oluşturulmuş mu? Varsa tabloyu kaldır ve aynı içeriği modül-referans-yöntem adımı ilişkilerini koruyan doğal teknik paragrafa dönüştür.
25. Gömülü şekil, grafik, ısı haritası ve diyagramlarda bulunan teknik sonuçlar ile açıklayıcı etiketler tamlık kontrolünde değerlendirilmiş mi?
26. ŞEKİLLERİN KISA AÇIKLAMASI içindeki Şekil 1, Şekil 2, Şekil 3... satırları aralarında boş paragraf gerektirmeyecek biçimde ardışık açıklamalar olarak verilmiş mi?
27. Buluş yazılım/algoritma ağırlıklıysa bağımsız sistem ve/veya yöntem istemi geniş bir donanımsal taşıyıcıya, tercihen elektronik cihaz üzerinde koşturulan yazılıma, açıkça dayandırılmış mı? Gereksiz sunucu/telefon/bilgisayar daraltması yapılmış mı?
28. Bağımlı istemler kaynakta geçen her ayrıntıyı ayrı isteme dönüştürmek yerine yalnız stratejik ve gerçek daraltma sağlayan özelliklerle kontrollü tutulmuş mu?
29. Bağımsız istemler yalnız “ne/sonuç” anlatımıyla mı kalıyor, yoksa teknikte uzman kişinin “nasıl gerçekleştiriliyor?” sorusuna cevap verecek biçimde kaynakta dayanaklı teknik taşıyıcı, girdi/veri, işlem mekanizması ve çıktı/ilişkiyi yeterince gösteriyor mu? Özellikle yazılım/algoritma istemlerinde yalnız elektronik cihaz demekle yetinilmiş mi, yoksa yazılımın cihaz üzerinde hangi teknik yapılar ve işlemler üzerinden sonucu ürettiği de istemden anlaşılabiliyor mu? Ana istem bu amaçla gereksiz tercihli ayrıntılarla daraltılmışsa sadeleştir.
30. Şekildeki mevcut referans numarası/okun fiziksel hedefi, referans listesi ve detaylı açıklamadaki unsur tanımıyla gerçekten uyuşuyor mu? Mevcut işaret sırf şeklin üzerinde bulunduğu için doğru kabul edilmemelidir.
31. Belirli bir alt parçaya ait referans genel tertibatı gösteriyor mu? Örneğin `9 = Travers` ise referans yalnız traversin fiziksel karşılığına yönelmelidir.
32. İlgili şekilde görünür ve tarifnamede gerçek referansla tanımlı bir unsur numarasızsa bu durum figure/reference audit notunda belirlenmiş mi? Görünmeyen veya konumu belirsiz unsurlar için uydurma işaretleme yapılmamalıdır.
33. BBF/BOM tablosunda “Diğer parçalar/Diğer elemanlar” gibi belirsiz satır nihai referans unsuru yapılmış mı? Yapılmışsa gerçek teknik parçalara ayır veya kaynakta açık karşılığı yoksa referans listesinden çıkar; teknik bilgiyi kaybetme.
34. Ana istemde bir referanslı unsur ilk kez tanımlanırken henüz tanımlanmamış daha sonraki referanslı unsur kullanılmış mı? Kural olarak her bullet tek yeni referanslı unsur tanımlıyor mu ve tanım sırası ana taşıyıcıdan bağlı unsurlara doğru mu?
35. Ürün/sistem/cihaz/tertibat/yapılanma istemlerinde “... bağlanması”, “... oluşturulması”, “... yapılması”, “... sağlanması” gibi yöntem/işlem isimleştirmesi var mı? Varsa unsur merkezli “... bağlanan ...”, “... yapısına sahip ...”, “... sağlayan ...” diline dönüştür.
36. Yöntem dışındaki tüm bağımlı istemler “olmasıdır.” veya “içermesidir.” mantığıyla mı bitiyor?
37. Aynı olmayan iki fiziksel unsur “ve/veya” ile tek unsur gibi bulanıklaştırılmış mı? Varsa teknik kimliklerini ve alternatif işlevlerini ayrı açık unsurlar olarak yaz.
38. Her zorunlu unsur yalnız konumla mı tanımlanmış, yoksa kaynak desteklediği ölçüde teknik işlevi de açıklanmış mı? Özellikle sızdırmazlık, kilitleme, ölçme, işleme gibi işlevsel unsurlarda “ne yapıyor?” cevabı görünür mü?
39. Ana istem buluşun mevcut tekniğe göre zorunlu/farklılaştırıcı çekirdeğini taşıyor mu? Genişletme uğruna esas teknik fark kaybolmuş mu?
40. Ana istemde “vidalanan/kaynaklanan/yapıştırılan” veya belirli çap/diş/ölçü gibi zorunlu olmayan kapsam daraltıcı ifade var mı? Kaynak bunları vazgeçilmez kılmıyorsa daha geniş ama dayanaklı dile dönüştür.
41. Ana istemin farklı maddelerinde aynı bağlantı/işlev özelliği gereksiz tekrar edilmiş mi? Sonraki maddeyi yalnız kendi unsur ilişkisi ve işleviyle sınırla.
42. Her bağımlı istem semantik olarak ana/üst isteme yeni bir teknik sınırlama ekliyor mu? Aynı özelliği farklı kelimelerle tekrar eden istemleri sil veya gerçek daraltıcı özelliğe dönüştür.
43. Bir istem silinmiş/değişmişse sonraki bağımlılık numaraları doğru isteme mi bağlı? Geçersiz bağımlılık kalmış mı?
44. Örnek mm/inç/çap/diş/ebat değerleri zorunlu teknik sınır olmadığı halde istemlere taşınmış mı? Taşınmışsa detaylı açıklamadaki örnek yapılanmaya bırak ve kaynak destekliyorsa genel boyutlandırılabilirlik dili kullan.
45. Referans unsur adı gereksiz biçimde özel örneğe kilitlenmiş mi? Kaynak destekliyorsa “O-ring” gibi özel gerçekleştirmeyi “sızdırmazlık elemanı” gibi genel teknik unsur altında açıklamaya taşı.
46. Tarifnamede her ayrıntıya zorla referans verilmiş mi? Yapıştırıcı/malzeme/kaplama gibi numarasız kalabilecek teknik özellikleri gereksiz referans unsuruna dönüştürme.
47. Kullanıcıya görünen metinde “müşteri tarafından iletilen teknik çizimde”, “müşteri bilgilerine göre”, “ek teknik belgede” gibi kaynak-atıf dili kalmış mı? Kalmışsa doğrudan teknik buluş anlatımına dönüştür.
48. Türkçe özet tek paragraf ve tek cümle mi? Buluş adı özet bölümünde ayrı başlık olarak kalacak, abstract alanına ikinci paragraf/cümle eklenmeyecek.
49. İstemlerde standart “olup, özelliği;” dışında noktalı virgül var mı? Varsa virgül veya noktayla düzelt.
50. Son kalite kapısında şu soruların tamamı EVET mi: kaynak tamlığı, açık referans adları, unsur tanımlama sırası, uzman-nasıl testi, farklılaştırıcı çekirdek, gereksiz kapsam daraltma yokluğu, bağımlı istem tekrarının olmaması, geçerli bağımlılık, örnek ölçülerin istemden uzak tutulması, ürün istem dili, referans senkronizasyonu ve tek paragraf/tek cümle özet?
51. YAPILANDIRILMIŞ ENVANTER içindeki HER `technical_facts` maddesinin `source_coverage_map` içinde tek tek karşılığı var mı? Teknik fact için mandatory=false kabul edilmez; Her kaydın `covered=true`, en az bir bölüm adı ve tarifname taslağında birebir geçen en az 20 karakterlik gerçek bir kanıt alıntısı var mı? Özellikle teknik avantajlar, ayırt edici yönler ve bağımsızlık/performans sonuçları “benzer anlam var” denilerek atlanmış mı?
52. Yazılım/modül ağırlıklı buluşta yalnız “işlemci/donanım” kelimesi geçmesiyle yetinilmiş mi, yoksa modül/yazılımın kaynakta dayanaklı teknik taşıyıcı üzerinde çalıştığı/koşturulduğu açık ilişkiyle yazılmış mı? Kaynak özel taşıyıcı veriyorsa genel elektronik cihaz ifadesi özel taşıyıcıyı silmiş mi?
53. BULUŞUN DETAYLI AÇIKLAMASI giriş cümlesinde buluş adı cümle içi normal yazımla mı kullanılmış? Başlıktaki Title Case düzeni cümle içine kopyalanmamış mı; SIM/eSIM gibi kısaltmalar korunmuş mu?
54. REFERANS NUMARALARI bölümündeki yöntem adımları `1001. ...` biçiminde önden yöntem numarasıyla mı yazılmış ve bu satırlarda sistem/cihaz `(1)`, `(2)` türü parantezli referans işaretleri kaldırılmış mı? Parantezli unsur referansları yalnız BULUŞUN DETAYLI AÇIKLAMASI bölümünden itibaren mi başlıyor? Sistem ve yöntem alt istemlerinin tamamı semantik tekrar kontrolünden geçti mi?
55. Aynı teknik taşıyıcı üzerinde aynı çalışma ilişkisine sahip ardışık yazılım modüllerinde gereksiz taşıyıcı tekrarı var mı? Uygunsa ortak numarasız üst bullet + ayrı gerçek alt bullet yazım stili kullanılmış mı; üst bullet hiçbir referans taşımıyor ve alt maddeler ilk-tanım sırasına uyuyor mu?
56. Referans listesinde ayrı olan iki unsur şekil üzerinde tek `2-3`/tek kutu/tek hedefte birleştirilmiş mi? Ayrı unsur kimlikleri ayrı kutucuk/çağrı/oklarla ayırt edilebilir mi?
57. Şekiller çıktı setinde REFERANS NUMARALARI bölümündeki bütün gerçek sistem unsur referansları en az bir kez gösteriliyor mu? Sistem+yöntem şekilleri varsa yöntem adımı referansları da akış şekillerinde tam mı?
58. Ortak `elektronik işlem birimi üzerinde koşturulan yazılım` üst maddesinin altında yalnız gerçekten yürütülebilir yazılım/modül/kontrolör/arayüz/yığın mı var? Veritabanı/bellek/veri deposu gibi pasif veri taşıyan unsur kaynakça açıkça yürütülebilir değilse ortak gruptan çıkarılmış mı?
59. BULUŞUN DETAYLI AÇIKLAMASI ve İSTEMLER içinde REFERANS NUMARALARI listesindeki bir unsur adı her geçtiğinde aynı/çekimli unsur adıyla doğru `(N)` referansını taşıyor mu? Özellikle bağımsız ve bağımlı yöntem istemlerindeki gNodeB, arayüz, kontrolör, veritabanı, yığın ve cihaz kullanımları numaralı mı?
60. Nihai Word üretildikten sonra ham kaynak zinciri ve beş son kapının tamamı tekrar geçmelidir: ham pasaj -> technical_fact -> nihai Word kanıtı; 1/5 BBF/KAYNAK TAMLIK, 2/5 ANA+ALT İSTEM kalite/tekrar/gereklilik, 3/5 DETAYLI AÇIKLAMA+İSTEMLER REFERANS NUMARASI tamlığı, 4/5 TAM ŞABLON, 5/5 UNSUR/YÖNTEM DİLİ. Sistem alt istemlerinde `bulunmasıdır` kullanılmışsa mutlaka unsur merkezli `olmasıdır/içermesidir` diline dönüştür.

JSON dışında hiçbir şey yazma. Çıktı, aşağıdaki şemaya tam uymalıdır:
{TARIFNAME_DRAFT_SCHEMA}

YAPILANDIRILMIŞ ENVANTER:
{json.dumps(extracted, ensure_ascii=False, indent=2)}

DOĞRULANMIŞ PATENT LİTERATÜRÜ:
{json.dumps(literature, ensure_ascii=False, indent=2)}

HAM BBF:
---
{source_text}
---

EK TEKNİK BELGELER:
---
{technical_supplement_text}
---

KONTROL EDİLECEK TASLAK:
{json.dumps(draft, ensure_ascii=False, indent=2)}
"""


def _tarifname_visible_draft_for_audit(draft: dict[str, Any]) -> dict[str, Any]:
    """Kaynak kapsam meta alanlarını dışarıda bırakarak kullanıcıya gidecek teknik taslak içeriğini döndürür."""
    return {
        key: value
        for key, value in draft.items()
        if key not in {"source_coverage_map", "coverage_audit"}
    }


def tarifname_final_raw_source_audit_prompt(
    source_passage_registry: list[dict[str, str]],
    extracted: dict[str, Any],
    draft: dict[str, Any],
    language: str = "Türkçe",
) -> str:
    visible_draft = _tarifname_visible_draft_for_audit(draft)
    return f"""{TARIFNAME_RULES}
Bu işlem tarifname TASLAĞI OLUŞTURULDUKTAN SONRA çalışan bağımsız SON HAM KAYNAK kontrolüdür. Tarifnameyi yeniden yazma.
Çıktı dili: {language}.

AMAÇ:
1. Ham kaynak pasaj auditinde `technical` olarak sınıflandırılmış HER passage_id'yi nihai tarifname taslağıyla yeniden karşılaştır.
2. `technical_facts` içindeki HER fact'i nihai tarifname taslağıyla bağımsız olarak yeniden karşılaştır.
3. Örnek senaryo, çalışma koşulu, avantaj, teknik etki, alternatif, performans sonucu, karşılaştırma, süreç ayrıntısı ve görselden çıkarılmış teknik bilgi dahil hiçbir teknik anlamı “benzer bir cümle var” diyerek geçme.
4. `source_coverage_map` ve `coverage_audit` bu kontrolde KANIT değildir; aşağıda bilerek verilmemiştir. Kanıt yalnız kullanıcıya gidecek tarifname taslağının gerçek metninden alınmalıdır.
5. Her evidence öğesi taslakta BİREBİR geçen en az 20 karakterlik bir alıntı olmalıdır.
6. En küçük teknik anlam kaybında `covered=false` yap ve `missing_detail` alanında neyin eksik olduğunu açıkça yaz.
7. passage_checks içinde yalnız source_passage_audit tarafından `technical` sınıflandırılmış pasajlar TAM BİR KEZ yer almalıdır. fact_checks içinde bütün technical_facts kimlikleri TAM BİR KEZ yer almalıdır.
8. Bütün satırlar covered=true değilse `all_pass=false` olmalıdır.

JSON dışında hiçbir şey yazma.
ŞEMA:
{{
  "passage_checks":[{{"passage_id":"B0001","covered":true,"evidence":["nihai taslaktan birebir alıntı"],"missing_detail":""}}],
  "fact_checks":[{{"fact_id":"T001","covered":true,"evidence":["nihai taslaktan birebir alıntı"],"missing_detail":""}}],
  "all_pass":true
}}

SOURCE_PASSAGE_REGISTRY:
{json.dumps(source_passage_registry, ensure_ascii=False, indent=2)}

KAYNAK PASAJ SINIFLANDIRMASI VE TECHNICAL FACT ENVANTERİ:
{json.dumps({"source_passage_audit": extracted.get("source_passage_audit") or [], "technical_facts": extracted.get("technical_facts") or []}, ensure_ascii=False, indent=2)}

KULLANICIYA GİDECEK NİHAİ TASLAK (coverage meta alanları hariç):
{json.dumps(visible_draft, ensure_ascii=False, indent=2)}
"""


def _visible_draft_text_for_audit(draft: dict[str, Any]) -> str:
    parts: list[str] = []

    def collect(value: Any) -> None:
        if isinstance(value, str):
            if value.strip():
                parts.append(value)
        elif isinstance(value, list):
            for item in value:
                collect(item)
        elif isinstance(value, dict):
            for item in value.values():
                collect(item)

    collect(_tarifname_visible_draft_for_audit(draft))
    return "\n".join(parts)


def _strip_claim_number(text: str) -> str:
    return re.sub(r"^\s*\d+\s*[.)-]\s*", "", str(text or "")).strip()


def add_numbered_claim(doc: Document, template: Document, text: str):
    """İstem numarasını şablondaki gerçek Word otomatik numaralandırmasıyla oluştur."""
    p = doc.add_paragraph()
    _copy_list_properties(p, template.paragraphs[85])
    _append_text_with_equations(p, _strip_claim_number(text))
    return format_paragraph(p)


def _replace_in_nested(value: Any, old: str, new: str) -> Any:
    if isinstance(value, str):
        return value.replace(old, new)
    if isinstance(value, list):
        return [_replace_in_nested(x, old, new) for x in value]
    if isinstance(value, dict):
        return {k: _replace_in_nested(v, old, new) for k, v in value.items()}
    return value


def _ensure_title_for_claim_mode(title: str, claim_mode: str, language: str = "Türkçe") -> str:
    text = str(title or "").strip()
    if claim_mode != "Sistem ve yöntem" or not text:
        return text
    if _english_spec(language):
        if "method" in text.casefold():
            return text
        if re.search(r"\bsystem\b", text, flags=re.IGNORECASE):
            return re.sub(r"\bsystem\b(?!.*\bsystem\b)", "System and Method", text, count=1, flags=re.IGNORECASE)
        return text + " System and Method"
    if "yöntem" in text.lower():
        return text
    matches = list(re.finditer(r"\bsistemi\b", text, flags=re.IGNORECASE))
    if matches:
        last = matches[-1]
        return text[: last.start()] + "Sistemi ve Yöntemi" + text[last.end() :]
    return text + " Sistemi ve Yöntemi"


_CONTINUATION_STARTERS = (
    "özellikle",
    "bununla birlikte",
    "bu nedenle",
    "ayrıca",
    "böylece",
    "bu kapsamda",
    "dolayısıyla",
)


def _merge_continuation_paragraphs(paragraphs: list[str]) -> list[str]:
    out: list[str] = []
    for raw in paragraphs or []:
        text = re.sub(r"\s+", " ", str(raw or "")).strip()
        if not text:
            continue
        if out and text.casefold().startswith(_CONTINUATION_STARTERS):
            out[-1] = out[-1].rstrip() + " " + text
        else:
            out.append(text)
    return out


def _merge_initial_element_paragraphs(draft: dict[str, Any]) -> None:
    """Detaylı açıklamadaki ilk referanslı unsur anlatımlarını tek sürekli paragrafta birleştir."""
    paragraphs = [str(x or "").strip() for x in (draft.get("detailed_paragraphs") or []) if str(x or "").strip()]
    if len(paragraphs) < 2:
        draft["detailed_paragraphs"] = paragraphs
        return
    element_numbers = {str(x.get("number", "")).strip() for x in (draft.get("elements") or []) if str(x.get("number", "")).strip()}
    if not element_numbers:
        draft["detailed_paragraphs"] = paragraphs
        return
    merged: list[str] = []
    prefix_parts: list[str] = []
    for para in paragraphs:
        lower = para.casefold()
        is_separate = lower.startswith((
            "buluşun bir yapılanmasında",
            "yöntemin gerçekleştirdiği",
            "sistemin çalışma",
            "buluşun çalışma",
            "yöntemin çalışma",
            "in one embodiment",
            "in an embodiment",
            "the method performs",
            "the operation of the system",
        ))
        element_ref = any(f"({n})" in para for n in element_numbers)
        if not merged and not is_separate and element_ref:
            prefix_parts.append(para)
            continue
        if prefix_parts:
            merged.append(" ".join(prefix_parts))
            prefix_parts = []
        merged.append(para)
    if prefix_parts:
        merged.append(" ".join(prefix_parts))
    draft["detailed_paragraphs"] = merged


def _ensure_literature_titles(draft: dict[str, Any], literature: list[dict[str, Any]] | None, language: str = "Türkçe") -> None:
    paragraphs = list(draft.get("literature_paragraphs") or [])
    for idx, doc_info in enumerate(literature or []):
        if idx >= len(paragraphs):
            break
        p = str(paragraphs[idx] or "").strip()
        en = str(doc_info.get("title_en", "") or "").strip()
        tr = str(doc_info.get("title_tr", "") or "").strip()
        if _english_spec(language):
            if not en or en in p:
                continue
            number = str(doc_info.get("application_number", "") or "").strip()
            paragraphs[idx] = f"Patent literature includes {number + ' ' if number else ''}entitled ‘{en}’. " + p
            continue
        if not en or not tr or (en in p and tr in p):
            continue
        if en in p:
            replacement = f"İngilizce başlığı “{en}” ve Türkçe karşılığı “{tr}” olan"
            patterns = [
                rf"ve\s+[“\"]{re.escape(en)}[”\"]\s+başlıklı",
                rf"[“\"]{re.escape(en)}[”\"]\s+başlıklı",
            ]
            changed = False
            for pattern in patterns:
                new_p, count = re.subn(pattern, replacement, p, count=1, flags=re.IGNORECASE)
                if count:
                    p = new_p
                    changed = True
                    break
            if not changed:
                quoted_en = f"“{en}”"
                if quoted_en in p:
                    p = p.replace(quoted_en, replacement, 1)
                else:
                    p = p.replace(en, replacement, 1)
        else:
            number = str(doc_info.get("application_number", "") or "").strip()
            lead = f"Literatürde yapılan araştırmalar sonucu {number + ' ' if number else ''}numaralı, İngilizce başlığı “{en}” ve Türkçe karşılığı “{tr}” olan patent dokümanına rastlanmıştır. "
            p = lead + p
        paragraphs[idx] = p
    draft["literature_paragraphs"] = paragraphs



def _normalize_technical_field_two_paragraphs(value: str, language: str = "Türkçe") -> str:
    """TEKNİK ALAN/TECHNICAL FIELD girişini iki zorunlu paragrafa ayır."""
    raw = str(value or "").strip()
    parts = [x.strip() for x in re.split(r"\n\s*\n", raw) if x.strip()]
    combined = " ".join(parts)
    if _english_spec(language):
        combined = re.sub(r"^The invention\s+(?!relates)", "The invention relates to ", combined, count=1, flags=re.IGNORECASE)
        match = re.match(r"^(The invention relates to .+?\.)\s*(.*)$", combined, flags=re.IGNORECASE | re.DOTALL)
        if not match:
            return raw
        first, rest = match.group(1).strip(), match.group(2).strip()
        if not rest:
            return first
        rest = re.sub(r"^(?:In particular,\s*)?the invention relates to\s*", "", rest, count=1, flags=re.IGNORECASE).strip()
        second = "In particular, the invention relates to " + (rest[:1].lower() + rest[1:] if rest else "")
        return first + "\n\n" + second
    raw = re.sub(r"^Buluş\s+(?!,)", "Buluş, ", raw, count=1, flags=re.IGNORECASE)
    parts = [x.strip() for x in re.split(r"\n\s*\n", raw) if x.strip()]
    combined = " ".join(parts)
    match = re.match(r"^(Buluş,\s+.+?ile ilgilidir\.)\s*(.*)$", combined, flags=re.IGNORECASE | re.DOTALL)
    if not match:
        return raw
    first = match.group(1).strip()
    rest = match.group(2).strip()
    if not rest:
        return first
    rest = re.sub(r"^Buluş,\s*özellikle\s*", "", rest, count=1, flags=re.IGNORECASE).strip()
    rest = re.sub(r"^Buluş\s+özellikle\s*", "", rest, count=1, flags=re.IGNORECASE).strip()
    second = "Buluş, özellikle " + (rest[:1].lower() + rest[1:] if rest else "")
    return first + "\n\n" + second


def _normalize_method_step_numbers(draft: dict[str, Any]) -> None:
    """Müşteri yöntem referanslarını koru; yalnız eksik referansları 1001... varsayılanıyla tamamla."""
    steps = draft.get("method_steps") or []
    if not steps:
        return
    used = {str(step.get("number", "") or "").strip() for step in steps if str(step.get("number", "") or "").strip()}
    next_default = 1001
    for step in steps:
        number = str(step.get("number", "") or "").strip()
        if number:
            step["number"] = number
            continue
        while str(next_default) in used:
            next_default += 1
        step["number"] = str(next_default)
        used.add(str(next_default))
        next_default += 1

    method_claim = draft.get("method_claim") or {}
    if method_claim.get("steps") is not None:
        claim_steps = list(method_claim.get("steps") or [])
        normalized: list[str] = []
        for i, step in enumerate(steps):
            text = str(step.get("text", "") or "").strip().rstrip(".,;:")
            number = str(step.get("number", "") or "").strip()
            # method_steps metni numarasızdır; istemde aynı metin + kaynak referansı kullanılır.
            normalized.append(f"{text} ({number})" if number else text)
        method_claim["steps"] = normalized
        draft["method_claim"] = method_claim


def _assign_missing_element_numbers(draft: dict[str, Any]) -> None:
    """Müşteri unsur referanslarını koru; yalnız eksik olanlara 1..N varsayılanı ata."""
    elements = draft.get("elements") or []
    if not elements:
        return
    used = {str(e.get("number", "") or "").strip() for e in elements if str(e.get("number", "") or "").strip()}
    next_default = 1
    for element in elements:
        number = str(element.get("number", "") or "").strip()
        if number:
            element["number"] = number
            continue
        while str(next_default) in used:
            next_default += 1
        element["number"] = str(next_default)
        used.add(str(next_default))
        next_default += 1


def _convert_mapping_tables_to_prose(draft: dict[str, Any], language: str = "Türkçe") -> None:
    """İşlem Adımı/Gerçekleştiren Unsur açıklama tablolarını teknik paragrafa çevir."""
    kept: list[dict[str, Any]] = []
    mapping_paragraphs: list[str] = []
    elements = draft.get("elements") or []
    steps = draft.get("method_steps") or []
    for table in draft.get("tables") or []:
        headers = [str(x or "").strip().casefold() for x in (table.get("headers") or [])]
        is_mapping = (any("işlem adımı" in h for h in headers) and any("gerçekleştiren unsur" in h for h in headers)) or (any("process step" in h for h in headers) and any("performing element" in h or "performing unit" in h for h in headers))
        if not is_mapping:
            kept.append(table)
            continue
        rows = table.get("rows") or []
        sentences: list[str] = []
        for idx, row in enumerate(rows):
            row = list(row or [])
            element = elements[idx] if idx < len(elements) else {}
            step = steps[idx] if idx < len(steps) else {}
            el_name = str(element.get("name", "") or (row[2] if len(row) > 2 else "ilgili modül")).strip()
            el_num = str(element.get("number", "") or "").strip()
            step_text = str(step.get("text", "") or (row[1] if len(row) > 1 else "ilgili işlem")).strip().rstrip(".,;:")
            step_num = str(step.get("number", "") or "").strip()
            explanation = str(row[3] if len(row) > 3 else "").strip().rstrip(".")
            if _english_spec(language):
                lead = f"The {el_name}{f' ({el_num})' if el_num else ''} performs the process step of {step_text[:1].lower() + step_text[1:]}{f' ({step_num})' if step_num else ''}"
                if explanation:
                    lead += f" and, in this context, {explanation[:1].lower() + explanation[1:]}"
            else:
                lead = f"{el_name}{f' ({el_num})' if el_num else ''}, {step_text[:1].lower() + step_text[1:]} işlem adımını{f' ({step_num})' if step_num else ''} gerçekleştirmektedir"
                if explanation:
                    lead += f" ve bu kapsamda {explanation[:1].lower() + explanation[1:]}"
            sentences.append(lead.rstrip(".") + ".")
        if sentences:
            prefix = "The technical relationship between the system and the method is established as follows. " if _english_spec(language) else "Sistem ile yöntem arasındaki teknik ilişki aşağıdaki şekilde kurulmaktadır. "
            mapping_paragraphs.append(prefix + " ".join(sentences))
    draft["tables"] = kept
    if mapping_paragraphs:
        detailed = [str(x or "").strip() for x in (draft.get("detailed_paragraphs") or []) if str(x or "").strip()]
        insert_at = 1 if detailed else 0
        for paragraph in reversed(mapping_paragraphs):
            detailed.insert(insert_at, paragraph)
        draft["detailed_paragraphs"] = detailed

def apply_tarifname_house_style(
    draft: dict[str, Any],
    claim_mode: str,
    literature: list[dict[str, Any]] | None = None,
    language: str = "Türkçe",
) -> dict[str, Any]:
    """Kullanıcının sabit terminoloji, paragraf ve başlık tercihlerini çıktıdan önce uygula."""
    draft = deepcopy(draft)
    if not _english_spec(language):
        for old, new in [
            ("Buluşun bir gerçekleştirilmesinde", "Buluşun bir yapılanmasında"),
            ("buluşun bir gerçekleştirilmesinde", "buluşun bir yapılanmasında"),
            ("Mevcut buluş", "Buluş"),
            ("mevcut buluş", "buluş"),
            ("Buluş özellikle", "Buluş, özellikle"),
        ]:
            draft = _replace_in_nested(draft, old, new)
    draft["title"] = _ensure_title_for_claim_mode(draft.get("title", ""), claim_mode, language)
    _assign_missing_element_numbers(draft)
    _normalize_method_step_numbers(draft)
    _convert_mapping_tables_to_prose(draft, language)
    draft["prior_art_general_paragraphs"] = _merge_continuation_paragraphs(draft.get("prior_art_general_paragraphs") or [])
    _merge_initial_element_paragraphs(draft)
    _ensure_literature_titles(draft, literature, language)
    draft["figure_descriptions"] = [
        re.sub(r"\b\d{4}\s*[-–]\s*\d{4}\s+numaralı\s+", "", str(x or ""), flags=re.IGNORECASE)
        for x in (draft.get("figure_descriptions") or [])
    ]
    draft["technical_field"] = _normalize_technical_field_two_paragraphs(draft.get("technical_field", ""), language)
    for step in draft.get("method_steps") or []:
        step["text"] = re.sub(r"[.,;:]+$", "", str(step.get("text", "") or "").strip())
    return draft



def _claim_refs(text: str) -> list[str]:
    """Parantezli referans işaretlerini görünme sırasıyla döndür."""
    return re.findall(r"\(\s*([A-Za-zÇĞİÖŞÜçğıöşü0-9_\-]+)\s*\)", str(text or ""))


def _system_claim_entries(system_claim: dict[str, Any] | None) -> list[Any]:
    """Ana sistem isteminin düz maddelerini ve ortak-taşıyıcı gruplarını sırasıyla döndür."""
    return list((system_claim or {}).get("elements") or [])


def _system_claim_entry_texts(entry: Any, include_group_lead: bool = True) -> list[str]:
    if isinstance(entry, dict):
        out: list[str] = []
        lead = str(entry.get("lead", "") or "").strip()
        if include_group_lead and lead:
            out.append(lead)
        out.extend(str(x or "").strip() for x in (entry.get("subelements") or []) if str(x or "").strip())
        return out
    text = str(entry or "").strip()
    return [text] if text else []


def _system_claim_all_texts(system_claim: dict[str, Any] | None, include_group_leads: bool = True) -> list[str]:
    out: list[str] = []
    for entry in _system_claim_entries(system_claim):
        out.extend(_system_claim_entry_texts(entry, include_group_leads))
    return out


def _reference_name_pattern(name: str) -> re.Pattern:
    """Referans adının (N)'den hemen önceki aynı/çekimli biçimini yaklaşık fakat sıkı eşleştir."""
    tokens = re.findall(r"[A-Za-zÇĞİÖŞÜçğıöşü0-9]+", str(name or ""))
    parts: list[str] = []
    for token in tokens:
        stem = token if len(token) <= 4 else token[:max(4, len(token) - 2)]
        parts.append(re.escape(stem) + r"\w*")
    return re.compile(r"\s+".join(parts) + r"\s*$", re.I)


def _validate_reference_identity(draft: dict[str, Any]) -> None:
    """(N) yalnız REFERANS NUMARALARI listesindeki aynı unsur adı/çekimiyle kullanılabilir."""
    element_map = {
        str(x.get("number", "") or "").strip(): str(x.get("name", "") or "").strip()
        for x in (draft.get("elements") or [])
        if str(x.get("number", "") or "").strip()
    }
    texts = [
        *map(str, draft.get("detailed_paragraphs") or []),
        str(draft.get("working_principle", "") or ""),
        *_system_claim_all_texts(draft.get("system_claim") or {}),
        *map(str, draft.get("dependent_system_claims") or []),
        *map(str, (draft.get("method_claim") or {}).get("steps") or []),
        *map(str, draft.get("dependent_method_claims") or []),
    ]
    for text in texts:
        for m in re.finditer(r"\(([^()]+)\)", str(text)):
            n = m.group(1).strip()
            if n not in element_map:
                continue
            before = str(text)[max(0, m.start() - 160):m.start()].rstrip()
            if not _reference_name_pattern(element_map[n]).search(before):
                raise ValueError(
                    f"Referans ({n}) '{element_map[n]}' unsurunun aynı/çekimli adıyla kullanılmalıdır; "
                    "kısaltma veya eş anlamlı ad referans numarasını taşıyamaz."
                )



def _reference_mention_pattern(name: str) -> re.Pattern:
    """Referans listesi unsur adını, son sözcükte Türkçe çekime izin vererek yakalar."""
    tokens = re.findall(r"[A-Za-zÇĞİÖŞÜçğıöşü0-9]+", str(name or ""))
    if not tokens:
        return re.compile(r"a^")
    parts = [re.escape(x) for x in tokens[:-1]]
    last = tokens[-1]
    stem = last if len(last) <= 4 else last[:max(4, len(last) - 2)]
    parts.append(re.escape(stem) + r"\w*")
    return re.compile(r"\b" + r"\s+".join(parts) + r"\b", re.I)


def _validate_reference_presence(draft: dict[str, Any]) -> None:
    """Detaylı açıklama ve istemlerde referans-listesi unsur adının numarasız kullanımını engeller."""
    elements = [x for x in (draft.get("elements") or []) if str(x.get("number", "") or "").strip() and str(x.get("name", "") or "").strip()]
    texts: list[tuple[str, str]] = []
    for i, text in enumerate(draft.get("detailed_paragraphs") or [], start=1):
        texts.append((f"BULUŞUN DETAYLI AÇIKLAMASI paragrafı {i}", str(text or "")))
    for i, step in enumerate(draft.get("method_steps") or [], start=1):
        texts.append((f"Yöntem işlem adımı {i}", str(step.get("text", "") or "")))
    if draft.get("working_principle"):
        texts.append(("Çalışma prensibi", str(draft.get("working_principle") or "")))
    sc = draft.get("system_claim") or {}
    texts.extend(("Ana sistem istemi", str(t or "")) for t in _system_claim_all_texts(sc))
    texts.extend((f"Bağımlı sistem istemi {i}", str(t or "")) for i, t in enumerate(draft.get("dependent_system_claims") or [], start=1))
    mc = draft.get("method_claim") or {}
    texts.extend(("Ana yöntem istemi", str(t or "")) for t in (mc.get("steps") or []))
    texts.extend((f"Bağımlı yöntem istemi {i}", str(t or "")) for i, t in enumerate(draft.get("dependent_method_claims") or [], start=1))

    for element in elements:
        number = str(element.get("number", "") or "").strip()
        name = str(element.get("name", "") or "").strip()
        mention_re = _reference_mention_pattern(name)
        ref_re = re.compile(r"^\s*(?:\([^)]{1,40}\)\s*)?\(\s*" + re.escape(number) + r"\s*\)")
        for label, text in texts:
            for m in mention_re.finditer(text):
                if not ref_re.match(text[m.end():m.end() + 70]):
                    raise ValueError(
                        f"{label}: '{name}' unsurunun kullanımı ({number}) referansını taşımıyor. "
                        "BULUŞUN DETAYLI AÇIKLAMASI ve İSTEMLER bölümünde referans-listesi unsurları her kullanımda doğru parantezli numarayla yazılmalıdır."
                    )


def _validate_common_carrier_scope(draft: dict[str, Any]) -> None:
    """Ortak yazılım taşıyıcısına pasif veri depolarının yanlışlıkla alınmasını engeller."""
    element_map = {str(x.get("number", "") or "").strip(): x for x in (draft.get("elements") or [])}
    passive_re = re.compile(r"veritaban|bellek|hafıza|veri depos|data store|profil tablos|kayıt tablos|veri yapıs", re.I)
    executable_re = re.compile(r"yazılım|modül|kontrolör|arayüz|yığın|stack|algoritma|koştur|çalıştır|yürüt", re.I)
    for entry in _system_claim_entries(draft.get("system_claim") or {}):
        if not isinstance(entry, dict):
            continue
        lead = str(entry.get("lead", "") or "")
        if not re.search(r"koşturulan|çalışan|yürütülen|executed|running", lead, re.I):
            continue
        for sub in entry.get("subelements") or []:
            text = str(sub or "")
            for ref in _claim_refs(text):
                info = element_map.get(str(ref).strip()) or {}
                name = str(info.get("name", "") or "")
                desc = str(info.get("description", "") or "")
                if passive_re.search(name) and not executable_re.search(desc):
                    raise ValueError(
                        f"Ortak elektronik işlem birimi/yazılım grubunda '{name} ({ref})' kullanılmış. "
                        "Veritabanı/bellek/veri deposu gibi pasif veri unsurları kaynakta açıkça yürütülebilir yazılım/modül olarak tanımlanmıyorsa ortak taşıyıcı grubunun dışında ayrı unsur olarak yazılmalıdır."
                    )

def _strip_known_element_reference_marks(text: str, element_numbers: list[str]) -> str:
    """REFERANS NUMARALARI yöntem satırında yalnız bilinen sistem/cihaz `(REF)` işaretlerini kaldır."""
    result = str(text or "")
    for number in sorted({str(x or "").strip() for x in element_numbers if str(x or "").strip()}, key=len, reverse=True):
        result = re.sub(rf"\s*\(\s*{re.escape(number)}\s*\)", "", result)
    return re.sub(r"\s{2,}", " ", result).strip()


def _normalize_claim_semantics(text: str) -> set[str]:
    """Bağımlı istem tekrarını kaba fakat deterministik biçimde yakalamak için içerik sözcükleri."""
    txt = str(text or "").casefold()
    txt = re.sub(r"^\s*istem\s+\d+(?:\s*(?:veya|ve|,)\s*\d+)*['’]?e\s+uygun\s+[^;]+;", " ", txt)
    txt = re.sub(r"\(\s*[a-z0-9_\-]+\s*\)", " ", txt)
    txt = re.sub(r"[^a-zçğıöşü0-9]+", " ", txt)
    stop = {"istem", "uygun", "olup", "özelliği", "bir", "ve", "veya", "ile", "olan", "olarak", "söz", "konusu", "şekilde", "şeklinde"}
    return {w for w in txt.split() if len(w) > 2 and w not in stop}


def _validate_system_claim_reference_order(system_claim: dict[str, Any], element_numbers: list[str]) -> None:
    """Düz veya ortak-taşıyıcı gruplu ana istemde ilk-tanım sırasını zorlar."""
    if not system_claim:
        return
    valid = set(element_numbers)
    seen = set(_claim_refs(system_claim.get("preamble", ""))) & valid

    def check_leaf(text: str, label: str) -> None:
        refs = [r for r in _claim_refs(str(text)) if r in valid]
        new_refs: list[str] = []
        for r in refs:
            if r not in seen and r not in new_refs:
                new_refs.append(r)
        if len(new_refs) > 1:
            raise ValueError(
                f"Ana istemin {label} birden fazla yeni referansı ({', '.join(new_refs)}) ilk kez birlikte tanımlıyor. "
                "Her düz/alt madde kural olarak tek yeni referanslı unsur tanımlamalıdır."
            )
        seen.update(new_refs)

    for idx, entry in enumerate(_system_claim_entries(system_claim), start=1):
        if isinstance(entry, dict):
            lead = str(entry.get("lead", "") or "").strip()
            subelements = [str(x or "").strip() for x in (entry.get("subelements") or []) if str(x or "").strip()]
            lead_refs = [r for r in _claim_refs(lead) if r in valid]
            if lead_refs:
                raise ValueError(
                    f"Ana istemin {idx}. ortak taşıyıcı üst maddesi referans numarası taşıyamaz ({', '.join(lead_refs)}). "
                    "Referanslı modülleri alt maddelerde ayrı ayrı tanımlayın."
                )
            if len(subelements) < 2:
                raise ValueError("Ortak taşıyıcı istem grubu en az iki ayrı referanslı alt madde içermelidir; aksi halde düz madde kullanın.")
            for sub_idx, sub in enumerate(subelements, start=1):
                check_leaf(sub, f"{idx}.{sub_idx}. alt maddesi")
        else:
            check_leaf(str(entry), f"{idx}. unsur maddesi")


def _validate_dependent_claim_semantic_repetition(system_claim: dict[str, Any], dependents: list[str]) -> None:
    """Birebir/çok yakın teknik tekrarları yerel kalite kapısında yakalar; nihai semantik kontrol AI kalite turunda da yapılır."""
    base = _normalize_claim_semantics(" ".join([str(system_claim.get("preamble", "")), *_system_claim_all_texts(system_claim)]))
    previous_sets: list[set[str]] = [base]
    for idx, claim in enumerate(dependents, start=2):
        words = _normalize_claim_semantics(claim)
        if not words:
            continue
        for prior in previous_sets:
            if len(words) >= 4:
                overlap = len(words & prior) / max(1, len(words))
                if overlap >= 0.92:
                    raise ValueError(
                        f"İstem {idx} üst istemde zaten bulunan teknik özelliği anlam olarak tekrar ediyor. "
                        "Bağımlı istem gerçek bir ek teknik sınırlama getirmelidir."
                    )
        previous_sets.append(words)


def _validate_dependent_method_claim_semantic_repetition(method_claim: dict[str, Any], dependents: list[str]) -> None:
    """Yöntem alt istemlerinde de ana/üst isteme karşı semantik tekrar kalite kapısı uygular."""
    if not method_claim:
        return
    base = _normalize_claim_semantics(" ".join([str(method_claim.get("preamble", "")), *map(str, method_claim.get("steps") or [])]))
    previous_sets: list[set[str]] = [base]
    for idx, claim in enumerate(dependents, start=1):
        words = _normalize_claim_semantics(claim)
        if not words:
            continue
        for prior in previous_sets:
            if len(words) >= 4:
                overlap = len(words & prior) / max(1, len(words))
                if overlap >= 0.92:
                    raise ValueError(
                        f"Bağımlı yöntem istemi {idx} ana/üst istemde zaten bulunan teknik özelliği anlam olarak tekrar ediyor. "
                        "Yöntem alt istemi gerçek bir ek teknik sınırlama getirmelidir."
                    )
        previous_sets.append(words)


def _validate_abstract_shape(abstract: str, language: str) -> None:
    if _english_spec(language):
        return
    text = str(abstract or "").strip()
    if not text:
        raise ValueError("ÖZET metni boş olamaz.")
    if re.search(r"\n\s*\n", text):
        raise ValueError("ÖZET tek paragraf olmalıdır.")
    endings = re.findall(r"[.!?](?=\s|$)", text)
    if len(endings) > 1:
        raise ValueError("ÖZET tek paragraf ve tek cümle olmalıdır; ayrı cümlelere bölünmemelidir.")




def _validate_method_step_action_language(draft: dict[str, Any], language: str = "Türkçe") -> None:
    if _english_spec(language):
        return
    action_end_re = re.compile(r"(?:ması|mesi)\s*$", re.IGNORECASE)
    for step in draft.get("method_steps") or []:
        number = str(step.get("number", "") or "").strip()
        text = str(step.get("text", "") or "").strip()
        clean = re.sub(rf"\s*\(\s*{re.escape(number)}\s*\)\s*$", "", text).strip().rstrip(".,;:") if number else text.rstrip(".,;:")
        if clean and not action_end_re.search(clean):
            raise ValueError(f"Yöntem işlem adımı {number or '?'} salt isimle veya işlem-sonu olmayan ifadeyle bitiyor: '{clean}'. '... yapılması/edilmesi/aktarılması/belirlenmesi' gibi gerçek işlem fiilimsisi kullanın.")
    method = draft.get("method_claim") or {}
    for raw in method.get("steps") or []:
        text = str(raw or "").strip().rstrip(".,;:")
        text = re.sub(r"\s*\(\s*[^()]+\s*\)\s*$", "", text).strip()
        if text and not action_end_re.search(text):
            raise ValueError(f"Bağımsız yöntem istemindeki adım gerçek işlem fiilimsisiyle bitmiyor: '{text}'.")


def _validate_no_generic_unsur_in_claims(draft: dict[str, Any], language: str = "Türkçe") -> None:
    if _english_spec(language):
        return
    texts = [*_system_claim_all_texts(draft.get("system_claim") or {}), *map(str, draft.get("dependent_system_claims") or []), *map(str, (draft.get("method_claim") or {}).get("steps") or []), *map(str, draft.get("dependent_method_claims") or [])]
    for text in texts:
        if re.search(r"\bbir\s+unsur\b|\bunsur\s+olmasıdır|\bunsur\s+içermesidir", str(text), re.IGNORECASE):
            raise ValueError("İstemlerde teknik eleman türü yerine belirsiz 'unsur' kullanılamaz; anten/modül/birim/eleman/sunucu/veritabanı gibi gerçek teknik tür yazılmalıdır.")


def _validate_main_claim_how_test(draft: dict[str, Any], extracted: dict[str, Any] | None = None) -> None:
    """Bağımsız sistem istemindeki modül yazımlarını teknikte uzman kişinin 'nasıl?' testiyle denetler.

    Özellikle yazılım/modül ağırlıklı unsurlarda İngilizce claim kalıbını andıran
    `X modülü (N), ... yapan bir modül` yapısı reddedilir. İşlev/mekanizma önce,
    teknik unsur adı ve referansı sonra kurulmalıdır. Ayrıca kaynak mekanizma
    açıklıyorsa salt sonuç fiili yeterli sayılmaz; girdi/ilişki + işlem + çıktı
    bağlantısından en azından yeniden üretilebilir teknik çekirdek görünmelidir.
    """
    system_claim = draft.get("system_claim") or {}
    if not system_claim:
        return
    all_items = _system_claim_all_texts(system_claim)
    elements = {str(x.get("number", "") or "").strip(): x for x in (draft.get("elements") or [])}
    source_functions = {
        str(x.get("number", "") or "").strip(): str(x.get("function", "") or "")
        for x in ((extracted or {}).get("elements") or [])
    }
    software_like = re.compile(r"(?:modül|birim|kontrolör|mekanizma|arayüz|simülatör|motor|üretici|alt sistem|yazılım|algoritma|yönetici)", re.I)
    action = re.compile(r"(?:tanımlayan|hesaplayan|belirleyen|oluşturan|seçen|güncelleyen|sınıflandıran|dönüştüren|değerlendiren|üreten|izleyen|sağlayan|uygulayan|kaydeden|sunan|örnekleyen|birleştiren|karşılaştıran|aktaran|yöneten|çalıştıran|gerçekleştiren|simüle\s+eden|haritalandıran|işleyen)", re.I)
    relation = re.compile(r"(?:kullanarak|üzerinden|göre|vasıtasıyla|birleştirerek|karşılaştırarak|parametre|değer|veri|çıktı|sonuç|eşik|koordinat|oran|indeks|sinyal|hedef|görev|enerji|karıştırma|aldatma|hız|mesafe|Q\s+değer)", re.I)

    for number, element in elements.items():
        name = str(element.get("name", "") or "").strip()
        if not number or not name or not software_like.search(name):
            continue
        ref_re = re.compile(r"\(\s*" + re.escape(number) + r"\s*\)")
        candidates = [t for t in all_items if ref_re.search(t)]
        if not candidates:
            continue  # referans varlığı başka sert kapıda denetlenir
        item = candidates[0].strip()

        # Kullanıcının istem dilinde istemediği İngilizce-claim benzeri sıra.
        if re.match(re.escape(name) + r"\s*\(\s*" + re.escape(number) + r"\s*\)\s*,", item, re.I):
            raise ValueError(
                f"Ana istemde {name} ({number}) İngilizce claim benzeri `X modülü (N), ... yapan bir modül` sırasıyla yazılmış. "
                "Önce kaynak destekli teknik işlev/mekanizma, sonra unsur adı ve referansı yazılmalıdır."
            )

        name_match = re.search(re.escape(name) + r"\s*\(\s*" + re.escape(number) + r"\s*\)", item, re.I)
        prefix = item[:name_match.start()] if name_match else item[:ref_re.search(item).start()]
        if not action.search(prefix):
            raise ValueError(
                f"Ana istemde {name} ({number}) için işlevi gerçekleştiren aktif teknik ilişki görünmüyor. "
                "Teknik uzman 'nasıl gerçekleştiriliyor?' sorusunun cevabını istemden görebilmelidir."
            )
        if not relation.search(prefix):
            raise ValueError(
                f"Ana istemde {name} ({number}) yalnız sonuç/fonksiyon düzeyinde kalmış. "
                "Kaynakta dayanak bulunduğu ölçüde girdi/veri, teknik işlem veya önceki/sonraki unsur ilişkisi yazılmalıdır."
            )

        low = item.casefold()
        # Sınıflandırma kaynağı kriter açıklıyorsa yalnız 'sınıflandıran' demek yeterli değildir.
        if re.search(r"\bsınıflandır(?:an|ır|ma)\b", low, re.I) and not re.search(r"(?:karşılaştır|kriter|koşul|eşik|\bgöre\b|mesafe|yarıçap|nedenine\s+göre)", prefix, re.I):
            raise ValueError(
                f"Ana istemde {name} ({number}) sınıflandırma sonucunu söylüyor fakat sınıflandırmanın hangi teknik kriter/karşılaştırma üzerinden yapıldığı görünmüyor."
            )
        # Hesaplama kaynağı bir matematik/ilişki açıklıyorsa hesaplamanın en az temel ilişkisi görünmeli.
        source_fn = source_functions.get(number, "")
        if "hesaplayan" in low:
            has_calc_relation = re.search(r"(?:oran|değişim|toplam|fark|ağırl|birleştir|üzerinden|\bgöre\b|kullan|bölün|bölerek|sayıs(?:ı|ının|al)|başarılı\s+görev|toplam\s+görev|parametre.*değişim|çıktılardan.*oran)", prefix, re.I)
            source_has_calc_detail = re.search(r"(?:formül|bağıntı|oran|ağırl|katsayı|değişim|eşik|=|/)", source_fn, re.I)
            # Kaynak fonksiyon alanı zayıf doldurulmuş olsa bile salt metrik isimlerini saymak
            # 'nasıl hesaplanıyor?' sorusunu cevaplamaz. En az temel ilişki görünmelidir.
            if not has_calc_relation and (source_has_calc_detail or re.search(r"(?:indeks|skor|metrik)", prefix, re.I)):
                raise ValueError(
                    f"Ana istemde {name} ({number}) için 'hesaplayan' sonucu var ancak hesaplamanın temel teknik ilişkisi görünmüyor."
                )


def _validate_claim_formula_markers(draft: dict[str, Any], language: str = "Türkçe") -> None:
    """İstemlerde yazılan açık matematik bağıntılarının düz metin bırakılmasını engeller."""
    if _english_spec(language):
        return
    texts = [
        *_system_claim_all_texts(draft.get("system_claim") or {}),
        *map(str, draft.get("dependent_system_claims") or []),
        *map(str, (draft.get("method_claim") or {}).get("steps") or []),
        *map(str, draft.get("dependent_method_claims") or []),
    ]
    formula_like = re.compile(r"\b[A-Za-zÇĞİÖŞÜçğıöşü][A-Za-zÇĞİÖŞÜçğıöşü0-9_]{0,30}\s*(?:=|≤|≥|<|>)\s*[-+()0-9A-Za-zÇĞİÖŞÜçğıöşü_]", re.I)
    for text in texts:
        raw = str(text or "")
        without_marked = EQ_MARKER_RE.sub("", raw)
        if formula_like.search(without_marked):
            raise ValueError(
                "İstemde açık matematik bağıntısı düz metin olarak yazılmış. Bağıntıyı `[[EQ: ...]]` biçiminde işaretleyin; Word çıktısında gerçek OMML denklem nesnesi olarak oluşturulacaktır."
            )


def _validate_word_math_format(data: bytes, draft: dict[str, Any]) -> None:
    """Nihai Word'de kaynak/draft formüllerinin gerçek OMML denklem nesnesi olmasını doğrular."""
    expected_display = sum(1 for f in (draft.get("formulas") or []) if str(f.get("expression", "") or "").strip())
    claim_texts = [
        *_system_claim_all_texts(draft.get("system_claim") or {}),
        *map(str, draft.get("dependent_system_claims") or []),
        *map(str, (draft.get("method_claim") or {}).get("steps") or []),
        *map(str, draft.get("dependent_method_claims") or []),
    ]
    expected_inline = sum(len(EQ_MARKER_RE.findall(str(t or ""))) for t in claim_texts)
    expected = expected_display + expected_inline
    if expected <= 0:
        return
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    m_ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    actual = len(root.findall(".//m:oMath", m_ns))
    if actual < expected:
        raise ValueError(
            f"ÇIKTI FORMÜL KONTROLÜ başarısız: {expected} matematiksel ifade beklenirken Word'de yalnız {actual} gerçek denklem nesnesi bulundu."
        )
    all_text = "".join((x.text or "") for x in root.iter())
    if "[[EQ:" in all_text or "[[FORMULA:" in all_text:
        raise ValueError("ÇIKTI FORMÜL KONTROLÜ başarısız: denklem işaretleyicisi nihai Word'de düz metin olarak kalmış.")


def validate_tarifname_draft(
    draft: dict[str, Any],
    claim_mode: str,
    literature: list[dict[str, Any]] | None = None,
    language: str = "Türkçe",
    extracted: dict[str, Any] | None = None,
) -> list[str]:
    warnings: list[str] = []
    technical_field_raw = str(draft.get("technical_field", "") or "").strip()
    tf_paragraphs = [x.strip() for x in re.split(r"\n\s*\n", technical_field_raw) if x.strip()]
    if _english_spec(language):
        if len(tf_paragraphs) < 2:
            raise ValueError('TECHNICAL FIELD must contain two paragraphs: “The invention relates to ... .” followed by “In particular, the invention relates to ... .”')
        if not re.fullmatch(r"The invention relates to .+?\.", tf_paragraphs[0], flags=re.IGNORECASE | re.DOTALL):
            raise ValueError('The first TECHNICAL FIELD paragraph must consist only of “The invention relates to ... .”')
        if not re.match(r"^In particular,\s*the invention relates to\b", tf_paragraphs[1], flags=re.IGNORECASE):
            raise ValueError('The second TECHNICAL FIELD paragraph must begin with “In particular, the invention relates to ...”.')
        first_cf = tf_paragraphs[0].casefold()
        if claim_mode == "Sistem ve yöntem" and not re.search(r"system\s+and\s+method\.?$", first_cf):
            raise ValueError('For System and Method claim mode, the first TECHNICAL FIELD paragraph must identify the invention as a system and method.')
        if claim_mode == "Yalnızca sistem" and not re.search(r"system\.?$", first_cf):
            raise ValueError('For system-only claim mode, the first TECHNICAL FIELD paragraph must identify the invention as a system.')
        if claim_mode == "Yalnızca yöntem" and not re.search(r"method\.?$", first_cf):
            raise ValueError('For method-only claim mode, the first TECHNICAL FIELD paragraph must identify the invention as a method.')
    else:
        technical_field = re.sub(r"\s+", " ", technical_field_raw).strip()
        if technical_field and not re.match(r"^Buluş,\s+.+?ile ilgilidir\.", technical_field, flags=re.IGNORECASE):
            raise ValueError('TEKNİK ALAN ilk cümlesi “Buluş, ... ile ilgilidir.” yapısında olmalıdır.')
        if re.search(r"\bBuluş özellikle\b", technical_field, flags=re.IGNORECASE):
            raise ValueError('TEKNİK ALAN içinde “Buluş özellikle” yerine “Buluş, özellikle” kullanılmalıdır.')
        if len(tf_paragraphs) < 2:
            raise ValueError('TEKNİK ALAN iki paragraf olmalıdır: ilk paragraf “Buluş, ... ile ilgilidir.”, ikinci paragraf “Buluş, özellikle ...” ile başlamalıdır.')
        if not re.fullmatch(r"Buluş,\s+.+?ile ilgilidir\.", tf_paragraphs[0], flags=re.IGNORECASE | re.DOTALL):
            raise ValueError('TEKNİK ALAN ilk paragrafı yalnız “Buluş, ... ile ilgilidir.” giriş cümlesinden oluşmalıdır.')
        if not re.match(r"^Buluş,\s*özellikle\b", tf_paragraphs[1], flags=re.IGNORECASE):
            raise ValueError('TEKNİK ALAN ikinci paragrafı “Buluş, özellikle ...” ile başlamalıdır.')
        first_cf = tf_paragraphs[0].casefold()
        if claim_mode == "Sistem ve yöntem" and not re.search(r"sistemi\s+ve\s+yöntemi\s+ile\s+ilgilidir\.$", first_cf):
            raise ValueError('Sistem ve yöntem istem yapısında TEKNİK ALAN ilk paragrafı “... sistemi ve yöntemi ile ilgilidir.” şeklinde bitmelidir.')
        if claim_mode == "Yalnızca sistem" and not re.search(r"sistemi\s+ile\s+ilgilidir\.$", first_cf):
            raise ValueError('Yalnızca sistem istem yapısında TEKNİK ALAN ilk paragrafı “... sistemi ile ilgilidir.” şeklinde bitmelidir.')
        if claim_mode == "Yalnızca yöntem" and not re.search(r"yöntemi\s+ile\s+ilgilidir\.$", first_cf):
            raise ValueError('Yalnızca yöntem istem yapısında TEKNİK ALAN ilk paragrafı “... yöntemi ile ilgilidir.” şeklinde bitmelidir.')

    steps = draft.get("method_steps") or []
    numbers = [str(x.get("number", "")).strip() for x in steps]
    if len(numbers) != len(set(numbers)):
        raise ValueError("REFERANS NUMARALARI bölümünde yinelenen yöntem adımı numarası bulundu.")
    if any(not n for n in numbers):
        raise ValueError("Numarası boş yöntem işlem adımı bulundu.")
    element_numbers = [str(x.get("number", "") or "").strip() for x in (draft.get("elements") or [])]
    if any(not n for n in element_numbers):
        raise ValueError("Sistem/cihaz unsurlarından en az birinin referans numarası boş.")
    overlap = sorted(set(element_numbers) & set(numbers))
    if overlap:
        warnings.append("Müşteri kaynaklı sistem unsuru ve yöntem adımı referansları çakışıyor: " + ", ".join(overlap) + ". Referanslar otomatik değiştirilmedi; teknik belirsizlik açısından kontrol edin.")

    for table in draft.get("tables") or []:
        headers = [str(x or "").casefold() for x in (table.get("headers") or [])]
        is_mapping = (any("işlem adımı" in h for h in headers) and any("gerçekleştiren unsur" in h for h in headers)) or (any("process step" in h for h in headers) and any("performing element" in h or "performing unit" in h for h in headers))
        if is_mapping:
            raise ValueError("Sistem-yöntem ilişki tablosu tarifname gövdesinde tablo olarak bırakılamaz; doğal teknik paragrafa dönüştürülmelidir.")

    normalized_to_numbers: dict[str, list[str]] = {}
    for step in steps:
        text = re.sub(r"\s+", " ", str(step.get("text", ""))).strip().lower()
        if not text:
            raise ValueError(f"{step.get('number','?')} numaralı yöntem adımının metni boş.")
        normalized_to_numbers.setdefault(text, []).append(str(step.get("number", "")))
    duplicate_groups = [nums for nums in normalized_to_numbers.values() if len(nums) > 1]
    if duplicate_groups:
        warnings.append(
            "Aynı metinle yazılmış farklı yöntem adımları bulundu: "
            + "; ".join(", ".join(group) for group in duplicate_groups)
            + ". Bu adımların farklı aşamaları temsil edip etmediğini kontrol edin."
        )

    all_claim_text = json.dumps(
        {
            "system_claim": draft.get("system_claim"),
            "dependent_system_claims": draft.get("dependent_system_claims"),
            "method_claim": draft.get("method_claim"),
            "dependent_method_claims": draft.get("dependent_method_claims"),
        },
        ensure_ascii=False,
    )
    missing = [n for n in numbers if n and f"({n})" not in all_claim_text and draft.get("method_claim")]
    if missing:
        raise ValueError("REFERANS NUMARALARI bölümündeki yöntem adımlarından istemde bulunmayan referanslar: " + ", ".join(missing))

    if claim_mode == "Yalnızca yöntem" and draft.get("system_claim"):
        raise ValueError("Yalnızca yöntem seçildiği halde sistem istemi üretildi.")
    if claim_mode == "Yalnızca sistem" and draft.get("method_claim"):
        raise ValueError("Yalnızca sistem seçildiği halde yöntem istemi üretildi.")
    if claim_mode in {"Yalnızca yöntem", "Sistem ve yöntem"} and not draft.get("method_claim"):
        raise ValueError("Seçilen istem yapısına rağmen bağımsız yöntem istemi üretilemedi.")
    if claim_mode in {"Yalnızca sistem", "Sistem ve yöntem"} and not draft.get("system_claim"):
        raise ValueError("Seçilen istem yapısına rağmen bağımsız sistem istemi üretilemedi.")

    hardware_anchor_re = re.compile(r"elektronik cihaz|elektronik işlem birimi|işlemci|donanım|bilgisayar|mikrodenetleyici|kontrol birimi|electronic device|processing unit|processor|hardware|computer|microcontroller|control unit", re.IGNORECASE)
    software_terms_re = re.compile(r"modül|birim|algoritma|yazılım|veri işleme|hesaplama|module|unit|algorithm|software|data processing|calculation", re.IGNORECASE)
    if draft.get("system_claim"):
        system_text = " ".join([
            str((draft.get("system_claim") or {}).get("preamble", "")),
            *_system_claim_all_texts(draft.get("system_claim") or {}),
        ])
        if len(software_terms_re.findall(system_text)) >= 2 and not hardware_anchor_re.search(system_text):
            raise ValueError("Yazılım/modül ağırlıklı bağımsız sistem istemi geniş bir donanımsal taşıyıcıya dayandırılmalıdır; örneğin elektronik cihaz üzerinde koşturulan yazılım vasıtasıyla.")
    if draft.get("method_claim"):
        method_text = " ".join([
            str((draft.get("method_claim") or {}).get("preamble", "")),
            *map(str, (draft.get("method_claim") or {}).get("steps") or []),
        ])
        if len(software_terms_re.findall(method_text)) >= 2 and not hardware_anchor_re.search(method_text):
            raise ValueError("Yazılım/algoritma ağırlıklı bağımsız yöntem istemi elektronik cihaz/işlemci gibi geniş bir donanımsal taşıyıcıya dayandırılmalıdır.")

    execution_relation_re = re.compile(r"üzerinde\s+(?:çalışan|koşturulan|yürütülen)|içerisinde\s+(?:çalışan|koşturulan|yürütülen)|vasıtasıyla|tarafından\s+(?:çalıştırılan|yürütülen)|executed\s+on|running\s+on|executed\s+by", re.IGNORECASE)
    if draft.get("system_claim"):
        system_text = " ".join([str((draft.get("system_claim") or {}).get("preamble", "")), *_system_claim_all_texts(draft.get("system_claim") or {})])
        if len(software_terms_re.findall(system_text)) >= 2 and not execution_relation_re.search(system_text):
            raise ValueError("Yazılım/modül ağırlıklı bağımsız sistem isteminde teknik taşıyıcı ile yazılım/modül arasında açık çalışma/koşturma ilişkisi bulunmalıdır; yalnız işlemci/donanım kelimesi yeterli değildir.")

    # BBF atomik teknik bilgi kapsam kapısı.
    if extracted is not None:
        mandatory_facts = list(extracted.get("technical_facts") or [])
        coverage = draft.get("source_coverage_map") or []
        coverage_by_id = {str(x.get("fact_id", "") or "").strip(): x for x in coverage if str(x.get("fact_id", "") or "").strip()}
        missing_ids: list[str] = []
        draft_content_for_evidence = {
            key: value for key, value in draft.items()
            if key not in {"source_coverage_map", "coverage_audit"}
        }
        searchable = re.sub(r"\s+", " ", json.dumps(draft_content_for_evidence, ensure_ascii=False)).casefold()
        invalid_evidence: list[str] = []
        for fact in mandatory_facts:
            fid = str(fact.get("id", "") or "").strip()
            if not fid:
                raise ValueError("Kaynak teknik bilgi envanterinde kimliği boş technical_facts maddesi bulundu.")
            row = coverage_by_id.get(fid)
            evidence = str((row or {}).get("evidence", "") or "").strip()
            if not row or row.get("covered") is not True or not (row.get("sections") or []) or not evidence:
                missing_ids.append(fid)
                continue
            evidence_norm = re.sub(r"\s+", " ", evidence).casefold()
            if len(evidence_norm) < 20 or evidence_norm not in searchable:
                invalid_evidence.append(fid)
        if missing_ids:
            raise ValueError("BBF tamlık kapısı başarısız: technical_facts maddelerinin tarifnamede kanıtlı karşılığı yok: " + ", ".join(missing_ids))
        if invalid_evidence:
            raise ValueError("BBF tamlık kapısı başarısız: source_coverage_map kanıtı nihai tarifname metninde birebir doğrulanamayan fact_id: " + ", ".join(invalid_evidence))

    # Tek-tuş istem kalite kapısı: ürün/sistem/yapılanma dili, unsur sırası, belirsiz referans ve özet.
    if not _english_spec(language):
        generic_names = {"diğer parçalar", "diğer parça", "diğer elemanlar", "çeşitli parçalar", "çeşitli elemanlar"}
        for element in draft.get("elements") or []:
            if str(element.get("name", "") or "").strip().casefold() in generic_names:
                raise ValueError("REFERANS NUMARALARI bölümünde ‘Diğer parçalar/Diğer elemanlar’ gibi belirsiz bir unsur adı kullanılamaz; teknik unsur net adlandırılmalıdır.")

        system_claim = draft.get("system_claim") or {}
        if system_claim:
            if str(system_claim.get("closing", "") or "").strip().casefold().rstrip(".") != "içermesidir":
                raise ValueError("Yöntem dışındaki bağımsız ürün/sistem/yapılanma istemi ‘içermesidir.’ ile kapanmalıdır.")
            action_noun_re = re.compile(
                r"\b(?:bağlanması|oluşturulması|yapılması|edilmesi|sağlanması|gerçekleştirilmesi|belirlenmesi|üretilmesi|hesaplanması|yerleştirilmesi|konumlandırılması|aktarılması|işlenmesi|tespit edilmesi)\b",
                re.IGNORECASE,
            )
            for entry in _system_claim_entries(system_claim):
                if isinstance(entry, dict):
                    lead = str(entry.get("lead", "") or "").strip()
                    subs = [str(x or "").strip() for x in (entry.get("subelements") or []) if str(x or "").strip()]
                    if not re.search(r"(?:\bve\s*;|\band\s*:)$", lead, re.IGNORECASE):
                        raise ValueError("Ortak taşıyıcı üst maddesi alt maddeleri başlatacak biçimde ‘... ve;’ (İngilizcede doğal iki nokta yapısı) ile bitmelidir.")
                    if lead.count(";") > 1:
                        raise ValueError("Ortak taşıyıcı üst maddesinde yalnız sondaki ‘ve;’ noktalı virgülüne izin verilir.")
                    for sub in subs:
                        if action_noun_re.search(sub):
                            raise ValueError("Ürün/sistem/yapılanma ana isteminin ortak taşıyıcı alt maddesinde yöntem/işlem isimleştirmesi bulundu; unsur merkezli dil kullanın.")
                        if ";" in sub:
                            raise ValueError("Ortak taşıyıcı alt maddelerinde noktalı virgül kullanılmamalıdır.")
                else:
                    item = str(entry)
                    if action_noun_re.search(item):
                        raise ValueError(
                            "Ürün/sistem/yapılanma ana isteminde işlem isimleştirmesi bulundu. ‘... bağlanması/oluşturulması’ yerine ‘... bağlanan/... yapısına sahip’ gibi unsur merkezli dil kullanın."
                        )
                    if ";" in item:
                        raise ValueError("Ana istem unsur maddelerinde noktalı virgül kullanılmamalıdır; ortak-taşıyıcı hiyerarşik grubundaki ‘ve;’ istisnası dışında kaldırın.")
            _validate_system_claim_reference_order(system_claim, element_numbers)
            _validate_reference_identity(draft)
            _validate_reference_presence(draft)
            _validate_common_carrier_scope(draft)
            _validate_main_claim_how_test(draft, extracted)
            _validate_claim_formula_markers(draft, language)

        dependents = [str(x or "").strip() for x in (draft.get("dependent_system_claims") or []) if str(x or "").strip()]
        bad_ending_re = re.compile(r"(?:yapmasıdır|etmesidir|belirlemesidir|bulunmasıdır|oluşturulmasıdır|bağlanmasıdır|sağlanmasıdır|gerçekleştirilmesidir|yapılmasıdır|edilmesidir)\.?$", re.IGNORECASE)
        for idx, claim in enumerate(dependents, start=2):
            if bad_ending_re.search(claim):
                raise ValueError(f"İstem {idx} ürün/sistem/yapılanma istem diline aykırı eylem sonucu ile bitiyor; ‘olmasıdır.’ veya ‘içermesidir.’ kullanın.")
            if not re.search(r"(?:olmasıdır|içermesidir)\.?$", claim, re.IGNORECASE):
                raise ValueError(f"İstem {idx} yöntem dışı bağımlı istemdir ve ‘olmasıdır.’ veya ‘içermesidir.’ ile bitmelidir.")
            semicolons = claim.count(";")
            if semicolons > 1 or (semicolons == 1 and not re.search(r"olup,\s*özelliği;", claim, re.IGNORECASE)):
                raise ValueError(f"İstem {idx} içinde standart ‘olup, özelliği;’ kalıbı dışında noktalı virgül kullanılmış.")
        method_dependents = [str(x or "").strip() for x in (draft.get("dependent_method_claims") or []) if str(x or "").strip()]
        for dep_index, claim in enumerate(method_dependents, start=1):
            if not re.search(r"işlem adım(?:ını|larını)\s+içermesidir\.?$", claim, re.IGNORECASE):
                raise ValueError(f"Bağımlı yöntem istemi {dep_index}, `işlem adımını içermesidir.` veya `işlem adımlarını içermesidir.` ile bitmelidir.")
        _validate_dependent_claim_semantic_repetition(system_claim, dependents)
        _validate_dependent_method_claim_semantic_repetition(
            draft.get("method_claim") or {},
            [str(x or "").strip() for x in (draft.get("dependent_method_claims") or []) if str(x or "").strip()],
        )
        _validate_no_generic_unsur_in_claims(draft, language)
        _validate_method_step_action_language(draft, language)
        _validate_claim_formula_markers(draft, language)
        _validate_abstract_shape(str(draft.get("abstract", "") or ""), language)

    audit = draft.get("coverage_audit") or {}
    mandatory_audit_flags = [
        "prior_art_complete", "reference_table_complete", "claims_consistent",
        "reference_names_clear", "reference_order_valid", "how_test_passed",
        "core_difference_present", "scope_not_overlimited", "dependent_claims_non_redundant",
        "dependent_claim_dependencies_valid", "example_dimensions_not_claim_limited",
        "product_claim_language_valid", "abstract_single_paragraph_sentence", "source_attribution_removed",
        "all_technical_facts_covered", "software_carrier_valid", "detail_intro_sentence_case",
    ]
    failed_flags = [key for key in mandatory_audit_flags if audit.get(key) is not True]
    if failed_flags:
        raise ValueError("Tarifname kalite denetiminde başarısız alanlar: " + ", ".join(failed_flags))

    user_facing_text = json.dumps(draft, ensure_ascii=False)
    if re.search(r"\bBBF\b|buluş bildirim formu|invention disclosure form|müşteri tarafından iletilen(?: teknik)? (?:çizim|belge)|müşteri bilgilerine göre|ek teknik belgede|iletilen teknik çizimde", user_facing_text, flags=re.IGNORECASE):
        raise ValueError("Tarifname taslağında kullanıcıya görünmemesi gereken kaynak/iletilen belge atfı bulundu; teknik bilgi doğrudan buluş anlatımı olarak yazılmalıdır.")
    if not _english_spec(language) and re.search(r"\bmevcut buluş\b", user_facing_text, flags=re.IGNORECASE):
        raise ValueError('Tarifname taslağında “mevcut buluş” ifadesi bulundu; “Buluş” dili kullanılmalıdır.')
    if claim_mode == "Sistem ve yöntem":
        required = "method" if _english_spec(language) else "yöntem"
        if required not in str(draft.get("title", "")).casefold():
            raise ValueError("Sistem ve yöntem istem yapısında başlık yöntem/method ifadesini içermiyor.")

    literature_paragraphs = [str(x or "").strip() for x in (draft.get("literature_paragraphs") or []) if str(x or "").strip()]
    literature_text = " ".join(literature_paragraphs)
    for doc_info in literature or []:
        en = str(doc_info.get("title_en", "") or "").strip()
        tr = str(doc_info.get("title_tr", "") or "").strip()
        if en and en not in literature_text:
            raise ValueError(f"Literatür paragrafında patent başlığı eksik: {en}")
        if not _english_spec(language) and tr and tr not in literature_text:
            raise ValueError(f"Literatür paragrafında Türkçe patent başlığı eksik: {tr}")
    if not _english_spec(language):
        for idx, paragraph in enumerate(literature_paragraphs, start=1):
            if not paragraph.startswith("Literatürde yapılan araştırmalar sonucu"):
                raise ValueError(f"Literatür paragrafı {idx}, bağlayıcı taslaktaki ‘Literatürde yapılan araştırmalar sonucu ...’ başlangıcını kullanmalıdır.")
            if re.search(r"\bBuluşta\s+ise\b", paragraph, flags=re.IGNORECASE):
                raise ValueError(f"Literatür paragrafı {idx} görüş/savunma dili içeriyor; ‘Buluşta ise ...’ yerine taslaktaki ‘Ancak ... emareye rastlanmamıştır.’ kalıbı kullanılmalıdır.")
            if not re.search(r"Ancak\s+.+?ile\s+ilgili\s+bir\s+emareye\s+rastlanmamıştır\.\s*$", paragraph, flags=re.IGNORECASE | re.DOTALL):
                raise ValueError(f"Literatür paragrafı {idx}, ‘Ancak ... ile ilgili bir emareye rastlanmamıştır.’ kalıbıyla bitmelidir.")
    return warnings


def _inline_invention_title(title: str) -> str:
    """Buluş başlığını cümle içinde normal küçük harf düzenine çevir; teknik kısaltmaları koru."""
    text = str(title or "").strip()

    def repl(match: re.Match[str]) -> str:
        token = match.group(0)
        if (len(token) > 1 and token.isupper()) or any(ch.isupper() for ch in token[1:]):
            return token
        return token.lower()

    return re.sub(r"[A-Za-zÇĞİÖŞÜçğıöşüÂâÎîÛû]+", repl, text)


def _reference_sentence_case(name: str) -> str:
    """Referans unsurunu başlık biçiminden çıkar; teknik kısaltmaları koru."""
    text = str(name or "").strip()
    first_seen = False

    def repl(match: re.Match[str]) -> str:
        nonlocal first_seen
        word = match.group(0)
        is_acronym = len(word) > 1 and word.isupper()
        if not first_seen:
            first_seen = True
            if is_acronym:
                return word
            return word[:1].upper() + word[1:].lower()
        if is_acronym:
            return word
        return word.lower()

    return re.sub(r"[A-Za-zÇĞİÖŞÜçğıöşü]+", repl, text)



def validate_tarifname_docx_structure(data: bytes, draft: dict[str, Any], language: str = "Türkçe") -> None:
    """Şablonun kullanıcı tarafından bağlayıcı kabul edilen Word yapılarını deterministik olarak denetler."""
    doc = Document(io.BytesIO(data))
    paras = doc.paragraphs
    texts = [p.text.strip() for p in paras]
    en = _english_spec(language)
    claims_label = "CLAIMS" if en else "İSTEMLER"
    abstract_label = "ABSTRACT" if en else "ÖZET"
    figures_label = "BRIEF DESCRIPTION OF THE FIGURES" if en else "ŞEKİLLERİN KISA AÇIKLAMASI"
    refs_label = "REFERENCE NUMERALS" if en else "REFERANS NUMARALARI"

    def index_of(label: str) -> int:
        try:
            return texts.index(label)
        except ValueError as exc:
            raise ValueError(f"Word şablon kontrolü: {label} başlığı bulunamadı.") from exc

    ci, ai = index_of(claims_label), index_of(abstract_label)
    fi, ri = index_of(figures_label), index_of(refs_label)
    conclusion_text = (
        "Consequently, the problems described above, which remain unresolved in view of the prior art, have created a need for an improvement in the relevant technical field."
        if en else
        "Sonuçta yukarıda bahsedilen ve mevcut teknik ışığında çözülemeyen sorunlar, ilgili teknik alanda bir yenilik yapmayı zorunlu kılmıştır."
    )
    conclusion_index = index_of(conclusion_text)
    if conclusion_index <= 0 or texts[conclusion_index - 1] != "":
        raise ValueError("Word şablon kontrolü: son önceki-teknik/literatür paragrafı ile ‘Sonuçta/Consequently’ paragrafı arasındaki fiziksel boş paragraf eksik.")
    if paras[ci].alignment != WD_ALIGN_PARAGRAPH.CENTER:
        raise ValueError("Word şablon kontrolü: İSTEMLER/CLAIMS başlığı ortalı değil.")
    if paras[ai].alignment != WD_ALIGN_PARAGRAPH.CENTER:
        raise ValueError("Word şablon kontrolü: ÖZET/ABSTRACT başlığı ortalı değil.")
    if paras[ci].paragraph_format.page_break_before is not True:
        raise ValueError("Word şablon kontrolü: İSTEMLER/CLAIMS yeni sayfadan başlamıyor.")
    if paras[ai].paragraph_format.page_break_before is not True:
        raise ValueError("Word şablon kontrolü: ÖZET/ABSTRACT yeni sayfadan başlamıyor.")
    if ci + 1 >= len(paras) or texts[ci + 1] != "":
        raise ValueError("Word şablon kontrolü: İSTEMLER başlığından sonraki şablon boş paragrafı korunmamış.")
    if fi + 1 >= len(paras) or texts[fi + 1] != "":
        raise ValueError("Word şablon kontrolü: ŞEKİLLERİN KISA AÇIKLAMASI başlığından sonraki boş paragraf korunmamış.")

    figure_descs = [str(x or "").strip() for x in (draft.get("figure_descriptions") or []) if str(x or "").strip()]
    if figure_descs:
        # Şekil açıklamaları kendi aralarında boşluksuz olmalı; sonrasında bir boş paragraf gelmeli.
        start = fi + 2
        for offset, expected in enumerate(figure_descs):
            if start + offset >= len(texts) or texts[start + offset] != expected:
                raise ValueError("Word şablon kontrolü: Şekil açıklamaları şablondaki ardışık düzende değil.")
        if start + len(figure_descs) >= len(texts) or texts[start + len(figure_descs)] != "":
            raise ValueError("Word şablon kontrolü: Son şekil açıklamasından sonraki boş paragraf eksik.")
    if ri <= fi:
        raise ValueError("Word şablon kontrolü: REFERANS NUMARALARI, şekil açıklamalarından sonra gelmelidir.")

    # REFERANS NUMARALARI yöntem satırları "1001. ..." biçiminde olmalı ve sistem/cihaz `(1)` vb. işaretleri içermemelidir.
    detail_label = "DETAILED DESCRIPTION OF THE INVENTION" if en else "BULUŞUN DETAYLI AÇIKLAMASI"
    di = index_of(detail_label)
    ref_segment = texts[ri + 1:di]
    element_numbers = [str(x.get("number", "") or "").strip() for x in (draft.get("elements") or []) if str(x.get("number", "") or "").strip()]
    for step in draft.get("method_steps") or []:
        number = str(step.get("number", "") or "").strip()
        raw = str(step.get("text", "") or "").strip()
        if number:
            raw = re.sub(rf"\s*\(\s*{re.escape(number)}\s*\)\s*$", "", raw).strip()
        expected_text = _strip_known_element_reference_marks(raw, element_numbers).rstrip(".,;:")
        expected = f"{number}. {expected_text}" if number else expected_text
        if expected not in ref_segment:
            raise ValueError(f"Word şablon kontrolü: REFERANS NUMARALARI yöntem satırı beklenen biçimde değil: {number or '?'}.")
    for ptxt in ref_segment:
        for number in element_numbers:
            if re.search(rf"\(\s*{re.escape(number)}\s*\)", ptxt):
                raise ValueError("Word şablon kontrolü: Parantezli sistem/cihaz referansları BULUŞUN DETAYLI AÇIKLAMASI bölümünden önce kullanılmamalıdır.")

    # Özet buluş adı başlığı: ÖZET + boş paragraf + kalın/ortalı başlık.
    title_idx = ai + 2
    if title_idx >= len(paras) or texts[title_idx] != str(draft.get("title", "") or "").strip():
        raise ValueError("Word şablon kontrolü: Özet içindeki buluş başlığı beklenen yerde değil.")
    title_p = paras[title_idx]
    if title_p.alignment != WD_ALIGN_PARAGRAPH.CENTER or not title_p.runs or not all(r.bold for r in title_p.runs if r.text):
        raise ValueError("Word şablon kontrolü: Özet içindeki buluş başlığı kalın ve ortalı olmalıdır.")

    # İstem paragrafları gerçek Word numaralandırmasıyla gelmeli.
    numbered_count = 0
    for p in paras[ci + 1:ai]:
        ppr = p._p.pPr
        if ppr is None:
            continue
        numpr = ppr.find(qn("w:numPr"))
        if numpr is None:
            continue
        numid = numpr.find(qn("w:numId"))
        if numid is not None and numid.get(qn("w:val")) == "2":
            numbered_count += 1
    expected_claims = (1 if draft.get("system_claim") else 0) + len(draft.get("dependent_system_claims") or []) + (1 if draft.get("method_claim") else 0) + len(draft.get("dependent_method_claims") or [])
    if numbered_count < expected_claims:
        raise ValueError("Word şablon kontrolü: İstemlerin tamamında gerçek Word otomatik numaralandırması uygulanmamış.")

    # Bağlayıcı şablonla istem kapanış konumu ve boşluk ritmi karşılaştırması.
    template = Document(str(TARIFNAME_TEMPLATE))
    template_close = template.paragraphs[93]
    def _indent_signature(p):
        pf = p.paragraph_format
        return (pf.left_indent, pf.right_indent, pf.first_line_indent, p.alignment, pf.line_spacing)

    claim_region = paras[ci + 1:ai]
    closings = [p for p in claim_region if p.text.strip() in {"içermesidir.", "işlem adımlarını içermesidir."}]
    expected_closings = (1 if draft.get("system_claim") else 0) + (1 if draft.get("method_claim") else 0)
    if len(closings) != expected_closings:
        raise ValueError("Word şablon kontrolü: bağımsız istem kapanış paragraflarının sayısı beklenenle uyuşmuyor.")
    for p in closings:
        if _indent_signature(p) != _indent_signature(template_close):
            raise ValueError("Word şablon kontrolü: 'içermesidir.' / 'işlem adımlarını içermesidir.' kapanış paragrafı şablondaki girinti ve hizaya uymuyor.")

    # Şablonda detaylı açıklama ile İSTEMLER arasında iki boş paragraf vardır.
    if ci < 2 or texts[ci-1] != "" or texts[ci-2] != "":
        raise ValueError("Word şablon kontrolü: İSTEMLER öncesindeki iki boş paragraf ritmi korunmamış.")
    # İstem açıklama metinleri ve istemler arasında tek boş paragraf korunur.
    note_texts = [template.paragraphs[79].text.strip(), template.paragraphs[81].text.strip(), template.paragraphs[83].text.strip()]
    for note in note_texts:
        if note in texts:
            ni = texts.index(note)
            if ni + 1 >= len(texts) or texts[ni + 1] != "":
                raise ValueError("Word şablon kontrolü: İSTEMLER açıklama paragrafları arasındaki boşluk şablona uymuyor.")

    # 4. kapının tam sürümü: bütün bölüm geçişleri, paragraf arketipleri, header/footer ve sayfa numarası konumu.
    validate_full_tarifname_template_fidelity(data, TARIFNAME_TEMPLATE, draft, language)


def validate_tarifname_post_generation_quality(
    data: bytes,
    draft: dict[str, Any],
    extracted: dict[str, Any],
    claim_mode: str,
    literature: list[dict[str, Any]] | None = None,
    language: str = "Türkçe",
    source_passage_registry: list[dict[str, str]] | None = None,
    final_raw_audit: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Word üretildikten sonra ham-kaynak zincirini ve 5 zorunlu kalite kapısını nihai çıktı üzerinde yeniden çalıştırır."""
    doc = Document(io.BytesIO(data))
    final_text = "\n".join(p.text for p in doc.paragraphs)

    if not source_passage_registry:
        raise ValueError("ÇIKTI SONRASI KAPI 1/5 — ham kaynak pasaj envanteri olmadan tarifname indirilemez.")
    if not final_raw_audit:
        raise ValueError("ÇIKTI SONRASI KAPI 1/5 — taslak sonrası bağımsız ham kaynak ikinci okuması yapılmadan tarifname indirilemez.")

    # 1A) Taslak üretildikten sonra yapılan bağımsız ham kaynak ikinci okumasını tekrar doğrula.
    validate_final_raw_source_audit(
        final_raw_audit,
        extracted,
        source_passage_registry,
        _visible_draft_text_for_audit(draft),
    )

    # 1B) Ham pasaj -> atomik technical_fact -> source_coverage_map -> nihai Word kanıt zinciri.
    source_stats = validate_final_source_coverage_chain(
        extracted,
        source_passage_registry,
        draft.get("source_coverage_map") or [],
        final_text,
    )

    # 2) Ana istem + alt istem kapısı: taslak denetimini Word üretiminden sonra tekrar çalıştır.
    validate_tarifname_draft(draft, claim_mode, literature or [], language, extracted)

    # 3) Referans kapısı: detaylı açıklama ve istemlerde canonical unsur kullanımları numaralı olmalı.
    _validate_reference_identity(draft)
    _validate_reference_presence(draft)

    # Nihai Word'de de DETAYLI AÇIKLAMA -> ÖZET arasında canonical isimlerin numarasız görünümünü ara.
    texts = [p.text.strip() for p in doc.paragraphs]
    detail_label = "DETAILED DESCRIPTION OF THE INVENTION" if _english_spec(language) else "BULUŞUN DETAYLI AÇIKLAMASI"
    abstract_label = "ABSTRACT" if _english_spec(language) else "ÖZET"
    try:
        di, ai = texts.index(detail_label), texts.index(abstract_label)
    except ValueError as exc:
        raise ValueError("ÇIKTI SONRASI KAPI 3/5 — referans denetimi için bölüm sınırları bulunamadı.") from exc
    segment = "\n".join(texts[di + 1:ai])
    for element in (draft.get("elements") or []):
        number = str(element.get("number", "") or "").strip()
        name = str(element.get("name", "") or "").strip()
        if not number or not name:
            continue
        mention_re = _reference_mention_pattern(name)
        ref_re = re.compile(r"^\s*(?:\([^)]{1,40}\)\s*)?\(\s*" + re.escape(number) + r"\s*\)")
        for m in mention_re.finditer(segment):
            if not ref_re.match(segment[m.end():m.end() + 70]):
                raise ValueError(f"ÇIKTI SONRASI KAPI 3/5 — '{name}' kullanımı ({number}) referansı olmadan nihai Word'e girmiş.")

    # 4) Şablon kapısı: nihai Word'ü bağlayıcı Tarifname_181176 şablon yapısıyla karşılaştır.
    validate_tarifname_docx_structure(data, draft, language)

    # 5) Unsur + işlem adımı dili kapısı: generic 'unsur' ve salt-isim yöntem adımları yasaktır.
    _validate_no_generic_unsur_in_claims(draft, language)
    _validate_method_step_action_language(draft, language)

    # Ek sert alt-kapı: formüller nihai .docx içinde düz metin değil gerçek Word matematik nesnesidir.
    _validate_word_math_format(data, draft)

    return {
        "source_completeness": True,
        "claims": True,
        "references": True,
        "template": True,
        "element_step_language": True,
        "formula_format": True,
        "how_test": True,
        **source_stats,
    }


def render_tarifname_docx_smoke_test(data: bytes) -> None:
    """Render ortamında DOCX'in LibreOffice ile PDF'e sorunsuz çevrilebildiğini doğrular."""
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        docx_path = td_path / "tarifname_qa.docx"
        docx_path.write_bytes(data)
        proc = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(td_path), str(docx_path)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=90,
            check=False,
        )
        pdf_path = td_path / "tarifname_qa.pdf"
        if proc.returncode != 0 or not pdf_path.exists() or pdf_path.stat().st_size == 0:
            detail = (proc.stderr or proc.stdout).decode("utf-8", errors="ignore")[-500:]
            raise ValueError("Word render kalite kontrolü başarısız oldu; dosya kullanıcıya sunulmadı. " + detail)
        if fitz is not None:
            pdf = fitz.open(pdf_path)
            try:
                if pdf.page_count < 1:
                    raise ValueError("Word render kalite kontrolü: PDF sayfası oluşmadı.")
                for page in pdf:
                    rect = page.rect
                    if rect.width <= 0 or rect.height <= 0:
                        raise ValueError("Word render kalite kontrolü: geçersiz sayfa geometrisi bulundu.")
            finally:
                pdf.close()


def build_tarifname_docx(draft: dict[str, Any], language: str = "Türkçe") -> bytes:
    """Bağlayıcı Tarifname_181176 şablonunu gövde arketipi olarak kullanarak DOCX üretir.

    Şablon yalnız font kaynağı değildir: başlık/boşluk ritmi, paragraf arketipleri,
    istem girintileri, header/footer ve sayfa numarası konumu şablondan korunur.
    """
    template = Document(str(TARIFNAME_TEMPLATE))
    doc = Document(str(TARIFNAME_TEMPLATE))
    clear_body(doc)
    en = _english_spec(language)
    labels = {
        "spec": "SPECIFICATION" if en else "TARİFNAME",
        "technical": "TECHNICAL FIELD" if en else "TEKNİK ALAN",
        "prior": "PRIOR ART" if en else "ÖNCEKİ TEKNİK",
        "short": "BRIEF DESCRIPTION OF THE INVENTION" if en else "BULUŞUN KISA AÇIKLAMASI",
        "figures": "BRIEF DESCRIPTION OF THE FIGURES" if en else "ŞEKİLLERİN KISA AÇIKLAMASI",
        "refs": "REFERENCE NUMERALS" if en else "REFERANS NUMARALARI",
        "detail": "DETAILED DESCRIPTION OF THE INVENTION" if en else "BULUŞUN DETAYLI AÇIKLAMASI",
        "claims": "CLAIMS" if en else "İSTEMLER",
        "abstract": "ABSTRACT" if en else "ÖZET",
    }

    # Section geometrisi ve header/footer şablondan aynen gelir; burada yeniden sayfa numarası eklenmez.
    # Özellikle footer'a PAGE alanı eklemek yasaktır.

    def tpl_text(index: int, text: str):
        copy_template_paragraph_with_text(doc, template, index, text)
        return doc.paragraphs[-1]

    def tpl_blank(index: int):
        copy_template_paragraph(doc, template, index)
        return doc.paragraphs[-1]

    def trim_trailing_blanks():
        body = doc._element.body
        while doc.paragraphs and not doc.paragraphs[-1].text.strip():
            body.remove(doc.paragraphs[-1]._p)

    # 0-5: tarifname başlığı, buluş adı ve sabit talimat; biçim şablondan birebir.
    tpl_text(0, labels["spec"])
    tpl_blank(1)
    tpl_text(2, str(draft.get("title", "") or ""))
    tpl_blank(3)
    tpl_text(4, "For preparation of the search report, the information provided in the specification and claims should be sufficiently clear and detailed to enable a person skilled in the art to carry out the subject product/method. If questions have been raised to clarify any issue, please provide the requested information. If you consider that an important element, process step or feature has not been stated before the claims section, please indicate your comments without changing the text, using highlighted text, and return them by e-mail." if en else template.paragraphs[4].text)
    tpl_blank(5)

    # TEKNİK ALAN: iki paragraf; aralarında ve bölüm sonunda şablon boşluğu.
    tpl_text(6, labels["technical"])
    tpl_blank(7)
    technical_field_parts = [x.strip() for x in re.split(r"\n\s*\n", str(draft.get("technical_field", "") or "")) if x.strip()]
    for idx, paragraph in enumerate(technical_field_parts):
        tpl_text(8, paragraph)
        tpl_blank(9)

    # ÖNCEKİ TEKNİK: her ana paragraf arasında bir boşluk; sonuç paragrafı şablonun space-after boşluğunu taşır.
    tpl_text(10, labels["prior"])
    tpl_blank(11)
    prior = [str(x or "").strip() for x in (draft.get("prior_art_general_paragraphs") or []) if str(x or "").strip()]
    literature = [str(x or "").strip() for x in (draft.get("literature_paragraphs") or []) if str(x or "").strip()]
    if not prior and not literature:
        # Eski/ara taslak şemalarıyla geriye uyumluluk; aktif şemada iki alan ayrıdır.
        prior = [str(x or "").strip() for x in (draft.get("prior_art_paragraphs") or []) if str(x or "").strip()]
    for paragraph in [*prior, *literature]:
        tpl_text(12, paragraph)
        tpl_blank(13)
    # Bağlayıcı şablon: son önceki-teknik/literatür paragrafı ile “Sonuçta...” paragrafı arasında
    # fiziksel olarak TAM BİR boş paragraf korunur. Bu boşluk trimlenmez.
    tpl_text(16, "Consequently, the problems described above, which remain unresolved in view of the prior art, have created a need for an improvement in the relevant technical field." if en else "Sonuçta yukarıda bahsedilen ve mevcut teknik ışığında çözülemeyen sorunlar, ilgili teknik alanda bir yenilik yapmayı zorunlu kılmıştır.")
    # Kullanıcı tarafından bağlayıcı hale getirilen ek şablon kuralı: sonuç paragrafı
    # ile BULUŞUN KISA AÇIKLAMASI arasında fiziksel olarak tam bir boş paragraf bulunur.
    tpl_blank(15)

    # BULUŞUN KISA AÇIKLAMASI
    tpl_text(17, labels["short"])
    tpl_blank(18)
    tpl_text(19, str(draft.get("short_description_intro", "") or ""))
    tpl_blank(20)
    for index, objective in enumerate(draft.get("objectives") or []):
        objective = str(objective or "").strip()
        if en:
            text = objective
        else:
            prefix = "Buluşun ana amacı, " if index == 0 else "Buluşun diğer bir amacı, "
            text = prefix + (objective[:1].lower() + objective[1:] if objective else "")
        tpl_text(21, text)
        tpl_blank(22)

    invention_definition = draft.get("unumbered_invention_definition") or draft.get("unumbered_system_definition")
    if invention_definition:
        tpl_text(29, str(invention_definition))
    invention_features = draft.get("unumbered_invention_features") or draft.get("unumbered_system_elements") or []
    for item in invention_features:
        add_template_list_item(doc, template, 30, str(item))
    if invention_features:
        if en:
            tpl_text(37, "The invention comprises the foregoing technical features." if draft.get("system_claim") else "The method comprises the foregoing process steps.")
        else:
            tpl_text(37, "işlem adımlarını içermesidir." if draft.get("method_claim") and not draft.get("system_claim") else "içermesidir.")
    tpl_blank(38)
    tpl_text(39, "For a better understanding of the configuration and advantages of the invention together with its additional elements, the invention should be considered with reference to the figures described below." if en else "Buluşun yapılanması ve ek elemanlarla birlikte avantajlarının en iyi şekilde anlaşılabilmesi için aşağıda açıklaması yapılan şekiller ile birlikte değerlendirilmesi gerekmektedir.")
    tpl_blank(40)

    # ŞEKİLLERİN KISA AÇIKLAMASI: başlıktan sonra boşluk; şekiller kendi aralarında bitişik; yalnız son şekilden sonra boşluk.
    tpl_text(41, labels["figures"])
    tpl_blank(42)
    figure_descriptions = [str(x or "").strip() for x in (draft.get("figure_descriptions") or []) if str(x or "").strip()]
    if not figure_descriptions:
        figure_descriptions = ["Figure 1 is a representative illustration of the configuration of the invention." if en else "Şekil 1, buluşa konu yapılanmanın temsili gösterimidir."]
    for figure in figure_descriptions:
        tpl_text(43, figure)
    tpl_blank(44)
    tpl_text(45, "The drawings are not necessarily to scale, and details that are not required for understanding the invention may be omitted. Elements that are substantially identical or perform substantially identical functions may be indicated by the same reference numeral or sign." if en else "Çizimlerin mutlaka ölçeklendirilmesi gerekmemektedir ve buluşu anlamak için gerekli olmayan detaylar ihmal edilmiş olabilmektedir. Bundan başka, en azından önemli ölçüde özdeş elemanlar veya benzer fonksiyonlara sahip olan elemanlar aynı numara ile gösterilmektedir.")
    tpl_blank(46)

    # REFERANS NUMARALARI: sistem referansları bitişik; yöntem varsa arada tek boşluk; detay başlığından önce tek boşluk.
    tpl_text(47, labels["refs"])
    tpl_blank(48)
    elements = draft.get("elements") or []
    for element in elements:
        # Referans listesinde unsur adı draft.elements ile karakter-karakter aynıdır;
        # sentence-case dönüşümü teknik kısaltma/proper-case adlarını (örn. Monte Carlo) bozabilir.
        tpl_text(49, f"{element.get('number','')}. {str(element.get('name','') or '').strip()}")
    method_steps = draft.get("method_steps") or []
    if elements and method_steps:
        tpl_blank(56)
    element_numbers_for_refs = [str(x.get("number", "") or "").strip() for x in elements]
    for step in method_steps:
        number = str(step.get("number", "") or "").strip()
        text = str(step.get("text", "") or "").strip()
        if number:
            text = re.sub(rf"\s*\(\s*{re.escape(number)}\s*\)\s*$", "", text).strip()
        text = _strip_known_element_reference_marks(text, element_numbers_for_refs).rstrip(".,;:")
        tpl_text(49, f"{number}. {text}" if number else text)
    tpl_blank(56)

    # BULUŞUN DETAYLI AÇIKLAMASI
    tpl_text(57, labels["detail"])
    tpl_blank(58)
    title = str(draft.get("title", "buluş") or "buluş").strip()
    inline_title = title if en else _inline_invention_title(title)
    tpl_text(59, f"In this detailed description, {title} is described by way of examples solely for a better understanding of the subject matter, without imposing any limiting effect." if en else f"Bu detaylı açıklamada, buluş konusu olan {inline_title} sadece konunun daha iyi anlaşılmasına yönelik hiçbir sınırlayıcı etki oluşturmayacak örneklerle açıklanmaktadır.")
    tpl_blank(60)

    detailed = [str(x or "").strip() for x in (draft.get("detailed_paragraphs") or []) if str(x or "").strip()]
    for paragraph in detailed:
        tpl_text(61, paragraph)
        tpl_blank(62)

    for formula in draft.get("formulas") or []:
        if formula.get("label"):
            p = tpl_text(61, str(formula.get("label", "")))
            for r in p.runs:
                if r.text:
                    r.bold = True
        if formula.get("expression"):
            add_display_equation(doc, template, 61, str(formula.get("expression", "")))
        if formula.get("explanation"):
            tpl_text(61, str(formula.get("explanation", "")))
        tpl_blank(62)

    for table_data in draft.get("tables") or []:
        caption = str(table_data.get("caption", "") or "").strip()
        if caption:
            p = tpl_text(61, caption)
            for r in p.runs:
                if r.text:
                    r.bold = True
        headers = [str(x) for x in table_data.get("headers") or []]
        rows = table_data.get("rows") or []
        column_count = max(len(headers), max((len(row) for row in rows), default=0))
        if column_count:
            table = doc.add_table(rows=1 if headers else 0, cols=column_count)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            if headers:
                for idx in range(column_count):
                    set_cell_text(table.rows[0].cells[idx], headers[idx] if idx < len(headers) else "", bold=True)
            for row_data in rows:
                cells = table.add_row().cells
                for idx in range(column_count):
                    set_cell_text(cells[idx], str(row_data[idx]) if idx < len(row_data) else "")
        tpl_blank(62)

    for paragraph in draft.get("experimental_results") or []:
        tpl_text(61, str(paragraph))
        tpl_blank(62)
    for paragraph in draft.get("alternatives") or []:
        tpl_text(61, str(paragraph))
        tpl_blank(62)

    if method_steps:
        tpl_text(65, "The process steps performed by the method are as follows:" if en else "Yöntemin gerçekleştirdiği işlem adımları aşağıdaki gibidir:")
        for idx, step in enumerate(method_steps):
            number = str(step.get("number", "") or "").strip()
            text = str(step.get("text", "") or "").strip()
            if number:
                text = re.sub(rf"\s*\(\s*{re.escape(number)}\s*\)\s*$", "", text).strip()
            text = text.rstrip(".,;:")
            punctuation = "." if idx == len(method_steps) - 1 else ","
            add_template_list_item(doc, template, 30, f"{text} ({number}){punctuation}")
        tpl_blank(73)
    if draft.get("working_principle"):
        # Şablonun 74. paragrafında metnin sonunda eski bir manuel sayfa sonu bulunur.
        # Aktif üretimde İSTEMLER başlığını page_break_before ile yeni sayfaya aldığımız için
        # bu manuel kırılmayı taşımak çift sayfa sonu ve boş sayfa üretir. Biçimi koruyup
        # yalnız sayfa sonu kontrolünü kaldırıyoruz.
        working_p = tpl_text(74, str(draft.get("working_principle", "")))
        for run in working_p.runs:
            for br in list(run._r.findall(qn("w:br"))):
                if br.get(qn("w:type")) == "page":
                    run._r.remove(br)

    # İSTEMLER öncesi TAM iki boş paragraf. Önceki içerikten kalan boşluklar temizlenir.
    trim_trailing_blanks()
    tpl_blank(75)
    tpl_blank(76)
    tpl_text(77, labels["claims"])
    claims_heading = doc.paragraphs[-1]
    claims_heading.paragraph_format.page_break_before = True
    claims_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tpl_blank(78)
    if en:
        english_claim_notes = {
            79: "If you consider that an important and novel element and/or process step is missing from the claims, please indicate your comments without changing the text, using highlighted text, and return them by e-mail.",
            81: "If the system can operate without one or more features recited in claim 1, or if a feature that must necessarily be present is not recited in claim 1, identifying that feature is important for determining the desired scope of protection.",
            83: "A system that implements the invention without using even one feature recited in claim 1 may fall outside the scope of that claim. Please evaluate the claims in light of this information.",
        }
        for index, blank_index in ((79,80),(81,82),(83,84)):
            tpl_text(index, english_claim_notes[index])
            tpl_blank(blank_index)
    else:
        for index, blank_index in ((79,80),(81,82),(83,84)):
            tpl_text(index, template.paragraphs[index].text)
            tpl_blank(blank_index)

    system_claim = draft.get("system_claim")
    if system_claim:
        entries = _system_claim_entries(system_claim)
        if en:
            add_numbered_claim(doc, template, f"{system_claim.get('preamble','').rstrip(' ,;:')}, comprising:")
            for idx, entry in enumerate(entries):
                if isinstance(entry, dict):
                    lead = str(entry.get("lead", "") or "").strip().rstrip(";:") + ":"
                    add_template_list_item(doc, template, 86, lead)
                    for sub in [str(x or "").strip() for x in (entry.get("subelements") or []) if str(x or "").strip()]:
                        add_nested_claim_list_item(doc, template, sub.rstrip(".,;:") + ";")
                else:
                    text = str(entry).rstrip(".,;:") + ("." if idx == len(entries) - 1 else ";")
                    add_template_list_item(doc, template, 86, text)
        else:
            add_numbered_claim(doc, template, f"{system_claim.get('preamble','')} olup, özelliği;")
            for entry in entries:
                if isinstance(entry, dict):
                    add_template_list_item(doc, template, 86, str(entry.get("lead", "") or ""))
                    for sub in entry.get("subelements") or []:
                        add_nested_claim_list_item(doc, template, str(sub))
                else:
                    add_template_list_item(doc, template, 86, str(entry))
            tpl_text(93, system_claim.get("closing", "içermesidir."))
        tpl_blank(94)
        for dependent in draft.get("dependent_system_claims") or []:
            add_numbered_claim(doc, template, str(dependent))
            tpl_blank(96)

    method_claim = draft.get("method_claim")
    if method_claim:
        if en:
            add_numbered_claim(doc, template, f"{method_claim.get('preamble','').rstrip(' ,;:')}, comprising:")
        else:
            add_numbered_claim(doc, template, f"{method_claim.get('preamble','')} olup, özelliği;")
        claim_steps = list(method_claim.get("steps") or [])
        for idx, item in enumerate(claim_steps):
            base_step = str(item).rstrip(".,;:")
            if en:
                step_text = base_step + ("." if idx == len(claim_steps) - 1 else ";")
            else:
                step_text = base_step if idx == len(claim_steps) - 1 else base_step + ","
            add_template_list_item(doc, template, 86, step_text)
        if not en:
            tpl_text(93, method_claim.get("closing", "işlem adımlarını içermesidir."))
        tpl_blank(94)
        for dependent in draft.get("dependent_method_claims") or []:
            add_numbered_claim(doc, template, str(dependent))
            # 98 numaralı şablon paragrafı ÖZET öncesindeki manuel sayfa sonunu içerir;
            # istemler arasında kullanılması her bağımlı yöntem istemini yeni sayfaya atar.
            # İstemler arasındaki normal boşluk arketipi 96 kullanılır.
            tpl_blank(96)

    # ÖZET öncesi tam bir boşluk ve yeni sayfa.
    trim_trailing_blanks()
    # Yeni sayfa kontrolü ÖZET başlığının page_break_before özelliğindedir; burada
    # manuel sayfa sonu içermeyen normal bir şablon boşluğu kullanılır.
    tpl_blank(96)
    tpl_text(99, labels["abstract"])
    summary_heading = doc.paragraphs[-1]
    summary_heading.paragraph_format.page_break_before = True
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tpl_blank(100)
    tpl_text(101, str(draft.get("title", "") or ""))
    tpl_blank(102)
    tpl_text(103, str(draft.get("abstract", "") or ""))
    tpl_blank(104)
    tpl_blank(105)

    out = io.BytesIO()
    doc.save(out)
    data = out.getvalue()
    # Üretim fonksiyonunun kendi içinde de tam şablon kapısı çalışır; dış akışta tekrar edilir.
    validate_full_tarifname_template_fidelity(data, TARIFNAME_TEMPLATE, draft, language)
    return data


# -----------------------------------------------------------------------------
# GÖRÜŞ MODÜLÜ
# -----------------------------------------------------------------------------
def gorus_analysis_prompt(
    report_type: str,
    reference: str,
    report_text: str,
    spec_text: str,
    prior_opinion_text: str,
    similar_text: str,
    customer_text: str,
) -> str:
    return f"""{GORUS_RULES}
Aşağıdaki dosyaları önce YALNIZCA analiz et. Bu aşamada görüş Word metni yazma.
Görüş türü: {report_type}
Ana dosya referansı: {reference}

Önce rapordaki itirazları, mevcut istemleri, tarifname dayanaklarını, varsa önceki görüşü ve müşteri bilgisini birlikte değerlendir. Türkiye araştırma raporuysa veya EP dosyası araştırma raporuysa savunma kapsamını yalnız X/Y kategorisi dokümanlarla sınırla. A kategorisini savunma dokümanı yapma. İnceleme/ofis aksiyonlarında yalnız uzmanın gerekçede fiilen kullandığı dokümanları esas al.
İstem değişikliği sırf daha iyi yazılabilir diye önerilmez. Yalnızca itirazı gidermek için gerçekten zorunluysa amendment_required=true yap.
Revizyon gerekiyorsa EN AZ DEĞİŞİKLİK ilkesini uygula. Her old_text, TARİFNAME içindeki tek bir paragrafta birebir bulunabilen mümkün olan en kısa ifade olsun; tüm istemi old_text olarak verme. Değişmeyen kelimeyi old_text/new_text içine alma: artikel değişiyorsa yalnız artikel, unsur adı değişiyorsa yalnız değişen unsur adı, eksik harf varsa yalnız gerekli karakter farkı öner.
EP raporunda Rule 42(1)(b) gereği önceki teknik dokümanlarının tarifnameye eklenmesi isteniyorsa description_prior_art_updates üret. Bu alanda D1/D2 etiketi kullanma. Mevcut `As a result of the research on the subject...` formatını izle. objective_summary yalnız ilgili X/Y kaynağın gerçek içeriği olsun. however_difference yalnız as-filed tarifnamede/istemlerde açıkça bulunan teknik farkı kullansın, yeni özellik veya yeni teknik etki eklemesin.
Her basis_quote tarifnamede birebir bulunan dayanak pasajı olsun. Kapsam aşımı/yeni konu yaratma.

JSON dışında yazma.
ŞEMA:
{{
  "analysis_summary":"",
  "examiner_issues":[""],
  "defense_direction":[""],
  "amendment_required":false,
  "amendment_reason":"",
  "no_amendment_reason":"",
  "amendments":[
    {{
      "claim_number":"1",
      "reason":"",
      "basis_quote":"",
      "old_text":"",
      "new_text":""
    }}
  ],
  "description_prior_art_updates":[
    {{
      "source_kind":"application|document",
      "publication_number":"",
      "document_title":"",
      "objective_summary":"",
      "however_difference":"",
      "basis_quote":""
    }}
  ]
}}

RAPOR:\n{report_text}\n
TARİFNAME:\n{spec_text}\n
ÖNCEKİ GÖRÜŞ:\n{prior_opinion_text}\n
BENZER DOKÜMANLAR:\n{similar_text}\n
MÜŞTERİ BİLGİLERİ:\n{customer_text}\n"""


def gorus_revision_refine_prompt(
    report_type: str,
    report_text: str,
    spec_text: str,
    prior_opinion_text: str,
    similar_text: str,
    customer_text: str,
    current_analysis: dict[str, Any],
    user_instruction: str,
) -> str:
    return f"""{GORUS_RULES}
Aşağıdaki mevcut istem-revizyon analizini kullanıcının talimatına göre yeniden değerlendir.
Yalnızca gerçekten zorunlu değişiklikleri bırak. En az değişiklik, açık tarifname dayanağı ve kapsam aşımı yapmama kuralları bağlayıcıdır.
old_text, kaynak TARİFNAME içindeki tek bir paragrafta birebir bulunabilen mümkün olan en kısa ifade olmalıdır. Değişmeyen ön/son kelimeleri old_text/new_text içine alma. Artikel değişimini yalnız artikel, unsur adı değişimini yalnız değişen unsur adı, eksik harf düzeltmesini yalnız karakter bazında öner.
Tüm istemi silip yeniden yazma. Kullanıcının talimatı teknik kaynaklarla çelişiyorsa uydurma yapma; güvenli olan minimum revizyonu seç.

JSON dışında yazma ve ilk analizle AYNI şemayı kullan.
Görüş türü: {report_type}
KULLANICI TALİMATI:\n{user_instruction}\n
MEVCUT ANALİZ:\n{json.dumps(current_analysis, ensure_ascii=False, indent=2)}\n
RAPOR:\n{report_text}\n
TARİFNAME:\n{spec_text}\n
ÖNCEKİ GÖRÜŞ:\n{prior_opinion_text}\n
BENZER DOKÜMANLAR:\n{similar_text}\n
MÜŞTERİ BİLGİLERİ:\n{customer_text}\n"""


def validate_gorus_analysis(analysis: dict[str, Any], spec_text: str) -> None:
    required = bool(analysis.get("amendment_required"))
    amendments = analysis.get("amendments") or []
    if required and not amendments:
        raise ValueError("Analiz istem revizyonu gerekli dedi ancak revizyon önerisi üretmedi.")
    normalized_spec = re.sub(r"\s+", " ", spec_text).strip()
    for item in amendments:
        old_text = re.sub(r"\s+", " ", str(item.get("old_text", ""))).strip()
        new_text = str(item.get("new_text", "")).strip()
        basis = re.sub(r"\s+", " ", str(item.get("basis_quote", ""))).strip()
        if not old_text or not new_text:
            raise ValueError("Revizyon önerisinde old_text/new_text boş bırakılamaz.")
        if required and not basis:
            raise ValueError("Her istem revizyonu için tarifnameden birebir dayanak gösterilmelidir.")
        if re.match(r"^\s*\d+\s*[.)-]", old_text) or len(old_text) > 600:
            raise ValueError("Revizyon old_text alanı tüm istemi kapsamamalı; mümkün olan en küçük ifade seçilmelidir.")
        if old_text not in normalized_spec:
            raise ValueError(f"Revize edilecek eski ifade tarifnamede birebir doğrulanamadı: {old_text[:120]}...")
        if basis and basis not in normalized_spec:
            raise ValueError(f"Revizyon dayanağı tarifnamede birebir doğrulanamadı: {basis[:120]}...")

    for upd in analysis.get("description_prior_art_updates") or []:
        kind=str(upd.get("source_kind", "")).strip().lower()
        objective=re.sub(r"\s+", " ", str(upd.get("objective_summary", ""))).strip()
        difference=re.sub(r"\s+", " ", str(upd.get("however_difference", ""))).strip()
        basis=re.sub(r"\s+", " ", str(upd.get("basis_quote", ""))).strip()
        if kind not in {"application","document"} or not objective or not difference or not basis:
            raise ValueError("EP önceki teknik güncellemesinde source_kind/objective_summary/however_difference/basis_quote zorunludur.")
        if re.search(r"\bD[1-9]\b", objective+" "+difference, flags=re.I):
            raise ValueError("EP tarifname literatüründe D1/D2 etiketi kullanılamaz.")
        if basis not in normalized_spec:
            raise ValueError("EP However fark cümlesinin tarifname dayanağı birebir doğrulanamadı.")


def gorus_prompt(
    report_type: str,
    reference: str,
    report_text: str,
    spec_text: str,
    prior_opinion_text: str,
    similar_text: str,
    customer_text: str,
    preanalysis: dict[str, Any] | None = None,
    revision_status: str = "Mevcut istemlerle devam",
    output_language: str = "Türkçe",
    applicant_override: str = "",
) -> str:
    language_instruction = (
        "Nihai görüşün tamamını İngilizce yaz; ancak bağlayıcı Word şablonunun kurum/metadata yerleşimini koru."
        if _english_spec(output_language)
        else "Nihai görüşün tamamını Türkçe yaz."
    )
    return f"""{GORUS_RULES}
Aşağıdaki dosyalara dayanarak seçilen görüş türüne uygun ilgili patent ofisine sunulacak ayrıntılı görüş metni hazırla.
Görüş türü: {report_type}
Ana dosya referansı: {reference}
Nihai istem durumu: {revision_status}
Görüş dili: {output_language}
Kullanıcı tarafından girilmiş başvuru sahibi (varsa bağlayıcı): {applicant_override}
{language_instruction}

Bu aşama ilk teknik analizden SONRA çalışmaktadır. İstemlerde kendiliğinden yeni revizyon önerme veya onaylanmış istem setini değiştirme.
Raporun sonucunu dürüstçe koru. Buluş basamağı itirazında genel değerlendirme bölümünü raporun gerçek doküman kapsamına göre kur. Uzman yalnız D1 kullanmışsa tek-D1 inceleme gerekçesi üzerinden değerlendir, birden fazla doküman fiilen kullanılmışsa birlikte değerlendirme yap.
Tarifname alıntılarında sayfa/satır numarası YAZMA; yalnız birebir quote text döndür. Sayfa ve satırlar Word üretiminden önce fiziksel tarifname üzerinden deterministik olarak eklenecektir.
Her D1/D2/D3 bölümünün `figure_caption` alanı şablondaki `D1 dokümanı - Şekil ...` biçiminde olsun; yalnız yüklenen özgün patentte gerçekten bulunan şekil numaralarını kullan.

JSON dışında yazma.
ŞEMA:
{{
 "application_no":"", "applicant":"", "reference":"{reference}", "report_date":"", "intro":"",
 "cited_documents":[{{"label":"D1","number":"","title":"","category":"","summary":""}}],
 "sections":[
   {{
     "label":"D1",
     "heading":"D1 (...) dokümanı:",
     "figure_caption":"D1 dokümanı - Şekil ...",
     "blocks":[
       {{"type":"paragraph","text":""}},
       {{"type":"quote","text":"","attach_to_previous":true}},
       {{"type":"paragraph","text":""}}
     ],
     "inventive_step_heading":"D1 karşısında buluş basamağı",
     "inventive_step_paragraphs":[""]
   }}
 ],
 "combined_assessment":{{"heading":"","paragraphs":[""]}},
 "conclusion":[""], "signoff":"Saygılarımızla,\nDESTEK PATENT A.Ş."
}}

ÖZEL:
- Şablon girişine uy: intro kısa olsun ve yalnız rapor tarihi/türü + istem/kriter sonucunu söylesin. Girişte D1/D2/D3 seçimini, `en yakın doküman` bilgisini veya hangi dokümana karşı savunma yapıldığını anlatma.
- `applicant_override` boş değilse JSON `applicant` alanını aynen bu değer yap; değiştirme veya kısaltma. Boşsa yalnız rapor/tarifnameden güvenilir biçimde çıkar.
- İnceleme raporunda X/Y etiketi yoksa category alanını boş bırak; uydurma kategori yazma.
- Uzman gerekçeli değerlendirmeyi yalnız D1 üzerinden kurmuşsa YALNIZ D1'i görüşe al. D2/D3 yalnız `ilgili dokümanlar` listesinde bulunuyor ancak gerekçede kullanılmıyorsa görüşe bölüm, şekil veya tamamlayıcı savunma olarak ekleme.
- Her dokümanın teknik öğretisini gerçekten yüklenen metinden çıkar. Patentte bulunmayan unsur/işlev yazma.
- Tarifname alıntıları spec metninde birebir geçen tam cümle/pasaj olsun.
- Buluş basamağı zincirinde çekirdek sıra teknik fark → teknik etki → objektif teknik problem şeklinde görünür olsun. Ayrıca ayırt edici teknik katkıyı, motivasyon/yönlendirmeyi, gereken ilave yapısal/işlevsel değişiklikleri ve geriye dönük değerlendirme riskini açıkça kur.
- Tek D1 gerekçesi varsa `combined_assessment.heading` alanını `D1 dokümanının inceleme gerekçesiyle birlikte değerlendirilmesi` mantığında kur. Birden çok doküman gerekçede kullanılmışsa `Dokümanların birlikte değerlendirilmesi` kullanılabilir.
- Tarifname quote bloğunu hemen önceki teknik savunmanın doğal devamı yap ve `attach_to_previous=true` döndür. `Tarifname sayfa...` ayrı paragraf olmayacak.
- `Bu teknik farkın...`, `Bu teknik etki...`, `Buna göre objektif teknik problem...` gibi bir önceki argümanın devamını gereksiz yeni paragrafa bölme.
- Model tarafından yazılan görüş metninde noktalı virgül (`;`) kullanma. Virgül veya nokta kullan. Birebir kaynak alıntısı noktalı virgül içeriyorsa alıntıyı değiştirme.
- Önceki teknik dokümanının unsur referans numaralarını (ör. piezoelektrik eleman 120, oturma tespit anahtarı 150) savunma için zorunlu olmadıkça yazma. Başvurunun kendi tarifname referansları gerektiğinde kullanılabilir.
- Genel değerlendirme en az birkaç güçlü paragraf olsun, yalnız dokümanı özetleme, uzmanı teknik katkı üzerinden ikna et.
- Müşteri bilgisinin tarifname dayanağı yoksa kullanma.
- Onaylı istem setine yeni değişiklik ekleme.
- Türkiye araştırma raporunda ve EP araştırma raporunda cited_documents ve sections alanlarında yalnız X/Y kategori dokümanları yer alsın. A kategorisi görüş bölümü değildir. İnceleme/ofis aksiyonunda yalnız uzmanın gerekçede fiilen kullandığı dokümanlara bölüm aç.
- EP ve İngilizce çıktıysa `intro` şu formatı izlesin: `In the Extended European Search Report dated [date], objections were raised under ... Applicant’s observations and amendments in response to the objections raised in the communication are hereby submitted below.` Doküman listesi intro içine gömülmez, cited_documents alanı üzerinden hemen arkasından gelir.
- EP ve İngilizce çıktıysa `conclusion` son paragrafı şu bağlayıcı formatta kur: `In the light of above explanations and defence, we believe that our claims meet the [applicable criteria] criteria. However, in case of probable rejection, we kindly request to be given opportunity for further amendments or at least oral proceedings.`
- EP bağımlı istem itirazında tüm itirazlı bağımlı istemleri veya teknik olarak anlamlı istem gruplarını ele al. Her grupta ek özelliğin bağımsız istemle birlikte teknik katkısını açıkla. Bluetooth, WiFi, QR vb. tekil teknolojileri tek başına inventive diye sunma.
- Article 84 için `the actor` gibi antecedent düzeltmesinde tarifnamede açıkça hangi rol gösterilmişse onu kullan. Açık dayanak yoksa otomatik `actors` çoğullaştırması yapma.

ÖN ANALİZ:\n{json.dumps(preanalysis or {}, ensure_ascii=False, indent=2)}\n
RAPOR:\n{report_text}\n
TARİFNAME / ONAYLI NİHAİ İSTEM SETİ:\n{spec_text}\n
ÖNCEKİ GÖRÜŞ:\n{prior_opinion_text}\n
BENZER DOKÜMANLAR:\n{similar_text}\n
MÜŞTERİ BİLGİLERİ:\n{customer_text}\n"""


def gorus_quality_audit_prompt(
    report_text: str,
    spec_text: str,
    prior_opinion_text: str,
    similar_text: str,
    customer_text: str,
    preanalysis: dict[str, Any],
    opinion: dict[str, Any],
) -> str:
    return f"""{GORUS_RULES}
Aşağıdaki oluşturulmuş GÖRÜŞ TASLAĞINI, ham kaynakların tamamına karşı bağımsız ikinci okuyucu olarak denetle. Metni yeniden yazma. Her kontrol için pass ve kısa note döndür. En küçük şüphede pass=false yap. Özellikle raporda sadece listelenen fakat gerekçede kullanılmayan dokümanın görüşe sızıp sızmadığını, uzmanın dayandığı her paragraf/istem gerekçesine cevap verilip verilmediğini, teknik katkının tarifnameye dayalı kurulup kurulmadığını, noktalı virgül bulunup bulunmadığını, tarifname dayanağının savunmanın aynı paragrafına bağlanıp bağlanmadığını ve önceki teknik referans numaralarının gereksiz kullanılıp kullanılmadığını kontrol et.

JSON dışında yazma.
ŞEMA:
{{
  "overall_pass": true,
  "checks": {{
    "report_scope": {{"pass":true,"note":""}},
    "examiner_ground_response": {{"pass":true,"note":""}},
    "technical_contribution": {{"pass":true,"note":""}},
    "technical_effect_problem": {{"pass":true,"note":""}},
    "motivation_hindsight": {{"pass":true,"note":""}},
    "spec_support_new_matter": {{"pass":true,"note":""}},
    "intro_concision": {{"pass":true,"note":""}},
    "paragraph_flow_inline_basis": {{"pass":true,"note":""}},
    "punctuation": {{"pass":true,"note":""}},
    "prior_art_reference_numbers": {{"pass":true,"note":""}},
    "conclusion_consistency": {{"pass":true,"note":""}}
  }},
  "required_fixes": []
}}

RAPOR:\n{report_text}\n
TARİFNAME / ONAYLI İSTEMLER:\n{spec_text}\n
ÖNCEKİ GÖRÜŞ:\n{prior_opinion_text}\n
SAVUNMA DOKÜMANLARI:\n{similar_text}\n
MÜŞTERİ BİLGİLERİ:\n{customer_text}\n
ÖN ANALİZ:\n{json.dumps(preanalysis or {}, ensure_ascii=False, indent=2)}\n
GÖRÜŞ TASLAĞI JSON:\n{json.dumps(opinion or {}, ensure_ascii=False, indent=2)}\n"""


def gorus_repair_prompt(
    report_text: str,
    spec_text: str,
    prior_opinion_text: str,
    similar_text: str,
    customer_text: str,
    preanalysis: dict[str, Any],
    opinion: dict[str, Any],
    audit: dict[str, Any],
) -> str:
    return f"""{GORUS_RULES}
Aşağıdaki görüş JSON'u ikinci kalite kontrolünde başarısız oldu. Yalnız belirtilen sorunları düzelt ve AYNI JSON ŞEMASIYLA eksiksiz görüş JSON'unu yeniden döndür. Metadata, onaylı istem seti, rapor sonucu ve kaynak dayanakları korunmalı. Yeni doküman veya yeni teknik özellik ekleme. Tarifname alıntıları birebir kalmalı. Model anlatımında noktalı virgül kullanma. Doğrudan tarifname dayanağını önceki savunma paragrafına `attach_to_previous=true` ile bağla.

JSON dışında yazma.
KALİTE RAPORU:\n{json.dumps(audit or {}, ensure_ascii=False, indent=2)}\n
MEVCUT GÖRÜŞ:\n{json.dumps(opinion or {}, ensure_ascii=False, indent=2)}\n
RAPOR:\n{report_text}\n
TARİFNAME / ONAYLI İSTEMLER:\n{spec_text}\n
ÖNCEKİ GÖRÜŞ:\n{prior_opinion_text}\n
SAVUNMA DOKÜMANLARI:\n{similar_text}\n
MÜŞTERİ BİLGİLERİ:\n{customer_text}\n
ÖN ANALİZ:\n{json.dumps(preanalysis or {}, ensure_ascii=False, indent=2)}\n"""

def validate_quotes(opinion: dict[str, Any], spec_text: str) -> None:
    normalized_spec = re.sub(r"\s+", " ", spec_text).strip()
    for section in opinion.get("sections") or []:
        for block in section.get("blocks") or []:
            if str(block.get("type", "")).lower() != "quote":
                continue
            text = re.sub(r"\s+", " ", str(block.get("text", ""))).strip()
            if text and text not in normalized_spec:
                raise ValueError(f"Tarifname alıntısı birebir doğrulanamadı: {text[:120]}...")
        for quote in section.get("quotes") or []:
            text = re.sub(r"\s+", " ", str(quote.get("text", ""))).strip()
            if text and text not in normalized_spec:
                raise ValueError(f"Tarifname alıntısı birebir doğrulanamadı: {text[:120]}...")


def _clone_paragraph_with_text(doc: Document, template_para, text: str, *, bold: bool | None = False, italic: bool | None = False):
    p_el = deepcopy(template_para._p)
    # Keep only paragraph properties; rebuild runs so dynamic content cannot inherit stale text.
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)
    body = doc._element.body
    body.insert(-1, p_el)
    p = doc.paragraphs[-1]
    r = p.add_run(str(text or ""))
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.bold = bold
    r.italic = italic
    # Opinion body is always 1.5 spaced and justified; institutional title archetypes are not passed here.
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _clone_blank(doc: Document, template_para):
    p_el = deepcopy(template_para._p)
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)
    doc._element.body.insert(-1, p_el)
    return doc.paragraphs[-1]


def _add_quote_with_location(doc: Document, template_para, block: dict[str, Any]):
    lead = str(block.get("lead", "")).strip()
    quote = str(block.get("text", "")).strip()
    p = _clone_paragraph_with_text(doc, template_para, "", bold=False)
    r1 = p.add_run((lead + " ") if lead else "")
    r1.font.name = "Arial"; r1.font.size = Pt(11); r1.bold = False
    r2 = p.add_run(f'“{quote}”')
    r2.font.name = "Arial"; r2.font.size = Pt(11); r2.bold = True
    return p



def _append_quote_with_location(paragraph, block: dict[str, Any]):
    lead = str(block.get("lead", "")).strip()
    quote = str(block.get("text", "")).strip()
    prefix = " " if paragraph.text and not paragraph.text.endswith((" ", "\n")) else ""
    r1 = paragraph.add_run(prefix + ((lead + " ") if lead else ""))
    r1.font.name = "Arial"; r1.font.size = Pt(11); r1.bold = False
    r2 = paragraph.add_run(f'“{quote}”')
    r2.font.name = "Arial"; r2.font.size = Pt(11); r2.bold = True
    return paragraph

def _add_original_figure_table(doc: Document, template: Document, caption: str, png_data: bytes):
    tbl_xml = deepcopy(template.tables[1]._tbl)
    doc._element.body.insert(-1, tbl_xml)
    table = doc.tables[-1]
    # Caption row: same table geometry and border archetype as binding template.
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(caption)
    r.font.name = "Arial"; r.font.size = Pt(11); r.bold = False
    # Figure row: original patent figure/page only.
    fig_cell = table.rows[1].cells[0]
    fig_cell.text = ""
    p2 = fig_cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.0
    img = Image.open(io.BytesIO(png_data))
    max_w = 15.2
    max_h = 15.8
    ratio = img.width / max(1, img.height)
    width_cm = min(max_w, max_h * ratio)
    height_cm = width_cm / ratio
    if height_cm > max_h:
        height_cm = max_h
        width_cm = height_cm * ratio
    run = p2.add_run()
    run.add_picture(io.BytesIO(png_data), width=Cm(width_cm), height=Cm(height_cm))
    return table


def build_gorus_docx(opinion: dict[str, Any], figure_images: dict[str, bytes] | None = None) -> bytes:
    """Build opinion by cloning binding 696809 paragraph/table archetypes.

    The body is rebuilt only from template archetypes so paragraph spacing, figure
    table geometry, header/footer and page setup remain bound to the template.
    """
    template = Document(str(GORUS_TEMPLATE))
    doc = Document(str(GORUS_TEMPLATE))
    title_1 = deepcopy(template.paragraphs[0]._p)
    title_2 = deepcopy(template.paragraphs[1]._p)
    metadata_table = deepcopy(template.tables[0]._tbl)
    blank_meta = deepcopy(template.paragraphs[2]._p)
    salutation = deepcopy(template.paragraphs[3]._p)
    clear_body(doc)
    body = doc._element.body
    body.insert(-1, title_1); body.insert(-1, title_2); body.insert(-1, metadata_table); body.insert(-1, blank_meta); body.insert(-1, salutation)

    table = doc.tables[0]
    values = [opinion.get("application_no", ""), opinion.get("applicant", ""), opinion.get("reference", "")]
    for row, value in zip(table.rows, values):
        set_cell_text(row.cells[2], str(value), bold=False, size=11)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True; run.font.name = "Arial"; run.font.size = Pt(11)
        for run in row.cells[1].paragraphs[0].runs:
            run.bold = False; run.font.name = "Arial"; run.font.size = Pt(11)

    # Template p4 intro, p5 blank, p6/7 cited-document rows, p8 blank.
    _clone_paragraph_with_text(doc, template.paragraphs[4], opinion.get("intro", ""), bold=False)
    _clone_blank(doc, template.paragraphs[5])
    docs = opinion.get("cited_documents") or []
    for i, d in enumerate(docs):
        cat = str(d.get("category", "")).strip()
        suffix = f" ({cat})" if cat else ""
        txt = f"{d.get('label','')}: {d.get('number','')}"
        title = str(d.get("title", "")).strip()
        if title:
            txt += f' - “{title}”'
        txt += suffix
        archetype = template.paragraphs[6 if i == 0 else 7]
        _clone_paragraph_with_text(doc, archetype, txt, bold=False)
    _clone_blank(doc, template.paragraphs[8])

    figure_images = figure_images or {}
    sections = opinion.get("sections") or []
    for si, section in enumerate(sections):
        heading_tpl = template.paragraphs[9 if si == 0 else 24]
        _clone_paragraph_with_text(doc, heading_tpl, section.get("heading", ""), bold=True)
        blocks = section.get("blocks") or []
        first_para_inserted = False
        figure_inserted = False
        for block in blocks:
            typ = str(block.get("type", "paragraph")).lower()
            if typ == "quote":
                if bool(block.get("attach_to_previous", True)) and doc.paragraphs and doc.paragraphs[-1].text.strip():
                    _append_quote_with_location(doc.paragraphs[-1], block)
                else:
                    _add_quote_with_location(doc, template.paragraphs[10], block)
            else:
                _clone_paragraph_with_text(doc, template.paragraphs[10], block.get("text", ""), bold=False)
                if not first_para_inserted:
                    first_para_inserted = True
                    label = str(section.get("label", "")).upper()
                    if label in figure_images:
                        _clone_blank(doc, template.paragraphs[11])
                        _clone_blank(doc, template.paragraphs[12])
                        caption = str(section.get("figure_caption", "")).strip() or f"{label} dokümanı - Şekil"
                        _add_original_figure_table(doc, template, caption, figure_images[label])
                        figure_inserted = True
        if not figure_inserted:
            label = str(section.get("label", "")).upper()
            if label in figure_images:
                _clone_blank(doc, template.paragraphs[11]); _clone_blank(doc, template.paragraphs[12])
                caption = str(section.get("figure_caption", "")).strip() or f"{label} dokümanı - Şekil"
                _add_original_figure_table(doc, template, caption, figure_images[label])
        _clone_blank(doc, template.paragraphs[20])
        inv_heading = str(section.get("inventive_step_heading", "")).strip()
        inv_paras = list(section.get("inventive_step_paragraphs") or [])
        if inv_heading:
            _clone_paragraph_with_text(doc, template.paragraphs[21], inv_heading, bold=True)
            for par in inv_paras:
                _clone_paragraph_with_text(doc, template.paragraphs[22], par, bold=False)
        _clone_blank(doc, template.paragraphs[23])

    combined = opinion.get("combined_assessment") or {}
    if combined:
        _clone_paragraph_with_text(doc, template.paragraphs[33], combined.get("heading", "Dokümanların birlikte değerlendirilmesi"), bold=True)
        for par in combined.get("paragraphs") or []:
            _clone_paragraph_with_text(doc, template.paragraphs[34], par, bold=False)
        _clone_blank(doc, template.paragraphs[35])

    for par in opinion.get("conclusion") or []:
        _clone_paragraph_with_text(doc, template.paragraphs[36], par, bold=False)
    _clone_blank(doc, template.paragraphs[37])
    # Exact signoff archetypes from template.
    lines = str(opinion.get("signoff", "Saygılarımızla,\nDESTEK PATENT A.Ş.")).splitlines()
    _clone_paragraph_with_text(doc, template.paragraphs[38], lines[0] if lines else "Saygılarımızla,", bold=False)
    _clone_paragraph_with_text(doc, template.paragraphs[39], lines[1] if len(lines) > 1 else "DESTEK PATENT A.Ş.", bold=False)

    out = io.BytesIO(); doc.save(out)
    return out.getvalue()



def _find_relaxed_phrase_span(text: str, phrase: str) -> tuple[int, int] | None:
    tokens = re.findall(r"\S+", str(phrase or "").strip())
    if not tokens:
        return None
    pattern = r"\s+".join(re.escape(token) for token in tokens)
    match = re.search(pattern, text)
    return (match.start(), match.end()) if match else None


def _claim_paragraph_map(doc: Document) -> dict[int, str | None]:
    """Her paragrafı İSTEMLER bölümü içinde en son görülen istem numarasıyla eşleştirir."""
    mapping: dict[int, str | None] = {}
    in_claims = False
    current_claim: str | None = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        upper = text.upper()
        if upper == "İSTEMLER":
            in_claims = True
            current_claim = None
            mapping[idx] = None
            continue
        if in_claims and upper == "ÖZET":
            in_claims = False
            current_claim = None
        if in_claims:
            match = re.match(r"^\s*(\d+)\s*[.)-]\s+", text)
            if match:
                current_claim = match.group(1)
            mapping[idx] = current_claim
        else:
            mapping[idx] = None
    return mapping


def _run_style_spans(paragraph) -> list[tuple[int, int, Any]]:
    spans: list[tuple[int, int, Any]] = []
    cursor = 0
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        rpr = run._r.find(qn("w:rPr"))
        spans.append((cursor, cursor + len(text), deepcopy(rpr) if rpr is not None else None))
        cursor += len(text)
    return spans


def _style_at(spans: list[tuple[int, int, Any]], position: int) -> Any:
    for start, end, rpr in spans:
        if start <= position < end:
            return deepcopy(rpr) if rpr is not None else None
    if spans:
        rpr = spans[-1][2]
        return deepcopy(rpr) if rpr is not None else None
    return None


def _append_unchanged_with_styles(parent, original: str, start: int, end: int, spans: list[tuple[int, int, Any]]) -> None:
    if end <= start:
        return
    cursor = start
    for span_start, span_end, rpr in spans:
        overlap_start = max(cursor, span_start)
        overlap_end = min(end, span_end)
        if overlap_end <= overlap_start:
            continue
        if overlap_start > cursor:
            _append_plain_run(parent, original[cursor:overlap_start], None)
        _append_plain_run(parent, original[overlap_start:overlap_end], rpr)
        cursor = overlap_end
        if cursor >= end:
            break
    if cursor < end:
        _append_plain_run(parent, original[cursor:end], _style_at(spans, cursor))


def _append_plain_run(parent, text: str, rpr: Any = None, *, deleted: bool = False) -> None:
    if not text:
        return
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(deepcopy(rpr))
    text_node = OxmlElement("w:delText" if deleted else "w:t")
    text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    parent.append(run)


def _append_revision(parent, text: str, *, kind: str, change_id: int, rpr: Any = None) -> None:
    if not text:
        return
    wrapper = OxmlElement("w:del" if kind == "delete" else "w:ins")
    wrapper.set(qn("w:id"), str(change_id))
    wrapper.set(qn("w:author"), "Patent Atölyesi")
    wrapper.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    _append_plain_run(wrapper, text, rpr, deleted=(kind == "delete"))
    parent.append(wrapper)


def _minimal_markup_parts(old_text: str, new_text: str) -> tuple[str, str, str, str]:
    """Return unchanged_prefix, deleted_mid, inserted_mid, unchanged_suffix.

    Common text is preserved positionally, not realigned from the middle. This produces
    patent-friendly minimal redlines: article changes mark only the article, noun
    replacements mark only the noun phrase, and a one-letter typo can become insertion-only.
    """
    tok_re = re.compile(r"\s+|[^\s]+")
    old_tokens = tok_re.findall(old_text)
    new_tokens = tok_re.findall(new_text)
    i = 0
    while i < min(len(old_tokens), len(new_tokens)) and old_tokens[i] == new_tokens[i]:
        i += 1
    j = 0
    while (
        j < len(old_tokens) - i
        and j < len(new_tokens) - i
        and old_tokens[len(old_tokens)-1-j] == new_tokens[len(new_tokens)-1-j]
    ):
        j += 1
    prefix = "".join(old_tokens[:i])
    old_mid = "".join(old_tokens[i:len(old_tokens)-j if j else len(old_tokens)])
    new_mid = "".join(new_tokens[i:len(new_tokens)-j if j else len(new_tokens)])
    suffix = "".join(old_tokens[len(old_tokens)-j:]) if j else ""

    # If the only changed unit is one lexical token, minimize further at character level.
    if old_mid and new_mid and not re.search(r"\s", old_mid) and not re.search(r"\s", new_mid):
        cp = 0
        while cp < min(len(old_mid), len(new_mid)) and old_mid[cp] == new_mid[cp]:
            cp += 1
        cs = 0
        while (
            cs < len(old_mid)-cp
            and cs < len(new_mid)-cp
            and old_mid[len(old_mid)-1-cs] == new_mid[len(new_mid)-1-cs]
        ):
            cs += 1
        prefix += old_mid[:cp]
        suffix = (old_mid[len(old_mid)-cs:] if cs else "") + suffix
        old_mid = old_mid[cp:len(old_mid)-cs if cs else len(old_mid)]
        new_mid = new_mid[cp:len(new_mid)-cs if cs else len(new_mid)]
    elif not old_mid and new_mid:
        pass
    elif old_mid and not new_mid:
        pass
    return prefix, old_mid, new_mid, suffix


def _rewrite_paragraph_with_changes(paragraph, operations: list[dict[str, Any]], *, track_changes: bool, id_start: int) -> int:
    original = paragraph.text
    if not original:
        raise ValueError("Revizyon uygulanacak paragraf boş.")

    spans: list[tuple[int, int, dict[str, Any]]] = []
    for op in operations:
        old_text = str(op.get("old_text", "")).strip()
        span = _find_relaxed_phrase_span(original, old_text)
        if span is None:
            raise ValueError(f"Markup için eski ifade aynı paragrafta bulunamadı: {old_text[:120]}...")
        spans.append((span[0], span[1], op))
    spans.sort(key=lambda x: x[0])
    for left, right in zip(spans, spans[1:]):
        if right[0] < left[1]:
            raise ValueError("Aynı istem paragrafında birbiriyle çakışan iki revizyon önerisi bulundu.")

    style_spans = _run_style_spans(paragraph)
    p_el = paragraph._p
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)

    cursor = 0
    change_id = id_start
    for start, end, op in spans:
        old_actual = original[start:end]
        new_text = str(op.get("new_text", ""))
        _append_unchanged_with_styles(p_el, original, cursor, start, style_spans)
        prefix, deleted_mid, inserted_mid, suffix = _minimal_markup_parts(old_actual, new_text)
        # Preserve unchanged prefix/suffix outside Track Changes. Only the minimum changed
        # character/word span is wrapped in w:del/w:ins.
        if prefix:
            _append_unchanged_with_styles(p_el, original, start, start + len(prefix), style_spans)
        change_pos = start + len(prefix)
        change_style = _style_at(style_spans, min(change_pos, max(start, end - 1)))
        if track_changes:
            if deleted_mid:
                _append_revision(p_el, deleted_mid, kind="delete", change_id=change_id, rpr=change_style)
                change_id += 1
            if inserted_mid:
                _append_revision(p_el, inserted_mid, kind="insert", change_id=change_id, rpr=change_style)
                change_id += 1
        else:
            if inserted_mid:
                _append_plain_run(p_el, inserted_mid, change_style)
        if suffix:
            _append_unchanged_with_styles(p_el, original, end - len(suffix), end, style_spans)
        cursor = end
    _append_unchanged_with_styles(p_el, original, cursor, len(original), style_spans)
    return change_id


def _enable_track_revisions(doc: Document) -> None:
    settings = doc.settings._element
    if settings.find(qn("w:trackRevisions")) is None:
        track = OxmlElement("w:trackRevisions")
        settings.insert(0, track)


def _ep_prior_art_paragraph_text(upd: dict[str, Any]) -> str:
    kind=str(upd.get("source_kind", "application")).strip().lower()
    objective=re.sub(r"\s+", " ", str(upd.get("objective_summary", ""))).strip().rstrip(".")
    difference=re.sub(r"\s+", " ", str(upd.get("however_difference", ""))).strip().rstrip(".")
    if kind == "document":
        title=str(upd.get("document_title", "")).strip()
        if not title:
            raise ValueError("EP document prior-art update requires document_title.")
        lead=f'As a result of the research on the subject, the document entitled "{title}" has been found. The document is related to {objective}.'
    else:
        pub=str(upd.get("publication_number", "")).strip()
        if not pub:
            raise ValueError("EP application prior-art update requires publication_number.")
        lead=f'As a result of the research on the subject, application numbered {pub} has been found. The application is related to {objective}.'
    return f"{lead} However, {difference}."


def _insert_ep_prior_art_updates(doc: Document, updates: list[dict[str, Any]], *, track_changes: bool, id_start: int) -> int:
    if not updates:
        return id_start
    anchor=None
    for p in doc.paragraphs:
        if p.text.strip().casefold().startswith("as a result of the research on the subject"):
            anchor=p
            break
    if anchor is None:
        raise ValueError("EP önceki teknik ekleme noktası bulunamadı.")
    parent=anchor._p.getparent(); insert_index=parent.index(anchor._p)+1
    base_rpr=None
    for r in anchor._p.findall(qn("w:r")):
        if "".join(t.text or "" for t in r.findall(qn("w:t"))).strip():
            rp=r.find(qn("w:rPr")); base_rpr=deepcopy(rp) if rp is not None else None; break
    change_id=id_start
    for upd in updates:
        text=_ep_prior_art_paragraph_text(upd)
        p_el=OxmlElement("w:p")
        ppr=anchor._p.find(qn("w:pPr"))
        if ppr is not None:
            p_el.append(deepcopy(ppr))
        if track_changes:
            _append_revision(p_el, text, kind="insert", change_id=change_id, rpr=base_rpr)
            change_id += 1
        else:
            _append_plain_run(p_el, text, base_rpr)
        parent.insert(insert_index,p_el); insert_index += 1
    return change_id


def build_claim_revision_docx(source_docx: bytes, amendments: list[dict[str, Any]], *, track_changes: bool, description_updates: list[dict[str, Any]] | None = None) -> bytes:
    description_updates = description_updates or []
    if not amendments and not description_updates:
        raise ValueError("Uygulanacak istem veya açıklama revizyonu bulunamadı.")
    doc = Document(io.BytesIO(source_docx))
    claim_map = _claim_paragraph_map(doc)

    assignments: dict[int, list[dict[str, Any]]] = {}
    for op in amendments:
        claim_no = str(op.get("claim_number", "")).strip()
        old_text = str(op.get("old_text", "")).strip()
        candidates: list[int] = []
        for idx, paragraph in enumerate(doc.paragraphs):
            if claim_map.get(idx) != claim_no:
                continue
            if _find_relaxed_phrase_span(paragraph.text, old_text) is not None:
                candidates.append(idx)
        if len(candidates) != 1:
            if not candidates:
                raise ValueError(
                    f"İstem {claim_no} için Markup uygulanacak ifade ilgili istem paragrafında bulunamadı: {old_text[:120]}..."
                )
            raise ValueError(
                f"İstem {claim_no} için aynı eski ifade birden fazla paragrafta bulundu. Revizyon ifadesini daha özgül hale getirin."
            )
        assignments.setdefault(candidates[0], []).append(op)

    if track_changes:
        _enable_track_revisions(doc)
    change_id = 1
    for paragraph_index in sorted(assignments):
        change_id = _rewrite_paragraph_with_changes(
            doc.paragraphs[paragraph_index],
            assignments[paragraph_index],
            track_changes=track_changes,
            id_start=change_id,
        )

    change_id = _insert_ep_prior_art_updates(doc, description_updates, track_changes=track_changes, id_start=change_id)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def build_claim_revision_pair(source_docx: bytes, amendments: list[dict[str, Any]], description_updates: list[dict[str, Any]] | None = None) -> tuple[bytes, bytes]:
    markup = build_claim_revision_docx(source_docx, amendments, track_changes=True, description_updates=description_updates)
    clean = build_claim_revision_docx(source_docx, amendments, track_changes=False, description_updates=description_updates)
    return markup, clean


# -----------------------------------------------------------------------------
# TİP 3 ÖN ARAŞTIRMA MODÜLÜ
# -----------------------------------------------------------------------------
def top10_research_prompt(bbf_text: str, cutoff_date: str) -> str:
    return f"""{ARASTIRMA_RULES}
Aşağıdaki BBF için araştırma kesim tarihi {cutoff_date} olacak şekilde global patent araştırması yap ve en benzer tam 10 patent dokümanını belirle.
Google Patents, Espacenet, PATENTSCOPE, TÜRKPATENT ve ulaşılabilir resmi/yarı resmi patent kaynaklarını kapsayacak geniş web araştırması yap.
Dokümanları teknik yakınlığa göre sırala. Numara, başlık, tarih ve kaynak URL doğrulanmış olsun. İlk araştırma çıktısında tam 10 dokümanla birlikte `totalpatent_query`, `proposed_d1` ve `proposed_d2` alanlarını mutlaka doldur. Her doküman için yayımlanmış özgün İngilizce Abstract metni doğrulanabiliyorsa `abstract_en` alanına doğrudan aktar; Türkçeye çevirme veya yeniden özetleme. Kullanıcı dokümanlarını bu aşamada sorma veya varsayma. JSON dışında yazma.
ŞEMA:
{{
 "subject_title":"",
 "technical_problem":"",
 "technical_effects":[""],
 "technical_features":[""],
 "method_steps":[{{"number":"1001","text":""}}],
 "keywords":[""],
 "ipc_cpc":[""],
 "documents":[{{
   "rank":1,"publication_number":"","application_number":"","title":"","date":"","jurisdiction":"","source_url":"",
   "summary":"","abstract_en":"","matching_features":[""],"missing_features":[""],"novelty_destroying":false,"novelty_reason":"","relevance_score":0
 }}],
 "totalpatent_query":"TotalPatent arama sorgusu: CN... or US...",
 "proposed_d1":"publication_number",
 "proposed_d2":"publication_number veya boş",
 "preliminary_novelty":"sağlanır/sağlanmaz",
 "preliminary_inventive_step":"sağlanır/sağlanmaz/belirsiz"
}}
BBF:\n{bbf_text}"""


# -----------------------------------------------------------------------------
# TİP 3 - ŞABLONA SIKI SADAKAT / ARAŞTIRMA GÜNCELLEME (v5.4)
# -----------------------------------------------------------------------------
def final_selection_prompt(
    bbf_text: str,
    top10: dict[str, Any],
    user_docs_text: str,
) -> str:
    return f"""{ARASTIRMA_RULES}
Aşağıdaki araştırma konusu, sistemin ilk global araştırmada bulduğu 10 doküman ve kullanıcının varsa yüklediği dokümanları birlikte incele.
İlk 10 dokümanı DEĞİŞTİRME veya yeniden sıralanmış yeni bir top-10 listesi üretme. Kullanıcı dokümanları arasından yalnız araştırma konusuna en ilgili birkaç belgeyi `10+` ek doküman olarak seç.
Nihai D1 ve gerekiyorsa D2'yi belirle. Kullanıcı dokümanı daha yakınsa D1/D2 değişebilir.
Tek doküman bütün esas teknik özellikleri doğrudan ve açık açıklıyorsa D1 ile yenilik sağlanmaz sonucuna git ve D2 seçme. Aksi halde en yakın D1 ve tamamlayıcı D2 ile buluş basamağını değerlendir.
ÖNEMLİ: Bu aşamada kullanıcı sonuç modu henüz alınmamıştır. Sonucu kullanıcı tercihine göre zorlama. Kendi teknik kanaatini bağımsız olarak ver ve `technical_opinion` alanını mutlaka `Bence buluş basamağı var.` veya `Bence buluş basamağı yok.` cümlesiyle başlat.
- `additional_relevant_user_documents` yalnız kullanıcı tarafından yüklenen en ilgili birkaç belgeyi içersin.
- `additional_query` alanı belge varsa tam olarak `10+ CN... or IN...` biçiminde olsun; ilgili kullanıcı belgesi yoksa boş bırak.
- İlk önerilen D1/D2 ile nihai D1/D2'yi karşılaştırarak `d1_changed` ve `d2_changed` alanlarını doğru doldur; değişiklik notunda eski ve yeni numarayı açık yaz.
- D1 ve D2 karşılaştırma satırları aynı özellik listesine ve aynı sıraya sahip olmalıdır.
- Sağ hücre metni çıplak + veya - olmasın. `+ Özet; İstem 1; Şekil 3 ...` veya `- Dokümanda ... açıklanmamaktadır.` mantığında somut dayanak yaz.
- Patent şekli için model tarafından üretilmiş veya temsili bir görsel verme. Kaynakta bulunan özgün şekle ait doğrudan URL bulunabiliyorsa `figure_image_url` alanına yaz; aksi halde boş bırak.
- Nihai D1/D2 için `abstract_en` alanına yalnız ilgili patentin yayımlanmış özgün İngilizce Abstract metnini koy. Kullanıcı yüklemesinde özgün İngilizce Abstract varsa onu öncelikle aynen kullan; yoksa doğrulanmış resmi/Google Patents kaynağından getir. Türkçe özet, çeviri veya model tarafından yeniden yazılmış abstract kullanma. Kaynak türünü `abstract_source` alanına yaz.
JSON dışında yazma.
ŞEMA:
{{
 "additional_relevant_user_documents":[{{"number":"","title":"","reason":""}}],
 "additional_query":"10+ CN... or IN...",
 "d1":{{"number":"","alternate_number":"","title":"","date":"","source":"system/user","source_url":"","summary":"","abstract_en":"","abstract_source":"user-file/official/espacenet/google-patents","figure_reference":"","figure_image_url":""}},
 "d2":null,
 "d1_changed":false,
 "d1_change_note":"D1 değişmedi: ... / D1 değişti: eski ...; yeni ...",
 "d2_changed":false,
 "d2_change_note":"D2 değişmedi: ... / D2 değişti: eski ...; yeni ...",
 "novelty_result":"sağlanır/sağlanmaz",
 "inventive_step_result":"sağlanır/sağlanmaz",
 "technical_opinion":"Bence buluş basamağı ... . Kısa teknik gerekçe...",
 "novelty_reasoning":[""],
 "inventive_step_reasoning":[""],
 "feature_list":[""],
 "comparison_rows_d1":[{{"feature":"","status_evidence":"+ Özet; İstem ...; Şekil ..."}}],
 "comparison_rows_d2":[{{"feature":"","status_evidence":"- ..."}}],
 "helper_documents":[{{"number":"","title":"","source_url":"","role":""}}],
 "warnings":[""]
}}
ARAŞTIRMA KONUSU:\n{bbf_text}\n
İLK TOP10 (AYNEN KORU):\n{json.dumps(top10, ensure_ascii=False, indent=2)}\n
KULLANICI DOKÜMANLARI:\n{user_docs_text}"""

def validate_research_selection(selection: dict[str, Any]) -> None:
    d1_rows = selection.get("comparison_rows_d1") or []
    d2_rows = selection.get("comparison_rows_d2") or []
    for label, rows in (("D1", d1_rows), ("D2", d2_rows)):
        for row in rows:
            evidence = str(row.get("status_evidence") or row.get("status") or "").strip()
            if not (evidence.startswith("+") or evidence.startswith("-")):
                raise ValueError(f"{label} karşılaştırma hücresi + veya - ile başlamalı ve dayanak içermelidir: {evidence!r}")
            if evidence.startswith("+") and len(evidence) < 5:
                raise ValueError(f"{label} karşılaştırma hücresinde '+' yanında dokümandaki somut yer/dayanak da belirtilmelidir.")
            row["status_evidence"] = evidence
    if selection.get("d2"):
        d1_features = [re.sub(r"\s+", " ", str(row.get("feature", ""))).strip() for row in d1_rows]
        d2_features = [re.sub(r"\s+", " ", str(row.get("feature", ""))).strip() for row in d2_rows]
        if d1_features != d2_features:
            raise ValueError("D1 ve D2 karşılaştırma tablolarındaki teknik özellik listeleri birebir aynı değildir.")
    for label, key in (("D1", "d1"), ("D2", "d2")):
        block = selection.get(key)
        if not block:
            continue
        abstract_en = str(block.get("abstract_en") or "").strip()
        if not abstract_en:
            raise ValueError(f"{label} için doğrulanmış özgün İngilizce Abstract bulunamadı.")
        if _contains_turkish_specific_chars(abstract_en):
            raise ValueError(f"{label} abstract_en alanı özgün İngilizce Abstract olmalıdır; Türkçe/çeviri metin kabul edilmez.")


def report_drafting_prompt(
    bbf_text: str,
    top10: dict[str, Any],
    selection: dict[str, Any],
    reference: str,
    cutoff_date: str,
    decision_mode: str,
) -> str:
    return f"""{ARASTIRMA_RULES}
Aşağıdaki verilere göre bağlayıcı `On_Arastirma_Raporu_181612_template.docx` içeriğine tam uyumlu Ön Araştırma Raporu metni oluştur.
DP referans numarası: {reference}
Araştırma kesim tarihi: {cutoff_date}
Kullanıcı sonuç modu: {decision_mode}
JSON dışında yazma.
ŞEMA:
{{
 "reference":"{reference}",
 "title":"",
 "report_date":"{date.today().strftime('%d.%m.%Y')}",
 "purpose":"Belirlenen konuda araştırmanın gerçekleştirilmesi",
 "scope":"Global (İlan edilmiş olan patent başvuruları)",
 "keywords":["English keyword or phrase"],
 "ipc_cpc":[{{"code":"","description":"ENGLISH IPC/CPC DESCRIPTION"}}],
 "evaluation_intro":"",
 "documents":[{{
   "label":"D1","number":"","alternate_number":"","title":"","date":"","source_url":"","figure_reference":"","figure_image_url":"",
   "description":["2-3 cümle"],"abstract":"ORIGINAL ENGLISH ABSTRACT VERBATIM","figure_caption":"D1- Şekil",
   "comparison_rows":[{{"feature":"","status_evidence":"+ Özet; İstem ...; Şekil ..."}}],
   "novelty_assessment":["5-10 satırlık değerlendirme"]
 }}],
 "inventive_step_paragraphs":[""],
 "conclusion_paragraphs":[""],
 "warnings":["uyarı paragrafı 1","uyarı paragrafı 2","uyarı paragrafı 3","uyarı paragrafı 4"],
 "attachments":["Benzer Dokümanlar","Ön İnceleme Raporu","Makine Tercümeleri"]
}}
ÖZEL:
- selection.d2 null ise yalnızca D1 bölümü oluştur.
- D1 ve D2 tablolarındaki feature alanları birebir aynı ve aynı sırada olmalıdır.
- comparison_rows sağ hücresi `status_evidence` alanıdır; + veya - ile başlar ve dokümandaki yeri açıkça yazar. Çıplak + / - kullanma.
- `documents` alanındaki D1/D2, NİHAİ SEÇİM içindeki D1/D2 ile aynı dokümanlar olmalı; yayın numarası, alternatif numara, başlık, tarih, source_url ve figure_image_url bilgilerini aynen taşı.
- `keywords` alanı yalnız İngilizce teknik anahtar kelimelerden oluşmalı; Türkçe kelime/ifade kullanma. Şablon 5x2 olduğundan en fazla 10 adet ver.
- `ipc_cpc[].description` yalnız İngilizce resmi sınıflandırma açıklaması olmalı; Türkçe açıklama yazma. Kodların kendisini değiştirme.
- `scope` alanını tam olarak `Global (İlan edilmiş olan patent başvuruları)` yaz; kesim tarihini bu sabit hücreye ekleme.
- D1/D2 `abstract` alanını ASLA yeniden yazma, Türkçeye çevirme veya özetleme. NİHAİ SEÇİM içindeki ilgili `abstract_en` değerini doğrudan kullan. `abstract_en` boşsa özgün İngilizce abstract olmadan raporu tamamlanmış gibi gösterme.
- `warnings` alanı şablondaki uyarı hücresinin dört ayrı paragraf yapısını koruyacak şekilde tam 4 paragraf olarak ver; tek paragrafta satır sonlarıyla birleştirme.
- Rapor metninde BBF/buluş bildirim formu ifadesi kullanma; `araştırma konusu` de.
- `→`, `=>` veya `özellik + özellik + özellik` gibi sembolik/yapay zekâ görünümlü anlatım kullanma.
- Yardımcı dokümanları yeni bir D3 başlığı açmadan yalnız buluş basamağı değerlendirmesinin doğal paragraf akışında kullan.
- Şablonda olmayan bölüm/başlık ekleme.
- Sonuçta yenilik ve buluş basamağı sonucunu açıkça yaz.
- Metni ikinci kez kontrol edip düzelt.
ARAŞTIRMA KONUSU:\n{bbf_text}\n
TOP10:\n{json.dumps(top10, ensure_ascii=False, indent=2)}\n
NİHAİ SEÇİM:\n{json.dumps(selection, ensure_ascii=False, indent=2)}"""


def research_update_analysis_prompt(first_text: str, revised_text: str, prior_report_text: str) -> str:
    return f"""{ARASTIRMA_GUNCELLEME_RULES}
İlk araştırma konusu, revize araştırma konusu ve ilk Ön Araştırma Raporunu birlikte analiz et.
Bu aşamada web araştırması yapma ve Word raporu yazma. İlk ve revize teknik içerik arasındaki gerçek teknik farkları çıkar; yalnız ifade/başlık değişikliklerini teknik katkı sayma.
İlk rapordaki D1/D2 ile yenilik/buluş basamağı gerekçelerini ayrıca çıkar ve her yeni farkın bu gerekçeler üzerindeki etkisini değerlendir.
JSON dışında yazma.
ŞEMA:
{{
 "first_title":"",
 "revised_title":"",
 "prior_d1":{{"number":"","alternate_number":"","title":"","date":""}},
 "prior_d2":{{"number":"","alternate_number":"","title":"","date":""}},
 "prior_novelty_result":"sağlanır/sağlanmaz/belirsiz",
 "prior_inventive_step_result":"sağlanır/sağlanmaz/belirsiz",
 "differences":[{{
   "old":"",
   "new":"",
   "technical_contribution":"evet/hayır/kısmen",
   "technical_effect":"",
   "effect_against_prior_d1_d2":""
 }}],
 "meaningful_change":true,
 "search_focus":[""],
 "preliminary_opinion":"",
 "reasons":[""]
}}
İLK ARAŞTIRMA KONUSU:\n{first_text}\n
REVİZE ARAŞTIRMA KONUSU:\n{revised_text}\n
İLK ÖN ARAŞTIRMA RAPORU:\n{prior_report_text}"""


def research_update_search_prompt(
    revised_text: str,
    prior_report_text: str,
    analysis: dict[str, Any],
    cutoff_date: str,
) -> str:
    return f"""{ARASTIRMA_GUNCELLEME_RULES}
Revize araştırma konusu için araştırma kesim tarihi {cutoff_date} olacak şekilde global patent araştırması yap.
İlk ön araştırma raporundaki D1/D2'yi başlangıç noktası olarak doğrula; özellikle revizyonda eklenen ayırt edici teknik özelliklere odaklan. Eski sorguyu aynen tekrarlama.
TR, EP, US, CN, KR, JP, GB, DE, WIPO ve ilgili diğer veri tabanlarında araştır. En yakın toplam 10 doğrulanmış patent dokümanını sırala. İlk D1/D2 bu 10 içinde olabilir.
Yeni bulunan ve ilk raporda bulunmayan dokümanları ayrıca işaretle. Nihai D1/D2'yi seç; yeni bir belge daha güçlü ise D1/D2'yi değiştir, değilse yardımcı doküman olarak tut.
D1/D2 karşılaştırma tablosunun sağ hücrelerinde yalnız + / - verme; somut yer/dayanak belirt.
Nihai D1/D2 için `abstract_en` alanına doğrulanmış yayımlanmış özgün İngilizce Abstract metnini doğrudan aktar; çeviri veya model özeti kullanma ve kaynağı `abstract_source` alanında belirt.
Özgün patent şekli için doğrulanabilir doğrudan `figure_image_url` bulunabiliyorsa yaz. Model tarafından şekil üretme.
JSON dışında yazma.
ŞEMA:
{{
 "documents":[{{"rank":1,"publication_number":"","alternate_number":"","title":"","date":"","jurisdiction":"","source_url":"","is_new_vs_prior_report":true,"summary":"","abstract_en":"","matching_revision_features":[""],"missing_revision_features":[""],"relevance_score":0}}],
 "new_documents":[{{"number":"","title":"","date":"","source_url":"","technical_relevance":""}}],
 "d1":{{"number":"","alternate_number":"","title":"","date":"","source_url":"","summary":"","abstract_en":"","abstract_source":"official/espacenet/google-patents","figure_reference":"","figure_image_url":""}},
 "d2":{{"number":"","alternate_number":"","title":"","date":"","source_url":"","summary":"","abstract_en":"","abstract_source":"official/espacenet/google-patents","figure_reference":"","figure_image_url":""}},
 "feature_list":[""],
 "comparison_rows_d1":[{{"feature":"","status_evidence":"+ Özet; İstem ...; Şekil ..."}}],
 "comparison_rows_d2":[{{"feature":"","status_evidence":"- ..."}}],
 "helper_documents":[{{"number":"","title":"","date":"","source_url":"","role":""}}],
 "novelty_result":"sağlanır/sağlanmaz",
 "inventive_step_result":"sağlanır/sağlanmaz",
 "technical_opinion":"Kanaatim: ...",
 "totalpatent_query":"TotalPatent arama sorgusu: ..."
}}
REVİZE ARAŞTIRMA KONUSU:\n{revised_text}\n
İLK ÖN ARAŞTIRMA RAPORU:\n{prior_report_text}\n
FARK ANALİZİ:\n{json.dumps(analysis, ensure_ascii=False, indent=2)}"""


def research_update_report_prompt(
    revised_text: str,
    prior_report_text: str,
    analysis: dict[str, Any],
    research: dict[str, Any],
    reference: str,
    cutoff_date: str,
    decision_mode: str,
) -> str:
    return f"""{ARASTIRMA_GUNCELLEME_RULES}
Aşağıdaki verilerden `On_Arastirma_Raporu_181612_template.docx` biçiminde nihai Ön Araştırma Raporu metni oluştur.
Bu bir güncelleme çalışması olsa da Word raporunda yeni bir format veya `Revizyon farkları` bölümü açma. Tam Tip 3 Ön Araştırma Raporu düzenini kullan.
DP referans numarası: {reference}
Araştırma kesim tarihi: {cutoff_date}
Kullanıcının buluş basamağı sonuç seçimi: {decision_mode}
Araştırmanın teknik yenilik sonucu: {research.get('novelty_result','')}

Kullanıcı 'Buluş basamağı sağlanıyor' seçtiyse inventive step sonucunu `sağlanır`, 'Buluş basamağı sağlanmıyor' seçtiyse `sağlanmaz` yaz. Ancak yenilik sonucu kaynaklarla çelişecek şekilde değiştirilmez.
İlk rapora atıf gerekiyorsa `ilk ön araştırma raporu`, teknik konuya atıf gerekiyorsa `revize araştırma konusu` de. `BBF` ifadesi kullanma.
Yeni bulunan yardımcı dokümanları şablonda olmayan D3/D4 başlığı açmadan buluş basamağı değerlendirmesinin doğal paragrafı içinde açıkla. Daha güçlü yeni doküman D1/D2 seçilmişse normal D1/D2 bölümünde kullan.
`documents` alanındaki D1 ve D2, YENİ ARAŞTIRMA içindeki `d1` ve `d2` ile aynı dokümanlar olmalı; yayın numarası, alternatif numara, başlık, tarih, source_url, figure_reference ve figure_image_url bilgilerini değiştirmeden taşı.
Karşılaştırma hücrelerinde + veya - işaretinin ardından dokümandaki somut yeri yaz.
`keywords` yalnız İngilizce olsun ve en fazla 10 adet ver; `ipc_cpc[].description` yalnız İngilizce resmi sınıflandırma açıklaması olsun.
`scope` tam olarak `Global (İlan edilmiş olan patent başvuruları)` olarak kalsın.
D1/D2 `abstract` alanına YENİ ARAŞTIRMA içindeki ilgili dokümanın `abstract_en` değerini doğrudan aktar; Türkçeye çevirme, özetleme veya yeniden yazma. Özgün İngilizce abstract yoksa raporu tamamlanmış gibi üretme.
`warnings` alanını şablondaki dört ayrı uyarı paragrafını koruyacak şekilde tam 4 paragraf olarak üret.
`→`, `=>`, oklar veya `özellik + özellik` gibi kısa sembolik anlatım kullanma.
JSON dışında yazma.
ŞEMA:
{{
 "reference":"{reference}",
 "title":"",
 "report_date":"{date.today().strftime('%d.%m.%Y')}",
 "purpose":"Belirlenen konuda araştırmanın gerçekleştirilmesi",
 "scope":"Global (İlan edilmiş olan patent başvuruları)",
 "keywords":["English keyword or phrase"],
 "ipc_cpc":[{{"code":"","description":"ENGLISH IPC/CPC DESCRIPTION"}}],
 "evaluation_intro":"",
 "documents":[{{
   "label":"D1","number":"","alternate_number":"","title":"","date":"","source_url":"","figure_reference":"","figure_image_url":"",
   "description":[""],"abstract":"ORIGINAL ENGLISH ABSTRACT VERBATIM","figure_caption":"D1- Şekil",
   "comparison_rows":[{{"feature":"","status_evidence":"+ Özet; İstem ...; Şekil ..."}}],
   "novelty_assessment":[""]
 }}],
 "inventive_step_paragraphs":[""],
 "conclusion_paragraphs":[""],
 "warnings":["uyarı paragrafı 1","uyarı paragrafı 2","uyarı paragrafı 3","uyarı paragrafı 4"],
 "attachments":["Benzer Dokümanlar","Ön İnceleme Raporu","Makine Tercümeleri"]
}}
REVİZE ARAŞTIRMA KONUSU:\n{revised_text}\n
İLK ÖN ARAŞTIRMA RAPORU:\n{prior_report_text}\n
FARK ANALİZİ:\n{json.dumps(analysis, ensure_ascii=False, indent=2)}\n
YENİ ARAŞTIRMA:\n{json.dumps(research, ensure_ascii=False, indent=2)}"""


def _iter_strings(value: Any):
    if isinstance(value, str):
        yield value
    elif isinstance(value, dict):
        for v in value.values():
            yield from _iter_strings(v)
    elif isinstance(value, (list, tuple)):
        for v in value:
            yield from _iter_strings(v)


def _contains_turkish_specific_chars(value: str) -> bool:
    return bool(re.search(r"[çğıöşüÇĞİÖŞÜ]", str(value or "")))


def _norm_ws(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def validate_research_report_language(report: dict[str, Any]) -> None:
    banned = [r"\bBBF\b", r"buluş bildirim formu", r"→", r"=>"]
    # Tablo durum/evidence hücrelerindeki + / - izinlidir; diğer alanlarda ok ve BBF dili yasaktır.
    for text in _iter_strings(report):
        for pattern in banned:
            if re.search(pattern, text, flags=re.IGNORECASE):
                raise ValueError(f"Ön araştırma raporu metninde kullanılmaması gereken ifade bulundu: {pattern}")

    keywords = [str(x).strip() for x in (report.get("keywords") or []) if str(x).strip()]
    if len(keywords) > 10:
        raise ValueError("Tip 3 anahtar kelime alanı şablondaki 5x2 yapı gereği en fazla 10 İngilizce ifade içerebilir.")
    if any(_contains_turkish_specific_chars(x) for x in keywords):
        raise ValueError("Tip 3 raporundaki anahtar kelimelerin tamamı İngilizce olmalıdır.")

    ipc_items = report.get("ipc_cpc") or []
    if len(ipc_items) > 4:
        raise ValueError("IPC/CPC alanı bağlayıcı şablonun dört satırlık yapısını aşmamalıdır.")
    for item in ipc_items:
        code = str(item.get("code") or "").strip()
        desc = str(item.get("description") or "").strip()
        if not code or not desc:
            raise ValueError("Her IPC/CPC satırında hem kod hem İngilizce açıklama bulunmalıdır.")
        if _contains_turkish_specific_chars(desc):
            raise ValueError("IPC/CPC açıklamalarının tamamı İngilizce olmalıdır.")

    scope = str(report.get("scope") or "").strip()
    if scope and scope != "Global (İlan edilmiş olan patent başvuruları)":
        raise ValueError("Tip 3 Kapsam hücresi bağlayıcı şablondaki sabit metinle aynı olmalıdır.")

    for block in report.get("documents") or []:
        abstract = str(block.get("abstract") or "").strip()
        if not abstract:
            raise ValueError(f"{block.get('label','D1/D2')} için özgün İngilizce Abstract bulunmadan rapor oluşturulamaz.")
        if _contains_turkish_specific_chars(abstract):
            raise ValueError(f"{block.get('label','D1/D2')} Abstract alanı özgün İngilizce metin olmalıdır; Türkçe/çeviri metin kabul edilmez.")


def validate_report_against_selection(report: dict[str, Any], selection: dict[str, Any]) -> None:
    """Nihai D1/D2 kimliği ve özgün İngilizce abstract metni rapora model tarafından değiştirilmeden taşınır."""
    selected_docs = [selection.get("d1"), selection.get("d2")]
    selected_docs = [d for d in selected_docs if d]
    report_docs = report.get("documents") or []
    if len(report_docs) != len(selected_docs):
        raise ValueError("Rapor D1/D2 sayısı nihai seçimle aynı değildir.")
    for idx, (src, dst) in enumerate(zip(selected_docs, report_docs), 1):
        src_num = str(src.get("number") or "").strip()
        dst_num = str(dst.get("number") or "").strip()
        if _norm_ws(src_num).lower() != _norm_ws(dst_num).lower():
            raise ValueError(f"Rapor D{idx} dokümanı nihai seçimdeki D{idx} ile aynı değildir.")
        abstract_en = str(src.get("abstract_en") or "").strip()
        if not abstract_en:
            raise ValueError(f"Nihai D{idx} için doğrulanmış özgün İngilizce Abstract bulunamadı.")
        # Modelin yeniden yazmasını engelle: rapor alanını doğrudan kaynaktaki özgün abstract ile üzerine yaz.
        dst["abstract"] = abstract_en


def _replace_paragraph_text_preserve_format(paragraph, text: str) -> None:
    proto = next((r for r in paragraph.runs if r.text or r._r.rPr is not None), None)
    proto_rpr = deepcopy(proto._r.rPr) if proto is not None and proto._r.rPr is not None else None
    paragraph.clear()
    run = paragraph.add_run(str(text or ""))
    if proto_rpr is not None:
        current = run._r.rPr
        if current is not None:
            run._r.remove(current)
        run._r.insert(0, proto_rpr)


def _replace_cell_text_preserve_format(cell, text: str) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    _replace_paragraph_text_preserve_format(p, text)
    # Hücrede önceki metinden kalan fazladan paragrafları temizle; hücre ölçüsü/kenarlıkları korunur.
    for extra in list(cell.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)


def _replace_cell_lines_preserve_format(cell, lines: list[str]) -> None:
    vals = [str(x).strip() for x in lines if str(x).strip()]
    protos = list(cell.paragraphs)
    # keep at least one paragraph
    while len(cell.paragraphs) > 1:
        extra = cell.paragraphs[-1]
        extra._element.getparent().remove(extra._element)
    if not vals:
        _replace_paragraph_text_preserve_format(cell.paragraphs[0], "")
        return
    _replace_paragraph_text_preserve_format(cell.paragraphs[0], vals[0])
    for i, value in enumerate(vals[1:], 1):
        new_p = cell.add_paragraph()
        if i < len(protos):
            src = protos[i]
            # paragraph properties
            if src._p.pPr is not None:
                new_p._p.insert(0, deepcopy(src._p.pPr))
            proto_run = next((r for r in src.runs if r.text or r._r.rPr is not None), None)
            run = new_p.add_run(value)
            if proto_run is not None and proto_run._r.rPr is not None:
                if run._r.rPr is not None:
                    run._r.remove(run._r.rPr)
                run._r.insert(0, deepcopy(proto_run._r.rPr))
        else:
            _replace_paragraph_text_preserve_format(new_p, value)


def _fill_keyword_table(cell, keywords: list[str]) -> None:
    if not cell.tables:
        return
    table = cell.tables[0]
    vals = [str(x).strip() for x in keywords if str(x).strip()][:10]
    vals += [""] * (10 - len(vals))
    k = 0
    for row in table.rows:
        for c in row.cells:
            _replace_cell_text_preserve_format(c, vals[k])
            k += 1


def _fill_ipc_cell(cell, ipc_items: list[dict[str, Any]]) -> None:
    """Şablondaki dört IPC paragrafını korur: kod kalın, İngilizce açıklama normal."""
    items = [x for x in (ipc_items or []) if str(x.get("code") or "").strip() and str(x.get("description") or "").strip()][:4]
    protos = list(cell.paragraphs)
    if not protos:
        protos = [cell.add_paragraph()]
    # Şablondaki dört paragraf geometrisini sabit tut.
    while len(cell.paragraphs) > 4:
        extra = cell.paragraphs[-1]
        extra._element.getparent().remove(extra._element)
    while len(cell.paragraphs) < 4:
        cell.add_paragraph()

    def fill_para(paragraph, item, proto):
        proto_ppr = deepcopy(proto._p.pPr) if proto is not None and proto._p.pPr is not None else None
        proto_run_rprs = [deepcopy(r._r.rPr) if r._r.rPr is not None else None for r in (list(proto.runs) if proto is not None else [])]
        if proto_ppr is not None:
            if paragraph._p.pPr is not None:
                paragraph._p.remove(paragraph._p.pPr)
            paragraph._p.insert(0, proto_ppr)
        paragraph.clear()
        if item is None:
            return
        code_run = paragraph.add_run(f"{str(item.get('code') or '').strip()}: ")
        desc_run = paragraph.add_run(str(item.get('description') or '').strip())
        if proto_run_rprs and proto_run_rprs[0] is not None:
            if code_run._r.rPr is not None:
                code_run._r.remove(code_run._r.rPr)
            code_run._r.insert(0, proto_run_rprs[0])
        code_run.bold = True
        if len(proto_run_rprs) > 1 and proto_run_rprs[1] is not None:
            if desc_run._r.rPr is not None:
                desc_run._r.remove(desc_run._r.rPr)
            desc_run._r.insert(0, proto_run_rprs[1])
        desc_run.bold = False

    for i in range(4):
        proto = protos[i] if i < len(protos) else protos[-1]
        fill_para(cell.paragraphs[i], items[i] if i < len(items) else None, proto)


def _fill_warning_cell(cell, warnings: list[str]) -> None:
    """Şablondaki dört ayrı uyarı paragrafını korur; yeni bölüm/paragraf geometrisi oluşturmaz."""
    vals = [str(x).strip() for x in (warnings or []) if str(x).strip()][:4]
    vals += [""] * (4 - len(vals))
    protos = list(cell.paragraphs)
    if not protos:
        protos = [cell.add_paragraph()]
    while len(cell.paragraphs) > 4:
        extra = cell.paragraphs[-1]
        extra._element.getparent().remove(extra._element)
    while len(cell.paragraphs) < 4:
        cell.add_paragraph()
    for i in range(4):
        p = cell.paragraphs[i]
        proto = protos[i] if i < len(protos) else protos[-1]
        proto_ppr = deepcopy(proto._p.pPr) if proto._p.pPr is not None else None
        if proto_ppr is not None:
            if p._p.pPr is not None:
                p._p.remove(p._p.pPr)
            p._p.insert(0, proto_ppr)
        _replace_paragraph_text_preserve_format(p, vals[i])

def _replace_comparison_table(table, rows: list[dict[str, Any]], label: str) -> None:
    # Şablon başlığını, satır yüksekliğini, hücre genişliklerini ve yazı biçimini koru.
    header = table.rows[0]
    _replace_cell_text_preserve_format(header.cells[0], "Araştırma konusu")
    _replace_cell_text_preserve_format(header.cells[1], f"{label} Dokümanı")
    prototype_tr = deepcopy(table.rows[1]._tr) if len(table.rows) > 1 else None
    for tr in list(table._tbl.tr_lst)[1:]:
        table._tbl.remove(tr)
    for item in rows:
        if prototype_tr is not None:
            table._tbl.append(deepcopy(prototype_tr))
            cells = table.rows[-1].cells
        else:
            cells = table.add_row().cells
        _replace_cell_text_preserve_format(cells[0], str(item.get("feature", "")))
        evidence = str(item.get("status_evidence") or item.get("status") or "").strip()
        _replace_cell_text_preserve_format(cells[1], evidence)


def _safe_remote_url(url: str) -> bool:
    url = str(url or "").strip().lower()
    return url.startswith("https://patents.google.com/") or url.startswith("https://patentimages.storage.googleapis.com/") or url.startswith("https://patentscope.wipo.int/") or url.startswith("https://worldwide.espacenet.com/")


def _fetch_remote_bytes(url: str, timeout: int = 30) -> bytes:
    if not _safe_remote_url(url):
        raise ValueError("Patent şekli için yalnız doğrulanabilir patent kaynağı URL'si kullanılabilir.")
    req = Request(url, headers={"User-Agent": "Mozilla/5.0 PatentAtolyesi/5.4"})
    with urlopen(req, timeout=timeout) as resp:
        data = resp.read(25 * 1024 * 1024 + 1)
    if len(data) > 25 * 1024 * 1024:
        raise ValueError("Patent şekli dosyası beklenenden büyük.")
    return data


def _google_patent_image_urls(source_url: str) -> list[str]:
    if not str(source_url).startswith("https://patents.google.com/"):
        return []
    try:
        html = _fetch_remote_bytes(source_url).decode("utf-8", errors="ignore")
    except Exception:
        return []
    urls = re.findall(r"https://patentimages\.storage\.googleapis\.com/[^\"'<>\s]+?\.(?:png|jpg|jpeg|webp)", html, flags=re.I)
    out: list[str] = []
    seen: set[str] = set()
    for url in urls:
        url = html_lib.unescape(url)
        if url in seen:
            continue
        seen.add(url)
        out.append(url)
    return out


def _normalize_image_bytes(data: bytes) -> bytes:
    with Image.open(io.BytesIO(data)) as im:
        im.load()
        if im.width < 300 or im.height < 180:
            raise ValueError("Patent şekli çözünürlüğü yetersiz.")
        if im.format == "PNG":
            return data
        out = io.BytesIO()
        if im.mode not in {"RGB", "RGBA", "L"}:
            im = im.convert("RGB")
        im.save(out, format="PNG")
        return out.getvalue()


def resolve_original_patent_figure(document_info: dict[str, Any]) -> bytes | None:
    """Yalnız özgün patent kaynağından şekil getirir; hiçbir zaman yapay şekil üretmez."""
    candidates: list[str] = []
    direct = str(document_info.get("figure_image_url") or "").strip()
    if direct:
        candidates.append(direct)
    source = str(document_info.get("source_url") or "").strip()
    if source:
        urls = _google_patent_image_urls(source)
        # D00000 çoğunlukla kapak/ilk çizim sayfasıdır; mümkünse ilk gerçek çizim sayfasını seç.
        preferred = [u for u in urls if "D00000" not in u.upper()]
        candidates.extend(preferred or urls)
    seen: set[str] = set()
    for url in candidates:
        if url in seen:
            continue
        seen.add(url)
        try:
            return _normalize_image_bytes(_fetch_remote_bytes(url))
        except Exception:
            continue
    return None


def _report_pdf_fallback_figures(asset: UploadedAsset | None) -> list[bytes]:
    """İlk raporda zaten bulunan özgün patent şekillerini son çare olarak yeniden kullanır."""
    if asset is None:
        return []
    imgs = extract_embedded_images(asset)
    unique: list[tuple[int, bytes]] = []
    seen: set[str] = set()
    import hashlib
    for img in imgs:
        try:
            with Image.open(io.BytesIO(img.data)) as im:
                area = im.width * im.height
                if im.width < 420 or im.height < 230:
                    continue
        except Exception:
            continue
        h = hashlib.sha1(img.data).hexdigest()
        if h in seen:
            continue
        seen.add(h)
        unique.append((area, img.data))
    # Patent şekilleri rapor logolarından belirgin biçimde daha büyük olur. Belge sırasını mümkün olduğunca korumak için
    # ilk dört büyük adayı alıp kaynak çıkarma sırasına göre kullanıyoruz.
    large_hashes = {hashlib.sha1(data).hexdigest() for _, data in sorted(unique, key=lambda x: x[0], reverse=True)[:4]}
    return [data for _, data in unique if hashlib.sha1(data).hexdigest() in large_hashes]


def _replace_figure_paragraph(paragraph, image_data: bytes | None) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not image_data:
        return
    try:
        with Image.open(io.BytesIO(image_data)) as im:
            w, h = im.size
        max_w_cm, max_h_cm = 15.8, 11.5
        ratio = min(max_w_cm / max(w, 1), max_h_cm / max(h, 1))
        width_cm = max(5.0, min(max_w_cm, w * ratio))
        paragraph.add_run().add_picture(io.BytesIO(image_data), width=Cm(width_cm))
    except Exception as exc:
        raise ValueError("Özgün patent şekli Word raporuna eklenemedi.") from exc


def _validate_research_template_fidelity(doc: Document) -> None:
    """Dinamik içerik dışında bağlayıcı Tip 3 şablon geometrisinin bozulmadığını doğrular."""
    template = Document(str(ARASTIRMA_TEMPLATE))
    if len(doc.sections) != len(template.sections):
        raise ValueError("Tip 3 çıktısında section sayısı bağlayıcı şablondan farklıdır.")
    for out_sec, tpl_sec in zip(doc.sections, template.sections):
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin", "page_width", "page_height"):
            if getattr(out_sec, attr) != getattr(tpl_sec, attr):
                raise ValueError(f"Tip 3 çıktısında {attr} bağlayıcı şablondan farklıdır.")
    if len(doc.paragraphs) != len(template.paragraphs):
        raise ValueError("Tip 3 çıktısında şablonda olmayan ana paragraf/başlık eklenmiş veya silinmiştir.")
    if len(doc.tables) != len(template.tables):
        raise ValueError("Tip 3 çıktısındaki ana tablo sayısı bağlayıcı şablondan farklıdır.")

    criteria = doc.tables[0]
    tpl_criteria = template.tables[0]
    if len(criteria.rows) != 5:
        raise ValueError("Tip 3 kriter tablosunun satır yapısı bozulmuştur.")
    for r in range(5):
        # Dinamik içerik hücreleri dışındaki sabit etiketleri koru.
        if _norm_ws(criteria.rows[r].cells[0].text) != _norm_ws(tpl_criteria.rows[r].cells[0].text):
            raise ValueError("Tip 3 kriter tablosundaki sabit alan etiketi şablondan farklıdır.")
        if _norm_ws(criteria.rows[r].cells[1].text) != _norm_ws(tpl_criteria.rows[r].cells[1].text):
            raise ValueError("Tip 3 kriter tablosundaki iki nokta/ayraç hücresi şablondan farklıdır.")

    if _norm_ws(criteria.rows[2].cells[2].text) != _norm_ws("Global (İlan edilmiş olan patent başvuruları)"):
        raise ValueError("Tip 3 Kapsam hücresi bağlayıcı şablondaki sabit metni korumuyor.")

    keyword_cell = criteria.rows[3].cells[2]
    if len(keyword_cell.tables) != 1 or len(keyword_cell.tables[0].rows) != 5 or len(keyword_cell.tables[0].columns) != 2:
        raise ValueError("Tip 3 anahtar kelime tablosunun 5x2 şablon geometrisi bozulmuştur.")

    ipc_cell = criteria.rows[4].cells[2]
    if len(ipc_cell.paragraphs) != 4:
        raise ValueError("Tip 3 IPC alanı şablondaki dört paragraf yapısını korumuyor.")
    for para in ipc_cell.paragraphs:
        if not para.text.strip():
            continue
        if len(para.runs) < 2 or para.runs[0].bold is not True or para.runs[1].bold is True:
            raise ValueError("Tip 3 IPC alanında kod kalın, İngilizce açıklama normal yazı biçimi korunmalıdır.")

    warning_cell = doc.tables[4].rows[0].cells[2]
    if len(warning_cell.paragraphs) != 4:
        raise ValueError("Tip 3 uyarı alanı şablondaki dört ayrı paragraf yapısını korumuyor.")


def build_research_docx(report: dict[str, Any], figure_fallbacks: list[bytes] | None = None) -> bytes:
    """Bağlayıcı Tip 3 şablonunu yerinde doldurur; gövdeyi yeniden kurmaz."""
    validate_research_report_language(report)
    docs = report.get("documents") or []
    if not docs:
        raise ValueError("Rapor için D1 dokümanı bulunamadı.")
    if len(docs) > 2:
        docs = docs[:2]
    d1 = docs[0]
    d2 = docs[1] if len(docs) > 1 else None

    doc = Document(str(ARASTIRMA_TEMPLATE))
    ps = doc.paragraphs
    if len(ps) < 107 or len(doc.tables) < 5:
        raise ValueError("Bağlayıcı Ön Araştırma Raporu şablonunun yapısı beklenen formatta değil.")

    # Kapak ve kriterler: paragraf/boşluk/section yapısı aynen korunur.
    _replace_paragraph_text_preserve_format(ps[1], report.get("reference", ""))
    _replace_paragraph_text_preserve_format(ps[6], report.get("title", ""))
    _replace_paragraph_text_preserve_format(ps[29], report.get("report_date", date.today().strftime('%d.%m.%Y')))

    criteria = doc.tables[0]
    _replace_cell_text_preserve_format(criteria.rows[0].cells[2], report.get("purpose", "Belirlenen konuda araştırmanın gerçekleştirilmesi"))
    _replace_cell_text_preserve_format(criteria.rows[1].cells[2], report.get("title", ""))
    _replace_cell_text_preserve_format(criteria.rows[2].cells[2], "Global (İlan edilmiş olan patent başvuruları)")
    _fill_keyword_table(criteria.rows[3].cells[2], report.get("keywords") or [])
    _fill_ipc_cell(criteria.rows[4].cells[2], report.get("ipc_cpc") or [])

    _replace_paragraph_text_preserve_format(ps[36], report.get("evaluation_intro", ""))

    def apply_doc(block: dict[str, Any], label: str, indices: dict[str, int], table_idx: int, fallback_idx: int):
        header = f"{label}- {block.get('number','')}"
        if block.get("alternate_number"):
            header += f" ({block.get('alternate_number')})"
        header += f"- {block.get('title','')}- {block.get('date','')}"
        _replace_paragraph_text_preserve_format(ps[indices['header']], header)
        _replace_paragraph_text_preserve_format(ps[indices['description']], " ".join(block.get("description") or []))
        _replace_paragraph_text_preserve_format(ps[indices['abstract']], block.get("abstract", ""))
        _replace_paragraph_text_preserve_format(ps[indices['figure_heading']], block.get("figure_caption", f"{label}- Şekil"))
        _replace_paragraph_text_preserve_format(ps[indices['comparison_heading']], f"Araştırma konusu ile {label} dokümanı arasında benzerlik değerlendirmesi:")
        _replace_comparison_table(doc.tables[table_idx], block.get("comparison_rows") or [], label)
        _replace_paragraph_text_preserve_format(ps[indices['assessment']], " ".join(block.get("novelty_assessment") or []))
        fig = resolve_original_patent_figure(block)
        if fig is None and figure_fallbacks and fallback_idx < len(figure_fallbacks):
            fig = figure_fallbacks[fallback_idx]
        if fig is None:
            raise ValueError(f"{label} için özgün patent şekli otomatik temin edilemedi. Yapay şekil oluşturulmadı; ilgili patentin orijinal PDF/şekil kaynağı gereklidir.")
        _replace_figure_paragraph(ps[indices['figure']], fig)

    apply_doc(d1, "D1", {"header":40,"description":42,"abstract":46,"figure_heading":47,"figure":48,"comparison_heading":49,"assessment":55}, 2, 0)

    if d2:
        apply_doc(d2, "D2", {"header":57,"description":59,"abstract":62,"figure_heading":64,"figure":65,"comparison_heading":66,"assessment":71}, 3, 1)
    else:
        # Şablonda D2 için ayrılmış alanı görünmez kıl; diğer format öğelerini bozma.
        for i in [57,59,61,62,64,65,66,68,69,71]:
            _replace_paragraph_text_preserve_format(ps[i], "")
        for tr in list(doc.tables[3]._tbl.tr_lst)[1:]:
            doc.tables[3]._tbl.remove(tr)
        _replace_cell_text_preserve_format(doc.tables[3].rows[0].cells[0], "")
        _replace_cell_text_preserve_format(doc.tables[3].rows[0].cells[1], "")

    inv = report.get("inventive_step_paragraphs") or []
    slots = [75, 77, 79, 81]
    for i, pidx in enumerate(slots):
        _replace_paragraph_text_preserve_format(ps[pidx], inv[i] if i < len(inv) else "")

    conclusion = report.get("conclusion_paragraphs") or []
    _replace_paragraph_text_preserve_format(ps[85], " ".join(conclusion))

    # Uyarılar tablosu şablonun kendi yerinde kalır; ayrı yeni bölüm açılmaz.
    warnings = [str(x).strip() for x in (report.get("warnings") or []) if str(x).strip()]
    _fill_warning_cell(doc.tables[4].rows[0].cells[2], warnings)

    _validate_research_template_fidelity(doc)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# -----------------------------------------------------------------------------
# ARAYÜZ
# -----------------------------------------------------------------------------
st.set_page_config(page_title=f"Patent Atölyesi {APP_VERSION}", page_icon="⚙️", layout="wide")
st.markdown(
    f"""
    <style>
      .block-container {{max-width: 1180px; padding-top: 1.5rem; padding-bottom: 3rem;}}
      .hero {{padding: 1.2rem 1.4rem; border:1px solid #e7e7e7; border-radius:16px; margin-bottom:1rem;}}
      .hero h1 {{margin:0; font-size:2rem;}}
      .hero p {{margin:.35rem 0 0 0; color:#666;}}
      .version {{font-size:.82rem; color:#888; margin-top:.45rem;}}
      div[data-testid="stDownloadButton"] button, div[data-testid="stFormSubmitButton"] button {{width:100%;}}
    </style>
    <div class="hero">
      <h1>Patent Atölyesi {APP_VERSION}</h1>
      <p>Tarifname, görüş, Tip 3 ön araştırma ve araştırma güncelleme çalışmalarını tek arayüzden oluşturun.</p>
      <div class="version">Kural sürümü: {RULESET_VERSION}</div>
    </div>
    """,
    unsafe_allow_html=True,
)


if not os.getenv("OPENAI_API_KEY", "").strip():
    st.warning("OPENAI_API_KEY henüz tanımlı değil. Arayüzü inceleyebilirsiniz; üretim düğmeleri API anahtarı olmadan çalışmaz.")

work_type = st.radio(
    "İş türü",
    ["Tarifname oluşturma", "Görüş hazırlama", "Tip 3 - Ön araştırma raporu", "Araştırma güncelleme - Tip 3"],
    horizontal=True,
)

# TARİFNAME
if work_type == "Tarifname oluşturma":
    st.subheader("Tarifname oluşturma")
    with st.form("tarifname_form"):
        c1, c2 = st.columns(2)
        with c1:
            bbf = st.file_uploader("BBF dosyası", type=["docx", "doc", "pdf", "txt"], key="tar_bbf")
            reference = st.text_input("DP referans numarası", value="")
            st.caption("Tarifname çıktı adı DP referansından otomatik oluşturulur: Tarifname_<DP referansı>.docx")
        with c2:
            language_choice = st.selectbox("Tarifname dili", ["Türkçe", "İngilizce"], index=0)
            claim_choice = st.selectbox(
                "İstem yapısı",
                ["BBF'ye göre otomatik belirle", "Yalnızca sistem", "Yalnızca yöntem", "Sistem ve yöntem"],
            )
            st.caption("Mevcut bir tarifnameyi değiştirme işlemi bu ekranda yapılmaz; tarifname düzenleme ayrı bir iş akışı olarak ele alınacaktır.")

        extra_technical_files = st.file_uploader(
            "Ek teknik müşteri belgeleri/notları (varsa)",
            type=["pdf", "docx", "doc", "txt", "md", "png", "jpg", "jpeg", "webp", "svg", "zip"],
            accept_multiple_files=True,
            key="tar_extra_technical",
        )
        example_files = st.file_uploader(
            "Örnek tarifnameler (yalnızca unsur/istem kurgusu için)",
            type=["pdf", "docx", "doc", "txt", "zip"],
            accept_multiple_files=True,
            key="tar_examples",
            help="Bu dosyaların teknik içeriği yeni tarifnameye aktarılmaz.",
        )

        separate_figures = st.checkbox("Şekilleri ayrı Word dosyası olarak oluştur")
        if separate_figures:
            st.caption("Şekiller çıktı adı DP referansından otomatik oluşturulur: Şekiller_<DP referansı>.docx")
        figure_files = st.file_uploader(
            "Ayrıca kullanılacak şekil dosyaları",
            type=["png", "jpg", "jpeg", "webp", "svg", "docx", "pdf"],
            accept_multiple_files=True,
            key="tar_figures",
            disabled=not separate_figures,
        )
        if separate_figures:
            st.caption(
                "Şekiller Word'e alınmadan önce nihai REFERANS NUMARALARI ile otomatik çapraz kontrol edilir. "
                "Eksik veya yanlış referans okları yalnız fiziksel karşılık güvenilir biçimde belirlenebiliyorsa düzeltilir; "
                "belirsiz referans varsa şekiller çıktısı durdurulur ve uydurma işaretleme yapılmaz."
            )

        literature = st.checkbox("Literatür araştırması yap ve önceki tekniğe ekle")
        lc1, lc2 = st.columns(2)
        with lc1:
            lit_count = st.number_input("Benzer patent sayısı", min_value=1, max_value=10, value=2, disabled=not literature)
        with lc2:
            jurisdiction = st.text_input("Tercih edilen ülke/veri tabanı", disabled=not literature)

        submit = st.form_submit_button("Tarifnameyi oluştur", type="primary")

    if submit:
        if bbf is None:
            st.error("BBF yükleyin.")
        elif not str(reference or "").strip():
            st.error("DP referans numarasını girin. Tarifname ve varsa Şekiller dosya adları bu numaradan otomatik oluşturulur.")
        else:
            try:
                output_name, figures_output_name = derive_tarifname_output_names(reference)
                progress = st.progress(0, text="Kaynak dosyalar okunuyor...")
                bbf_asset = UploadedAsset(bbf.name, bbf.getvalue(), bbf.type)
                source = extract_text_from_asset(bbf_asset)

                technical_assets = assets_from_uploads(extra_technical_files)
                technical_text, technical_images = combine_asset_text("EK TEKNİK BELGE", technical_assets)
                technical_figure_assets: list[UploadedAsset] = []
                for asset in technical_assets:
                    technical_figure_assets.extend(extract_embedded_images(asset))
                example_assets = assets_from_uploads(example_files)
                example_text, _ = combine_asset_text("ÖRNEK TARİFNAME - YALNIZCA KURGU", example_assets)

                embedded_images = extract_embedded_images(bbf_asset)
                provided_figure_assets: list[UploadedAsset] = []
                if separate_figures:
                    for uploaded in figure_files or []:
                        fig_asset = UploadedAsset(uploaded.name, uploaded.getvalue(), uploaded.type)
                        provided_figure_assets.extend(extract_embedded_images(fig_asset))

                # Nihai şekil dosyaları ayrıca yüklenmişse model bunları da görerek referans senkronizasyonunu denetler.
                model_images = [*provided_figure_assets, *technical_images, *embedded_images][:24]

                progress.progress(15, text="BBF içeriği, referans tablosu, şekiller ve istem çekirdeği çıkarılıyor...")
                extracted = ask_json(
                    tarifname_extraction_prompt(source, technical_text, example_text, language_choice),
                    images=model_images,
                )

                progress.progress(22, text="BBF teknik bilgi envanteri ham kaynak pasajlarıyla ikinci kez, madde madde doğrulanıyor...")
                source_passage_registry = build_source_passage_registry(source, technical_text)
                extracted = ask_json(
                    tarifname_extraction_quality_prompt(source, technical_text, extracted, source_passage_registry, language_choice),
                    images=model_images,
                )
                _validate_technical_fact_inventory(extracted)
                validate_source_passage_audit(extracted, source_passage_registry)

                progress.progress(28, text="İstem yapısı belirleniyor...")
                mode = resolve_tarifname_claim_mode(extracted, claim_choice)
                st.info(f"Kullanılan istem yapısı: {mode}")

                lit_docs: list[dict[str, Any]] = []
                if literature:
                    progress.progress(38, text="Patent literatürü araştırılıyor...")
                    lit_docs = (
                        ask_json(
                            tarifname_literature_prompt(extracted, int(lit_count), jurisdiction, language_choice),
                            web_search=True,
                        ).get("documents")
                        or []
                    )

                progress.progress(55, text="Tarifname ve istemlerin tam taslağı hazırlanıyor...")
                draft = ask_json(
                    tarifname_drafting_prompt(
                        extracted,
                        mode,
                        lit_docs,
                        source,
                        technical_text,
                        example_text,
                        language_choice,
                    ),
                    images=model_images,
                )

                progress.progress(73, text="BBF ile tamlık, istem mantığı ve şablon kuralları ikinci kez kontrol ediliyor...")
                draft = ask_json(
                    tarifname_quality_prompt(
                        source,
                        technical_text,
                        extracted,
                        draft,
                        mode,
                        lit_docs,
                        language_choice,
                    ),
                    images=model_images,
                )

                validation_feedback = ""
                warnings: list[str] = []
                final_raw_audit: dict[str, Any] = {}
                for repair_round in range(3):
                    if mode == "Yalnızca sistem":
                        draft["method_claim"] = None
                        draft["dependent_method_claims"] = []
                        draft["method_steps"] = []
                    elif mode == "Yalnızca yöntem":
                        draft["system_claim"] = None
                        draft["dependent_system_claims"] = []

                    draft = apply_tarifname_house_style(draft, mode, lit_docs, language_choice)
                    try:
                        warnings = validate_tarifname_draft(draft, mode, lit_docs, language_choice, extracted)
                        progress.progress(80 + repair_round * 3, text="Taslak oluşturuldu; ham BBF ve ek teknik kaynaklarla bağımsız son ikinci okuma yapılıyor...")
                        final_raw_audit = ask_json(
                            tarifname_final_raw_source_audit_prompt(
                                source_passage_registry, extracted, draft, language_choice
                            ),
                            images=model_images,
                        )
                        validate_final_raw_source_audit(
                            final_raw_audit,
                            extracted,
                            source_passage_registry,
                            _visible_draft_text_for_audit(draft),
                        )
                        break
                    except ValueError as validation_exc:
                        validation_feedback = str(validation_exc)
                        if repair_round >= 2:
                            raise
                        progress.progress(78 + repair_round * 3, text=f"Kalite kapısı düzeltme turu {repair_round + 1}: {validation_feedback[:120]}...")
                        draft = ask_json(
                            tarifname_quality_prompt(
                                source, technical_text, extracted, draft, mode, lit_docs, language_choice, validation_feedback
                            ),
                            images=model_images,
                        )

                for warning in warnings:
                    st.warning(warning)

                technical_facts = list(extracted.get("technical_facts") or [])
                coverage_rows = draft.get("source_coverage_map") or []
                with st.expander(f"BBF teknik bilgi kapsam kontrolü ({len(technical_facts)}/{len(technical_facts)} teknik bilgi eşleştirildi)", expanded=False):
                    by_id = {str(r.get("fact_id", "")): r for r in coverage_rows}
                    for fact in technical_facts:
                        fid = str(fact.get("id", ""))
                        row = by_id.get(fid, {})
                        sections = ", ".join(map(str, row.get("sections") or []))
                        st.write(f"✅ {fid}: {fact.get('statement','')} — {sections}")

                progress.progress(88, text="Word dosyası hazırlanıyor; ham kaynak zinciri ve beş son kalite kapısı nihai Word üzerinde tekrar doğrulanıyor...")
                data = build_tarifname_docx(draft, language_choice)
                validate_tarifname_docx_structure(data, draft, language_choice)
                final_gates = validate_tarifname_post_generation_quality(
                    data,
                    draft,
                    extracted,
                    mode,
                    lit_docs,
                    language_choice,
                    source_passage_registry=source_passage_registry,
                    final_raw_audit=final_raw_audit,
                )
                render_tarifname_docx_smoke_test(data)
                required_final_gates = ["source_completeness", "claims", "references", "template", "element_step_language"]
                if not all(final_gates.get(key) is True for key in required_final_gates):
                    raise ValueError("Nihai tarifname kalite kapılarının tamamı doğrulanmadan indirme açılamaz.")
                st.success(
                    f"Ham veri kontrolü yapıldı: {final_gates['raw_passages_total']} ham kaynak pasajı tam bir kez sınıflandırıldı; "
                    f"{final_gates['technical_passages']} teknik ham pasaj ve {final_gates['technical_facts']} atomik teknik bilgi "
                    "nihai Word metnindeki gerçek kanıtlarla doğrulandı."
                )
                st.success("Kontrol kapıları tamamlandı: ✅ 1/5 Ham kaynak/BBF tamlığı  ✅ 2/5 Ana + alt istemler  ✅ 3/5 Referanslar  ✅ 4/5 Tam şablon  ✅ 5/5 Unsur/yöntem dili")
                st.caption("Tarifname Word indirmesi yalnız ham veri ikinci okuması ve bütün son kalite kapıları başarıyla tamamlandığında açılır.")

                figure_data = None
                figure_reports: list[dict[str, Any]] = []
                figure_unresolved: list[str] = []
                if separate_figures:
                    # BBF içindeki kullanılabilir özgün teknik şekiller ZORUNLU kaynak şekildir.
                    # Ayrıca yüklenen müşteri şekilleri bunlara eklenir; salt ayrı şekil yüklenmiş olması BBF şekillerini düşürmez.
                    # Aynı görsel bayt düzeyinde yineleniyorsa aşağıdaki deduplikasyon tek kopya bırakır.
                    all_figure_assets: list[UploadedAsset] = [*embedded_images, *technical_figure_assets, *provided_figure_assets]
                    source_figure_inventory = [asset.name for asset in [*embedded_images, *technical_figure_assets, *provided_figure_assets]]
                    # Aynı görselin birden fazla kez eklenmesini engelle.
                    deduplicated: list[UploadedAsset] = []
                    seen_images: set[int] = set()
                    for asset in all_figure_assets:
                        marker = hash(asset.data)
                        if marker not in seen_images:
                            seen_images.add(marker)
                            deduplicated.append(asset)

                    # Kaynak şekil kalite kapısı: BBF'den çıkarılan kullanılabilir teknik görseller
                    # seçim listesinden sessizce düşürülemez.
                    selected_source_names = {asset.name for asset in deduplicated}
                    omitted_source_figures = [name for name in source_figure_inventory if name not in selected_source_names]
                    if omitted_source_figures:
                        figure_unresolved.append(
                            "BBF içindeki zorunlu kaynak teknik şekiller seçimden düştü: " + ", ".join(omitted_source_figures)
                        )

                    if not deduplicated:
                        figure_unresolved.append("Şekiller Word dosyası için kullanılabilir müşteri görseli bulunamadı.")
                    else:
                        def figure_progress(index: int, total: int, stage: str) -> None:
                            stage_text = {
                                "audit": "referansları tarifnameyle karşılaştırıyor",
                                "edit": "eksik/yanlış referans oklarını düzeltiyor",
                                "verify": "düzeltmeyi ve müşteri geometrisini doğruluyor",
                            }.get(stage, "şekilleri kontrol ediyor")
                            ratio = (index - 1) / max(total, 1)
                            percent = min(97, 89 + int(ratio * 8))
                            progress.progress(percent, text=f"ŞEKİL {index}/{total}: {stage_text}...")

                        prepared_figures, figure_reports, figure_unresolved = prepare_figures_with_reference_audit(
                            deduplicated,
                            draft,
                            language_choice,
                            progress_callback=figure_progress,
                        )
                        if not figure_unresolved:
                            progress.progress(98, text="Referansları doğrulanmış şekiller Word dosyasına yerleştiriliyor...")
                            figure_data = build_figures_docx(prepared_figures, language_choice)

                progress.progress(100, text="Hazır")
                st.success("Tarifname oluşturuldu ve BBF tamlık kontrolü tamamlandı.")
                st.download_button(
                    "Tarifname Word dosyasını indir",
                    data=data,
                    file_name=safe_output_name(output_name, f"Tarifname_{str(reference).strip()}.docx"),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                )
                if separate_figures and figure_reports:
                    with st.expander("Şekil referans kontrolü", expanded=bool(figure_unresolved)):
                        for report in figure_reports:
                            idx = report.get("figure_index", "?")
                            status = report.get("final_status", "unresolved")
                            if status == "ok":
                                st.write(f"✅ ŞEKİL {idx}: referanslar doğru; özgün müşteri şekli korundu.")
                            elif status == "corrected":
                                st.write(f"🛠️ ŞEKİL {idx}: eksik/yanlış referans numarası veya oku düzeltildi ve ikinci kontrolden geçti.")
                            else:
                                st.write(f"⚠️ {report.get('message', f'ŞEKİL {idx}: referans belirsizliği giderilemedi.')}")
                if figure_unresolved:
                    st.error(
                        "Şekiller dosyası oluşturulmadı. En az bir referansın fiziksel karşılığı güvenilir biçimde "
                        "doğrulanamadı veya otomatik düzeltme müşteri geometrisi kontrolünü geçemedi. "
                        "Tarifname Word dosyası kullanılabilir; şekilleri düzeltip yeniden yükleyin."
                    )
                    for message in figure_unresolved:
                        st.warning(message)
                if figure_data is not None:
                    st.download_button(
                        "Şekiller Word dosyasını indir",
                        data=figure_data,
                        file_name=safe_output_name(figures_output_name, f"Şekiller_{str(reference).strip()}.docx"),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
            except Exception as exc:
                st.exception(exc)

# GÖRÜŞ
elif work_type == "Görüş hazırlama":
    st.subheader("Görüş hazırlama")
    st.caption("Akış: görüş türü → rapor ve önceki görüş → tarifname → araştırma raporunda X/Y veya ofis aksiyonunda gerekçede fiilen kullanılan savunma dokümanları → teknik analiz/revizyon kararı → ham-kaynak ikinci okuma → son Markup fiziksel sayfa/satır kontrolü → Word kalite kapıları → görüş çıktısı.")

    opinion_modes = [
        "Araştırma raporuna karşı",
        "İnceleme raporuna karşı",
        "EP araştırma raporu veya ofis aksiyon",
        "Yurtdışı ofis aksiyon",
    ]
    opinion_case_mode = st.radio(
        "Görüş çalışma sekmesi",
        opinion_modes,
        horizontal=True,
        key="gor_case_mode",
    )
    report_type_map = {
        "Araştırma raporuna karşı": "Türkiye araştırma raporuna karşı görüş",
        "İnceleme raporuna karşı": "Türkiye inceleme raporuna karşı görüş",
        "EP araştırma raporu veya ofis aksiyon": "EP araştırma raporu veya ofis aksiyonuna karşı görüş",
        "Yurtdışı ofis aksiyon": "Yurtdışı ofis aksiyonuna karşı görüş",
    }
    report_label_map = {
        "Araştırma raporuna karşı": "Türkiye araştırma raporu",
        "İnceleme raporuna karşı": "Türkiye inceleme raporu",
        "EP araştırma raporu veya ofis aksiyon": "EP araştırma raporu / ofis aksiyon",
        "Yurtdışı ofis aksiyon": "Yurtdışı ofis aksiyon",
    }
    report_type = report_type_map[opinion_case_mode]
    opinion_language = st.selectbox("Görüş dili", ["Türkçe", "İngilizce"], key="gor_language")
    report_file = st.file_uploader(report_label_map[opinion_case_mode], type=["pdf", "docx", "doc", "txt"], key="gor_report")

    prior_file = None
    prior_yes = st.radio("Önceki sunulan görüş var mı?", ["Hayır", "Evet"], horizontal=True, key="gor_prior_yes")
    if prior_yes == "Evet":
        prior_file = st.file_uploader("Önceki sunulan görüş", type=["pdf", "docx", "doc", "txt"], key="gor_prior")

    customer_yes = st.radio("Müşteriden bilgi var mı?", ["Hayır", "Evet"], horizontal=True)
    customer_files = []
    if customer_yes == "Evet":
        customer_files = st.file_uploader(
            "Müşteri bilgileri",
            type=["pdf", "docx", "doc", "txt", "png", "jpg", "jpeg", "webp", "zip"],
            accept_multiple_files=True,
            key="gor_customer",
        )

    spec_file = st.file_uploader("Tarifname", type=["pdf", "docx", "doc", "txt"], key="gor_spec")
    reference = st.text_input("Ana dosya referansı nedir?", value="")
    applicant_override = st.text_input("Başvuru sahibi (raporda yoksa girin)", value="", key="gor_applicant")
    output_name = st.text_input("Çıktı dosyasının adı", value="Görüş Metni_XXXXXX.docx")

    for key, default in {
        "gorus_required_docs": None,
        "gorus_scope_report_text": None,
        "gorus_analysis": None,
        "gorus_source": None,
        "gorus_markup_data": None,
        "gorus_clean_data": None,
        "gorus_final_spec_text": None,
        "gorus_opinion_data": None,
        "gorus_opinion_status": None,
        "gorus_quality_report": None,
    }.items():
        if key not in st.session_state:
            st.session_state[key] = default

    if st.button("1. Raporu analiz et ve savunmada gerekli dokümanları belirle", type="primary", use_container_width=True):
        if not all([reference.strip(), report_file, spec_file]):
            st.error("Referans, rapor ve tarifnameyi yükleyin.")
        elif prior_yes == "Evet" and prior_file is None:
            st.error("Önceki sunulan görüş var seçildi; önceki görüşü yükleyin.")
        elif customer_yes == "Evet" and not customer_files:
            st.error("Müşteriden bilgi var seçildi; müşteri bilgilerini yükleyin.")
        else:
            try:
                report_text_scope = extract_text_from_asset(UploadedAsset(report_file.name, report_file.getvalue(), report_file.type))
                xy_scope = (
                    opinion_case_mode == "Araştırma raporuna karşı"
                    or (opinion_case_mode == "EP araştırma raporu veya ofis aksiyon" and is_ep_search_report(report_text_scope))
                )
                required_docs = detect_ep_xy_documents(report_text_scope) if xy_scope else detect_examiner_reasoned_documents(report_text_scope)
                if not required_docs:
                    raise ValueError("Savunmada kullanılacak doküman otomatik kesinleştirilemedi. Türkiye/EP araştırma raporlarında yalnız X/Y kategorileri, inceleme ve ofis aksiyonlarında ise gerekçede fiilen kullanılan dokümanlar kabul edilir.")
                st.session_state.gorus_required_docs = required_docs
                st.session_state.gorus_scope_report_text = report_text_scope
                st.session_state.gorus_analysis = None
                st.session_state.gorus_source = None
                st.session_state.gorus_opinion_data = None
                st.session_state.gorus_quality_report = None
            except Exception as exc:
                st.exception(exc)

    required_docs = st.session_state.gorus_required_docs
    similar_files = []
    if required_docs:
        st.success("Savunmada kullanılacak doküman(lar) tespit edildi.")
        for item in required_docs:
            st.markdown(f"- **{item.get('label','')} — {item.get('number','')}**")
        required_text = ", ".join(f"{x.get('label','')} {x.get('number','')}" for x in required_docs)
        similar_files = st.file_uploader(
            f"Savunma için gerekli doküman(lar): {required_text}. Lütfen yalnız bu dokümanları yükleyin.",
            type=["pdf", "docx", "doc", "txt", "zip"],
            accept_multiple_files=True,
            key="gor_sim",
        )

    if required_docs and st.button("2. Raporu, istemleri ve savunma dokümanlarını teknik olarak analiz et", type="primary", use_container_width=True):
        if not similar_files:
            st.error("Yukarıda tespit edilen savunma dokümanlarını yükleyin.")
        else:
            try:
                progress = st.progress(0, text="Dosyalar okunuyor...")
                report_text = extract_text_from_asset(UploadedAsset(report_file.name, report_file.getvalue(), report_file.type))
                spec_bytes = spec_file.getvalue()
                spec_text = extract_text_from_asset(UploadedAsset(spec_file.name, spec_bytes, spec_file.type))
                prior_text = ""
                if prior_file:
                    prior_text = extract_text_from_asset(UploadedAsset(prior_file.name, prior_file.getvalue(), prior_file.type))
                sim_assets = assets_from_uploads(similar_files)
                sim_text, sim_images = combine_asset_text("BENZER DOKÜMAN", sim_assets)
                cust_assets = assets_from_uploads(customer_files)
                cust_text, cust_images = combine_asset_text("MÜŞTERİ BİLGİSİ", cust_assets)
                model_images = [*sim_images, *cust_images]

                progress.progress(35, text="Rapor itirazları, savunma dokümanları ve mevcut istemler analiz ediliyor...")
                analysis = ask_json(
                    gorus_analysis_prompt(
                        report_type,
                        reference,
                        report_text,
                        spec_text,
                        prior_text,
                        sim_text,
                        cust_text,
                    ),
                    images=model_images,
                )
                validate_gorus_analysis(analysis, spec_text)

                st.session_state.gorus_analysis = analysis
                st.session_state.gorus_source = {
                    "report_type": report_type,
                    "opinion_case_mode": opinion_case_mode,
                    "language": opinion_language,
                    "reference": reference,
                    "applicant_override": applicant_override.strip(),
                    "output_name": output_name,
                    "required_docs": required_docs,
                    "report_text": report_text,
                    "spec_text": spec_text,
                    "spec_name": spec_file.name,
                    "spec_bytes": spec_bytes,
                    "prior_text": prior_text,
                    "sim_text": sim_text,
                    "sim_assets": sim_assets,
                    "cust_text": cust_text,
                    "model_images": model_images,
                }
                st.session_state.gorus_markup_data = None
                st.session_state.gorus_clean_data = None
                st.session_state.gorus_final_spec_text = None
                st.session_state.gorus_opinion_data = None
                st.session_state.gorus_opinion_status = None
                st.session_state.gorus_quality_report = None
                progress.progress(100, text="Teknik analiz tamamlandı")
            except Exception as exc:
                st.exception(exc)

    analysis = st.session_state.gorus_analysis
    source_state = st.session_state.gorus_source

    ready_to_generate = False
    final_spec_text = None
    revision_status = ""
    opinion_step = 3

    if analysis and source_state:
        st.success("Rapor ve savunma dokümanlarının teknik analizi tamamlandı.")
        if analysis.get("analysis_summary"):
            st.write(analysis.get("analysis_summary"))

        issues = analysis.get("examiner_issues") or []
        if issues:
            st.markdown("**Raporda odaklanılması gereken hususlar:**")
            for item in issues:
                st.markdown(f"- {item}")

        directions = analysis.get("defense_direction") or []
        if directions:
            st.markdown("**Önerilen savunma yönü:**")
            for item in directions:
                st.markdown(f"- {item}")

        amendment_required = bool(analysis.get("amendment_required"))
        if amendment_required:
            st.warning("Analiz, istemlerde sınırlı bir revizyonun gerekli olabileceğini belirledi. Görüş henüz oluşturulmayacak.")
            if analysis.get("amendment_reason"):
                st.write(analysis.get("amendment_reason"))

            amendments = analysis.get("amendments") or []
            for idx, amendment in enumerate(amendments, 1):
                claim_no = amendment.get("claim_number", "")
                with st.expander(f"Revizyon {idx} – İstem {claim_no}", expanded=True):
                    st.markdown(f"**Gerekçe:** {amendment.get('reason','')}")
                    if amendment.get("basis_quote"):
                        st.markdown("**Tarifname dayanağı:**")
                        st.code(amendment.get("basis_quote", ""), language=None)
                    st.markdown("**Mevcut ifade:**")
                    st.code(amendment.get("old_text", ""), language=None)
                    st.markdown("**Önerilen ifade:**")
                    st.code(amendment.get("new_text", ""), language=None)

            refine_instruction = st.text_area(
                "Revizyon önerileri için düzeltme talimatı (varsa)",
                key="gor_revision_instruction",
                placeholder="Örneğin: yalnızca istem 1'i değiştir; istem 4'e dokunma; eklenen ifadeyi tarifnamedeki şu dayanakla sınırla...",
            )
            if st.button("Revizyon önerilerini yeniden analiz et", use_container_width=True):
                if not refine_instruction.strip():
                    st.error("Yeniden analiz için düzeltme talimatını yazın.")
                else:
                    try:
                        progress = st.progress(0, text="Revizyon önerileri yeniden değerlendiriliyor...")
                        revised_analysis = ask_json(
                            gorus_revision_refine_prompt(
                                source_state["report_type"],
                                source_state["report_text"],
                                source_state["spec_text"],
                                source_state["prior_text"],
                                source_state["sim_text"],
                                source_state["cust_text"],
                                analysis,
                                refine_instruction,
                            ),
                            images=source_state.get("model_images") or [],
                        )
                        validate_gorus_analysis(revised_analysis, source_state["spec_text"])
                        st.session_state.gorus_analysis = revised_analysis
                        st.session_state.gorus_markup_data = None
                        st.session_state.gorus_clean_data = None
                        st.session_state.gorus_final_spec_text = None
                        st.session_state.gorus_opinion_data = None
                        st.session_state.gorus_opinion_status = None
                        progress.progress(100, text="Revizyon önerileri güncellendi")
                        st.rerun()
                    except Exception as exc:
                        st.exception(exc)

            decision = st.radio(
                "İstem revizyonu kararı",
                ["Henüz karar vermedim", "Önerilen revizyonları uygula", "Revizyon yapmadan mevcut istemlerle devam et"],
                key="gor_revision_decision",
            )

            if decision == "Önerilen revizyonları uygula":
                if Path(source_state.get("spec_name", "")).suffix.lower() != ".docx":
                    st.error("Gerçek Word Track Changes/Markup üretimi için tarifnameyi .docx olarak yükleyin.")
                else:
                    if st.button("3. Onaylı revizyonlardan Markup ve temiz tarifnameyi oluştur", type="primary", use_container_width=True):
                        try:
                            progress = st.progress(0, text="İstem revizyonları Word dosyasına uygulanıyor...")
                            markup_data, clean_data = build_claim_revision_pair(
                                source_state["spec_bytes"],
                                analysis.get("amendments") or [],
                                analysis.get("description_prior_art_updates") or [],
                            )
                            validate_minimal_tracked_changes(markup_data)
                            if analysis.get("description_prior_art_updates"):
                                clean_doc=Document(io.BytesIO(clean_data))
                                validate_ep_prior_art_markup_text([p.text for p in clean_doc.paragraphs], source_state["spec_text"])
                            st.session_state.gorus_markup_data = markup_data
                            st.session_state.gorus_clean_data = clean_data
                            st.session_state.gorus_final_spec_text = docx_text(clean_data)
                            st.session_state.gorus_opinion_data = None
                            st.session_state.gorus_opinion_status = None
                            progress.progress(100, text="Markup ve temiz sürüm hazır")
                        except Exception as exc:
                            st.exception(exc)

                    if st.session_state.gorus_markup_data and st.session_state.gorus_clean_data:
                        ref_name = re.sub(r"[^0-9A-Za-zÇĞİÖŞÜçğıöşü_-]+", "_", source_state.get("reference", "")) or "rev"
                        c1, c2 = st.columns(2)
                        with c1:
                            st.download_button(
                                "Markup (Track Changes) tarifnameyi indir",
                                data=st.session_state.gorus_markup_data,
                                file_name=f"Düzenlenen_tarifname_track_changes_{ref_name}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                            )
                        with c2:
                            st.download_button(
                                "Temiz revize tarifnameyi indir",
                                data=st.session_state.gorus_clean_data,
                                file_name=f"Düzenlenen_tarifname_temiz_{ref_name}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                            )

                        confirmed = st.checkbox(
                            "Revize istemleri kontrol ettim ve bu istem setini görüş için onaylıyorum.",
                            key="gor_revised_claims_confirmed",
                        )
                        if confirmed:
                            ready_to_generate = True
                            final_spec_text = st.session_state.gorus_final_spec_text
                            revision_status = "Kullanıcı tarafından onaylanmış revize istem seti"
                            opinion_step = 4

            elif decision == "Revizyon yapmadan mevcut istemlerle devam et":
                no_revision_confirm = st.checkbox(
                    "Revizyon önerisini gördüm; mevcut istemlerle görüş hazırlanmasını onaylıyorum.",
                    key="gor_no_revision_confirmed",
                )
                if no_revision_confirm:
                    ready_to_generate = True
                    final_spec_text = source_state["spec_text"]
                    revision_status = "Kullanıcının açık kararıyla mevcut istemlerle revizyonsuz devam"
                    opinion_step = 3
        else:
            st.info(analysis.get("no_amendment_reason") or "İlk analizde istem revizyonu gerekli görülmedi. Mevcut istemlerle görüş hazırlanabilir.")
            ready_to_generate = True
            final_spec_text = source_state["spec_text"]
            revision_status = "İlk analizde revizyon gerekmiyor; mevcut istemler esas alındı"
            opinion_step = 3

        if ready_to_generate and final_spec_text:
            if st.button(f"{opinion_step}. Görüş metnini oluştur", type="primary", use_container_width=True):
                try:
                    progress = st.progress(0, text="Onaylı istem seti üzerinden görüş hazırlanıyor...")
                    opinion = ask_json(
                        gorus_prompt(
                            source_state["report_type"],
                            source_state["reference"],
                            source_state["report_text"],
                            final_spec_text,
                            source_state["prior_text"],
                            source_state["sim_text"],
                            source_state["cust_text"],
                            preanalysis=analysis,
                            revision_status=revision_status,
                            output_language=source_state.get("language") or "Türkçe",
                            applicant_override=source_state.get("applicant_override") or "",
                        ),
                        images=source_state.get("model_images") or [],
                    )
                    if source_state.get("applicant_override"):
                        opinion["applicant"] = source_state["applicant_override"]
                    validate_quotes(opinion, final_spec_text)
                    validate_opinion_against_raw_sources(
                        opinion, source_state["report_text"], final_spec_text,
                        source_state["prior_text"], source_state["sim_text"], source_state["cust_text"],
                        allowed_documents=source_state.get("required_docs"),
                    )
                    progress.progress(58, text="Ham kaynaklara karşı bağımsız ikinci okuma yapılıyor...")
                    quality_audit = ask_json(
                        gorus_quality_audit_prompt(
                            source_state["report_text"], final_spec_text, source_state["prior_text"],
                            source_state["sim_text"], source_state["cust_text"], analysis, opinion,
                        ),
                        images=source_state.get("model_images") or [],
                    )
                    try:
                        validate_ai_quality_audit(quality_audit)
                    except Exception:
                        progress.progress(68, text="İkinci okuma bulgularına göre taslak bir kez düzeltiliyor...")
                        opinion = ask_json(
                            gorus_repair_prompt(
                                source_state["report_text"], final_spec_text, source_state["prior_text"],
                                source_state["sim_text"], source_state["cust_text"], analysis, opinion, quality_audit,
                            ),
                            images=source_state.get("model_images") or [],
                        )
                        if source_state.get("applicant_override"):
                            opinion["applicant"] = source_state["applicant_override"]
                        validate_quotes(opinion, final_spec_text)
                        validate_opinion_against_raw_sources(
                            opinion, source_state["report_text"], final_spec_text,
                            source_state["prior_text"], source_state["sim_text"], source_state["cust_text"],
                            allowed_documents=source_state.get("required_docs"),
                        )
                        quality_audit = ask_json(
                            gorus_quality_audit_prompt(
                                source_state["report_text"], final_spec_text, source_state["prior_text"],
                                source_state["sim_text"], source_state["cust_text"], analysis, opinion,
                            ),
                            images=source_state.get("model_images") or [],
                        )
                        validate_ai_quality_audit(quality_audit)
                    # Sayfa/satır numaraları modelden alınmaz. Markup üretildiyse TEK otorite kullanıcıya
                    # verilecek son Markup Word dosyasının fiziksel render'ıdır; clean/orijinal sürüm kullanılmaz.
                    if revision_status.startswith("Kullanıcı tarafından onaylanmış revize") and st.session_state.gorus_markup_data:
                        final_spec_bytes = st.session_state.gorus_markup_data
                        final_spec_name = "son_markup_tarifname.docx"
                    else:
                        final_spec_bytes = source_state["spec_bytes"]
                        final_spec_name = source_state["spec_name"]
                    annotate_quote_locations(
                        opinion, final_spec_name, final_spec_bytes, source_state.get("language") or "Türkçe"
                    )
                    validate_quote_locations_against_spec(
                        opinion, final_spec_name, final_spec_bytes, source_state.get("language") or "Türkçe"
                    )
                    figure_images = extract_cited_original_figure_pages(
                        opinion.get("cited_documents") or [],
                        source_state.get("sim_assets") or [],
                    )
                    required_figure_labels = [str(d.get("label", "")) for d in opinion.get("cited_documents") or [] if d.get("number")]
                    missing_figs = [x for x in required_figure_labels if str(x).upper() not in figure_images]
                    if missing_figs:
                        raise ValueError("Özgün patent şekli çıkarılamayan dokümanlar: " + ", ".join(missing_figs))
                    data = build_gorus_docx(opinion, figure_images=figure_images)
                    validate_gorus_template_fidelity(data, GORUS_TEMPLATE, opinion, required_figure_labels)
                    validate_gorus_docx_content_flow(data)
                    render_gorus_docx_smoke_test(data)
                    st.session_state.gorus_opinion_data = data
                    st.session_state.gorus_opinion_status = revision_status
                    st.session_state.gorus_quality_report = build_gorus_quality_report()
                    progress.progress(100, text="Görüş metni ve tüm kalite kapıları hazır")
                except Exception as exc:
                    st.exception(exc)

        if st.session_state.gorus_opinion_data and st.session_state.gorus_opinion_status == revision_status:
            st.success("Görüş metni oluşturuldu ve kalite kapılarının tamamı geçti.")
            qr = st.session_state.get("gorus_quality_report") or {}
            if qr.get("checks"):
                with st.expander("Çıktı kalite kontrolü", expanded=True):
                    for check in qr.get("checks") or []:
                        st.write(f"✅ {check.get('name','')}")
            st.download_button(
                "Word görüş metnini indir",
                data=st.session_state.gorus_opinion_data,
                file_name=safe_output_name(source_state.get("output_name") or output_name, "Görüş Metni.docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )

# ARAŞTIRMA
elif work_type == "Tip 3 - Ön araştırma raporu":
    st.subheader("Tip 3 - Ön araştırma raporu")
    c1, c2 = st.columns(2)
    with c1:
        bbf = st.file_uploader("BBF dosyası", type=["docx", "doc", "pdf", "txt"], key="res_bbf")
        reference = st.text_input("DP referans numarası", value="")
    with c2:
        output_name = st.text_input("Çıktı dosyasının adı", value="Ön Araştırma Raporu_XXXXXX.docx")
        cutoff = st.date_input("Araştırma kesim tarihi", value=date.today())

    for key, default in {
        "top10_result": None,
        "research_bbf_text": None,
        "research_cutoff": None,
        "research_selection": None,
        "research_user_images": [],
    }.items():
        if key not in st.session_state:
            st.session_state[key] = default

    if st.button("1. Global araştırmayı yap ve en benzer 10 dokümanı bul", type="primary", use_container_width=True):
        if bbf is None:
            st.error("BBF yükleyin.")
        elif not reference.strip():
            st.error("DP referans numarasını girin.")
        else:
            try:
                progress = st.progress(0, text="BBF okunuyor ve teknik çekirdek çıkarılıyor...")
                bbf_text = extract_text_from_asset(UploadedAsset(bbf.name, bbf.getvalue(), bbf.type))
                cutoff_text = cutoff.strftime("%d.%m.%Y")
                progress.progress(20, text="Global patent veritabanlarında araştırma yapılıyor...")
                top10 = ask_json(top10_research_prompt(bbf_text, cutoff_text), web_search=True)
                docs = top10.get("documents") or []
                if len(docs) != 10:
                    raise ValueError(f"Tam 10 doküman yerine {len(docs)} doküman döndü. Araştırmayı tekrar çalıştırın.")
                st.session_state.top10_result = top10
                st.session_state.research_bbf_text = bbf_text
                st.session_state.research_cutoff = cutoff_text
                st.session_state.research_selection = None
                st.session_state.research_user_images = []
                progress.progress(100, text="Araştırma tamamlandı")
            except Exception as exc:
                st.exception(exc)

    if st.session_state.top10_result:
        st.success("En benzer 10 doküman bulundu.")
        docs = st.session_state.top10_result.get("documents") or []
        totalpatent_query = st.session_state.top10_result.get("totalpatent_query") or (
            "TotalPatent arama sorgusu: " + " or ".join(str(d.get("publication_number", "")) for d in docs)
        )
        st.code(totalpatent_query, language=None)
        rows = [{
            "Sıra": d.get("rank"), "Yayın no": d.get("publication_number"), "Başlık": d.get("title"),
            "Tarih": d.get("date"), "Yakınlık": d.get("relevance_score"),
            "Yeniliği bozar mı?": "Evet" if d.get("novelty_destroying") else "Hayır",
        } for d in docs]
        st.dataframe(rows, use_container_width=True, hide_index=True)
        st.caption(f"Önerilen D1: {st.session_state.top10_result.get('proposed_d1','')} | Önerilen D2: {st.session_state.top10_result.get('proposed_d2','') or '-'}")

        own_docs = st.radio("Sizin araştırdığınız benzer dokümanlar var mı?", ["Hayır", "Evet"], horizontal=True)
        user_files = []
        if own_docs == "Evet":
            user_files = st.file_uploader("Sizin bulduğunuz benzer dokümanlar", type=["pdf", "zip", "docx", "doc", "txt", "png", "jpg", "jpeg", "webp"], accept_multiple_files=True, key="res_user_docs")

        if st.button("2. Kullanıcı dokümanlarını incele ve nihai D1/D2'yi belirle", type="primary", use_container_width=True):
            if own_docs == "Evet" and not user_files:
                st.error("Benzer dokümanları yükleyin.")
            else:
                try:
                    progress = st.progress(0, text="Kullanıcı dokümanları inceleniyor...")
                    user_assets = assets_from_uploads(user_files)
                    user_text, user_images = combine_asset_text("KULLANICI BENZER DOKÜMANI", user_assets)
                    progress.progress(45, text="10+ dokümanları ve nihai D1/D2 belirleniyor...")
                    selection = ask_json(final_selection_prompt(
                        st.session_state.research_bbf_text,
                        st.session_state.top10_result,
                        user_text,
                    ), images=user_images, web_search=True)
                    validate_research_selection(selection)
                    st.session_state.research_selection = selection
                    st.session_state.research_user_images = user_images
                    progress.progress(100, text="Teknik değerlendirme tamamlandı")
                except Exception as exc:
                    st.exception(exc)

        selection = st.session_state.research_selection
        if selection:
            additional_query = str(selection.get("additional_query") or "").strip()
            if additional_query:
                st.code(additional_query, language=None)
            elif own_docs == "Evet":
                st.caption("10+ olarak eklenecek yeterince ilgili kullanıcı dokümanı belirlenmedi.")

            d1_note = selection.get("d1_change_note") or (
                "D1 değişti." if selection.get("d1_changed") else "D1 değişmedi."
            )
            d2_note = selection.get("d2_change_note") or (
                "D2 değişti." if selection.get("d2_changed") else "D2 değişmedi."
            )
            st.write(f"**{d1_note}**")
            st.write(f"**{d2_note}**")
            st.success(f"Nihai D1: {selection.get('d1',{}).get('number','')} | Nihai D2: {(selection.get('d2') or {}).get('number','-')}")
            st.info(selection.get("technical_opinion") or f"Bence buluş basamağı {'var' if selection.get('inventive_step_result') == 'sağlanır' else 'yok'}.")

            decision_mode = st.selectbox(
                "Rapor sonucunu nasıl oluşturayım?",
                ["Otomatik belirle", "Buluş basamağı var", "Buluş basamağı yok"],
                help="Bu seçim ancak 10+, nihai D1/D2 ve sistem kanaati gösterildikten sonra yapılır.",
            )

            if st.button("3. Ön Araştırma Raporunu oluştur", type="primary", use_container_width=True):
                try:
                    final_selection = deepcopy(selection)
                    if decision_mode == "Buluş basamağı var":
                        final_selection["inventive_step_result"] = "sağlanır"
                    elif decision_mode == "Buluş basamağı yok":
                        final_selection["inventive_step_result"] = "sağlanmaz"
                    progress = st.progress(0, text="Yenilik ve buluş basamağı raporu hazırlanıyor...")
                    report = ask_json(report_drafting_prompt(
                        st.session_state.research_bbf_text,
                        st.session_state.top10_result,
                        final_selection,
                        reference,
                        st.session_state.research_cutoff or cutoff.strftime("%d.%m.%Y"),
                        decision_mode,
                    ))
                    validate_report_against_selection(report, final_selection)
                    validate_research_report_language(report)
                    progress.progress(75, text="Bağlayıcı Word şablonu dolduruluyor...")
                    data = build_research_docx(report)
                    progress.progress(100, text="Hazır")
                    st.info(f"Yenilik: {final_selection.get('novelty_result','')} | Buluş basamağı: {final_selection.get('inventive_step_result','')}")
                    effective_output_name = output_name.replace("XXXXXX", reference.strip()) if reference.strip() else output_name
                    st.download_button("Word raporunu indir", data=data, file_name=safe_output_name(effective_output_name, "Ön Araştırma Raporu.docx"), mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
                except Exception as exc:
                    st.exception(exc)

# ARAŞTIRMA GÜNCELLEME
else:
    st.subheader("Araştırma güncelleme - Tip 3")
    st.caption("İlk araştırma konusu ile revize araştırma konusu karşılaştırılır; ilk rapordaki D1/D2 dikkate alınarak yeni araştırma yapılır ve nihai rapor standart Tip 3 Ön Araştırma Raporu formatında oluşturulur.")

    c1, c2 = st.columns(2)
    with c1:
        first_bbf = st.file_uploader("1. İlk BBF", type=["docx", "doc", "pdf", "txt"], key="upd_first_bbf")
        revised_bbf = st.file_uploader("2. Revize BBF", type=["docx", "doc", "pdf", "txt"], key="upd_revised_bbf")
        prior_report = st.file_uploader("3. İlk Ön Araştırma Raporu", type=["pdf", "docx", "doc", "txt"], key="upd_prior_report")
    with c2:
        update_reference = st.text_input("DP referans numarası", value="", key="upd_reference")
        update_output_name = st.text_input("Çıktı dosyasının adı", value="Ön Araştırma Raporu_XXXXXX_rev.docx", key="upd_output_name")
        update_cutoff = st.date_input("Araştırma kesim tarihi", value=date.today(), key="upd_cutoff")

    for key, default in {
        "update_analysis": None,
        "update_research": None,
        "update_first_text": None,
        "update_revised_text": None,
        "update_prior_report_text": None,
        "update_prior_report_asset": None,
    }.items():
        if key not in st.session_state:
            st.session_state[key] = default

    if st.button("1. Farkları ve teknik katkıyı analiz et", type="primary", use_container_width=True):
        if first_bbf is None:
            st.error("İlk BBF dosyasını yükleyin.")
        elif revised_bbf is None:
            st.error("Revize BBF dosyasını yükleyin.")
        elif prior_report is None:
            st.error("İlk Ön Araştırma Raporunu yükleyin.")
        else:
            try:
                progress = st.progress(0, text="İlk araştırma konusu okunuyor...")
                first_asset = UploadedAsset(first_bbf.name, first_bbf.getvalue(), first_bbf.type)
                revised_asset = UploadedAsset(revised_bbf.name, revised_bbf.getvalue(), revised_bbf.type)
                report_asset = UploadedAsset(prior_report.name, prior_report.getvalue(), prior_report.type)
                first_text = extract_text_from_asset(first_asset)
                progress.progress(25, text="Revize araştırma konusu okunuyor...")
                revised_text = extract_text_from_asset(revised_asset)
                progress.progress(45, text="İlk ön araştırma raporu ve D1/D2 gerekçeleri okunuyor...")
                report_text = extract_text_from_asset(report_asset)
                progress.progress(65, text="Teknik farklar ve katkılar karşılaştırılıyor...")
                analysis = ask_json(research_update_analysis_prompt(first_text, revised_text, report_text))
                st.session_state.update_analysis = analysis
                st.session_state.update_research = None
                st.session_state.update_first_text = first_text
                st.session_state.update_revised_text = revised_text
                st.session_state.update_prior_report_text = report_text
                st.session_state.update_prior_report_asset = report_asset
                progress.progress(100, text="Fark analizi tamamlandı")
            except Exception as exc:
                st.exception(exc)

    if st.session_state.update_analysis:
        analysis = st.session_state.update_analysis
        st.markdown("### Fark analizi")
        diff_rows = []
        for i, d in enumerate(analysis.get("differences") or [], 1):
            diff_rows.append({
                "No": i,
                "İlk araştırma konusu": d.get("old", ""),
                "Revize araştırma konusu": d.get("new", ""),
                "Teknik katkı": d.get("technical_contribution", ""),
                "Teknik etki": d.get("technical_effect", ""),
                "İlk D1/D2 karşısındaki etkisi": d.get("effect_against_prior_d1_d2", ""),
            })
        if diff_rows:
            st.dataframe(diff_rows, use_container_width=True, hide_index=True)
        st.write(f"**İlk rapordaki D1:** {(analysis.get('prior_d1') or {}).get('number','-')}  |  **D2:** {(analysis.get('prior_d2') or {}).get('number','-')}")
        st.info(analysis.get("preliminary_opinion", ""))

        if st.button("2. Revize konu için yeni patent araştırmasını yap", type="primary", use_container_width=True):
            try:
                progress = st.progress(0, text="Revize teknik farklara göre global araştırma yapılıyor...")
                research = ask_json(
                    research_update_search_prompt(
                        st.session_state.update_revised_text or "",
                        st.session_state.update_prior_report_text or "",
                        analysis,
                        update_cutoff.strftime("%d.%m.%Y"),
                    ),
                    web_search=True,
                )
                docs = research.get("documents") or []
                if len(docs) != 10:
                    raise ValueError(f"Araştırma güncellemede tam 10 doğrulanmış doküman beklenirken {len(docs)} doküman döndü. Araştırmayı tekrar çalıştırın.")
                temp_selection = {
                    "d1": research.get("d1") or {},
                    "d2": research.get("d2"),
                    "comparison_rows_d1": research.get("comparison_rows_d1") or [],
                    "comparison_rows_d2": research.get("comparison_rows_d2") or [],
                }
                validate_research_selection(temp_selection)
                st.session_state.update_research = research
                progress.progress(100, text="Yeni araştırma tamamlandı")
            except Exception as exc:
                st.exception(exc)

    if st.session_state.update_research:
        research = st.session_state.update_research
        st.markdown("### Yeni araştırma sonucu")
        new_docs = research.get("new_documents") or []
        if new_docs:
            st.write("**İlk raporda bulunmayan yeni yakın dokümanlar:**")
            st.dataframe([
                {
                    "Yayın no": d.get("number", ""),
                    "Başlık": d.get("title", ""),
                    "Tarih": d.get("date", ""),
                    "Teknik ilgisi": d.get("technical_relevance", ""),
                }
                for d in new_docs
            ], use_container_width=True, hide_index=True)
        else:
            st.caption("İlk rapordaki dokümanlardan daha yakın yeni bir doküman tespit edilmedi.")

        st.code(research.get("totalpatent_query", ""), language=None)
        st.write(f"**Önerilen nihai D1:** {(research.get('d1') or {}).get('number','-')}  |  **D2:** {(research.get('d2') or {}).get('number','-')}")
        st.write(f"**Yenilik ön sonucu:** {research.get('novelty_result','-')}  |  **Buluş basamağı ön sonucu:** {research.get('inventive_step_result','-')}")
        st.info(research.get("technical_opinion", ""))

        recommended = "Buluş basamağı sağlanmıyor" if str(research.get("inventive_step_result", "")).strip() == "sağlanmaz" else "Buluş basamağı sağlanıyor"
        update_decision = st.radio(
            "Raporu hangi sonuçla hazırlayayım?",
            ["Buluş basamağı sağlanıyor", "Buluş basamağı sağlanmıyor"],
            index=1 if recommended == "Buluş basamağı sağlanmıyor" else 0,
            horizontal=True,
            key="upd_decision",
        )
        st.caption(f"Sistem önerisi: {recommended}")

        if st.button("3. Ön Araştırma Raporunu oluştur", type="primary", use_container_width=True):
            if not update_reference.strip():
                st.error("DP referans numarasını girin.")
            else:
                try:
                    progress = st.progress(0, text="Standart Tip 3 rapor metni hazırlanıyor...")
                    report = ask_json(
                        research_update_report_prompt(
                            st.session_state.update_revised_text or "",
                            st.session_state.update_prior_report_text or "",
                            st.session_state.update_analysis or {},
                            research,
                            update_reference.strip(),
                            update_cutoff.strftime("%d.%m.%Y"),
                            update_decision,
                        )
                    )
                    validate_report_against_selection(report, {"d1": research.get("d1"), "d2": research.get("d2")})
                    validate_research_report_language(report)
                    progress.progress(55, text="D1/D2 özgün patent şekilleri temin ediliyor...")
                    fallbacks = _report_pdf_fallback_figures(st.session_state.update_prior_report_asset)
                    progress.progress(75, text="Bağlayıcı Ön Araştırma Raporu şablonu dolduruluyor...")
                    data = build_research_docx(report, figure_fallbacks=fallbacks)
                    progress.progress(100, text="Rapor hazır")
                    effective_output = update_output_name.replace("XXXXXX", update_reference.strip())
                    st.success("Güncelleme raporu standart Ön Araştırma Raporu formatında oluşturuldu.")
                    st.download_button(
                        "Word raporunu indir",
                        data=data,
                        file_name=safe_output_name(effective_output, "Ön Araştırma Raporu_rev.docx"),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True,
                    )
                except Exception as exc:
                    st.exception(exc)

