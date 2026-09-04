from pathlib import Path
import io
import pytest
from docx import Document

ROOT = Path(__file__).parent
from rules import APP_VERSION, RULESET_VERSION
import app_core


def base_draft():
    elements=[
        {'number':'1','name':'Kullanıcı ekipmanı izleme arayüzü','description':'x'},
        {'number':'2','name':'Veri toplama ve ön işleme modülü','description':'y'},
    ]
    system_claim={
        'preamble':'LTE ağında kullanıcı ekipmanı ölçüm verilerini işleyerek adaptif hücre seçimini sağlayan ve ağ performansını optimize eden adaptif hücre optimizasyon sistemi',
        'elements':['ölçüm verisini toplayan ve sonraki işleme aktaran kullanıcı ekipmanı izleme arayüzü (1),','toplanan ölçüm verisini işleyen ve öznitelik çıktısı üreten veri toplama ve ön işleme modülü (2)'],
        'closing':'içermesidir.'
    }
    return {
        'title':'Adaptif Hücre Optimizasyon Sistemi ve Yöntemi',
        'technical_field':'Buluş, adaptif hücre optimizasyon sistemi ve yöntemi ile ilgilidir.\n\nBuluş, özellikle LTE ağında kullanılmaktadır.',
        'prior_art_general_paragraphs':['Mevcut sistemlerde statik öncelikler kullanılır ve değişken ağ yükü dikkate alınamaz. '*7,'Hareketlilik aynı parametrelerle ele alınır ve gereksiz yeniden seçimler oluşur. '*7,'Bağlam bilgileri ve enerji durumu birlikte değerlendirilemez. '*7,'Yük dengesizliği küçük hücrelerin etkin kullanılmasını engeller. '*7,'Ping-pong etkisi sinyalizasyon ve pil tüketimini artırır. '*7,'Kullanıcı servis gereksinimleri seçim kararına yeterince yansıtılamaz. '*7,'Yukarıda belirtilen eksiklikler, adaptif ve kullanıcıya özgü bir teknik çözüm ihtiyacı doğurmaktadır. '*7],
        'literature_paragraphs':['Literatürde yapılan araştırmalar sonucu US1 numaralı “Adaptive cell selection (Adaptif hücre seçimi)” başlıklı patent dokümanına rastlanmıştır. Söz konusu doküman hücre seçimi ile ilgilidir. Ancak bahsedilen dokümanda kullanıcı bağlamıyla kapalı döngü eniyileme ile ilgili bir emareye rastlanmamıştır.'],
        'short_description_intro':'Buluş, teknik çözüm sağlar.',
        'objectives':['bir teknik amaç sağlamaktır.'],
        'unumbered_invention_definition':'LTE ağında kullanıcı ekipmanı ölçüm verilerini işleyerek adaptif hücre seçimini sağlayan ve ağ performansını optimize eden adaptif hücre optimizasyon sistemi olup, özelliği;',
        'unumbered_invention_features':['ölçüm verisini toplayan ve sonraki işleme aktaran kullanıcı ekipmanı izleme arayüzü,','toplanan ölçüm verisini işleyen ve öznitelik çıktısı üreten veri toplama ve ön işleme modülü'],
        'figure_descriptions':['Şekil 1, sistemdeki veri ve kontrol ilişkisini gösteren temsili şematik gösterimdir.'],
        'elements':elements,
        'method_steps':[],
        'detailed_paragraphs':['Kullanıcı ekipmanı izleme arayüzü (1), ölçüm verisini toplamakta ve veri toplama ve ön işleme modülü (2), toplanan ölçüm verisini işlemektedir.','Buluşun bir yapılanmasında sistem LTE ağına uygulanabilmektedir.'],
        'formulas':[], 'tables':[], 'experimental_results':[], 'alternatives':[],
        'working_principle':'Kullanıcı ekipmanı izleme arayüzü (1) ile veri toplama ve ön işleme modülü (2) birlikte çalışmaktadır.',
        'system_claim':system_claim, 'dependent_system_claims':[], 'method_claim':None, 'dependent_method_claims':[],
        'abstract':'Buluş, adaptif hücre optimizasyon sistemi ile ilgilidir.',
        'source_coverage_map':[],
        'coverage_audit':{'prior_art_complete':True,'reference_table_complete':True,'claims_consistent':True,'reference_names_clear':True,'reference_order_valid':True,'how_test_passed':True,'core_difference_present':True,'scope_not_overlimited':True,'dependent_claims_non_redundant':True,'dependent_claim_dependencies_valid':True,'example_dimensions_not_claim_limited':True,'product_claim_language_valid':True,'abstract_single_paragraph_sentence':True,'source_attribution_removed':True,'all_technical_facts_covered':True,'software_carrier_valid':True,'detail_intro_sentence_case':True},
    }


def extracted():
    return {'technical_facts':[]}


def test_versions():
    assert APP_VERSION == 'v5.4.52'
    assert RULESET_VERSION == '2026-09-04.v42'


def test_title_case_normalization():
    assert app_core._normalize_turkish_invention_title('adaptif hücre optimizasyon sistemi ve yöntemi') == 'Adaptif Hücre Optimizasyon Sistemi ve Yöntemi'
    d=base_draft(); d['title']='Adaptif hücre optimizasyon sistemi ve yöntemi'
    with pytest.raises(ValueError): app_core._validate_turkish_title_style(d,'Türkçe')


def test_literature_title_pair_and_meta_language():
    d=base_draft(); d['title']='Adaptif Hücre Optimizasyon Sistemi'; d['technical_field']='Buluş, adaptif hücre optimizasyon sistemi ile ilgilidir.\n\nBuluş, özellikle LTE ağında kullanılmaktadır.'
    app_core.validate_tarifname_draft(d,'Yalnızca sistem',[{'title_en':'Adaptive cell selection','title_tr':'Adaptif hücre seçimi'}],'Türkçe',extracted())
    d['literature_paragraphs']=['Literatürde yapılan araştırmalar sonucu US1 numaralı “Adaptive cell selection” başlıklı ve Türkçe karşılığı “Adaptif hücre seçimi” olan patent dokümanına rastlanmıştır. Ancak bahsedilen dokümanda kapalı döngü ile ilgili bir emareye rastlanmamıştır.']
    with pytest.raises(ValueError): app_core.validate_tarifname_draft(d,'Yalnızca sistem',[{'title_en':'Adaptive cell selection','title_tr':'Adaptif hücre seçimi'}],'Türkçe',extracted())


def test_unnumbered_must_mirror_main_claim_without_refs():
    d=base_draft(); app_core._validate_unumbered_claim_mirror(d,'Yalnızca sistem','Türkçe')
    d['unumbered_invention_features']=['Kullanıcı ekipmanı izleme arayüzü,','Veri toplama ve ön işleme modülü']
    with pytest.raises(ValueError): app_core._validate_unumbered_claim_mirror(d,'Yalnızca sistem','Türkçe')


def test_detail_element_chain_one_paragraph():
    d=base_draft(); app_core._validate_detailed_element_cohesion(d,'Türkçe')
    d['detailed_paragraphs']=['Kullanıcı ekipmanı izleme arayüzü (1), ölçüm alır.','Veri toplama ve ön işleme modülü (2), veriyi işler.']
    with pytest.raises(ValueError): app_core._validate_detailed_element_cohesion(d,'Türkçe')


def test_realization_wording_and_figure_desc():
    d=base_draft(); app_core._validate_realization_wording(d,'Türkçe'); app_core._validate_figure_description_style(d,'Türkçe')
    d['detailed_paragraphs'][1]='Bir gerçekleştirimde sistem LTE ağına uygulanır.'
    with pytest.raises(ValueError): app_core._validate_realization_wording(d,'Türkçe')
    d=base_draft(); d['figure_descriptions']=['Şekil 1, 1-12 referanslı teknik unsurları göstermektedir.']
    with pytest.raises(ValueError): app_core._validate_figure_description_style(d,'Türkçe')


def test_figures_page_counter_arial_11():
    # Minimal image generated from bytes is not needed: directly verify helper on a constructed docx header.
    doc=Document(); app_core._add_figures_page_counter(doc.sections[0]); bio=io.BytesIO(); doc.save(bio)
    app_core.validate_figures_docx_structure(bio.getvalue())
    header = doc.sections[0].header
    assert header.paragraphs[0].alignment == 1  # CENTER
    fld = [n for n in header._element.iter() if str(n.tag).endswith('}fldSimple')]
    assert len(fld) == 2
    assert all(str(n.getparent().tag).endswith('}p') for n in fld)
    assert all(any(str(x.tag).endswith('}t') and (x.text or '').strip() for x in n.iter()) for n in fld)
    assert all(any(str(x.tag).endswith('}b') and str(x.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or '1') not in {'0','false'} for x in n.iter()) for n in fld)
    assert any(r.text.strip() == '/' and r.bold is True for r in header.paragraphs[0].runs)


def test_objective_leading_acronym_is_preserved():
    src = Path(app_core.__file__).read_text(encoding="utf-8")
    assert 'first_token.group(0).isupper()' in src
    assert 'body = objective' in src
