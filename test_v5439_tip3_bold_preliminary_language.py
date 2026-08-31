from __future__ import annotations

import io

import pytest
from docx import Document
from PIL import Image

from app_core import build_research_docx, validate_research_report_language


def _png_bytes() -> bytes:
    im = Image.new('L', (900, 500), 255)
    out = io.BytesIO()
    im.save(out, format='PNG')
    return out.getvalue()


def _report() -> dict:
    common_rows = [
        {"feature": "çoklu veri kaynağından risk değerlendirmesi", "status_evidence": "+ Abstract; Claim 1"},
        {"feature": "olay kaydının oluşturulması", "status_evidence": "- Açık ve doğrudan açıklama bulunmamaktadır"},
    ]
    return {
        "reference": "182009",
        "title": "Araç İçi Risk Değerlendirme Sistemi",
        "report_date": "31.08.2026",
        "purpose": "Belirlenen konuda araştırmanın gerçekleştirilmesi",
        "scope": "Global (İlan edilmiş olan patent başvuruları)",
        "keywords": ["vehicle risk assessment", "driver monitoring"],
        "ipc_cpc": [
            {"code": "B60R 25/00", "description": "Fittings or systems for preventing or indicating unauthorised use or theft of vehicles"},
        ],
        "evaluation_intro": "Bu alan Word üreticisi tarafından deterministik oluşturulur.",
        "documents": [
            {
                "label": "D1", "number": "CN120088929A", "alternate_number": "", "title": "Driver emergency calling method", "date": "03.06.2025", "source_url": "", "figure_reference": "Fig. 1", "figure_image_url": "",
                "description": ["D1 dokümanı araç içi acil durum tespiti ile ilgilidir."],
                "abstract": "A vehicle emergency method is disclosed for detecting an emergency and sending information.",
                "figure_caption": "D1- Şekil 1",
                "comparison_rows": common_rows,
                "novelty_assessment": ["Bu nedenle araştırma konusu, D1 dokümanı karşısında bütün teknik özellikleri ve aralarındaki ilişki bakımından doğrudan ve açık biçimde açıklanmadığından yenilik kriterini sağlamaktadır."],
            },
            {
                "label": "D2", "number": "GB2608795A", "alternate_number": "", "title": "Vehicle monitoring system", "date": "18.01.2023", "source_url": "", "figure_reference": "Fig. 1", "figure_image_url": "",
                "description": ["D2 dokümanı araç izleme sistemi ile ilgilidir."],
                "abstract": "A vehicle monitoring system includes sensors and an alerting arrangement.",
                "figure_caption": "D2- Şekil 1",
                "comparison_rows": common_rows,
                "novelty_assessment": ["Bu nedenle araştırma konusu D2 karşısında yenilik kriterini sağladığı düşünülmektedir."],
            },
        ],
        "inventive_step_paragraphs": ["D1 ve D2 birlikte değerlendirildiğinde araştırma konusu buluş basamağı kriterini sağladığı düşünülmektedir."],
        "conclusion_paragraphs": ["Araştırma konusu yenilik kriterini sağladığı ve buluş basamağı kriterini sağladığı düşünülmektedir."],
        "warnings": ["Uyarı 1", "Uyarı 2", "Uyarı 3", "Uyarı 4"],
        "attachments": ["Benzer Dokümanlar", "Ön İnceleme Raporu", "Makine Tercümeleri"],
    }


def test_direct_language_gate_rejects_categorical_patentability_statement():
    report = _report()
    with pytest.raises(ValueError, match="ön araştırma raporudur"):
        validate_research_report_language(report)


def test_build_softens_result_language_and_bolds_d1_d2_identity():
    report = _report()
    data = build_research_docx(report, figure_fallbacks=[_png_bytes(), _png_bytes()])
    doc = Document(io.BytesIO(data))

    intro = doc.paragraphs[36]
    bold_text = "".join(r.text for r in intro.runs if r.bold is True)
    assert bold_text == "CN120088929A (D1) ve GB2608795A (D2)"
    assert "teknik yakınlığı en yüksek dokümanlar" in intro.text

    d1_assessment = doc.paragraphs[55].text
    assert "yenilik kriterini sağladığı düşünülmektedir" in d1_assessment
    assert "yenilik kriterini sağlamaktadır" not in d1_assessment
