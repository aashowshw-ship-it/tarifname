from __future__ import annotations

import io
import json
import re
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

from gorus_audit import validate_minimal_tracked_changes


ALLOWED_DECISIONS = {
    "apply",
    "partial",
    "explain",
    "clarification",
    "figure_action",
    "procedural_action",
}
ALLOWED_OPERATION_TYPES = {"replace_text", "insert_paragraph_after", "insert_paragraph_before"}
ALLOWED_BASIS_SOURCES = {"existing_spec", "customer_request"}


def derive_markup_output_name(source_name: str) -> str:
    """Return a stable markup filename while removing browser duplicate suffixes such as (1)."""
    stem = Path(source_name or "Description").stem.strip() or "Description"
    stem = re.sub(r"\s*\(\d+\)\s*$", "", stem).strip()
    stem = re.sub(r"_markup(?:_final)?$", "", stem, flags=re.I).strip()
    return f"{stem}_markup.docx"


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _find_relaxed_phrase_span(text: str, phrase: str) -> tuple[int, int] | None:
    tokens = re.findall(r"\S+", str(phrase or "").strip())
    if not tokens:
        return None
    pattern = r"\s+".join(re.escape(token) for token in tokens)
    match = re.search(pattern, text)
    return (match.start(), match.end()) if match else None


def _ooxml_run_is_customer_red(run: Any, W: str) -> bool:
    """Recognize explicit red-font customer notes without treating ordinary black text as review content."""
    color = run.find(f"{{{W}}}rPr/{{{W}}}color")
    if color is None:
        return False
    value = str(color.get(f"{{{W}}}val", "") or "").strip().upper().lstrip("#")
    # Strict red shades only. Theme/automatic/black text is never stripped.
    return value in {"FF0000", "C00000", "E60000", "D90000", "CC0000"}


def extract_docx_review_context(data: bytes) -> str:
    """Extract comments, tracked changes and explicit red-font customer notes from a returned DOCX.

    This is intentionally review-only context; it does not replace the clean specification text.
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": W}
    blocks: list[str] = []
    try:
        from lxml import etree

        with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
            comments: dict[str, str] = {}
            if "word/comments.xml" in zf.namelist():
                croot = etree.fromstring(zf.read("word/comments.xml"))
                for c in croot.xpath(".//w:comment", namespaces=ns):
                    cid = c.get(f"{{{W}}}id", "")
                    ctext = _norm(" ".join(c.xpath(".//w:t/text()", namespaces=ns)))
                    if cid and ctext:
                        comments[cid] = ctext
            root = etree.fromstring(zf.read("word/document.xml"))
            for p in root.xpath(".//w:p", namespaces=ns):
                normal_parts = p.xpath("./w:r/w:t/text()", namespaces=ns)
                inserted = [_norm(x) for x in p.xpath(".//w:ins//w:t/text()", namespaces=ns) if _norm(x)]
                deleted = [_norm(x) for x in p.xpath(".//w:del//w:delText/text()", namespaces=ns) if _norm(x)]
                cids = [x.get(f"{{{W}}}id", "") for x in p.xpath(".//w:commentRangeStart", namespaces=ns)]
                red_parts: list[str] = []
                for run in p.xpath(".//w:r", namespaces=ns):
                    if not _ooxml_run_is_customer_red(run, W):
                        continue
                    # Tracked insertions/deletions are already represented separately.
                    if run.xpath("ancestor::w:ins|ancestor::w:del", namespaces=ns):
                        continue
                    text = _norm("".join(run.xpath(".//w:t/text()", namespaces=ns)))
                    if text:
                        red_parts.append(text)
                if not (inserted or deleted or red_parts or any(cid in comments for cid in cids)):
                    continue
                para_text = _norm("".join(normal_parts))
                if para_text:
                    blocks.append(f"PARAGRAPH: {para_text}")
                for text in red_parts:
                    blocks.append(f"RED CUSTOMER NOTE: {text}")
                for text in deleted:
                    blocks.append(f"TRACK DELETE: {text}")
                for text in inserted:
                    blocks.append(f"TRACK INSERT: {text}")
                for cid in cids:
                    if cid in comments:
                        blocks.append(f"WORD COMMENT: {comments[cid]}")
    except Exception:
        return ""
    return "\n".join(blocks)


def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    seen_cells: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            marker = id(cell._tc)
            if marker in seen_cells:
                continue
            seen_cells.add(marker)
            for block in cell.iter_inner_content():
                if isinstance(block, Paragraph):
                    yield block
                elif isinstance(block, Table):
                    yield from _iter_table_paragraphs(block)


def iter_document_paragraphs(doc: Document) -> list[Paragraph]:
    out: list[Paragraph] = []
    for block in doc.iter_inner_content():
        if isinstance(block, Paragraph):
            out.append(block)
        elif isinstance(block, Table):
            out.extend(_iter_table_paragraphs(block))
    return out


def document_text(doc_or_bytes: Document | bytes) -> str:
    doc = Document(io.BytesIO(doc_or_bytes)) if isinstance(doc_or_bytes, (bytes, bytearray)) else doc_or_bytes
    return "\n".join(p.text for p in iter_document_paragraphs(doc) if p.text.strip())


def _paragraph_matches(paragraphs: list[Paragraph], phrase: str) -> list[Paragraph]:
    return [p for p in paragraphs if _find_relaxed_phrase_span(p.text, phrase) is not None]


def _quote_is_supported(quote: str, source_text: str) -> bool:
    if not _norm(quote):
        return False
    return _find_relaxed_phrase_span(source_text, quote) is not None


def _claims_section_text(source_text: str) -> str:
    text = str(source_text or "")
    match = re.search(r"\bİSTEMLER\b", text, flags=re.I)
    if not match:
        return ""
    tail = text[match.end():]
    end = re.search(r"\bÖZET\b", tail, flags=re.I)
    return tail[: end.start()] if end else tail


def _validate_explicit_claim_visibility(plan: dict[str, Any], source_text: str, customer_text: str) -> None:
    """Deterministic backstop for 'make it explicit in the claims' customer requests.

    The AI may correctly notice that a test/function is semantically present yet still fail to make
    the customer-requested terminology visible. If the customer explicitly asks for claim visibility,
    every technical acronym that is both requested and already supported somewhere in the source
    specification must either already occur in the claims or be introduced by a planned operation.
    Unsupported acronyms are deliberately ignored here; they remain subject to the normal basis gate.
    """
    customer = str(customer_text or "")
    trigger = bool(
        re.search(r"istem(?:ler|lere|lerde|e|in)?", customer, flags=re.I)
        and re.search(r"açıkça|vurgula|vurgulan|görün|ekle|eklen", customer, flags=re.I)
    )
    if not trigger:
        return
    acronyms = sorted(set(re.findall(r"(?<![A-Za-zÇĞİÖŞÜçğıöşü0-9])([A-ZÇĞİÖŞÜ]{2,8})(?![A-Za-zÇĞİÖŞÜçğıöşü0-9])", customer)))
    if not acronyms:
        return
    claims_text = _claims_section_text(source_text)
    planned_new = "\n".join(str(op.get("new_text", "")) for op in (plan.get("operations") or []))
    missing: list[str] = []
    for acronym in acronyms:
        # Only enforce terminology that has technical basis in the existing specification.
        if re.search(rf"(?<![A-Za-z0-9]){re.escape(acronym)}(?![A-Za-z0-9])", source_text) is None:
            continue
        already_visible = re.search(rf"(?<![A-Za-z0-9]){re.escape(acronym)}(?![A-Za-z0-9])", claims_text) is not None
        planned_visible = re.search(rf"(?<![A-Za-z0-9]){re.escape(acronym)}(?![A-Za-z0-9])", planned_new) is not None
        if not (already_visible or planned_visible):
            missing.append(acronym)
    if missing:
        raise ValueError(
            "Tarifname düzenleme istem-görünürlük kapısı: müşteri istemlerde açıkça vurgulanmasını istediği ve "
            "mevcut tarifnamede dayanağı bulunan terminolojinin bir kısmı planlanan istem değişikliklerinde görünmüyor: "
            + ", ".join(missing)
        )


def validate_update_plan(plan: dict[str, Any], source_docx: bytes, customer_text: str, filing_status: str) -> None:
    requests = list(plan.get("requests") or [])
    if not requests:
        raise ValueError("Tarifname düzenleme kalite kapısı: müşteri talep/soru envanteri boş.")
    ids = [str(r.get("id", "")).strip() for r in requests]
    if any(not x for x in ids) or len(ids) != len(set(ids)):
        raise ValueError("Tarifname düzenleme kalite kapısı: her müşteri talebinin benzersiz bir id değeri olmalıdır.")
    by_id = {str(r.get("id", "")).strip(): r for r in requests}

    operations = list(plan.get("operations") or [])
    comments = list(plan.get("comments") or [])
    figure_actions = list(plan.get("figure_actions") or [])
    open_questions = list(plan.get("blocking_clarifications") or [])

    for r in requests:
        rid = str(r.get("id", "")).strip()
        decision = str(r.get("decision", "")).strip()
        if decision not in ALLOWED_DECISIONS:
            raise ValueError(f"Tarifname düzenleme kalite kapısı: {rid} için geçersiz karar `{decision}`.")
        answer = _norm(r.get("answer_for_customer", ""))
        linked_ops = [x for x in operations if str(x.get("request_id", "")).strip() == rid]
        linked_comments = [x for x in comments if str(x.get("request_id", "")).strip() == rid]
        linked_fig = [x for x in figure_actions if str(x.get("request_id", "")).strip() == rid]
        if decision in {"apply", "partial"} and not linked_ops:
            raise ValueError(f"Tarifname düzenleme kalite kapısı: {rid} uygulanacak/kısmen uygulanacak denmiş fakat Word revizyon işlemi yok.")
        if decision == "explain" and not (answer or linked_comments):
            raise ValueError(f"Tarifname düzenleme kalite kapısı: {rid} yalnız açıklanacak denmiş fakat müşteri cevabı/Word yorumu yok.")
        if decision == "figure_action" and not (linked_fig or answer):
            raise ValueError(f"Tarifname düzenleme kalite kapısı: {rid} şekil aksiyonu gerektiriyor fakat şekil önerisi/cevap yok.")
        if decision == "clarification" and not (answer or open_questions):
            raise ValueError(f"Tarifname düzenleme kalite kapısı: {rid} için gerekli netleştirme sorusu bulunmuyor.")
        if decision == "procedural_action" and not answer:
            raise ValueError(f"Tarifname düzenleme kalite kapısı: {rid} usuli/stratejik konu için müşteri cevabı bulunmuyor.")

    doc = Document(io.BytesIO(source_docx))
    paragraphs = iter_document_paragraphs(doc)
    source_text = document_text(doc)
    _validate_explicit_claim_visibility(plan, source_text, customer_text)

    for op in operations:
        rid = str(op.get("request_id", "")).strip()
        if rid not in by_id:
            raise ValueError(f"Tarifname düzenleme kalite kapısı: revizyon işleminin bilinmeyen request_id değeri var: {rid}")
        op_type = str(op.get("type", "")).strip()
        if op_type not in ALLOWED_OPERATION_TYPES:
            raise ValueError(f"Tarifname düzenleme kalite kapısı: geçersiz Word işlemi `{op_type}`.")
        basis_source = str(op.get("basis_source", "")).strip()
        if basis_source not in ALLOWED_BASIS_SOURCES:
            raise ValueError(f"Tarifname düzenleme kalite kapısı: {rid} için basis_source geçersiz.")
        basis_quote = str(op.get("basis_quote", "")).strip()
        basis_pool = source_text if basis_source == "existing_spec" else customer_text
        if not _quote_is_supported(basis_quote, basis_pool):
            raise ValueError(f"Tarifname düzenleme kaynak kapısı: {rid} dayanak alıntısı `{basis_source}` içinde doğrulanamadı.")

        if filing_status == "Başvuru yapıldı" and basis_source == "customer_request":
            raise ValueError(
                f"Tarifname düzenleme new-matter kapısı: {rid} yalnız müşteri dönüşündeki yeni teknik bilgiye dayanıyor; "
                "başvuru yapılmış dosyada otomatik olarak metne eklenemez."
            )
        if filing_status.startswith("Rüçhan başvurusu") and basis_source == "customer_request":
            raise ValueError(
                f"Tarifname düzenleme rüçhan kapısı: {rid} yalnız sonraki müşteri bilgisine dayanıyor. "
                "Bu özelliğin sonraki başvuruya eklenmesi rüçhan etkisi bakımından kullanıcı kararı gerektirir."
            )

        if op_type == "replace_text":
            locator = str(op.get("locator_text", "")).strip()
            old_text = str(op.get("old_text", "")).strip()
            if not locator or not old_text:
                raise ValueError(f"Tarifname düzenleme kalite kapısı: {rid} replace_text için locator_text ve old_text zorunludur.")
            matches = _paragraph_matches(paragraphs, locator)
            if len(matches) != 1:
                raise ValueError(
                    f"Tarifname düzenleme hedef kapısı: {rid} locator_text tek bir paragrafı göstermiyor (eşleşme={len(matches)})."
                )
            if _find_relaxed_phrase_span(matches[0].text, old_text) is None:
                raise ValueError(f"Tarifname düzenleme hedef kapısı: {rid} old_text hedef paragrafta bulunamadı.")
        else:
            anchor = str(op.get("anchor_text", "")).strip()
            if not anchor or not _norm(op.get("new_text", "")):
                raise ValueError(f"Tarifname düzenleme kalite kapısı: {rid} paragraf ekleme için anchor_text ve new_text zorunludur.")
            matches = _paragraph_matches(paragraphs, anchor)
            if len(matches) != 1:
                raise ValueError(
                    f"Tarifname düzenleme hedef kapısı: {rid} anchor_text tek bir paragrafı göstermiyor (eşleşme={len(matches)})."
                )

    for fa in figure_actions:
        rid = str(fa.get("request_id", "")).strip()
        if rid not in by_id:
            raise ValueError(f"Tarifname düzenleme şekil kapısı: bilinmeyen request_id: {rid}")
        figure = str(fa.get("figure", "")).strip()
        recommended_change = _norm(fa.get("recommended_change", ""))
        if not figure or not recommended_change:
            raise ValueError(f"Tarifname düzenleme şekil kapısı: {rid} için figure/recommended_change zorunludur.")
        if bool(fa.get("safe_auto_edit")):
            basis_source = str(fa.get("basis_source", "")).strip()
            if basis_source not in ALLOWED_BASIS_SOURCES:
                raise ValueError(f"Tarifname düzenleme şekil kapısı: {rid} otomatik şekil revizyonu için basis_source geçersiz.")
            basis_quote = str(fa.get("basis_quote", "")).strip()
            basis_pool = source_text if basis_source == "existing_spec" else customer_text
            if not _quote_is_supported(basis_quote, basis_pool):
                raise ValueError(f"Tarifname düzenleme şekil kaynak kapısı: {rid} dayanak alıntısı `{basis_source}` içinde doğrulanamadı.")
            if not _norm(fa.get("edit_instructions", "")):
                raise ValueError(f"Tarifname düzenleme şekil kapısı: {rid} safe_auto_edit=true ise edit_instructions zorunludur.")
            if filing_status == "Başvuru yapıldı" and basis_source == "customer_request":
                raise ValueError(
                    f"Tarifname düzenleme şekil new-matter kapısı: {rid} yalnız müşteri dönüşündeki yeni teknik bilgiye dayanıyor; "
                    "başvuru yapılmış dosyada otomatik şekil revizyonu yapılamaz."
                )
            if filing_status.startswith("Rüçhan başvurusu") and basis_source == "customer_request":
                raise ValueError(
                    f"Tarifname düzenleme şekil rüçhan kapısı: {rid} yalnız sonraki müşteri bilgisine dayanıyor; "
                    "otomatik şekil revizyonu için kullanıcı kararı gerekir."
                )

    for c in comments:
        rid = str(c.get("request_id", "")).strip()
        if rid not in by_id:
            raise ValueError(f"Tarifname düzenleme yorum kapısı: bilinmeyen request_id: {rid}")
        anchor = str(c.get("anchor_text", "")).strip()
        text = _norm(c.get("text", ""))
        if not anchor or not text:
            raise ValueError(f"Tarifname düzenleme yorum kapısı: {rid} yorumunda anchor_text/text eksik.")
        matches = _paragraph_matches(paragraphs, anchor)
        if len(matches) != 1:
            raise ValueError(f"Tarifname düzenleme yorum kapısı: {rid} yorum anchor'ı tek paragraf göstermiyor.")

    if plan.get("coverage_complete") is not True:
        raise ValueError("Tarifname düzenleme ikinci okuma kapısı: müşteri taleplerinin tamamının karşılandığı doğrulanmadı.")


def _run_style_spans(paragraph: Paragraph) -> list[tuple[int, int, Any]]:
    spans: list[tuple[int, int, Any]] = []
    cursor = 0
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        rpr = run._r.find(qn("w:rPr"))
        spans.append((cursor, cursor + len(text), deepcopy(rpr) if rpr is not None else None))
        cursor += len(text)
    return spans


def _style_at(spans: list[tuple[int, int, Any]], position: int) -> Any:
    for start, end, rpr in spans:
        if start <= position < end:
            return deepcopy(rpr) if rpr is not None else None
    if spans:
        rpr = spans[-1][2]
        return deepcopy(rpr) if rpr is not None else None
    return None


def _append_plain_run(parent, text: str, rpr: Any = None, *, deleted: bool = False) -> None:
    if not text:
        return
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(deepcopy(rpr))
    text_node = OxmlElement("w:delText" if deleted else "w:t")
    text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    parent.append(run)


def _append_revision(parent, text: str, *, kind: str, change_id: int, rpr: Any = None) -> None:
    if not text:
        return
    wrapper = OxmlElement("w:del" if kind == "delete" else "w:ins")
    wrapper.set(qn("w:id"), str(change_id))
    wrapper.set(qn("w:author"), "Destek Patent")
    wrapper.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    _append_plain_run(wrapper, text, rpr, deleted=(kind == "delete"))
    parent.append(wrapper)


def _append_unchanged_with_styles(parent, original: str, start: int, end: int, spans: list[tuple[int, int, Any]]) -> None:
    if end <= start:
        return
    cursor = start
    for span_start, span_end, rpr in spans:
        overlap_start = max(cursor, span_start)
        overlap_end = min(end, span_end)
        if overlap_end <= overlap_start:
            continue
        if overlap_start > cursor:
            _append_plain_run(parent, original[cursor:overlap_start], None)
        _append_plain_run(parent, original[overlap_start:overlap_end], rpr)
        cursor = overlap_end
        if cursor >= end:
            break
    if cursor < end:
        _append_plain_run(parent, original[cursor:end], _style_at(spans, cursor))


def _minimal_markup_parts(old_text: str, new_text: str) -> tuple[str, str, str, str]:
    tok_re = re.compile(r"\s+|[^\s]+")
    old_tokens = tok_re.findall(old_text)
    new_tokens = tok_re.findall(new_text)
    i = 0
    while i < min(len(old_tokens), len(new_tokens)) and old_tokens[i] == new_tokens[i]:
        i += 1
    j = 0
    while j < len(old_tokens) - i and j < len(new_tokens) - i and old_tokens[-1-j] == new_tokens[-1-j]:
        j += 1
    prefix = "".join(old_tokens[:i])
    old_mid = "".join(old_tokens[i:len(old_tokens)-j if j else len(old_tokens)])
    new_mid = "".join(new_tokens[i:len(new_tokens)-j if j else len(new_tokens)])
    suffix = "".join(old_tokens[len(old_tokens)-j:]) if j else ""
    if old_mid and new_mid and not re.search(r"\s", old_mid) and not re.search(r"\s", new_mid):
        cp = 0
        while cp < min(len(old_mid), len(new_mid)) and old_mid[cp] == new_mid[cp]:
            cp += 1
        cs = 0
        while cs < len(old_mid)-cp and cs < len(new_mid)-cp and old_mid[-1-cs] == new_mid[-1-cs]:
            cs += 1
        prefix += old_mid[:cp]
        suffix = (old_mid[len(old_mid)-cs:] if cs else "") + suffix
        old_mid = old_mid[cp:len(old_mid)-cs if cs else len(old_mid)]
        new_mid = new_mid[cp:len(new_mid)-cs if cs else len(new_mid)]
    return prefix, old_mid, new_mid, suffix


def _rewrite_paragraph_with_changes(paragraph: Paragraph, operations: list[dict[str, Any]], *, track_changes: bool, id_start: int) -> int:
    original = paragraph.text
    style_spans = _run_style_spans(paragraph)
    targets: list[tuple[int, int, dict[str, Any]]] = []
    for op in operations:
        old_text = str(op.get("old_text", "")).strip()
        span = _find_relaxed_phrase_span(original, old_text)
        if span is None:
            raise ValueError(f"Markup hedefi uygulanamadı: {old_text[:100]}")
        targets.append((span[0], span[1], op))
    targets.sort(key=lambda x: x[0])
    for left, right in zip(targets, targets[1:]):
        if right[0] < left[1]:
            raise ValueError("Aynı paragrafta çakışan iki tarifname revizyonu bulundu.")

    p_el = paragraph._p
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)
    cursor = 0
    change_id = id_start
    for start, end, op in targets:
        old_actual = original[start:end]
        new_text = str(op.get("new_text", ""))
        _append_unchanged_with_styles(p_el, original, cursor, start, style_spans)
        prefix, deleted_mid, inserted_mid, suffix = _minimal_markup_parts(old_actual, new_text)
        if prefix:
            _append_unchanged_with_styles(p_el, original, start, start + len(prefix), style_spans)
        change_pos = start + len(prefix)
        change_style = _style_at(style_spans, min(change_pos, max(start, end - 1)))
        if track_changes:
            if deleted_mid:
                _append_revision(p_el, deleted_mid, kind="delete", change_id=change_id, rpr=change_style)
                change_id += 1
            if inserted_mid:
                _append_revision(p_el, inserted_mid, kind="insert", change_id=change_id, rpr=change_style)
                change_id += 1
        elif inserted_mid:
            _append_plain_run(p_el, inserted_mid, change_style)
        if suffix:
            _append_unchanged_with_styles(p_el, original, end - len(suffix), end, style_spans)
        cursor = end
    _append_unchanged_with_styles(p_el, original, cursor, len(original), style_spans)
    return change_id


def _enable_track_revisions(doc: Document) -> None:
    settings = doc.settings._element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.insert(0, OxmlElement("w:trackRevisions"))


def _insert_paragraph_relative(anchor: Paragraph, text: str, *, before: bool, track_changes: bool, change_id: int) -> tuple[Paragraph, int]:
    new_p = OxmlElement("w:p")
    ppr = anchor._p.find(qn("w:pPr"))
    if ppr is not None:
        new_p.append(deepcopy(ppr))
    base_rpr = None
    for run in anchor.runs:
        if run.text:
            rpr = run._r.find(qn("w:rPr"))
            base_rpr = deepcopy(rpr) if rpr is not None else None
            break
    if track_changes:
        _append_revision(new_p, text, kind="insert", change_id=change_id, rpr=base_rpr)
        change_id += 1
    else:
        _append_plain_run(new_p, text, base_rpr)
    if before:
        anchor._p.addprevious(new_p)
    else:
        anchor._p.addnext(new_p)
    return Paragraph(new_p, anchor._parent), change_id


def _add_comment(doc: Document, anchor_para: Paragraph, text: str) -> None:
    runs = [r for r in anchor_para.runs if (r.text or "").strip()]
    if not runs:
        raise ValueError("Word yorumu için değişmeden kalmış yorum anchor run'ı bulunamadı.")
    if not hasattr(doc, "add_comment"):
        raise ValueError("Kurulu python-docx sürümü Word comment API'sini desteklemiyor.")
    doc.add_comment(runs=[runs[0]], text=text, author="Destek Patent", initials="DP")


def build_updated_spec_docx(source_docx: bytes, plan: dict[str, Any], *, track_changes: bool, add_comments: bool = False) -> bytes:
    doc = Document(io.BytesIO(source_docx))
    if track_changes:
        _enable_track_revisions(doc)
    change_id = 1

    original_paragraphs = iter_document_paragraphs(doc)
    replace_assignments: dict[int, tuple[Paragraph, list[dict[str, Any]]]] = {}
    insert_ops: list[tuple[Paragraph, dict[str, Any]]] = []

    for op in plan.get("operations") or []:
        op_type = str(op.get("type", "")).strip()
        if op_type == "replace_text":
            matches = _paragraph_matches(original_paragraphs, str(op.get("locator_text", "")))
            if len(matches) != 1:
                raise ValueError("Tarifname markup üretiminde replace_text hedefi benzersiz değil.")
            key = id(matches[0]._p)
            if key not in replace_assignments:
                replace_assignments[key] = (matches[0], [])
            replace_assignments[key][1].append(op)
        else:
            matches = _paragraph_matches(original_paragraphs, str(op.get("anchor_text", "")))
            if len(matches) != 1:
                raise ValueError("Tarifname markup üretiminde paragraf ekleme anchor'ı benzersiz değil.")
            insert_ops.append((matches[0], op))

    # Existing text is revised before structural insertions. This keeps locator matching tied to the source file.
    for paragraph, ops in [value for value in replace_assignments.values()]:
        change_id = _rewrite_paragraph_with_changes(paragraph, ops, track_changes=track_changes, id_start=change_id)

    # Multiple insertions after one anchor preserve plan order; before insertions preserve plan order as well.
    after_tail: dict[int, Paragraph] = {}
    before_head: dict[int, Paragraph] = {}
    for anchor, op in insert_ops:
        op_type = str(op.get("type", "")).strip()
        text = str(op.get("new_text", ""))
        key = id(anchor._p)
        if op_type == "insert_paragraph_after":
            actual_anchor = after_tail.get(key, anchor)
            new_para, change_id = _insert_paragraph_relative(
                actual_anchor, text, before=False, track_changes=track_changes, change_id=change_id
            )
            after_tail[key] = new_para
        else:
            # Inserting before the same original anchor repeatedly naturally retains plan order
            # by inserting each new paragraph after the previous inserted head.
            if key in before_head:
                new_para, change_id = _insert_paragraph_relative(
                    before_head[key], text, before=False, track_changes=track_changes, change_id=change_id
                )
            else:
                new_para, change_id = _insert_paragraph_relative(
                    anchor, text, before=True, track_changes=track_changes, change_id=change_id
                )
            before_head[key] = new_para

    if add_comments and plan.get("comments"):
        paragraphs = iter_document_paragraphs(doc)
        for comment in plan.get("comments") or []:
            matches = _paragraph_matches(paragraphs, str(comment.get("anchor_text", "")))
            if len(matches) != 1:
                raise ValueError("Word comment anchor'ı revizyon sonrası benzersiz olarak bulunamadı.")
            _add_comment(doc, matches[0], str(comment.get("text", "")))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def validate_update_result(source_docx: bytes, markup_docx: bytes, accepted_docx: bytes, plan: dict[str, Any]) -> None:
    if plan.get("operations"):
        validate_minimal_tracked_changes(markup_docx)

    original = Document(io.BytesIO(source_docx))
    marked = Document(io.BytesIO(markup_docx))
    if len(original.sections) != len(marked.sections):
        raise ValueError("Tarifname düzenleme format kapısı: bölüm/section sayısı değişti.")
    attrs = ["page_width", "page_height", "top_margin", "bottom_margin", "left_margin", "right_margin"]
    for idx, (a, b) in enumerate(zip(original.sections, marked.sections), 1):
        for attr in attrs:
            if getattr(a, attr) != getattr(b, attr):
                raise ValueError(f"Tarifname düzenleme format kapısı: section {idx} `{attr}` değişti.")

    accepted_text = document_text(accepted_docx)
    for op in plan.get("operations") or []:
        new_text = _norm(op.get("new_text", ""))
        if new_text and _find_relaxed_phrase_span(accepted_text, new_text) is None:
            raise ValueError(
                "Tarifname düzenleme çıktı kapısı: planlanan yeni metin clean kabul görünümünde doğrulanamadı: "
                + new_text[:120]
            )

    expected_comments = len(plan.get("comments") or [])
    if expected_comments:
        with zipfile.ZipFile(io.BytesIO(markup_docx), "r") as zf:
            if "word/comments.xml" not in zf.namelist():
                raise ValueError("Tarifname düzenleme yorum kapısı: planlanan Word yorumları çıktı dosyasında bulunamadı.")
            from lxml import etree
            root = etree.fromstring(zf.read("word/comments.xml"))
            W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            actual_comments = len(root.xpath(".//w:comment", namespaces={"w": W}))
            if actual_comments < expected_comments:
                raise ValueError("Tarifname düzenleme yorum kapısı: Word yorumlarının bir kısmı çıktı dosyasına yazılmadı.")


def tarifname_update_analysis_prompt(
    rules: str,
    spec_text: str,
    customer_text: str,
    filing_status: str,
    supplemental_text: str = "",
    user_instruction: str = "",
) -> str:
    return f"""{rules}

TARİFNAME DÜZENLEME İŞ AKIŞI — İLK ANALİZ
Başvuru durumu: {filing_status}
Kullanıcının ek yönlendirmesi: {user_instruction or '(yok)'}

Görevin mevcut tarifnameyi sıfırdan yeniden yazmak değildir. Müşterinin HER talep, soru, itiraz ve önerisini ayrı ayrı çıkar; teknik ve hukuki açıdan değerlendir; yalnız gerçekten uygun olan değişiklikleri öner. Bir kelime/ibare yeterliyse tüm cümleyi değiştirme. Her metinsel değişiklik mevcut Word paragrafında minimum Track Changes farkı üretmeye uygun olmalıdır.

ÖNEMLİ:
- `locator_text`, mevcut tarifnamedeki TEK bir paragrafı benzersiz gösterecek yeterli ama gereksiz uzun olmayan birebir/çok yakın pasajdır.
- `old_text`, yalnız değiştirilecek en küçük mevcut ifade olmalıdır. Yeni paragraf gerçekten gerekliyse replace_text ile tüm paragrafı yeniden yazma; insert_paragraph_after/before kullan.
- `basis_source=existing_spec` ise basis_quote mevcut tarifnamede; `basis_source=customer_request` ise basis_quote müşteri dönüşünde birebir bulunmalıdır.
- Başvuru yapılmışsa yalnız müşteri dönüşünde ortaya çıkan yeni teknik içeriği otomatik ekleme; clarification olarak işaretle.
- Rüçhan başvurusu yapılmış ve sonraki başvuru hazırlanıyorsa yalnız sonraki müşteri bilgisindeki yeni teknik içeriği otomatik ekleme; rüçhan etkisi için clarification olarak işaretle.
- Müşterinin önerdiği claim wording bağlayıcı değildir. Teknik niyeti koruyup patent dilini sen belirle.
- Müşteri bir teknik işlev/terim/kısaltmanın istemlerde AÇIKÇA görünmesini veya vurgulanmasını istiyorsa ve aynı teknik içerik mevcut tarifnamede zaten anlam olarak destekleniyorsa bunu `zaten var` diyerek sessizce geçme. En küçük kelime/ibare değişikliğiyle görünür hale getir. Tam teknik ad ile kısaltmayı birlikte kullanmak uygunsa `tam ad (KISALTMA)` biçimini tercih et. Müşteri iki mevcut alternatifi slash ile görünür istemişse teknik belirsizlik yaratmıyorsa `tam ad (A) / tam ad (B)` biçimi kullanılabilir.
- Aynı müşteri maddesinde birden çok test/işlev/özellik sayılmışsa her birini ayrı ayrı dayanak açısından kontrol et. Bazıları istemde, bazıları yalnız detaylı açıklamada destekleniyor olabilir; desteklenenleri uygun isteme minimum müdahaleyle açıkça taşı, desteklenmeyeni ekleme.
- Bağımsız istem gereksiz daraltılmaz; tercihli implementation ayrıntıları dayanak varsa bağımlı isteme taşınabilir.
- Buluş bütünlüğü, PCT/EP stratejisi, ISA, rüçhan, maliyet gibi usuli soruları tarifnameye zorla yazma; mail cevabı/stratejik açıklama olarak ele al.
- Şekiller verilmişse tarifnameyle teknik uyumunu değerlendir ve figure_actions üret. Yalnız kaynakla desteklenen, hedef şekli açıkça belirli ve sınırlı bir revizyon (ör. mevcut iki unsur arasında istenen kablo/bağlantı çizgisi, mevcut monitör üzerinde istenen PLR/MFC düğmesi, referans/ok düzeltmesi) güvenle otomatik uygulanabiliyorsa `safe_auto_edit=true` ver; aksi halde false ver.
- `safe_auto_edit=true` için `basis_source`, doğrulanabilir `basis_quote` ve tek anlamlı `edit_instructions` zorunludur. Başvuru sonrası yalnız müşteri talebinden gelen yeni teknik geometri safe_auto_edit olamaz.
- Word comment yalnız değişiklik yapılmayan veya stratejik açıklamanın dokümanda görünmesi gerçekten faydalıysa öner. comment.anchor_text mevcut ve DEĞİŞMEYECEK bir ifadeyi hedeflemelidir.

JSON dışında yazma.
ŞEMA:
{{
  "analysis_summary":"",
  "coverage_complete":false,
  "requests":[
    {{
      "id":"R1",
      "source_locator":"müşterinin talep/soru başlığı veya kısa birebir ifadesi",
      "customer_request":"",
      "category":"technical_text|claim_scope|question|figure|procedural",
      "decision":"apply|partial|explain|clarification|figure_action|procedural_action",
      "reason":"",
      "answer_for_customer":""
    }}
  ],
  "operations":[
    {{
      "request_id":"R1",
      "type":"replace_text|insert_paragraph_after|insert_paragraph_before",
      "section":"",
      "locator_text":"",
      "anchor_text":"",
      "old_text":"",
      "new_text":"",
      "basis_source":"existing_spec|customer_request",
      "basis_quote":"",
      "reason":""
    }}
  ],
  "comments":[
    {{"request_id":"R1","anchor_text":"mevcut ve değişmeden kalacak kısa ifade","text":""}}
  ],
  "figure_actions":[
    {{"request_id":"R1","figure":"Şekil/Figure X veya genel","issue":"","recommended_change":"","safe_auto_edit":false,"basis_source":"existing_spec|customer_request","basis_quote":"","edit_instructions":""}}
  ],
  "blocking_clarifications":[""],
  "open_procedural_items":[""],
  "coverage_notes":[""]
}}

MEVCUT TARİFNAME:
{spec_text}

MÜŞTERİ REVİZYONLARI / SORULARI:
{customer_text}

EK TEKNİK / ŞEKİL BAĞLAMI:
{supplemental_text}
"""


def tarifname_update_quality_prompt(
    rules: str,
    spec_text: str,
    customer_text: str,
    filing_status: str,
    current_plan: dict[str, Any],
    mail_language: str,
    supplemental_text: str = "",
    user_instruction: str = "",
) -> str:
    return f"""{rules}

TARİFNAME DÜZENLEME — BAĞIMSIZ İKİNCİ OKUMA VE NİHAİ PLAN
Başvuru durumu: {filing_status}
Mail dili: {mail_language}
Kullanıcının ek yönlendirmesi: {user_instruction or '(yok)'}

Aşağıdaki müşteri dönüşünü ve mevcut tarifnameyi SIFIRDAN yeniden oku. İlk planın doğru olduğunu varsayma. Müşterinin her ana maddesini, alt talebini ve doğrudan sorusunu tek tek karşılaştır. Eksik kalan talep/soru varsa plana ekle. Uygun olmayan müşteri talebini sırf müşteri istedi diye uygulama; nedenini `answer_for_customer` ve gerekiyorsa Word comment ile açıkla. Minimum metin farkı kuralını tekrar uygula.

`coverage_complete=true` yalnız bütün müşteri talepleri bir sonuç durumuna bağlandıysa verilebilir. Her request şu sonuçlardan birine sahip olmalı: apply, partial, explain, clarification, figure_action, procedural_action. Hiçbiri sessizce atlanamaz.

İKİNCİ OKUMADA ÖZELLİKLE:
- Müşteri `istemlerde açıkça vurgulansın/görülsün` dediği bir teknik içeriği mevcut istem veya tarifnamede semantik olarak buluyorsan, `zaten var` sonucu yeterli değildir; requested terminology/kısaltma görünürlüğünü minimum Track Changes ile gerçekten sağla.
- Aynı talepte sayılan her teknik test/işlevi tek tek kontrol et; bir kısmını uygulayıp diğerlerini sessizce atlama.
- Şekil aksiyonunda otomatik düzenleme güvenliyse `safe_auto_edit=true` + dayanak + edit_instructions ver; yalnız öneri bırakılması gerekiyorsa false ver.

Mail:
- Markup dosyasının ekte olduğunu belirt.
- Yapılan ana teknik/istem değişikliklerini gruplayarak kısa açıkla; kelime kelime change-log çıkarma.
- Uygulanmayan/kısmen uygulanan müşteri taleplerine doğrudan cevap ver.
- Müşterinin doğrudan sorularını cevapla; kaynaklar kesin hükmü desteklemiyorsa kesin konuşma.
- Şekil revizyonu gerekiyorsa somut olarak hangi şekillerde neyin güncellenebileceğini ve gerekiyorsa editable kaynakları iste.
- Açık usuli/stratejik konuları belirt.
- Mail `{mail_language}` dilinde olsun.

JSON dışında yazma ve aşağıdaki şemayı aynen koru. `requests`, `operations`, `comments`, `figure_actions` alanlarını düzeltilmiş nihai içerikle tekrar döndür.
ŞEMA:
{{
  "analysis_summary":"",
  "coverage_complete":true,
  "requests":[],
  "operations":[],
  "comments":[],
  "figure_actions":[],
  "blocking_clarifications":[],
  "open_procedural_items":[],
  "coverage_notes":[],
  "mail":{{"subject":"","body":""}}
}}

MEVCUT TARİFNAME:
{spec_text}

MÜŞTERİ REVİZYONLARI / SORULARI:
{customer_text}

EK TEKNİK / ŞEKİL BAĞLAMI:
{supplemental_text}

İLK PLAN:
{json.dumps(current_plan or {}, ensure_ascii=False, indent=2)}
"""


def prepare_review_baseline_docx(data: bytes) -> bytes:
    """Reject pre-existing tracked revisions and remove old Word comments.

    This lets a customer-returned file serve as both the request carrier and the baseline source:
    customer insertions are treated as requests rather than silently accepted, customer deletions are
    restored, and Destek Patent then writes a fresh markup layer after substantive review.
    """
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R = "http://schemas.openxmlformats.org/package/2006/relationships"
    CT = "http://schemas.openxmlformats.org/package/2006/content-types"
    ns = {"w": W}
    src = io.BytesIO(data)
    out = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        names = set(zin.namelist())
        for name in zin.namelist():
            # Old review-comment auxiliary parts are deliberately discarded.
            if name in {"word/comments.xml", "word/commentsExtended.xml", "word/people.xml"}:
                continue
            payload = zin.read(name)
            if name == "word/document.xml":
                root = etree.fromstring(payload)
                # Reject insertions.
                for ins in list(root.xpath(".//w:ins", namespaces=ns)):
                    parent = ins.getparent()
                    if parent is not None:
                        parent.remove(ins)
                # Restore deletions as ordinary runs, preserving run properties.
                for dele in list(root.xpath(".//w:del", namespaces=ns)):
                    parent = dele.getparent()
                    if parent is None:
                        continue
                    idx = parent.index(dele)
                    for child in list(dele):
                        if child.tag != f"{{{W}}}r":
                            continue
                        run = deepcopy(child)
                        for dt in run.xpath(".//w:delText", namespaces=ns):
                            dt.tag = f"{{{W}}}t"
                            dt.set(qn("xml:space"), "preserve")
                        parent.insert(idx, run)
                        idx += 1
                    parent.remove(dele)
                # Remove explicit red-font customer note runs from the baseline. They remain available
                # through extract_docx_review_context as request text and are not silently accepted into the patent body.
                red_note_paragraphs: list[Any] = []
                for run in list(root.xpath(".//w:r", namespaces=ns)):
                    if not _ooxml_run_is_customer_red(run, W):
                        continue
                    paragraph = next(iter(run.xpath("ancestor::w:p[1]", namespaces=ns)), None)
                    if paragraph is not None and paragraph not in red_note_paragraphs:
                        red_note_paragraphs.append(paragraph)
                    parent = run.getparent()
                    if parent is not None:
                        parent.remove(run)
                # Only paragraphs touched by red-note stripping may be removed; unrelated template blank
                # paragraphs are part of the original Word geometry and must remain intact.
                for p in red_note_paragraphs:
                    has_text = bool(p.xpath(".//w:t[normalize-space(.) != '']", namespaces=ns))
                    has_object = bool(p.xpath(".//w:drawing|.//w:pict|.//w:object|.//w:br|.//w:fldSimple", namespaces=ns))
                    if not has_text and not has_object:
                        parent = p.getparent()
                        if parent is not None:
                            parent.remove(p)

                # Remove old comment anchors/references.
                for node in list(root.xpath(".//w:commentRangeStart|.//w:commentRangeEnd", namespaces=ns)):
                    parent = node.getparent()
                    if parent is not None:
                        parent.remove(node)
                for cref in list(root.xpath(".//w:commentReference", namespaces=ns)):
                    run = cref.getparent()
                    if run is not None and run.tag == f"{{{W}}}r":
                        parent = run.getparent()
                        if parent is not None:
                            parent.remove(run)
                payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            elif name == "word/settings.xml":
                root = etree.fromstring(payload)
                for node in list(root.xpath(".//w:trackRevisions", namespaces=ns)):
                    parent = node.getparent()
                    if parent is not None:
                        parent.remove(node)
                payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            elif name == "word/_rels/document.xml.rels":
                root = etree.fromstring(payload)
                for rel in list(root):
                    typ = str(rel.get("Type", ""))
                    target = str(rel.get("Target", ""))
                    if typ.endswith("/comments") or typ.endswith("/commentsExtended") or target.endswith("comments.xml") or target.endswith("commentsExtended.xml") or target.endswith("people.xml"):
                        root.remove(rel)
                payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            elif name == "[Content_Types].xml":
                root = etree.fromstring(payload)
                for node in list(root):
                    part = str(node.get("PartName", ""))
                    if part in {"/word/comments.xml", "/word/commentsExtended.xml", "/word/people.xml"}:
                        root.remove(node)
                payload = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            zout.writestr(name, payload)
    return out.getvalue()
