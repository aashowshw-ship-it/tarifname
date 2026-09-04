from pathlib import Path
from docx import Document
import io

from app_core import build_gorus_docx
from gorus_audit import validate_gorus_template_fidelity
from rules import APP_VERSION, RULESET_VERSION, GORUS_RULES

ROOT = Path(__file__).resolve().parent


def _opinion():
    return {
        "application_no": "2026/000001",
        "applicant": "TEST",
        "reference": "698833",
        "report_date": "01.09.2026",
        "intro": "Sayın Uzman girişinden sonra kullanılan örnek değerlendirme metnidir.",
        "amendment_assessment": {"heading": "", "blocks": []},
        "cited_documents": [
            {"label": "D1", "number": "US20250014043A1", "title": "THIS TITLE MUST NOT RENDER", "category": "X"},
            {"label": "D2", "number": "CN117993919A", "title": "SECOND TITLE MUST NOT RENDER", "category": "X"},
            {"label": "D3", "number": "CN117708880A", "title": "THIRD TITLE MUST NOT RENDER", "category": "X"},
        ],
        "sections": [],
        "combined_assessment": {"heading": "", "paragraphs": []},
        "conclusion": [],
        "signoff": "Saygılarımızla,\nDESTEK PATENT A.Ş.",
    }


def test_v552_versions_and_rule_text():
    assert APP_VERSION == "v5.4.52"
    assert RULESET_VERSION == "2026-09-04.v42"
    low = GORUS_RULES.casefold()
    assert "bibliyografik satırlarında patent/doküman başlığı yazılmaz" in low
    assert "satırın tamamı" in low


def test_cited_document_rows_are_number_only_and_fully_bold():
    opinion = _opinion()
    data = build_gorus_docx(opinion)
    doc = Document(io.BytesIO(data))
    expected = [
        "D1: US20250014043A1",
        "D2: CN117993919A",
        "D3: CN117708880A",
    ]
    texts = [p.text.strip() for p in doc.paragraphs]
    for text in expected:
        assert text in texts
        p = next(p for p in doc.paragraphs if p.text.strip() == text)
        runs = [r for r in p.runs if r.text.strip()]
        assert runs and all(bool(r.bold) for r in runs)
    assert all("THIS TITLE" not in t and "SECOND TITLE" not in t and "THIRD TITLE" not in t for t in texts)


def test_gate_rejects_title_suffix_or_nonbold_row():
    opinion = _opinion()
    data = build_gorus_docx(opinion)
    # The clean generated document must pass the D-list bibliography gate. No figures/quotes are required here.
    validate_gorus_template_fidelity(data, ROOT / "Gorus_metni_696809_template.docx", opinion)
