from __future__ import annotations

import io
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def _emu(v):
    return None if v is None else int(v)


def _line(v):
    if v is None:
        return None
    try:
        return round(float(v), 6)
    except Exception:
        return str(v)


def paragraph_format_signature(p, *, ignore_page_break: bool = False) -> tuple:
    pf = p.paragraph_format
    ppr = p._p.pPr
    num_id = ilvl = None
    if ppr is not None:
        numpr = ppr.find(qn("w:numPr"))
        if numpr is not None:
            n = numpr.find(qn("w:numId"))
            l = numpr.find(qn("w:ilvl"))
            num_id = n.get(qn("w:val")) if n is not None else None
            ilvl = l.get(qn("w:val")) if l is not None else None
    return (
        p.style.style_id if p.style is not None else None,
        p.alignment,
        _emu(pf.left_indent), _emu(pf.right_indent), _emu(pf.first_line_indent),
        _emu(pf.space_before), _emu(pf.space_after), _line(pf.line_spacing),
        None if ignore_page_break else pf.page_break_before,
        pf.keep_with_next, pf.keep_together, pf.widow_control,
        num_id, ilvl,
    )


def _run_signature(p) -> tuple:
    runs = [r for r in p.runs if r.text]
    if not runs:
        return ()
    out = []
    for r in runs:
        out.append((
            r.bold, r.italic, r.underline,
            r.font.name,
            None if r.font.size is None else int(r.font.size),
            None if r.font.color is None or r.font.color.rgb is None else str(r.font.color.rgb),
            None if r.font.highlight_color is None else str(r.font.highlight_color),
        ))
    return tuple(out)


def _xml_part_signature(data: bytes, prefix: str) -> dict[str, tuple]:
    out: dict[str, tuple] = {}
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for name in sorted(n for n in zf.namelist() if n.startswith(prefix) and n.endswith(".xml")):
            root = ET.fromstring(zf.read(name))
            texts = tuple((x.text or "") for x in root.findall(".//w:t", NS))
            instr = tuple(" ".join((x.text or "").split()) for x in root.findall(".//w:instrText", NS))
            drawings = len(root.findall(".//w:drawing", NS)) + len(root.findall(".//w:pict", NS))
            paragraphs = len(root.findall(".//w:p", NS))
            out[name] = (texts, instr, drawings, paragraphs)
    return out


def _field_count(parts: dict[str, tuple], field: str = "PAGE") -> int:
    target = field.casefold()
    total = 0
    for _, sig in parts.items():
        instr = sig[1]
        total += sum(1 for x in instr if x.casefold().startswith(target))
    return total


def _count_blanks_before(texts: list[str], index: int) -> int:
    n = 0
    i = index - 1
    while i >= 0 and texts[i] == "":
        n += 1
        i -= 1
    return n


def _count_blanks_after(texts: list[str], index: int) -> int:
    n = 0
    i = index + 1
    while i < len(texts) and texts[i] == "":
        n += 1
        i += 1
    return n


def _assert_fmt(out_p, tpl_p, message: str, *, ignore_page_break: bool = False, compare_runs: bool = True) -> None:
    if paragraph_format_signature(out_p, ignore_page_break=ignore_page_break) != paragraph_format_signature(tpl_p, ignore_page_break=ignore_page_break):
        raise ValueError(f"TAM ŞABLON KONTROLÜ: {message} paragraf biçimi Tarifname_181176_template.docx ile aynı değil.")
    if compare_runs and out_p.text.strip() and tpl_p.text.strip():
        # Metin değişebilir; ancak paragrafın ilk görünür run karakter biçimi şablonla aynı olmalı.
        a = _run_signature(out_p)
        b = _run_signature(tpl_p)
        if a and b and a[0] != b[0]:
            raise ValueError(f"TAM ŞABLON KONTROLÜ: {message} yazı karakteri/kalınlık/renk biçimi şablonla aynı değil.")


def validate_full_tarifname_template_fidelity(
    data: bytes,
    template_path: str | Path,
    draft: dict[str, Any],
    language: str = "Türkçe",
) -> None:
    """Nihai tarifnamenin bağlayıcı Word şablonuna yapısal ve görsel-geometri düzeyinde uyumunu denetler.

    Bu kontrol yalnız birkaç başlığı değil; bölüm geçiş boşluklarını, paragraf arketiplerini,
    istem/özet sayfa kırımlarını, referans-listesi ritmini, header/footer ve sayfa numarası
    konumunu birlikte denetler. Çıktı bu kapıyı geçmeden indirilemez.
    """
    out = Document(io.BytesIO(data))
    tpl = Document(str(template_path))
    paras = out.paragraphs
    texts = [p.text.strip() for p in paras]
    en = str(language or "").casefold().startswith("ing") or str(language or "").casefold().startswith("en")
    labels = {
        "spec": "SPECIFICATION" if en else "TARİFNAME",
        "technical": "TECHNICAL FIELD" if en else "TEKNİK ALAN",
        "prior": "PRIOR ART" if en else "ÖNCEKİ TEKNİK",
        "short": "BRIEF DESCRIPTION OF THE INVENTION" if en else "BULUŞUN KISA AÇIKLAMASI",
        "figures": "BRIEF DESCRIPTION OF THE FIGURES" if en else "ŞEKİLLERİN KISA AÇIKLAMASI",
        "refs": "REFERENCE NUMERALS" if en else "REFERANS NUMARALARI",
        "detail": "DETAILED DESCRIPTION OF THE INVENTION" if en else "BULUŞUN DETAYLI AÇIKLAMASI",
        "claims": "CLAIMS" if en else "İSTEMLER",
        "abstract": "ABSTRACT" if en else "ÖZET",
    }

    def idx(label: str) -> int:
        try:
            return texts.index(label)
        except ValueError as exc:
            raise ValueError(f"TAM ŞABLON KONTROLÜ: {label} başlığı bulunamadı.") from exc

    indices = {k: idx(v) for k, v in labels.items()}
    if list(indices.values()) != sorted(indices.values()):
        raise ValueError("TAM ŞABLON KONTROLÜ: tarifname bölüm sırası bağlayıcı şablonla uyuşmuyor.")

    # Sayfa/section geometrisi doğrudan bağlayıcı şablondan gelmeli.
    if len(out.sections) != len(tpl.sections):
        raise ValueError("TAM ŞABLON KONTROLÜ: section sayısı şablonla aynı değil.")
    geom_attrs = (
        "top_margin", "bottom_margin", "left_margin", "right_margin",
        "page_width", "page_height", "header_distance", "footer_distance", "gutter",
    )
    for n, (os, ts) in enumerate(zip(out.sections, tpl.sections), start=1):
        for attr in geom_attrs:
            if _emu(getattr(os, attr)) != _emu(getattr(ts, attr)):
                raise ValueError(f"TAM ŞABLON KONTROLÜ: {n}. section {attr} değeri şablondan sapmış.")
        if os.orientation != ts.orientation:
            raise ValueError(f"TAM ŞABLON KONTROLÜ: {n}. section sayfa yönü şablondan sapmış.")

    # Header/footer birebir aynı kaynak yapıyı korumalı. Özellikle PAGE footer'a eklenemez.
    tpl_bytes = Path(template_path).read_bytes()
    tpl_headers = _xml_part_signature(tpl_bytes, "word/header")
    tpl_footers = _xml_part_signature(tpl_bytes, "word/footer")
    out_headers = _xml_part_signature(data, "word/header")
    out_footers = _xml_part_signature(data, "word/footer")
    if out_headers != tpl_headers:
        raise ValueError("TAM ŞABLON KONTROLÜ: header içeriği/sayfa numarası şablondan farklı. Header şablondan aynen korunmalıdır.")
    if out_footers != tpl_footers:
        raise ValueError("TAM ŞABLON KONTROLÜ: footer şablondan farklı. Tarifnameye alt sayfa numarası veya yeni footer eklenemez.")
    if _field_count(out_footers, "PAGE") != _field_count(tpl_footers, "PAGE"):
        raise ValueError("TAM ŞABLON KONTROLÜ: footer içinde fazladan PAGE alanı bulundu; sayfa numarası yalnız şablondaki üst konumda kullanılmalıdır.")
    if _field_count(out_headers, "PAGE") != _field_count(tpl_headers, "PAGE"):
        raise ValueError("TAM ŞABLON KONTROLÜ: header PAGE alanları şablonla uyuşmuyor.")

    # Sabit başlıkların paragraf arketipleri.
    archetypes = {
        "spec": 0, "technical": 6, "prior": 10, "short": 17,
        "figures": 41, "refs": 47, "detail": 57,
        "claims": 77, "abstract": 99,
    }
    for key, ti in archetypes.items():
        _assert_fmt(paras[indices[key]], tpl.paragraphs[ti], labels[key], ignore_page_break=key in {"claims", "abstract"})

    # İSTEMLER ve ÖZET ek kural: yeni sayfa ve ortalı.
    for key in ("claims", "abstract"):
        p = paras[indices[key]]
        if p.alignment != WD_ALIGN_PARAGRAPH.CENTER or p.paragraph_format.page_break_before is not True:
            raise ValueError(f"TAM ŞABLON KONTROLÜ: {labels[key]} ortalı ve yeni sayfadan başlamalıdır.")

    # Başlangıç ritmi: başlık, boşluk, buluş adı, boşluk, talimat, boşluk, teknik alan.
    si = indices["spec"]
    if si != 0:
        raise ValueError("TAM ŞABLON KONTROLÜ: TARİFNAME/SPECIFICATION belgenin ilk paragrafı olmalıdır.")
    if len(paras) < 7 or texts[1] != "" or texts[3] != "" or texts[5] != "":
        raise ValueError("TAM ŞABLON KONTROLÜ: tarifname başlangıcındaki boş paragraf ritmi şablonla aynı değil.")
    _assert_fmt(paras[1], tpl.paragraphs[1], "TARİFNAME sonrası boşluk", compare_runs=False)
    _assert_fmt(paras[2], tpl.paragraphs[2], "buluş başlığı")
    _assert_fmt(paras[3], tpl.paragraphs[3], "buluş başlığı sonrası boşluk", compare_runs=False)
    _assert_fmt(paras[4], tpl.paragraphs[4], "giriş talimatı")
    _assert_fmt(paras[5], tpl.paragraphs[5], "giriş talimatı sonrası boşluk", compare_runs=False)

    # Her ana başlıktan sonra şablon boşluğu zorunlu.
    blank_after_archetype = {
        "technical": 7, "prior": 11, "short": 18,
        "figures": 42, "refs": 48, "detail": 58,
        "claims": 78, "abstract": 100,
    }
    for key, ti in blank_after_archetype.items():
        i = indices[key]
        if i + 1 >= len(texts) or texts[i + 1] != "":
            raise ValueError(f"TAM ŞABLON KONTROLÜ: {labels[key]} başlığından sonraki boş paragraf eksik.")
        _assert_fmt(paras[i + 1], tpl.paragraphs[ti], f"{labels[key]} sonrası boş paragraf", compare_runs=False)

    # TEKNİK ALAN iki paragraf + aralarında ve sonrasında birer şablon boşluğu.
    ti, pi = indices["technical"], indices["prior"]
    tech_segment = texts[ti + 2:pi]
    nonblank_tech = [x for x in tech_segment if x]
    if len(nonblank_tech) != 2:
        raise ValueError("TAM ŞABLON KONTROLÜ: TEKNİK ALAN iki ana paragraftan oluşmalıdır.")
    if tech_segment != [nonblank_tech[0], "", nonblank_tech[1], ""]:
        raise ValueError("TAM ŞABLON KONTROLÜ: TEKNİK ALAN paragraf/boşluk ritmi şablon düzenine uygun değil.")
    _assert_fmt(paras[ti + 2], tpl.paragraphs[8], "TEKNİK ALAN ilk paragrafı")
    _assert_fmt(paras[ti + 3], tpl.paragraphs[9], "TEKNİK ALAN ara boşluğu", compare_runs=False)
    _assert_fmt(paras[ti + 4], tpl.paragraphs[8], "TEKNİK ALAN ikinci paragrafı")
    _assert_fmt(paras[ti + 5], tpl.paragraphs[9], "TEKNİK ALAN son boşluğu", compare_runs=False)

    # Önceki teknik: paragraflar arasında boşluk; sonuç paragrafı şablondaki space-after boşluğunu taşımalı.
    pri, shi = indices["prior"], indices["short"]
    if shi <= pri + 2:
        raise ValueError("TAM ŞABLON KONTROLÜ: ÖNCEKİ TEKNİK içeriği eksik.")
    conclusion_p = paras[shi - 1]
    if texts[shi - 1] == "":
        raise ValueError("TAM ŞABLON KONTROLÜ: BULUŞUN KISA AÇIKLAMASI öncesindeki boşluk şablonda fiziksel boş paragraf değil, sonuç paragrafının space-after değeridir; şablon paragrafı kullanılmalıdır.")
    if _emu(conclusion_p.paragraph_format.space_after) != _emu(tpl.paragraphs[16].paragraph_format.space_after):
        raise ValueError("TAM ŞABLON KONTROLÜ: BULUŞUN KISA AÇIKLAMASI öncesindeki görsel boşluk şablondaki sonuç paragrafı kadar değil.")
    _assert_fmt(conclusion_p, tpl.paragraphs[16], "ÖNCEKİ TEKNİK sonuç paragrafı")
    # Önceki teknik gövdesindeki bütün normal paragraflar arasında bir boş paragraf olmalı.
    prior_body = texts[pri + 2:shi - 1]
    for j in range(len(prior_body) - 1):
        if prior_body[j] and prior_body[j + 1]:
            raise ValueError("TAM ŞABLON KONTROLÜ: ÖNCEKİ TEKNİK ana paragrafları arasında şablon boşluğu eksik.")

    # ŞEKİLLER başlığından önce tam bir boş paragraf.
    fi = indices["figures"]
    if _count_blanks_before(texts, fi) != 1:
        raise ValueError("TAM ŞABLON KONTROLÜ: ŞEKİLLERİN KISA AÇIKLAMASI öncesinde tam bir şablon boş paragrafı bulunmalıdır.")
    _assert_fmt(paras[fi - 1], tpl.paragraphs[40], "ŞEKİLLERİN KISA AÇIKLAMASI öncesi boşluk", compare_runs=False)

    # Şekil açıklamaları ardışık; son şekilden sonra boşluk; sonra Çizimlerin...; sonra boşluk; REFERANSLAR.
    ri = indices["refs"]
    boilerplate_idx = None
    boilerplate_prefix = "The drawings are not necessarily" if en else "Çizimlerin mutlaka ölçeklendirilmesi"
    for j in range(fi + 2, ri):
        if texts[j].startswith(boilerplate_prefix):
            boilerplate_idx = j
            break
    if boilerplate_idx is None:
        raise ValueError("TAM ŞABLON KONTROLÜ: şekil açıklamalarından sonraki sabit 'Çizimlerin...' paragrafı bulunamadı.")
    fig_region = texts[fi + 2:boilerplate_idx]
    if not fig_region or fig_region[-1] != "" or any(x == "" for x in fig_region[:-1]):
        raise ValueError("TAM ŞABLON KONTROLÜ: Şekil açıklamaları aralarında boşluk olmadan ardışık olmalı ve yalnız son şekilden sonra bir boş paragraf gelmelidir.")
    for j in range(fi + 2, boilerplate_idx - 1):
        _assert_fmt(paras[j], tpl.paragraphs[43], "Şekil açıklaması")
    _assert_fmt(paras[boilerplate_idx - 1], tpl.paragraphs[44], "son Şekil açıklaması sonrası boşluk", compare_runs=False)
    _assert_fmt(paras[boilerplate_idx], tpl.paragraphs[45], "Çizimlerin... paragrafı")
    if _count_blanks_before(texts, ri) != 1:
        raise ValueError("TAM ŞABLON KONTROLÜ: REFERANS NUMARALARI öncesindeki şablon boşluğu eksik/fazla.")
    _assert_fmt(paras[ri - 1], tpl.paragraphs[46], "REFERANS NUMARALARI öncesi boşluk", compare_runs=False)

    # Referans listesi: sistem referansları bitişik; yöntem varsa arada tam bir boşluk; detaylı açıklama öncesinde tam bir boşluk.
    di = indices["detail"]
    ref_lines = texts[ri + 2:di]
    if not ref_lines or ref_lines[-1] != "":
        raise ValueError("TAM ŞABLON KONTROLÜ: BULUŞUN DETAYLI AÇIKLAMASI öncesindeki şablon boş paragrafı eksik.")
    if _count_blanks_before(texts, di) != 1:
        raise ValueError("TAM ŞABLON KONTROLÜ: BULUŞUN DETAYLI AÇIKLAMASI öncesinde tam bir boş paragraf bulunmalıdır.")
    _assert_fmt(paras[di - 1], tpl.paragraphs[56], "BULUŞUN DETAYLI AÇIKLAMASI öncesi boşluk", compare_runs=False)
    ref_nonfinal = ref_lines[:-1]
    blanks = [i for i, x in enumerate(ref_nonfinal) if x == ""]
    system_refs = [str(x.get("number", "") or "").strip() for x in (draft.get("elements") or []) if str(x.get("number", "") or "").strip()]
    method_refs = [str(x.get("number", "") or "").strip() for x in (draft.get("method_steps") or []) if str(x.get("number", "") or "").strip()]
    if system_refs and method_refs:
        if len(blanks) != 1:
            raise ValueError("TAM ŞABLON KONTROLÜ: sistem referansları ile yöntem referansları arasında tam bir boş paragraf olmalıdır.")
    elif blanks:
        raise ValueError("TAM ŞABLON KONTROLÜ: REFERANS NUMARALARI içinde gereksiz boş paragraf bulundu.")
    for j in range(ri + 2, di - 1):
        if texts[j]:
            _assert_fmt(paras[j], tpl.paragraphs[49], "REFERANS NUMARALARI satırı")

    # Detaylı açıklama başlığından sonra giriş + boşluk; ana açıklama paragraflarında boşluk ritmi korunur.
    if di + 3 >= len(texts) or not texts[di + 2] or texts[di + 3] != "":
        raise ValueError("TAM ŞABLON KONTROLÜ: detaylı açıklama giriş paragrafı ve sonrasındaki boşluk şablona uymuyor.")
    _assert_fmt(paras[di + 2], tpl.paragraphs[59], "detaylı açıklama giriş paragrafı")
    _assert_fmt(paras[di + 3], tpl.paragraphs[60], "detaylı açıklama girişinden sonraki boşluk", compare_runs=False)

    # İSTEMLER öncesi tam iki boş paragraf; şablonun 75/76 arketipleri.
    ci = indices["claims"]
    if _count_blanks_before(texts, ci) != 2:
        raise ValueError("TAM ŞABLON KONTROLÜ: İSTEMLER öncesinde şablondaki gibi tam iki boş paragraf olmalıdır.")
    _assert_fmt(paras[ci - 2], tpl.paragraphs[75], "İSTEMLER öncesi ilk boşluk", compare_runs=False)
    _assert_fmt(paras[ci - 1], tpl.paragraphs[76], "İSTEMLER öncesi ikinci boşluk", compare_runs=False)

    # İstem açıklama blokları arasında tek boşluk.
    tpl_notes = [tpl.paragraphs[79].text.strip(), tpl.paragraphs[81].text.strip(), tpl.paragraphs[83].text.strip()]
    for note in tpl_notes:
        if note in texts:
            ni = texts.index(note)
            if _count_blanks_after(texts, ni) != 1:
                raise ValueError("TAM ŞABLON KONTROLÜ: İSTEMLER açıklama paragrafları arasındaki boşluk ritmi şablona uymuyor.")

    # Bağımsız istem kapanışları arketip 93; istemler arasında boşluk.
    ai = indices["abstract"]
    for i, p in enumerate(paras[ci + 1:ai], start=ci + 1):
        if p.text.strip() in {"içermesidir.", "işlem adımlarını içermesidir."}:
            _assert_fmt(p, tpl.paragraphs[93], "bağımsız istem kapanışı")
    # Numaralı istem paragrafları numId=2 olmalı; listeler numId=3 veya hiyerarşik alt liste.
    numbered_claims = 0
    for p in paras[ci + 1:ai]:
        sig = paragraph_format_signature(p)
        if sig[-2] == "2":
            numbered_claims += 1
    expected_claims = (1 if draft.get("system_claim") else 0) + len(draft.get("dependent_system_claims") or []) + (1 if draft.get("method_claim") else 0) + len(draft.get("dependent_method_claims") or [])
    if numbered_claims < expected_claims:
        raise ValueError("TAM ŞABLON KONTROLÜ: tüm istemler gerçek Word otomatik numaralandırmasıyla oluşturulmamış.")

    # ÖZET ritmi: öncesinde tek boşluk; başlık + boşluk + ortalı/kalın buluş adı + boşluk + tek paragraf özet + şablon son boşlukları.
    if _count_blanks_before(texts, ai) != 1:
        raise ValueError("TAM ŞABLON KONTROLÜ: ÖZET öncesindeki istem boşluğu şablona uymuyor.")
    if ai + 4 >= len(paras):
        raise ValueError("TAM ŞABLON KONTROLÜ: ÖZET bölümü eksik.")
    if texts[ai + 1] != "" or texts[ai + 3] != "":
        raise ValueError("TAM ŞABLON KONTROLÜ: ÖZET başlık/buluş adı boşlukları şablona uymuyor.")
    if texts[ai + 2] != str(draft.get("title", "") or "").strip():
        raise ValueError("TAM ŞABLON KONTROLÜ: ÖZET içindeki buluş adı beklenen yerde değil.")
    _assert_fmt(paras[ai + 2], tpl.paragraphs[101], "ÖZET buluş adı")
    if paras[ai + 2].alignment != WD_ALIGN_PARAGRAPH.CENTER or not any(r.bold for r in paras[ai + 2].runs if r.text):
        raise ValueError("TAM ŞABLON KONTROLÜ: ÖZET içindeki buluş adı kalın ve ortalı olmalıdır.")
    # Özet tek paragraf olmalı; sonraki iki paragraf şablondaki gibi boş olmalı.
    abstract_text = str(draft.get("abstract", "") or "").strip()
    if texts[ai + 4] != abstract_text:
        raise ValueError("TAM ŞABLON KONTROLÜ: ÖZET metni tek paragraf halinde beklenen yerde değil.")
    if ai + 6 >= len(texts) or texts[ai + 5] != "" or texts[ai + 6] != "":
        raise ValueError("TAM ŞABLON KONTROLÜ: ÖZET sonundaki iki şablon boş paragrafı korunmamış.")
    _assert_fmt(paras[ai + 4], tpl.paragraphs[103], "ÖZET metni")
    _assert_fmt(paras[ai + 5], tpl.paragraphs[104], "ÖZET ilk son boşluğu", compare_runs=False)
    _assert_fmt(paras[ai + 6], tpl.paragraphs[105], "ÖZET ikinci son boşluğu", compare_runs=False)
