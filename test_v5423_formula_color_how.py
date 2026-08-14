from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import RGBColor

from word_math import add_display_equation
from template_audit import _assert_exact_colored_runs
from validators import validate_draft

ROOT = Path(__file__).resolve().parent
TPL = ROOT / "Tarifname_181176_template.docx"


def test_active_generator_preserves_fixed_template_paragraph_when_text_is_identical():
    src = (ROOT / "app.py").read_text(encoding="utf-8")
    assert 'if str(text) == template.paragraphs[index].text:' in src
    assert 'copy_template_paragraph(doc, template, index)' in src
    tpl = Document(str(TPL))
    visible = [(r.text, str(r.font.color.rgb) if r.font.color.rgb else None) for r in tpl.paragraphs[4].runs if r.text]
    assert any(color == "0000FF" for _, color in visible)
    assert any(color == "FF0000" for _, color in visible)


def test_template_color_guard_rejects_collapsed_blue_run():
    tpl = Document(str(TPL))
    out = Document(str(TPL))
    p = out.paragraphs[4]
    # Mavi bir run'ı kırmızıya çevirerek tam şablon renk guard'ını boz.
    blue = next(r for r in p.runs if r.text and str(r.font.color.rgb) == "0000FF")
    blue.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    with pytest.raises(ValueError, match="kırmızı/mavi"):
        _assert_exact_colored_runs(p, tpl.paragraphs[4], "giriş talimatı")


def test_display_formula_is_real_omml_with_subscript_and_fraction():
    tpl = Document(str(TPL))
    out = Document()
    add_display_equation(out, tpl, 61, "Durum = 0,30·energy_used/start_energy + 0,20·pn_sat_ratio")
    buf = io.BytesIO()
    out.save(buf)
    with zipfile.ZipFile(io.BytesIO(buf.getvalue())) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    assert "<m:oMath" in xml
    assert "<m:sSub>" in xml
    assert "<m:f>" in xml
    assert "energy" in xml and "used" in xml and "start" in xml


def _claim_draft(item: str):
    return {
        "technical_field": "Buluş, bir kontrol sistemi ile ilgilidir.\n\nBuluş, özellikle veri işleyen bir kontrol sistemi ile ilgilidir.",
        "elements": [{"number": "1", "name": "Kontrol modülü"}],
        "method_steps": [],
        "system_claim": {
            "preamble": "Bir elektronik işlem birimi üzerinde koşturulan yazılım vasıtasıyla çalışan bir sistem",
            "elements": [item],
            "closing": "içermesidir.",
        },
        "dependent_system_claims": [],
        "method_claim": None,
        "dependent_method_claims": [],
        "tables": [],
        "objectives": ["veri işleme güvenilirliğini artırmaktır."],
    }


def test_english_like_module_order_is_rejected_by_how_guard():
    findings = validate_draft(_claim_draft("Kontrol modülü (1), giriş verisini kullanarak sonucu hesaplayan bir modül,"))
    assert any("İngilizce claim sırasıyla" in f["message"] for f in findings)


def test_turkish_function_first_module_order_passes_order_guard():
    findings = validate_draft(_claim_draft("giriş verisini eşik değer ile karşılaştırarak kontrol sonucunu belirleyen kontrol modülü (1),"))
    assert not any("İngilizce claim sırasıyla" in f["message"] for f in findings)
    assert not any("'nasıl?'" in f["message"] for f in findings)


def test_plain_formula_in_claim_is_rejected_and_eq_marker_is_allowed():
    bad = _claim_draft("giriş verisini kullanarak x = 0,4·J + 0,6·S değerini hesaplayan kontrol modülü (1),")
    findings = validate_draft(bad)
    assert any("matematik bağıntısı düz metin" in f["message"] for f in findings)

    good = _claim_draft("giriş verisini kullanarak [[EQ: x = 0,4·J + 0,6·S]] değerini hesaplayan kontrol modülü (1),")
    findings = validate_draft(good)
    assert not any("matematik bağıntısı düz metin" in f["message"] for f in findings)


def _multi_claim_draft(items, names):
    return {
        "technical_field": "Buluş, bir kontrol sistemi ile ilgilidir.\n\nBuluş, özellikle veri işleyen bir kontrol sistemi ile ilgilidir.",
        "elements": [{"number": str(i + 1), "name": name} for i, name in enumerate(names)],
        "method_steps": [],
        "system_claim": {
            "preamble": "Bir elektronik işlem birimi üzerinde koşturulan yazılım vasıtasıyla çalışan bir sistem",
            "elements": items,
            "closing": "içermesidir.",
        },
        "dependent_system_claims": [],
        "method_claim": None,
        "dependent_method_claims": [],
        "tables": [],
        "objectives": ["veri işleme güvenilirliğini artırmaktır."],
    }


def test_how_guard_rejects_result_only_classification_module():
    draft = _multi_claim_draft(
        ["her görev çıktısını başarı veya tanımlı başarısızlık nedenlerinden biri olarak sınıflandıran görev sonucu sınıflandırma modülü (1),"],
        ["Görev sonucu sınıflandırma modülü"],
    )
    findings = validate_draft(draft)
    assert any("sınıflandırma kriterini/mekanizmasını" in f["message"] for f in findings)


def test_how_guard_rejects_metric_names_without_calculation_relation():
    draft = _multi_claim_draft(
        ["sınıflandırılmış görev sonuçlarından görev güvenilirlik indeksini ve elektronik harp dayanım skorunu hesaplayan metrik hesaplama modülü (1),"],
        ["Metrik hesaplama modülü"],
    )
    findings = validate_draft(draft)
    assert any("temel hesaplama ilişkisini" in f["message"] for f in findings)


def test_how_guard_accepts_source_backed_basic_calculation_relation():
    draft = _multi_claim_draft(
        ["başarılı görev sayısını toplam görev sayısına oranlayarak görev güvenilirlik indeksini hesaplayan metrik hesaplama modülü (1),"],
        ["Metrik hesaplama modülü"],
    )
    findings = validate_draft(draft)
    assert not any("temel hesaplama ilişkisini" in f["message"] for f in findings)
