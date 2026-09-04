import io

import pytest
from docx import Document

from app_core import build_gorus_docx
from gorus_audit import validate_opinion_narrative_rules, validate_gorus_template_fidelity
from rules import APP_VERSION, RULESET_VERSION, GORUS_RULES

def base_opinion():
    return {
        "application_no": "2026/000001",
        "applicant": "TEST A.Ş.",
        "reference": "699999",
        "intro": "Araştırma raporunda istem 1 bakımından yenilik ve buluş basamağı itirazı bildirilmiştir. Başvuru sahibinin görüşleri aşağıda sunulmaktadır.",
        "cited_documents": [{"label": "D1", "number": "WO0000000001A1", "category": "X", "summary": "D1 teknik bir sensör yapısını açıklamaktadır."}],
        "sections": [{
            "label": "D1",
            "heading": "D1 (WO0000000001A1) dokümanı:",
            "blocks": [{"type": "paragraph", "text": "D1 kısa olarak açıklanmıştır ve ayırt edici teknik fark başvurudaki sensör ile işleme biriminin doğrudan işlevsel ilişkisidir."}],
            "novelty_heading": "",
            "novelty_paragraphs": ["D1 istem 1'deki teknik özelliklerin tamamını ve aralarındaki işlevsel ilişkiyi doğrudan ve açık biçimde açıklamamaktadır."],
            "inventive_step_heading": "",
            "inventive_step_paragraphs": ["Teknik etki, sensör çıktısının istemde tanımlanan işlem sırasına göre değerlendirilmesidir. Objektif teknik problem, bu teknik işlevsel ilişkinin nasıl sağlanacağıdır. D1 bu değişiklik için motivasyon veya yönlendirme sağlamamaktadır."],
        }],
        "combined_assessment": {"heading": "", "paragraphs": []},
        "conclusion": ["Bu nedenlerle istem 1'in yenilik ve buluş basamağı kriterlerini sağladığı değerlendirilmektedir."],
        "signoff": "Saygılarımızla,\nDESTEK PATENT A.Ş.",
    }

def test_version_and_new_binding_rules():
    assert APP_VERSION == "v5.4.51"
    assert RULESET_VERSION == "2026-09-04.v41"
    low = GORUS_RULES.casefold()
    for phrase in ["ara başlıklar oluşturulmaz", "devralmaktadır", "mimari", "saygılarımızla", "bu farklardan"]:
        assert phrase in low

def test_individual_d_subheadings_are_rejected():
    op = base_opinion()
    validate_opinion_narrative_rules(op, "yenilik buluş basamağı", "teknik katkı teknik etki objektif teknik problem motivasyon")
    bad = base_opinion()
    bad["sections"][0]["novelty_heading"] = "D1 karşısında yenilik"
    with pytest.raises(ValueError, match="başlık"):
        validate_opinion_narrative_rules(bad, "yenilik buluş basamağı", "teknik katkı teknik etki objektif teknik problem motivasyon")

@pytest.mark.parametrize("bad_word", [
    "İstem 2, İstem 1'in teknik katkısını devralmaktadır.",
    "Bu mimari teknik etki sağlamaktadır.",
    "The claimed architecture provides the technical effect.",
    "Claim 2 inherits the technical contribution of claim 1.",
])
def test_forbidden_opinion_diction_is_rejected(bad_word):
    op = base_opinion()
    op["sections"][0]["inventive_step_paragraphs"][0] += " " + bad_word
    with pytest.raises(ValueError, match="dil kapısı"):
        validate_opinion_narrative_rules(op, "yenilik buluş basamağı", "teknik katkı teknik etki objektif teknik problem motivasyon")

def test_continuation_paragraph_is_rejected():
    op = base_opinion()
    op["sections"][0]["inventive_step_paragraphs"].append("Bu farklardan kaynaklanan teknik etki ayrıca değerlendirilmiştir.")
    with pytest.raises(ValueError, match="paragraf devamlılığı"):
        validate_opinion_narrative_rules(op, "yenilik buluş basamağı", "teknik katkı teknik etki objektif teknik problem motivasyon")

def test_docx_has_no_individual_subheadings_and_bold_signoff():
    op = base_opinion()
    data = build_gorus_docx(op)
    doc = Document(io.BytesIO(data))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "D1 karşısında yenilik" not in texts
    assert "D1 karşısında buluş basamağı" not in texts
    for signoff in ("Saygılarımızla,", "DESTEK PATENT A.Ş."):
        p = next(p for p in doc.paragraphs if p.text.strip() == signoff)
        runs = [r for r in p.runs if r.text.strip()]
        assert runs and all(r.bold for r in runs)
