from __future__ import annotations

import io

import pytest
from docx import Document

from rules import APP_VERSION, RULESET_VERSION, TARIFNAME_DUZENLEME_RULES
from tarifname_update import tarifname_update_analysis_prompt, validate_update_plan
from tarifname_figure_update import parse_figure_number


def _docx_bytes(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def test_v5435_versions_and_visibility_rule():
    assert APP_VERSION == "v5.4.49"
    assert RULESET_VERSION == "2026-09-01.v34"
    assert "zaten semantik olarak var" in tarifname_update_analysis_prompt(
        TARIFNAME_DUZENLEME_RULES,
        "mini sıvı yükleme testi ve pasif bacak kaldırma testi",
        "istemlerde açıkça vurgulayın",
        "Henüz başvuru yapılmadı",
    )
    assert "tam ad (KISALTMA)" in TARIFNAME_DUZENLEME_RULES
    assert "safe_auto_edit" in tarifname_update_analysis_prompt(
        TARIFNAME_DUZENLEME_RULES,
        "Şekil 1'de monitör vardır.",
        "Şekil 1'de kablo gösterilsin.",
        "Henüz başvuru yapılmadı",
    )


def test_figure_number_parser_requires_explicit_target():
    assert parse_figure_number("Şekil 1") == 1
    assert parse_figure_number("Figure 12") == 12
    assert parse_figure_number("genel") is None


def test_safe_auto_edit_requires_supported_basis_and_blocks_post_filing_customer_only_geometry():
    spec = _docx_bytes("İzleme birimi kablolu olarak kateter ile veri iletişimi sağlar.")
    plan = {
        "coverage_complete": True,
        "requests": [
            {
                "id": "R1",
                "decision": "figure_action",
                "answer_for_customer": "Şekil 1 güncellenecektir.",
            }
        ],
        "operations": [],
        "comments": [],
        "figure_actions": [
            {
                "request_id": "R1",
                "figure": "Şekil 1",
                "issue": "Kablo görünmüyor",
                "recommended_change": "İzleme birimi ile kateter arasında kablo göster.",
                "safe_auto_edit": True,
                "basis_source": "existing_spec",
                "basis_quote": "İzleme birimi kablolu olarak kateter ile veri iletişimi sağlar.",
                "edit_instructions": "Mevcut izleme birimi ile mevcut kateter arasında sade bir kablo çizgisi ekle.",
            }
        ],
        "blocking_clarifications": [],
    }
    validate_update_plan(plan, spec, "Şekil 1'e kablo ekleyin.", "Başvuru yapıldı")

    plan["figure_actions"][0]["basis_source"] = "customer_request"
    plan["figure_actions"][0]["basis_quote"] = "Şekil 1'e kablo ekleyin."
    with pytest.raises(ValueError, match="şekil new-matter"):
        validate_update_plan(plan, spec, "Şekil 1'e kablo ekleyin.", "Başvuru yapıldı")


def test_claim_visibility_gate_catches_supported_acronym_silently_omitted():
    spec = _docx_bytes(
        "BULUŞUN DETAYLI AÇIKLAMASI\nPLR ve MFC ile PPV ve SVV testleri sistemde kullanılabilir.\n"
        "İSTEMLER\n1. Sistem olup mini sıvı yükleme testi ve pasif bacak kaldırma testi hesaplar.\nÖZET"
    )
    customer = "İSTEMLERE EKLEME YAPIN; PLR, MFC, PPV ve SVV istemlerde açıkça vurgulansın."
    plan = {
        "coverage_complete": True,
        "requests": [{"id": "R1", "decision": "apply", "answer_for_customer": "Uygulandı."}],
        "operations": [
            {
                "request_id": "R1",
                "type": "replace_text",
                "section": "İSTEMLER",
                "locator_text": "Sistem olup mini sıvı yükleme testi",
                "old_text": "mini sıvı yükleme testi",
                "new_text": "mini sıvı yükleme testi (MFC)",
                "basis_source": "existing_spec",
                "basis_quote": "MFC",
                "reason": "MFC görünür kılındı",
            }
        ],
        "comments": [],
        "figure_actions": [],
        "blocking_clarifications": [],
    }
    with pytest.raises(ValueError, match="istem-görünürlük"):
        validate_update_plan(plan, spec, customer, "Henüz başvuru yapılmadı")


def test_red_customer_notes_are_extracted_and_removed_from_baseline():
    from docx.shared import RGBColor
    from tarifname_update import extract_docx_review_context, prepare_review_baseline_docx, document_text

    doc = Document()
    doc.add_paragraph("İSTEMLER")
    note = doc.add_paragraph()
    red = note.add_run("İSTEMLERE EKLEME YAPILMASINI RİCA EDİYORUM. PLR ve MFC açıkça yazılsın.")
    red.font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph("1. Sistem olup, özelliği; mevcut teknik unsurları içermesidir.")
    out = io.BytesIO()
    doc.save(out)
    data = out.getvalue()

    context = extract_docx_review_context(data)
    assert "RED CUSTOMER NOTE" in context
    assert "PLR ve MFC" in context

    baseline = prepare_review_baseline_docx(data)
    text = document_text(baseline)
    assert "İSTEMLERE EKLEME" not in text
    assert "1. Sistem olup" in text
