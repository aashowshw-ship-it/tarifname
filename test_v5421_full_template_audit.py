import io
from copy import deepcopy
from pathlib import Path

import pytest
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
from template_audit import validate_full_tarifname_template_fidelity

TPL = ROOT / "Tarifname_181176_template.docx"

DRAFT = {
    "title": "Test Sistemi",
    "elements": [{"number": "1", "name": "Test modülü"}],
    "method_steps": [],
    "figure_descriptions": ["Şekil 1, test sisteminin temsili gösterimidir."],
    "system_claim": {"preamble": "Bir test sistemi", "elements": ["test modülü (1)"], "closing": "içermesidir."},
    "dependent_system_claims": [],
    "method_claim": None,
    "dependent_method_claims": [],
    "abstract": "Buluş, bir test sistemi ile ilgilidir.",
}


def _clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _copy(doc, tpl, idx, text=None):
    el = deepcopy(tpl.paragraphs[idx]._p)
    if text is not None:
        nodes = list(el.iter(qn("w:t")))
        nodes[0].text = text
        for n in nodes[1:]:
            n.text = ""
    doc._element.body.insert(-1, el)
    return doc.paragraphs[-1]


def _make_good_bytes():
    tpl = Document(str(TPL))
    doc = Document(str(TPL))
    _clear_body(doc)
    for idx, text in [
        (0, "TARİFNAME"), (1, None), (2, DRAFT["title"]), (3, None), (4, tpl.paragraphs[4].text), (5, None),
        (6, "TEKNİK ALAN"), (7, None), (8, "Buluş, test sistemleri ile ilgilidir."), (9, None),
        (8, "Buluş, özellikle test haberleşmesi alanına yöneliktir."), (9, None),
        (10, "ÖNCEKİ TEKNİK"), (11, None), (12, "Mevcut sistemlerde teknik sorun bulunmaktadır."), (13, None),
        (16, "Sonuçta yukarıda bahsedilen ve mevcut teknik ışığında çözülemeyen sorunlar, ilgili teknik alanda bir yenilik yapmayı zorunlu kılmıştır."),
        (17, "BULUŞUN KISA AÇIKLAMASI"), (18, None), (19, "Buluş, test sistemi ile ilgilidir."), (20, None),
        (21, "Buluşun ana amacı, test işlevini sağlamaktır."), (22, None),
        (29, "Yukarıdaki amaçları gerçekleştirmek üzere buluş, bir test sistemi olup, özelliği;"),
        (30, "test modülü,"), (37, "içermesidir."), (38, None),
        (39, "Buluşun yapılanması ve ek elemanlarla birlikte avantajlarının en iyi şekilde anlaşılabilmesi için aşağıda açıklaması yapılan şekiller ile birlikte değerlendirilmesi gerekmektedir."), (40, None),
        (41, "ŞEKİLLERİN KISA AÇIKLAMASI"), (42, None), (43, DRAFT["figure_descriptions"][0]), (44, None),
        (45, tpl.paragraphs[45].text), (46, None), (47, "REFERANS NUMARALARI"), (48, None),
        (49, "1. Test modülü"), (56, None), (57, "BULUŞUN DETAYLI AÇIKLAMASI"), (58, None),
        (59, "Bu detaylı açıklamada, buluş konusu olan test sistemi sadece konunun daha iyi anlaşılmasına yönelik hiçbir sınırlayıcı etki oluşturmayacak örneklerle açıklanmaktadır."), (60, None),
        (61, "Test modülü (1), test işlevini gerçekleştiren teknik modüldür."),
        (75, None), (76, None), (77, "İSTEMLER"), (78, None),
        (79, tpl.paragraphs[79].text), (80, None), (81, tpl.paragraphs[81].text), (82, None), (83, tpl.paragraphs[83].text), (84, None),
        (85, "Bir test sistemi olup, özelliği;"), (86, "test modülü (1)"), (93, "içermesidir."), (98, None),
        (99, "ÖZET"), (100, None), (101, DRAFT["title"]), (102, None), (103, DRAFT["abstract"]), (104, None), (105, None),
    ]:
        p = _copy(doc, tpl, idx, text)
        if idx in (77, 99):
            p.paragraph_format.page_break_before = True
    out = io.BytesIO(); doc.save(out); return out.getvalue()


def test_good_template_structure_passes():
    validate_full_tarifname_template_fidelity(_make_good_bytes(), TPL, DRAFT, "Türkçe")


def test_footer_page_number_is_rejected():
    doc = Document(io.BytesIO(_make_good_bytes()))
    p = doc.sections[0].footer.paragraphs[0]
    run = p.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    out = io.BytesIO(); doc.save(out)
    with pytest.raises(ValueError, match="footer"):
        validate_full_tarifname_template_fidelity(out.getvalue(), TPL, DRAFT, "Türkçe")


def test_missing_figure_heading_gap_is_rejected():
    doc = Document(io.BytesIO(_make_good_bytes()))
    texts = [p.text.strip() for p in doc.paragraphs]
    i = texts.index("ŞEKİLLERİN KISA AÇIKLAMASI")
    doc._element.body.remove(doc.paragraphs[i-1]._p)
    out = io.BytesIO(); doc.save(out)
    with pytest.raises(ValueError, match="ŞEKİLLERİN KISA AÇIKLAMASI öncesinde"):
        validate_full_tarifname_template_fidelity(out.getvalue(), TPL, DRAFT, "Türkçe")


def test_short_description_visual_gap_is_rejected():
    doc = Document(io.BytesIO(_make_good_bytes()))
    texts = [p.text.strip() for p in doc.paragraphs]
    i = texts.index("BULUŞUN KISA AÇIKLAMASI")
    doc.paragraphs[i-1].paragraph_format.space_after = 0
    out = io.BytesIO(); doc.save(out)
    with pytest.raises(ValueError, match="BULUŞUN KISA AÇIKLAMASI öncesindeki görsel boşluk"):
        validate_full_tarifname_template_fidelity(out.getvalue(), TPL, DRAFT, "Türkçe")
