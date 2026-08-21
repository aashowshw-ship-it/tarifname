import io
import zipfile
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pytest

from rules import APP_VERSION, RULESET_VERSION, TARIFNAME_DUZENLEME_RULES
from tarifname_update import (
    derive_markup_output_name,
    extract_docx_review_context,
    prepare_review_baseline_docx,
    validate_update_plan,
    build_updated_spec_docx,
    validate_update_result,
    document_text,
)


def _base_docx() -> bytes:
    d = Document()
    d.add_paragraph('SPECIFICATION')
    d.add_paragraph('The codec is a transparent gzip codec used in the response path.')
    d.add_paragraph('CLAIMS')
    d.add_paragraph('A system comprising the transparent gzip codec.')
    b = io.BytesIO(); d.save(b); return b.getvalue()


def _plan(basis_source='existing_spec'):
    quote = 'transparent gzip codec' if basis_source == 'existing_spec' else 'Please generalize gzip to content coding.'
    return {
        'coverage_complete': True,
        'requests': [{
            'id':'R1','customer_request':'Generalize gzip wording','category':'technical_text',
            'decision':'apply','reason':'broader supported terminology','answer_for_customer':'Updated.'
        }],
        'operations': [{
            'request_id':'R1','type':'replace_text','section':'DETAILED DESCRIPTION',
            'locator_text':'The codec is a transparent gzip codec used in the response path.',
            'anchor_text':'','old_text':'gzip','new_text':'content-coding',
            'basis_source':basis_source,'basis_quote':quote,'reason':'minimum terminology revision'
        }],
        'comments': [], 'figure_actions': [], 'blocking_clarifications': [], 'open_procedural_items': []
    }


def test_versions_and_rules():
    assert APP_VERSION == 'v5.4.32'
    assert RULESET_VERSION == '2026-08-21.v22'
    assert 'EN AZ DEĞİŞİKLİK' in TARIFNAME_DUZENLEME_RULES
    assert 'Track Changes' in TARIFNAME_DUZENLEME_RULES
    assert 'Müşteriye gönderilecek mail zorunlu çıktıdır' in TARIFNAME_DUZENLEME_RULES


def test_markup_filename_removes_browser_duplicate_suffix():
    assert derive_markup_output_name('Description_181569(1).docx') == 'Description_181569_markup.docx'
    assert derive_markup_output_name('Description_181569.docx') == 'Description_181569_markup.docx'


def test_minimum_markup_and_format_gate():
    src = _base_docx(); plan = _plan()
    validate_update_plan(plan, src, 'Please generalize gzip to content coding.', 'Henüz başvuru yapılmadı')
    markup = build_updated_spec_docx(src, plan, track_changes=True, add_comments=False)
    clean = build_updated_spec_docx(src, plan, track_changes=False, add_comments=False)
    validate_update_result(src, markup, clean, plan)
    assert 'content-coding codec used in the response path' in document_text(clean)
    with zipfile.ZipFile(io.BytesIO(markup)) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    assert '<w:delText xml:space="preserve">gzip</w:delText>' in xml
    assert '>content-coding</w:t>' in xml
    assert 'transparent gzip codec used in the response path' not in document_text(clean)


def test_existing_customer_track_changes_are_read_then_rejected_for_baseline():
    d = Document(); p = d.add_paragraph(); r = p.add_run('Original ')
    dele = OxmlElement('w:del'); dele.set(qn('w:id'),'1'); dr=OxmlElement('w:r'); dt=OxmlElement('w:delText'); dt.text='word'; dr.append(dt); dele.append(dr); p._p.append(dele)
    ins = OxmlElement('w:ins'); ins.set(qn('w:id'),'2'); ir=OxmlElement('w:r'); it=OxmlElement('w:t'); it.text='replacement'; ir.append(it); ins.append(ir); p._p.append(ins)
    d.add_comment(runs=[r], text='Please review this change', author='Client', initials='C')
    b=io.BytesIO(); d.save(b); raw=b.getvalue()
    ctx=extract_docx_review_context(raw)
    assert 'TRACK DELETE: word' in ctx and 'TRACK INSERT: replacement' in ctx and 'WORD COMMENT:' in ctx
    baseline=prepare_review_baseline_docx(raw)
    assert 'Original word' in document_text(baseline)
    assert 'replacement' not in document_text(baseline)
    with zipfile.ZipFile(io.BytesIO(baseline)) as z:
        assert 'word/comments.xml' not in z.namelist()


def test_post_filing_customer_only_new_matter_is_blocked():
    src=_base_docx(); plan=_plan('customer_request')
    with pytest.raises(ValueError, match='new-matter'):
        validate_update_plan(plan, src, 'Please generalize gzip to content coding.', 'Başvuru yapıldı')


def test_comments_are_written_for_explanation():
    src=_base_docx(); plan=_plan(); plan['comments']=[{
        'request_id':'R1','anchor_text':'The codec is a transparent','text':'This wording was generalized while preserving the existing embodiment.'
    }]
    markup=build_updated_spec_docx(src, plan, track_changes=True, add_comments=True)
    clean=build_updated_spec_docx(src, plan, track_changes=False, add_comments=False)
    validate_update_result(src, markup, clean, plan)
    with zipfile.ZipFile(io.BytesIO(markup)) as z:
        assert 'word/comments.xml' in z.namelist()
