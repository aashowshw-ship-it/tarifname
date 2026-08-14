from __future__ import annotations

import io
import re
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


def _clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _fmt(p, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, center=False):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else align
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(11)
        if bold:
            r.bold = True
    return p


def _add_text(doc, text: str, *, bold=False, center=False):
    p = doc.add_paragraph()
    p.add_run(text)
    return _fmt(p, bold=bold, center=center)


def _add_heading(doc, text: str):
    return _add_text(doc, text, bold=True)


def _add_bullet(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.905)
    p.paragraph_format.first_line_indent = Cm(-0.635)
    p.paragraph_format.line_spacing = 1.5
    p.add_run("•\t" + text)
    return _fmt(p)


def _add_nested_dash(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(3.75)
    p.paragraph_format.first_line_indent = Cm(-0.65)
    p.paragraph_format.line_spacing = 1.5
    p.add_run("-\t" + text)
    return _fmt(p)


def _copy_template_paragraph(doc: Document, template: Document, index: int):
    doc._element.body.insert(-1, deepcopy(template.paragraphs[index]._p))


def _copy_template_paragraph_with_text(doc: Document, template: Document, index: int, text: str):
    p_el = deepcopy(template.paragraphs[index]._p)
    text_nodes = list(p_el.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
    if text_nodes:
        text_nodes[0].text = text
        for node in text_nodes[1:]:
            node.text = ''
    doc._element.body.insert(-1, p_el)


def _numbered_claim(doc: Document, number: int, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{number}. {text}")
    r.font.name = "Arial"
    r.font.size = Pt(11)


def _add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType", "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType", "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def _strip_known_element_reference_marks(text: str, element_numbers: list[str]) -> str:
    result = str(text or "")
    for number in sorted({str(x or "").strip() for x in element_numbers if str(x or "").strip()}, key=len, reverse=True):
        result = re.sub(rf"\s*\(\s*{re.escape(number)}\s*\)", "", result)
    return re.sub(r"\s{2,}", " ", result).strip()


def build_docx(draft: dict[str, Any], template_path: str | Path) -> bytes:
    template = Document(str(template_path))
    doc = Document(str(template_path))
    _clear_body(doc)

    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        _add_page_number(section)

    _add_text(doc, "TARİFNAME", bold=True, center=True)
    doc.add_paragraph()
    _add_text(doc, draft["title"], bold=True, center=True)
    doc.add_paragraph()
    _copy_template_paragraph(doc, template, 4)
    doc.add_paragraph()

    _add_heading(doc, "TEKNİK ALAN")
    doc.add_paragraph()
    technical_field_parts = [
        x.strip() for x in re.split(r"\n\s*\n", str(draft.get("technical_field", "") or "")) if x.strip()
    ]
    for idx, paragraph in enumerate(technical_field_parts):
        _add_text(doc, paragraph)
        if idx < len(technical_field_parts) - 1:
            doc.add_paragraph()
    doc.add_paragraph()

    _add_heading(doc, "ÖNCEKİ TEKNİK")
    doc.add_paragraph()
    for text in draft.get("prior_art_paragraphs") or []:
        _add_text(doc, text)
        doc.add_paragraph()
    _add_text(doc, "Sonuçta yukarıda bahsedilen ve mevcut teknik ışığında çözülemeyen sorunlar, ilgili teknik alanda bir yenilik yapmayı zorunlu kılmıştır.")

    _add_heading(doc, "BULUŞUN KISA AÇIKLAMASI")
    doc.add_paragraph()
    _add_text(doc, draft.get("short_description_intro", ""))
    for i, obj in enumerate(draft.get("objectives") or []):
        prefix = "Buluşun ana amacı, " if i == 0 else "Buluşun diğer bir amacı, "
        _add_text(doc, prefix + obj[:1].lower() + obj[1:] if obj else prefix)
        doc.add_paragraph()

    if draft.get("unumbered_system_definition"):
        _add_text(doc, draft["unumbered_system_definition"])
    for el in draft.get("unumbered_system_elements") or []:
        _add_bullet(doc, el)
    if draft.get("unumbered_system_elements"):
        _add_text(doc, "içermesidir.")
    doc.add_paragraph()
    _add_text(doc, "Buluşun yapılanması ve ek elemanlarla birlikte avantajlarının en iyi şekilde anlaşılabilmesi için aşağıda açıklaması yapılan şekiller ile birlikte değerlendirilmesi gerekmektedir.")

    _add_heading(doc, "ŞEKİLLERİN KISA AÇIKLAMASI")
    doc.add_paragraph()
    figures = draft.get("figure_descriptions") or ["Şekil 1, buluşa konu sistemin temsili bir gösterimidir."]
    for x in figures:
        _add_text(doc, x)
    _add_text(doc, "Çizimlerin mutlaka ölçeklendirilmesi gerekmemektedir ve buluşu anlamak için gerekli olmayan detaylar ihmal edilmiş olabilmektedir. Bundan başka, en azından büyük ölçüde özdeş olan veya en azından büyük ölçüde özdeş işlevleri olan elemanlar, aynı numara ile gösterilmektedir.")

    _add_heading(doc, "REFERANS NUMARALARI")
    doc.add_paragraph()
    for e in draft.get("elements") or []:
        _add_text(doc, f"{e['number']}. {e['name']}")
    if draft.get("elements") and draft.get("method_steps"):
        doc.add_paragraph()
    element_numbers = [str(x.get("number", "") or "").strip() for x in (draft.get("elements") or [])]
    for s in draft.get("method_steps") or []:
        number = str(s.get("number", "") or "").strip()
        text = str(s.get("text", "") or "").strip()
        if number:
            text = re.sub(rf"\s*\(\s*{re.escape(number)}\s*\)\s*$", "", text).strip()
        text = _strip_known_element_reference_marks(text, element_numbers).rstrip(".,;:")
        _add_text(doc, f"{number}. {text}" if number else text)

    _add_heading(doc, "BULUŞUN DETAYLI AÇIKLAMASI")
    doc.add_paragraph()
    _add_text(doc, f"Bu detaylı açıklamada, buluş konusu olan {draft['title'].lower()} sadece konunun daha iyi anlaşılmasına yönelik hiçbir sınırlayıcı etki oluşturmayacak örneklerle açıklanmaktadır.")
    for para in draft.get("detailed_paragraphs") or []:
        doc.add_paragraph()
        _add_text(doc, para)

    _copy_template_paragraph(doc, template, 75)
    _copy_template_paragraph(doc, template, 76)
    _copy_template_paragraph_with_text(doc, template, 77, "İSTEMLER")
    claims_p = doc.paragraphs[-1]
    claims_p.paragraph_format.page_break_before = True
    claims_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _copy_template_paragraph(doc, template, 78)
    for idx in (79, 81, 83):
        _copy_template_paragraph(doc, template, idx)
        doc.add_paragraph()

    claim_no = 1
    sc = draft.get("system_claim") or {}
    preamble = sc.get("preamble", "")
    _numbered_claim(doc, claim_no, preamble + " olup, özelliği;")
    for el in sc.get("elements") or []:
        if isinstance(el, dict):
            _add_bullet(doc, str(el.get("lead", "") or ""))
            for sub in el.get("subelements") or []:
                _add_nested_dash(doc, str(sub))
        else:
            _add_bullet(doc, str(el))
    _copy_template_paragraph_with_text(doc, template, 93, sc.get("closing", "içermesidir."))
    claim_no += 1

    for dep in draft.get("dependent_system_claims") or []:
        _numbered_claim(doc, claim_no, dep)
        claim_no += 1

    mc = draft.get("method_claim")
    if mc:
        _numbered_claim(doc, claim_no, mc.get("preamble", "") + " olup, özelliği;")
        for step in mc.get("steps") or []:
            step_text = str(step).rstrip().rstrip(".;,") + ","
            _add_bullet(doc, step_text)
        _copy_template_paragraph_with_text(doc, template, 93, mc.get("closing", "işlem adımlarını içermesidir."))
        claim_no += 1
        for dep in draft.get("dependent_method_claims") or []:
            _numbered_claim(doc, claim_no, dep)
            claim_no += 1

    _copy_template_paragraph_with_text(doc, template, 99, "ÖZET")
    summary_p = doc.paragraphs[-1]
    summary_p.paragraph_format.page_break_before = True
    summary_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _copy_template_paragraph(doc, template, 100)
    _copy_template_paragraph_with_text(doc, template, 101, draft["title"])
    _copy_template_paragraph(doc, template, 102)
    _copy_template_paragraph_with_text(doc, template, 103, draft.get("abstract", ""))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
